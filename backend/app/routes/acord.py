import base64
import asyncio
import io
import logging
import os
import re
import uuid
from pathlib import Path
from typing import Any, List, Optional

from fastapi import APIRouter, BackgroundTasks, File, Header, HTTPException, Query, UploadFile

from Models.acord_form_understanding import AcordFormSummary, parse_acord_form
from Models.acord_form_understanding.azure_document_intelligence import try_azure_document_intelligence_to_uir
from Models.acord_form_understanding.bytescout_cli import try_extract_pdf_to_uir
from Models.acord_form_understanding.extraction_pipeline import (
    build_summary_from_uir,
    build_uir_from_pdf_text,
    openai_compat_extract_structured,
)
from Models.acord_form_understanding.ocr_runtime import configure_tesseract_runtime
from Models.acord_form_understanding.pdf_acroform import extract_acroform_fields, pdf_has_acroform_widgets
from Models.acord_form_understanding.pdf_ocr import ocr_pdf_pages
from Models.acord_form_understanding.uir import KeyValue, Table, TextBlock, UnifiedIntermediateRepresentation
from app.core.supabase import postgrest_get, postgrest_insert, postgrest_patch, verify_user
from app.services.webhook_engine import (
    WEBHOOK_EVENT_INFERENCE_COMPLETE,
    resolve_tenant_id_for_user,
    try_emit_webhook_event,
)
from app.services.runpod_orchestrator import ensure_runpod_ml_ready
from app.services.vectorstore_ingestion import ingest_text_into_vectorstore
from app.schemas.acord_workflow import (
    AcordAdminReviewRequest,
    AcordBatchReviewRequest,
    AcordExtractJobStatusResponse,
    AcordExtractStartResponse,
    AcordExtractResponse,
    AcordSubmitRequest,
    PreviewSftTrainingRecordBody,
    ReExtractRequest,
)
from app.services.acord_training import create_job_row, spawn_job_runner
from app.services.nl_summary import generate_nl_summary as _generate_nl_summary
# NOTE: importing fine_tuning / torch at module import time can make API startup
# extremely slow. Import hardened inference lazily only when needed.

from urllib.parse import quote

try:
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None  # type: ignore[assignment]

try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover - optional dependency
    fitz = None  # type: ignore[assignment]

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover - optional dependency
    docx = None  # type: ignore[assignment]

try:
    import openpyxl
except ImportError:  # pragma: no cover - optional dependency
    openpyxl = None  # type: ignore[assignment]

try:
    from PIL import Image
except ImportError:  # pragma: no cover - optional dependency
    Image = None  # type: ignore[assignment]

try:
    import pytesseract
except ImportError:  # pragma: no cover - optional dependency
    pytesseract = None  # type: ignore[assignment]

try:
    import pdfplumber  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - optional dependency
    pdfplumber = None  # type: ignore[assignment]


router = APIRouter(prefix="/api/acord", tags=["acord"])
logger = logging.getLogger("fideon.acord")

ACORD_POD_ID = os.getenv("ACORD_POD_ID", "acord_form_understanding")
_POD_COLLECTION_CACHE: dict[str, str] = {}
_EXTRACT_JOBS: dict[str, dict[str, Any]] = {}
_EXTRACT_JOBS_LOCK = asyncio.Lock()
_EXTRACT_JOB_TABLE = "acord_extract_jobs"

configure_tesseract_runtime()


def _auth_debug_meta(authorization: str | None) -> dict[str, Any]:
    has_auth = bool((authorization or "").strip())
    is_bearer = bool((authorization or "").lower().startswith("bearer "))
    return {"has_auth": has_auth, "is_bearer": is_bearer}


async def _ensure_runpod_ready_for_acord() -> None:
    """
    Ensure RunPod pod is resumed and ML HTTP is ready before ACORD extraction.
    This restores the expected flow for `/api/acord/extract/start` where clicking
    Parse ACORD Form can cold-start the pod automatically.
    """
    # If pod config is absent, keep existing behavior (local-only extraction paths).
    if not (os.getenv("RUNPOD_POD_ID") or os.getenv("POD_ID")):
        return
    try:
        await ensure_runpod_ml_ready()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"RunPod readiness failed: {exc}") from exc

def _confidence_threshold() -> float:
    """Read threshold from env so it can be tuned without a deploy."""
    try:
        return float(os.getenv("ACORD_CONFIDENCE_THRESHOLD", "0.85"))
    except (TypeError, ValueError):
        return 0.85


def _clamp01(v: float) -> float:
    return max(0.0, min(1.0, float(v)))


def _hardened_form_type_hint(form_type_hint: Optional[str], extracted_form_type: Optional[str]) -> str:
    raw = (form_type_hint or extracted_form_type or "").strip().lower()
    if raw in {"125", "acord_125", "acord125"}:
        return "acord_125"
    if raw in {"25", "acord_25", "acord25"}:
        return "acord_25"
    return "acord_125"


def _apply_hardened_overlay(extracted: dict, hardened: dict) -> dict:
    data = hardened.get("data") or {}
    out = dict(extracted or {})
    producer = dict(out.get("producer") or {})
    policy_info = dict(out.get("policy_info") or {})
    carrier = dict(policy_info.get("carrier") or {})

    if data.get("agency_name") is not None:
        producer["name"] = data.get("agency_name")
    if data.get("contact_name") is not None:
        producer["contact_name"] = data.get("contact_name")
    if data.get("email") is not None:
        producer["email"] = data.get("email")
    if data.get("phone") is not None:
        producer["phone"] = data.get("phone")

    if data.get("policy_number") is not None:
        policy_info["policy_number"] = data.get("policy_number")
    if data.get("carrier") is not None:
        carrier["name"] = data.get("carrier")
    if carrier:
        policy_info["carrier"] = carrier
    if policy_info:
        out["policy_info"] = policy_info
    if producer:
        out["producer"] = producer

    lob = data.get("lines_of_business")
    if lob:
        if isinstance(lob, str):
            out["lines_of_business_indicated"] = [x.strip() for x in lob.split(",") if x.strip()]
        elif isinstance(lob, list):
            out["lines_of_business_indicated"] = [str(x).strip() for x in lob if str(x).strip()]

    out["hardened_inference"] = {
        "status": hardened.get("status"),
        "final_score": hardened.get("final_score"),
        "confidence": hardened.get("confidence"),
        "trust_score": hardened.get("trust_score"),
        "field_confidence": hardened.get("field_confidence"),
        "review_fields": hardened.get("review_fields"),
    }
    return out


def _evaluate_confidence_and_feedback(
    *,
    base_confidence: float,
    run_status: str,
    original_json: dict,
    edited_json: dict,
    feedback_rows: list[dict],
) -> dict:
    """
    Lightweight confidence calibration for UI/ops visibility.
    This does NOT overwrite stored extraction confidence; it provides
    an additional derived signal that combines model confidence + human feedback.
    """
    status = (run_status or "").strip().lower()
    latest_thumbs: Optional[bool] = None
    latest_actor: Optional[str] = None
    corrections_count = 0
    for fb in feedback_rows or []:
        cj = fb.get("corrected_json")
        if cj is not None:
            corrections_count += 1
        tv = fb.get("thumbs_up")
        if latest_thumbs is None and isinstance(tv, bool):
            latest_thumbs = tv
            latest_actor = str(fb.get("actor_role") or "")

    has_user_or_admin_edits = edited_json != original_json
    adjustment = 0.0
    reasons: list[str] = []
    if latest_thumbs is True:
        adjustment += 0.05
        reasons.append("latest_feedback_positive")
    elif latest_thumbs is False:
        adjustment -= 0.08
        reasons.append("latest_feedback_negative")

    if has_user_or_admin_edits:
        adjustment -= 0.10
        reasons.append("json_was_manually_corrected")

    if status == "approved":
        adjustment += 0.04
        reasons.append("admin_approved")
    elif status in {"rejected", "needs_admin_review"}:
        adjustment -= 0.06
        reasons.append("requires_or_failed_admin_review")

    calibrated = _clamp01(float(base_confidence or 0.0) + adjustment)
    return {
        "base_confidence": _clamp01(base_confidence),
        "calibrated_confidence": calibrated,
        "adjustment": round(adjustment, 4),
        "reasons": reasons,
        "feedback_signals": {
            "total_feedback_entries": len(feedback_rows or []),
            "corrections_count": corrections_count,
            "has_manual_edits": has_user_or_admin_edits,
            "latest_thumbs_up": latest_thumbs,
            "latest_feedback_actor_role": latest_actor,
        },
    }


def _eval_weighted(rows_by_set: dict[str, dict], metric_key: str, sets: list[str]) -> Optional[float]:
    total_n = 0.0
    weighted = 0.0
    for s in sets:
        row = rows_by_set.get(s) or {}
        val = row.get(metric_key)
        if val is None:
            continue
        m = row.get("metrics_json") or {}
        n = float((m.get("n") if isinstance(m, dict) else 0) or 0)
        if n <= 0:
            continue
        total_n += n
        weighted += float(val) * n
    if total_n <= 0:
        return None
    return weighted / total_n


def _quality_gate_snapshot_from_eval_rows(eval_rows: list[dict]) -> dict:
    """
    Compute a gate pass/fail snapshot from persisted acord_eval_results rows.
    """
    by_set: dict[str, dict] = {str(r.get("eval_set") or ""): r for r in (eval_rows or [])}
    # DB uses "oos" key.
    seen_key = "seen"
    para_key = "paraphrased"
    oos_key = "oos"

    json_valid = _eval_weighted(by_set, "exact_match", [seen_key, para_key])  # backward-compatible fallback
    json_exact = _eval_weighted(by_set, "exact_match", [seen_key, para_key])
    json_recall = None
    json_extra = None
    oos_hall = _eval_weighted(by_set, "hallucination_rate", [oos_key])

    # Prefer richer JSON metrics from metrics_json when available.
    def _weighted_from_metrics_json(metric_name: str) -> Optional[float]:
        total_n = 0.0
        weighted = 0.0
        for s in (seen_key, para_key):
            row = by_set.get(s) or {}
            m = row.get("metrics_json") or {}
            if not isinstance(m, dict):
                continue
            val = m.get(metric_name)
            n = float(m.get("n") or 0)
            if val is None or n <= 0:
                continue
            total_n += n
            weighted += float(val) * n
        if total_n <= 0:
            return None
        return weighted / total_n

    json_valid_m = _weighted_from_metrics_json("json_valid_rate")
    json_exact_m = _weighted_from_metrics_json("json_exact_match")
    json_recall_m = _weighted_from_metrics_json("json_field_recall")
    json_extra_m = _weighted_from_metrics_json("json_extra_field_rate")
    if json_valid_m is not None:
        json_valid = json_valid_m
    if json_exact_m is not None:
        json_exact = json_exact_m
    if json_recall_m is not None:
        json_recall = json_recall_m
    if json_extra_m is not None:
        json_extra = json_extra_m

    min_json_valid = float(os.getenv("FT_ACORD_QG_MIN_JSON_VALID_RATE", os.getenv("FT_QG_MIN_JSON_VALID_RATE", "0.90")))
    min_json_exact = float(os.getenv("FT_ACORD_QG_MIN_JSON_EXACT_MATCH", os.getenv("FT_QG_MIN_JSON_EXACT_MATCH", "0.70")))
    min_field_recall = float(os.getenv("FT_ACORD_QG_MIN_JSON_FIELD_RECALL", os.getenv("FT_QG_MIN_JSON_FIELD_RECALL", "0.80")))
    max_extra_field = float(
        os.getenv("FT_ACORD_QG_MAX_JSON_EXTRA_FIELD_RATE", os.getenv("FT_QG_MAX_JSON_EXTRA_FIELD_RATE", "0.10"))
    )
    max_oos_halluc = float(
        os.getenv("FT_ACORD_QG_MAX_OOS_HALLUCINATION_RATE", os.getenv("FT_QG_MAX_OOS_HALLUCINATION_RATE", "0.25"))
    )
    require_oos = (os.getenv("FT_ACORD_QG_REQUIRE_OOS", "false").strip().lower() in {"1", "true", "yes", "on"})

    checks: list[dict] = []

    def _check(name: str, val: Optional[float], op: str, th: float) -> bool:
        if val is None:
            checks.append({"metric": name, "value": None, "threshold": th, "operator": op, "ok": False})
            return False
        ok = (val >= th) if op == ">=" else (val <= th)
        checks.append({"metric": name, "value": val, "threshold": th, "operator": op, "ok": ok})
        return ok

    ok_all = True
    ok_all &= _check("json_valid_rate", json_valid, ">=", min_json_valid)
    ok_all &= _check("json_exact_match", json_exact, ">=", min_json_exact)
    ok_all &= _check("json_field_recall", json_recall, ">=", min_field_recall)
    ok_all &= _check("json_extra_field_rate", json_extra, "<=", max_extra_field)
    if require_oos or oos_hall is not None:
        ok_all &= _check("oos_hallucination_rate", oos_hall, "<=", max_oos_halluc)

    return {
        "pass": bool(ok_all),
        "checks": checks,
    }


async def _latest_corrected_json_for_run(run_id: str) -> Optional[dict]:
    """
    Return most recent non-null corrected_json from feedback history for a run.
    """
    rows = await postgrest_get(
        "acord_extraction_feedback",
        f"select=corrected_json,created_at&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=50",
    )
    for row in rows or []:
        cj = row.get("corrected_json")
        if isinstance(cj, dict):
            return cj
    return None


async def _pod_collection_name(pod_id: str) -> str:
    if pod_id in _POD_COLLECTION_CACHE:
        return _POD_COLLECTION_CACHE[pod_id]

    agents = await postgrest_get(
        "agent_catalog",
        f"id=eq.{quote(pod_id, safe='')}&is_active=eq.true&select=domain_id,rag_collection_override",
    )
    if not agents:
        raise RuntimeError(f"pod agent not found: {pod_id}")
    agent = agents[0]

    domain_id = agent.get("domain_id")
    if not domain_id:
        raise RuntimeError(f"pod agent missing domain_id: {pod_id}")

    domains = await postgrest_get(
        "domain_catalog",
        f"id=eq.{quote(str(domain_id), safe='')}&is_active=eq.true&select=id,rag_collection",
    )
    if not domains:
        raise RuntimeError(f"domain not found: {domain_id} for pod_id={pod_id}")
    domain = domains[0]

    collection = agent.get("rag_collection_override") or domain.get("rag_collection") or f"{domain['id']}_index"
    _POD_COLLECTION_CACHE[pod_id] = collection
    return collection


async def _ingest_acord_into_vectorstore(doc_id: str, text: str) -> None:
    """
    Persist ACORD text into the correct pod pgvector collection.
    """
    if not text.strip():
        logger.info("ACORD[ingest] doc_id=%s produced empty text; skipping.", doc_id)
        return

    collection_name = await _pod_collection_name(ACORD_POD_ID)
    wrote = ingest_text_into_vectorstore(
        collection_name=collection_name,
        doc_id=doc_id,
        text=text,
        pod_id=ACORD_POD_ID,
        source="acord-upload",
    )
    logger.info("ACORD[ingest] doc_id=%s wrote %d chunks into '%s'.", doc_id, wrote, collection_name)


async def _read_text_from_file(file: UploadFile) -> str:
    contents = await file.read()
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    logger.info(
        "ACORD[upload] filename=%s size=%d content_type=%s",
        filename or "<unknown>",
        len(contents),
        content_type or "<unknown>",
    )

    # Simple text handling
    if filename.endswith(".txt") or "text" in content_type:
        try:
            return contents.decode("utf-8", errors="ignore")
        except UnicodeDecodeError as exc:  # pragma: no cover - unlikely
            raise HTTPException(status_code=400, detail=str(exc)) from exc

    # CSV handling (simple UTF-8 decode)
    if filename.endswith(".csv") or content_type in {"text/csv", "application/csv"}:
        return contents.decode("utf-8", errors="ignore")

    # Basic PDF handling via PyPDF2
    if filename.endswith(".pdf") or content_type in {"application/pdf", "pdf"}:
        if PdfReader is None:
            raise HTTPException(
                status_code=500,
                detail="PDF support not installed on server (PyPDF2 missing).",
            )
        try:
            from io import BytesIO

            reader = PdfReader(BytesIO(contents))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()

            # Robust fallback: scanned PDFs often have no embedded text.
            # If text is too small, attempt OCR via PyMuPDF + Tesseract.
            if len(text) >= 200:
                return text

            if fitz is None:
                logger.warning("ACORD[pdf] low extracted text and PyMuPDF missing; skipping OCR fallback.")
                return text
            if Image is None or pytesseract is None:
                logger.warning("ACORD[pdf] low extracted text and OCR deps missing; skipping OCR fallback.")
                return text

            logger.info("ACORD[pdf] low extracted text (%d chars). Running OCR fallback.", len(text))
            doc = fitz.open(stream=contents, filetype="pdf")
            ocr_pages: list[str] = []
            for i in range(min(len(doc), 20)):  # safety cap
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                img = Image.open(BytesIO(img_bytes))
                ocr_text = pytesseract.image_to_string(img) or ""
                if ocr_text.strip():
                    ocr_pages.append(ocr_text)
            ocr = "\n".join(ocr_pages).strip()
            return ocr or text
        except Exception as exc:  # pragma: no cover - defensive
            raise HTTPException(status_code=400, detail=f"Failed to read PDF: {exc}") from exc

    # Word (.docx) handling via python-docx
    if filename.endswith(".docx") or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }:
        if docx is None:
            raise HTTPException(
                status_code=500,
                detail="DOCX support not installed on server (python-docx missing).",
            )
        try:
            from io import BytesIO

            doc = docx.Document(BytesIO(contents))
            paras = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
            return "\n".join(paras)
        except Exception as exc:  # pragma: no cover - defensive
            raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {exc}") from exc

    # Excel (.xlsx) handling via openpyxl
    if filename.endswith(".xlsx") or content_type in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    }:
        if openpyxl is None:
            raise HTTPException(
                status_code=500,
                detail="XLSX support not installed on server (openpyxl missing).",
            )
        try:
            from io import BytesIO

            wb = openpyxl.load_workbook(BytesIO(contents), data_only=True)
            lines: List[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"[SHEET] {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    # Convert row to a tab-separated string (skip fully empty rows)
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        lines.append("\t".join(cells).strip())
            return "\n".join(lines).strip()
        except Exception as exc:  # pragma: no cover - defensive
            raise HTTPException(status_code=400, detail=f"Failed to read XLSX: {exc}") from exc

    # Images (.png/.jpg/.jpeg) via OCR (optional: requires pytesseract + Tesseract installed)
    if filename.endswith((".png", ".jpg", ".jpeg", ".webp")) or content_type.startswith("image/"):
        if Image is None:
            raise HTTPException(
                status_code=500,
                detail="Image support not installed on server (pillow missing).",
            )
        if pytesseract is None:
            raise HTTPException(
                status_code=500,
                detail="OCR support not installed on server (pytesseract missing).",
            )
        try:
            from io import BytesIO

            img = Image.open(BytesIO(contents))
            text = pytesseract.image_to_string(img)
            if not text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="OCR produced empty text. Image may be too low quality or Tesseract not configured.",
                )
            return text
        except HTTPException:
            raise
        except Exception as exc:  # pragma: no cover - defensive
            raise HTTPException(status_code=400, detail=f"Failed to OCR image: {exc}") from exc

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Please upload PDF, TXT, CSV, DOCX, XLSX, or an image (PNG/JPG).",
    )


def _uir_has_content(uir: UnifiedIntermediateRepresentation, min_chars: int = 200) -> bool:
    """Return True if the UIR has at least min_chars of usable text."""
    total = sum(len(tb.text.strip()) for tb in uir.text_blocks)
    total += sum(len(kv.key) + len(kv.value) for kv in uir.key_values)

    # Many PDFs yield a lot of *template* text (labels, headings) with very few/no filled values.
    # When that happens, the LLM tends to do label-bleed (e.g., "QUOTE ISSUE POLICY RENEW").
    #
    # So we accept the engine only if we have (a) enough text volume and (b) evidence
    # of *real filled values* like dates, money, or policy-number-like tokens.
    all_text = " ".join(tb.text for tb in (uir.text_blocks or []))
    has_date = bool(re.search(r"\b\d{2}/\d{2}/\d{4}\b", all_text))
    has_money = bool(re.search(r"\$\s*[\d,]+", all_text))
    has_policy_token = bool(re.search(r"\b[A-Z]{1,5}[-\s]?\d[\d\-]{4,20}\b", all_text))
    has_any_filled_value = has_date or has_money or has_policy_token

    if total < min_chars:
        return False
    return has_any_filled_value


def _uir_non_empty(uir: UnifiedIntermediateRepresentation) -> bool:
    """True if UIR has any non-empty text, key_values, or tables."""
    if not uir:
        return False
    t = sum(len((tb.text or "").strip()) for tb in (uir.text_blocks or []))
    t += sum(len((kv.key or "")) + len((kv.value or "")) for kv in (uir.key_values or []))
    if t > 0:
        return True
    return bool(uir.tables)


def _merge_acord_uirs(
    parts: list[tuple[str, UnifiedIntermediateRepresentation]],
) -> UnifiedIntermediateRepresentation | None:
    """
    Staging / production-quality: combine all PDF engines into one UIR so the LLM
    sees layout KVs + tables (pdfplumber) plus alternate text layers (pymupdf/pypdf2/ocr).
    """
    if not parts:
        return None
    merged_blocks: list[TextBlock] = []
    seen_kv: set[tuple[str, str]] = set()
    merged_kvs: list[KeyValue] = []
    merged_tables: list[Table] = []
    engine_order: list[str] = []

    for engine_name, uir in parts:
        engine_order.append(engine_name)
        block_text = "\n".join((tb.text or "") for tb in (uir.text_blocks or [])).strip()
        if block_text:
            merged_blocks.append(
                TextBlock(
                    text=f"===== [{engine_name.upper()}] =====\n{block_text}",
                    page=1,
                    bbox=None,
                    source="pdf_text",
                ),
            )
        for kv in uir.key_values or []:
            sig = ((kv.key or "").strip().lower(), (kv.value or "").strip()[:400])
            if sig in seen_kv:
                continue
            seen_kv.add(sig)
            merged_kvs.append(kv)
        merged_tables.extend(uir.tables or [])

    if not merged_blocks and not merged_kvs and not merged_tables:
        return None
    return UnifiedIntermediateRepresentation(
        text_blocks=merged_blocks,
        tables=merged_tables,
        key_values=merged_kvs,
        layout={
            "extraction_engine": "merged",
            "extraction_engines": engine_order,
        },
    )


def _ocr_engine_from_parts(parts: list[tuple[str, UnifiedIntermediateRepresentation]]) -> Optional[str]:
    for eng_name, u in parts:
        if eng_name == "ocr" and u.layout:
            oe = u.layout.get("ocr_engine_used")
            if isinstance(oe, str) and oe.strip():
                return oe.strip()
    return None


def _patch_uir_pdf_classification(
    uir: UnifiedIntermediateRepresentation,
    has_widgets: bool,
    parts: list[tuple[str, UnifiedIntermediateRepresentation]],
) -> None:
    """Record fillable vs flattened + which OCR engine produced raster text (if any)."""
    lay = dict(uir.layout or {})
    lay["pdf_form_classification"] = "fillable" if has_widgets else "flattened"
    if not lay.get("ocr_engine_used"):
        oe = _ocr_engine_from_parts(parts)
        if oe:
            lay["ocr_engine_used"] = oe
    uir.layout = lay


def _attach_acroform_fields_if_enabled(uir: UnifiedIntermediateRepresentation, pdf_bytes: bytes) -> None:
    """
    Fillable ACORD PDFs: read embedded AcroForm field values (ground truth for key slots).
    No-op for flattened / print-only PDFs (e.g. ReportLab). Controlled by ACORD_USE_ACROFORM_FIRST.
    """
    if os.getenv("ACORD_USE_ACROFORM_FIRST", "true").strip().lower() not in {"1", "true", "yes", "on"}:
        return
    try:
        flat = extract_acroform_fields(pdf_bytes)
    except Exception as exc:
        logger.warning("ACORD[pdf] AcroForm extraction failed: %s", exc)
        return
    if not flat:
        return
    lay = dict(uir.layout or {})
    lay["acroform_fields"] = flat
    uir.layout = lay
    logger.info("ACORD[pdf] AcroForm widgets: %s fields (hybrid with VL/LLM)", len(flat))


def _attach_acord_vl_page_images(uir: UnifiedIntermediateRepresentation, pdf_bytes: bytes) -> None:
    """
    Rasterize PDF pages to JPEG (base64) for Qwen2.5-VL multimodal extraction.
    Controlled by ACORD_VL_ENABLED, ACORD_VL_MAX_PAGES, ACORD_VL_RENDER_DPI, ACORD_VL_JPEG_QUALITY.
    """
    if os.getenv("ACORD_VL_ENABLED", "true").strip().lower() not in {"1", "true", "yes", "on"}:
        return
    if fitz is None or Image is None:
        return
    try:
        max_pages = max(1, int(os.getenv("ACORD_VL_MAX_PAGES", "12")))
        dpi = float(os.getenv("ACORD_VL_RENDER_DPI", "144"))
        quality = int(os.getenv("ACORD_VL_JPEG_QUALITY", "85"))
    except ValueError:
        max_pages, dpi, quality = 12, 144.0, 85
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    images_b64: list[str] = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        n = min(len(doc), max_pages)
        for i in range(n):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            images_b64.append(base64.b64encode(buf.getvalue()).decode("ascii"))
    finally:
        doc.close()
    lay = dict(uir.layout or {})
    lay["acord_vl_page_images_base64"] = images_b64
    uir.layout = lay
    logger.info("ACORD[pdf] attached VL page rasters: pages=%s dpi=%s", len(images_b64), dpi)


async def _extract_from_pdf(
    contents: bytes,
    filename: str = "",
    form_type_hint: Optional[str] = None,
) -> AcordFormSummary:
    """
    Staging-style PDF extraction: run every available engine, merge into one UIR,
    then a single LLM structured pass (best field coverage). Falls back to the
    best single engine, then legacy heuristics.

    Engines (all collected when available):
      ByteScout → pdfplumber (layout + tables) → PyMuPDF → PyPDF2 → Azure DI (optional) → OCR
    """
    from io import BytesIO

    # When all engines fail, we currently fall back to the legacy heuristic
    # parser. For quality-focused runs, allow disabling that behavior.
    no_legacy_fallback = os.getenv("ACORD_NO_LEGACY_FALLBACK", "").strip().lower() in {
        "1", "true", "yes", "on",
    }
    merge_all = os.getenv("ACORD_PDF_MERGE_ALL", "1").strip().lower() not in {
        "0", "false", "no", "off",
    }
    engine_errors: list[str] = []
    parts: list[tuple[str, UnifiedIntermediateRepresentation]] = []

    has_widgets = pdf_has_acroform_widgets(contents)
    logger.info(
        "ACORD[pdf] PDF input: %s",
        "fillable (AcroForm widgets)" if has_widgets else "flattened (no AcroForm — OCR + VL)",
    )

    # ── Engine 1: ByteScout CLI ────────────────────────────────────────────
    try:
        bs_uir = try_extract_pdf_to_uir(contents, filename=filename)
        if bs_uir is not None and _uir_non_empty(bs_uir):
            lay = dict(bs_uir.layout or {})
            lay["extraction_engine"] = "bytescout"
            bs_uir.layout = lay
            parts.append(("bytescout", bs_uir))
            logger.info("ACORD[pdf] collected engine=bytescout")
    except Exception as exc:
        logger.exception("ACORD[pdf/bytescout] failed")
        engine_errors.append(f"bytescout: {type(exc).__name__}: {exc}")

    # ── Engine 2: pdfplumber (layout + tables) ─────────────────────────────
    if pdfplumber is not None:
        try:
            with pdfplumber.open(BytesIO(contents)) as pdf:
                pages_words: list[list[dict]] = []
                pages_text_: list[str] = []
                pages_tables_: list[list[list[list[str]]]] = []
                for p in pdf.pages[:20]:
                    raw_text = (p.extract_text() or "").strip()
                    import re as _re
                    raw_text = _re.sub(r"\bn\b(?=\s+[A-Z])", "X", raw_text)
                    pages_text_.append(raw_text)
                    pages_words.append(p.extract_words() or [])
                    try:
                        pages_tables_.append(p.extract_tables() or [])
                    except Exception:
                        pages_tables_.append([])
            uir = build_uir_from_pdf_text(
                pages_words=pages_words,
                pages_text=pages_text_,
                pages_tables=pages_tables_,
                extraction_engine="pdfplumber",
            )
            if _uir_non_empty(uir):
                parts.append(("pdfplumber", uir))
                logger.info("ACORD[pdf] collected engine=pdfplumber")
            else:
                logger.info(
                    "ACORD[pdf/pdfplumber] empty UIR (%d chars raw)",
                    sum(len(tb.text) for tb in uir.text_blocks),
                )
        except Exception as exc:
            logger.exception("ACORD[pdf/pdfplumber] failed")
            engine_errors.append(f"pdfplumber: {type(exc).__name__}: {exc}")

    # ── Engine 3: PyMuPDF text extraction ──────────────────────────────────
    if fitz is not None:
        try:
            doc = fitz.open(stream=contents, filetype="pdf")
            fitz_pages = [doc.load_page(i).get_text("text") or "" for i in range(min(len(doc), 20))]
            full_text = "\n".join(fitz_pages).strip()
            if len(full_text) >= 20:
                uir = UnifiedIntermediateRepresentation(
                    text_blocks=[TextBlock(text=full_text, page=1, bbox=None, source="pdf_text")],
                    layout={"extraction_engine": "pymupdf"},
                )
                parts.append(("pymupdf", uir))
                logger.info("ACORD[pdf] collected engine=pymupdf (%d chars)", len(full_text))
        except Exception as exc:
            logger.exception("ACORD[pdf/pymupdf] failed")
            engine_errors.append(f"pymupdf: {type(exc).__name__}: {exc}")

    # ── Engine 4: PyPDF2 basic text extraction ─────────────────────────────
    if PdfReader is not None:
        try:
            reader = PdfReader(BytesIO(contents))
            pdf2_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
            if len(pdf2_text) >= 20:
                uir = UnifiedIntermediateRepresentation(
                    text_blocks=[TextBlock(text=pdf2_text, page=1, bbox=None, source="pdf_text")],
                    layout={"extraction_engine": "pypdf2"},
                )
                parts.append(("pypdf2", uir))
                logger.info("ACORD[pdf] collected engine=pypdf2 (%d chars)", len(pdf2_text))
        except Exception as exc:
            logger.exception("ACORD[pdf/pypdf2] failed")
            engine_errors.append(f"pypdf2: {type(exc).__name__}: {exc}")

    # ── Engine: Azure Document Intelligence (optional cloud layout + OCR) ──
    if os.getenv("ACORD_SKIP_AZURE_DI", "").strip().lower() not in {"1", "true", "yes", "on"}:
        try:
            azure_uir = await try_azure_document_intelligence_to_uir(contents)
            if azure_uir is not None and _uir_non_empty(azure_uir):
                parts.append(("azure_di", azure_uir))
                logger.info("ACORD[pdf] collected engine=azure_di")
        except Exception as exc:
            logger.warning("ACORD[pdf/azure_di] failed: %s", exc)
            engine_errors.append(f"azure_di: {type(exc).__name__}: {exc}")

    # ── Engine: OCR (rasterize pages — Tesseract and/or optional PaddleOCR) ─
    try:
        ocr_text, ocr_engine_used = ocr_pdf_pages(contents, flattened_pdf=not has_widgets)
        if ocr_text:
            uir = UnifiedIntermediateRepresentation(
                text_blocks=[TextBlock(text=ocr_text, page=1, bbox=None, source="ocr")],
                layout={
                    "extraction_engine": "ocr",
                    "ocr_engine_used": ocr_engine_used,
                    "pdf_form_classification": "fillable" if has_widgets else "flattened",
                },
            )
            parts.append(("ocr", uir))
            logger.info(
                "ACORD[pdf] collected engine=ocr (%d chars, backend=%s)",
                len(ocr_text),
                ocr_engine_used or "?",
            )
    except Exception as exc:
        logger.warning("ACORD[pdf/ocr] failed: %s", exc)
        engine_errors.append(f"ocr: {type(exc).__name__}: {exc}")

    def _uir_char_len(u: UnifiedIntermediateRepresentation) -> int:
        return sum(len((tb.text or "")) for tb in (u.text_blocks or []))

    _ENGINE_PRIORITY = ("bytescout", "pdfplumber", "pymupdf", "pypdf2", "azure_di", "ocr")

    # ── Merge all engines → one LLM call (staging default) ────────────────
    if merge_all:
        merged = _merge_acord_uirs(parts)
        merged_len = _uir_char_len(merged) if merged else 0
        if merged and merged_len >= 80:
            logger.info(
                "ACORD[pdf] merged engines=%s total_chars=%d",
                (merged.layout or {}).get("extraction_engines"),
                merged_len,
            )
            try:
                _patch_uir_pdf_classification(merged, has_widgets, parts)
                _attach_acroform_fields_if_enabled(merged, contents)
                _attach_acord_vl_page_images(merged, contents)
                llm_json = await openai_compat_extract_structured(merged, form_type_hint=form_type_hint)
                if llm_json:
                    return build_summary_from_uir(merged, llm_json)
            except Exception as exc:
                logger.exception("ACORD[pdf/merged-llm] failed")
                engine_errors.append(f"merged_llm: {type(exc).__name__}: {exc}")

    # ── Fallback: single engine (priority order if merge off, else by text length) ─
    if not merge_all:
        ordered_parts = []
        for eng in _ENGINE_PRIORITY:
            for n, u in parts:
                if n == eng:
                    ordered_parts.append((n, u))
                    break
        iter_parts = ordered_parts
    else:
        iter_parts = sorted(parts, key=lambda x: _uir_char_len(x[1]), reverse=True)

    for eng_name, uir in iter_parts:
        if _uir_has_content(uir):
            logger.info("ACORD[pdf] fallback single engine=%s", eng_name)
            try:
                _patch_uir_pdf_classification(uir, has_widgets, parts)
                _attach_acroform_fields_if_enabled(uir, contents)
                _attach_acord_vl_page_images(uir, contents)
                llm_json = await openai_compat_extract_structured(uir, form_type_hint=form_type_hint)
                if llm_json:
                    return build_summary_from_uir(uir, llm_json)
            except Exception as exc:
                logger.exception("ACORD[pdf/fallback-%s] LLM failed", eng_name)
                engine_errors.append(f"{eng_name}_llm: {type(exc).__name__}: {exc}")

    # ── Legacy heuristic parser (absolute last resort) ─────────────────────
    logger.warning("ACORD[pdf] all engines exhausted; using legacy heuristic parser")
    if no_legacy_fallback:
        # Quality-focused mode: do not silently return low-quality heuristic output.
        raise HTTPException(
            status_code=422,
            detail={
                "error": "ACORD extraction failed across all PDF engines",
                "engine_errors": engine_errors[-10:],  # cap size
                "hint": "Check pdf engine logs + ensure OCR TESSDATA_PREFIX/eng.traineddata exists.",
            },
        )
    try:
        if PdfReader is not None:
            from io import BytesIO as _BytesIO2
            fallback_text = "\n".join(
                p.extract_text() or "" for p in PdfReader(_BytesIO2(contents)).pages
            ).strip()
        else:
            fallback_text = ""
    except Exception:
        fallback_text = ""
    return parse_acord_form(fallback_text)


async def _extract_summary_from_file(
    file: UploadFile,
    form_type_hint: Optional[str] = None,
) -> AcordFormSummary:
    """
    Dispatch extraction to the multi-engine PDF chain or plain-text pipeline.

    form_type_hint: user-selected ACORD form number (e.g. '25', '125') — prepended
    to the LLM prompt so the model focuses on the right schema.
    """
    contents = await file.read()
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    # PDF → multi-engine fallback chain
    if filename.endswith(".pdf") or content_type in {"application/pdf", "pdf"}:
        return await _extract_from_pdf(contents, filename=filename, form_type_hint=form_type_hint)

    # Non-PDF: read text via existing handler then run through LLM pipeline
    from io import BytesIO as _BIO
    file.file = _BIO(contents)  # reset without seeking on a potentially consumed stream
    file.file.seek(0)
    text = await _read_text_from_file(file)
    uir = UnifiedIntermediateRepresentation(
        text_blocks=[TextBlock(text=text, page=1, bbox=None, source="txt")],
        layout={"extraction_engine": "txt"},
    )
    llm_json = await openai_compat_extract_structured(uir, form_type_hint=form_type_hint)
    return build_summary_from_uir(uir, llm_json) if llm_json else parse_acord_form(text)


async def _extract_summary_from_raw_text(
    raw_text: str,
    form_type_hint: Optional[str] = None,
) -> AcordFormSummary:
    """Re-run LLM extraction over stored raw text (no layout info available)."""
    uir = UnifiedIntermediateRepresentation(
        text_blocks=[TextBlock(text=raw_text, page=1, bbox=None, source="txt")],
        layout={"extraction_engine": "txt"},
    )
    llm_json = await openai_compat_extract_structured(uir, form_type_hint=form_type_hint)
    return build_summary_from_uir(uir, llm_json) if llm_json else parse_acord_form(raw_text)


async def _run_extract_and_persist(
    *,
    file: UploadFile,
    user_id: str,
    form_type_hint: Optional[str],
    source_mime: Optional[str] = None,
) -> AcordExtractResponse:
    summary = await _extract_summary_from_file(file, form_type_hint=form_type_hint)
    text = summary.raw_text or ""

    extracted = summary.model_dump(mode="json", exclude_none=True)
    extracted.pop("raw_text", None)
    overall_confidence = float(summary.overall_confidence or 0.0)

    hardened_enabled = (os.getenv("ACORD_HARDENED_INFERENCE_ENABLED", "false").strip().lower() in {"1", "true", "yes", "on"})
    if hardened_enabled and (summary.raw_text or "").strip():
        try:
            from fine_tuning.acord_form_pipeline.inference_production import run_default_extraction

            hardened_form = _hardened_form_type_hint(form_type_hint, str(extracted.get("form_type") or ""))
            hardened = run_default_extraction(
                text=summary.raw_text or "",
                form_type=hardened_form,
            )
            if hardened.get("status") in {"ok", "needs_review", "low_confidence"}:
                extracted = _apply_hardened_overlay(extracted, hardened)
                final_score = hardened.get("final_score")
                if isinstance(final_score, (float, int)):
                    overall_confidence = _clamp01(float(final_score) / 100.0)
        except Exception as exc:
            logger.warning("ACORD[hardened] overlay failed: %s", exc)

    try:
        row = (await postgrest_insert(
            "acord_extraction_runs",
            {
                "created_by": user_id,
                "source_filename": file.filename,
                "source_mime": source_mime or file.content_type,
                "form_type_detected": extracted.get("form_type"),
                "raw_text": summary.raw_text,
                "original_extracted_json": extracted,
                "extracted_json": extracted,
                "overall_confidence": overall_confidence,
                "status": "draft",
            },
        ))[0]
    except HTTPException as persist_exc:
        logger.exception(
            "ACORD[persist] could not save run — returning extracted fields only: %s",
            persist_exc.detail,
        )
        detail_txt = persist_exc.detail if isinstance(persist_exc.detail, str) else str(persist_exc.detail)
        nl_summary_partial = await _generate_nl_summary(extracted, text)
        return AcordExtractResponse(
            run_id="",
            status="draft",
            overall_confidence=overall_confidence,
            extracted=extracted,
            partial=True,
            persist_error=detail_txt[:2000],
            warning=(
                "Extraction completed but the draft could not be saved to the database. "
                "You can still review and copy the JSON; submit/feedback needs a saved run."
            ),
            natural_language_summary=nl_summary_partial,
        )

    doc_id = (file.filename or "acord-upload").rsplit(".", 1)[0]
    try:
        await _ingest_acord_into_vectorstore(doc_id, summary.raw_text or text)
    except Exception as exc:  # pragma: no cover - defensive
        logger.exception("ACORD[ingest] failed for doc_id=%s: %s", doc_id, exc)

    tenant_id = await resolve_tenant_id_for_user(user_id)
    await try_emit_webhook_event(
        tenant_id,
        WEBHOOK_EVENT_INFERENCE_COMPLETE,
        {
            "run_id": str(row["id"]),
            "workflow": "acord",
            "form_type": extracted.get("form_type"),
            "overall_confidence": overall_confidence,
        },
    )

    nl_summary = await _generate_nl_summary(extracted, text)

    return AcordExtractResponse(
        run_id=str(row["id"]),
        status=row.get("status") or "draft",
        overall_confidence=float(row.get("overall_confidence") or 0.0),
        extracted=row.get("extracted_json") or {},
        partial=False,
        natural_language_summary=nl_summary,
    )


async def _set_extract_job(job_id: str, patch: dict[str, Any]) -> None:
    persisted: dict[str, Any] = {}
    async with _EXTRACT_JOBS_LOCK:
        cur = dict(_EXTRACT_JOBS.get(job_id) or {})
        cur.update(patch)
        _EXTRACT_JOBS[job_id] = cur
        persisted = dict(cur)
        if len(_EXTRACT_JOBS) > 500:
            # Keep memory bounded for long-running processes.
            for old_id in list(_EXTRACT_JOBS.keys())[:100]:
                if old_id != job_id:
                    _EXTRACT_JOBS.pop(old_id, None)
    try:
        await postgrest_patch(
            _EXTRACT_JOB_TABLE,
            f"job_id=eq.{quote(job_id, safe='')}",
            {
                "status": persisted.get("status"),
                "phase": persisted.get("phase"),
                "error": persisted.get("error"),
                "result": persisted.get("result"),
            },
        )
    except Exception as exc:
        # Keep requests non-blocking, but include reason for operators.
        logger.warning("ACORD[extract/job] db patch failed: job_id=%s error=%s", job_id, exc)


async def _run_extract_job(
    *,
    job_id: str,
    user_id: str,
    filename: str,
    content_type: str,
    contents: bytes,
    form_type_hint: Optional[str],
) -> None:
    started_at = asyncio.get_event_loop().time()
    logger.info(
        "ACORD[extract/job] started: job_id=%s user_id=%s filename=%s content_type=%s form_type_hint=%s",
        job_id,
        user_id,
        filename,
        content_type,
        form_type_hint,
    )
    await _set_extract_job(job_id, {"status": "running", "phase": "warming_model"})
    try:
        logger.info("ACORD[extract/job] phase=warming_model job_id=%s", job_id)
        await _ensure_runpod_ready_for_acord()
        logger.info("ACORD[extract/job] phase=generate_extracting job_id=%s", job_id)
        await _set_extract_job(job_id, {"status": "running", "phase": "generate_extracting"})
        upload = UploadFile(filename=filename, file=io.BytesIO(contents))
        resp = await _run_extract_and_persist(
            file=upload,
            user_id=user_id,
            form_type_hint=form_type_hint,
            source_mime=content_type,
        )
        await _set_extract_job(
            job_id,
            {"status": "succeeded", "phase": "completed", "result": resp.model_dump(mode="json"), "error": None},
        )
        elapsed = round(asyncio.get_event_loop().time() - started_at, 2)
        logger.info("ACORD[extract/job] succeeded: job_id=%s elapsed_s=%s", job_id, elapsed)
    except Exception as exc:
        elapsed = round(asyncio.get_event_loop().time() - started_at, 2)
        logger.exception("ACORD[extract/job] failed: job_id=%s elapsed_s=%s", job_id, elapsed)
        await _set_extract_job(job_id, {"status": "failed", "phase": "failed", "error": str(exc)[:2000]})






@router.post("/parse", response_model=AcordFormSummary)
async def parse_acord_form_endpoint(file: UploadFile = File(...)):
    """
    Upload an ACORD form (PDF or text), extract text, and return a normalized
    structured summary.

    This endpoint performs the "model understanding" step. The frontend can then
    pass the extracted text into the generic `acord_form_understanding` agent
    to obtain a grounded LLM explanation using the shared RAG pipeline.
    """
    summary = await _extract_summary_from_file(file)
    text = summary.raw_text or ""

    # Persist the raw text into the insurance_index so that other insurance
    # agents (e.g. policy comparison, multi-document analysis) can retrieve
    # this upload via the shared vector store.
    doc_id = (file.filename or "acord-upload").rsplit(".", 1)[0]
    try:
        await _ingest_acord_into_vectorstore(doc_id, summary.raw_text or text)
    except Exception as exc:  # pragma: no cover - defensive
        logger.exception("ACORD[ingest] failed for doc_id=%s: %s", doc_id, exc)

    return summary


@router.post("/extract", response_model=AcordExtractResponse)
async def extract_acord_endpoint(
    file: UploadFile = File(...),
    authorization: str | None = Header(default=None),
    form_type_hint: Optional[str] = Query(default=None, description="e.g. '25', '125', '140' — user-selected ACORD form type"),
):
    """
    Upload a document, extract text + ACORD fields, and persist a draft run row.

    This is the workflow-aware version of `/api/acord/parse`.
    Pass `?form_type_hint=125` when the user has selected a form type in the UI.
    """
    logger.info(
        "ACORD[extract] request: filename=%s content_type=%s form_type_hint=%s auth=%s",
        file.filename,
        file.content_type,
        form_type_hint,
        _auth_debug_meta(authorization),
    )
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        await _ensure_runpod_ready_for_acord()
        return await _run_extract_and_persist(
            file=file,
            user_id=str(user_id),
            form_type_hint=form_type_hint,
            source_mime=file.content_type,
        )
    except HTTPException as exc:
        logger.warning(
            "ACORD[extract] failed: filename=%s content_type=%s status=%s detail=%s",
            file.filename,
            file.content_type,
            exc.status_code,
            exc.detail,
        )
        raise


@router.post("/extract/start", response_model=AcordExtractStartResponse)
async def start_extract_acord_endpoint(
    file: UploadFile = File(...),
    authorization: str | None = Header(default=None),
    form_type_hint: Optional[str] = Query(default=None, description="e.g. '25', '125', '140'"),
):
    """
    Start asynchronous extraction and return immediately with a job_id.
    Frontend should poll /api/acord/extract/status/{job_id}.
    """
    logger.info(
        "ACORD[extract/start] request: filename=%s content_type=%s form_type_hint=%s auth=%s",
        file.filename,
        file.content_type,
        form_type_hint,
        _auth_debug_meta(authorization),
    )
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")

    job_id = str(uuid.uuid4())
    logger.info("ACORD[extract/start] creating job: job_id=%s user_id=%s", job_id, user_id)
    try:
        await postgrest_insert(
            _EXTRACT_JOB_TABLE,
            {
                "job_id": job_id,
                "user_id": str(user_id),
                "status": "queued",
                "phase": "queued",
                "error": None,
                "result": None,
            },
        )
    except Exception as exc:
        # DB persistence is required for cross-instance status reliability.
        logger.exception("ACORD[extract/job] db insert failed: job_id=%s", job_id)
        raise HTTPException(
            status_code=500,
            detail=f"Could not persist async extraction job. Verify Supabase service config. ({exc})",
        ) from exc
    await _set_extract_job(
        job_id,
        {
            "job_id": job_id,
            "status": "queued",
            "phase": "queued",
            "user_id": str(user_id),
            "error": None,
            "result": None,
        },
    )
    asyncio.create_task(
        _run_extract_job(
            job_id=job_id,
            user_id=str(user_id),
            filename=file.filename or "upload.bin",
            content_type=file.content_type or "",
            contents=contents,
            form_type_hint=form_type_hint,
        )
    )
    logger.info("ACORD[extract/start] queued job: job_id=%s", job_id)
    return AcordExtractStartResponse(job_id=job_id, status="queued")


@router.get("/extract/status/{job_id}", response_model=AcordExtractJobStatusResponse)
async def get_extract_acord_status(
    job_id: str,
    authorization: str | None = Header(default=None),
):
    logger.debug("ACORD[extract/status] request: job_id=%s auth=%s", job_id, _auth_debug_meta(authorization))
    user = await verify_user(authorization)
    user_id = str(user.get("id") or "")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")
    job: dict[str, Any] = {}
    try:
        rows = await postgrest_get(
            _EXTRACT_JOB_TABLE,
            f"select=job_id,user_id,status,phase,error,result&job_id=eq.{quote(job_id, safe='')}&limit=1",
        )
        if rows:
            row = rows[0]
            job = {
                "job_id": row.get("job_id"),
                "user_id": row.get("user_id"),
                "status": row.get("status"),
                "phase": row.get("phase"),
                "error": row.get("error"),
                "result": row.get("result"),
            }
            # Re-warm local cache for subsequent polls.
            async with _EXTRACT_JOBS_LOCK:
                _EXTRACT_JOBS[job_id] = dict(job)
    except Exception as exc:
        logger.warning("ACORD[extract/job] db read failed: job_id=%s error=%s", job_id, exc)
    if not job:
        async with _EXTRACT_JOBS_LOCK:
            job = dict(_EXTRACT_JOBS.get(job_id) or {})
    if not job:
        logger.warning("ACORD[extract/status] job not found: job_id=%s user_id=%s", job_id, user_id)
        raise HTTPException(status_code=404, detail="Job not found")
    if str(job.get("user_id") or "") != user_id:
        logger.warning(
            "ACORD[extract/status] forbidden: job_id=%s owner_user_id=%s requester_user_id=%s",
            job_id,
            job.get("user_id"),
            user_id,
        )
        raise HTTPException(status_code=403, detail="Forbidden")

    result = job.get("result")
    return AcordExtractJobStatusResponse(
        job_id=job_id,
        status=str(job.get("status") or "queued"),
        phase=(str(job.get("phase")) if job.get("phase") else None),
        result=(AcordExtractResponse(**result) if isinstance(result, dict) else None),
        error=(str(job.get("error")) if job.get("error") else None),
    )


@router.get("/runs/{run_id}")
async def get_acord_run(run_id: str, authorization: str | None = Header(default=None)):
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=*&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Run not found")

    run = rows[0]
    # Enforce access: owner or admin (RLS in DB also applies if using anon key, but backend uses service role).
    if run.get("created_by") != user_id:
        roles = await postgrest_get("user_roles", f"select=role&user_id=eq.{quote(user_id, safe='')}&limit=1")
        role = roles[0].get("role") if roles else None
        if role not in {"admin", "global_admin"}:
            raise HTTPException(status_code=403, detail="Forbidden")
    feedback_rows = await postgrest_get(
        "acord_extraction_feedback",
        f"select=corrected_json,thumbs_up,actor_role,created_at&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=20",
    )
    latest_corrected = None
    for fb in feedback_rows:
        if fb.get("corrected_json") is not None:
            latest_corrected = fb.get("corrected_json")
            break

    original_json = run.get("original_extracted_json")
    if original_json is None:
        original_json = run.get("extracted_json") or {}

    edited_json = latest_corrected if latest_corrected is not None else (run.get("extracted_json") or {})
    has_edits = latest_corrected is not None and latest_corrected != original_json

    run["original_extracted_json"] = original_json
    run["edited_extracted_json"] = edited_json
    run["has_edits"] = has_edits
    run["confidence_evaluation"] = _evaluate_confidence_and_feedback(
        base_confidence=float(run.get("overall_confidence") or 0.0),
        run_status=str(run.get("status") or ""),
        original_json=original_json if isinstance(original_json, dict) else {},
        edited_json=edited_json if isinstance(edited_json, dict) else {},
        feedback_rows=feedback_rows if isinstance(feedback_rows, list) else [],
    )
    return {"run": run}


@router.post("/runs/{run_id}/preview-training-jsonl")
async def preview_training_jsonl(
    run_id: str,
    body: PreviewSftTrainingRecordBody,
    authorization: str | None = Header(default=None),
):
    """
    Return the exact JSONL object shape produced by `export_approved_acord_dataset` for this run:
    six-field `output` via `build_sft_label_json`, metadata without confidence scores.
    """
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=created_by&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Run not found")
    run_row = rows[0]
    if run_row.get("created_by") != user_id:
        roles = await postgrest_get("user_roles", f"select=role&user_id=eq.{quote(user_id, safe='')}&limit=1")
        role = roles[0].get("role") if roles else None
        if role not in {"admin", "global_admin"}:
            raise HTTPException(status_code=403, detail="Forbidden")

    from fine_tuning.export_approved_acord_dataset import build_training_jsonl_record

    record = build_training_jsonl_record(
        extracted_json=body.extracted_json,
        raw_text=body.raw_text or "",
        run_id=run_id,
        source_filename=body.source_filename,
    )
    return {"record": record}


@router.post("/runs/{run_id}/re-extract", response_model=AcordExtractResponse)
async def re_extract_run(
    run_id: str,
    body: ReExtractRequest,
    authorization: str | None = Header(default=None),
    file: Optional[UploadFile] = File(default=None),
):
    """
    Re-run the extraction pipeline for an existing run.

    - If a new `file` is provided, the full multi-engine pipeline is used (highest quality).
    - If no file is provided, the stored `raw_text` is re-processed through the LLM with
      the new (or same) form_type_hint. Useful when the model or prompt improves.

    The run row is updated in-place (extracted_json, overall_confidence, form_type_detected).
    The run status is reset to 'draft' so the user can validate and resubmit.
    """
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # Fetch existing run (owner or admin)
    rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=*&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Run not found")
    run = rows[0]
    if run.get("created_by") != user_id:
        roles = await postgrest_get("user_roles", f"select=role&user_id=eq.{quote(user_id, safe='')}&limit=1")
        role = (roles[0].get("role") if roles else None)
        if role not in {"admin", "global_admin"}:
            raise HTTPException(status_code=403, detail="Forbidden")

    hint = body.form_type_hint or run.get("form_type_detected")

    # Re-extract
    if file is not None:
        summary = await _extract_summary_from_file(file, form_type_hint=hint)
        logger.info("ACORD[re-extract] run_id=%s engine=file hint=%s", run_id, hint)
    else:
        raw_text = run.get("raw_text") or ""
        if not raw_text.strip():
            raise HTTPException(
                status_code=400,
                detail="No raw text stored for this run. Upload the original file to re-extract.",
            )
        summary = await _extract_summary_from_raw_text(raw_text, form_type_hint=hint)
        logger.info("ACORD[re-extract] run_id=%s engine=raw_text hint=%s", run_id, hint)

    extracted = summary.model_dump(mode="json", exclude_none=True)
    extracted.pop("raw_text", None)
    overall_confidence = float(summary.overall_confidence or 0.0)

    hardened_enabled = (os.getenv("ACORD_HARDENED_INFERENCE_ENABLED", "false").strip().lower() in {"1", "true", "yes", "on"})
    if hardened_enabled and (summary.raw_text or "").strip():
        try:
            from fine_tuning.acord_form_pipeline.inference_production import run_default_extraction

            hardened_form = _hardened_form_type_hint(hint, str(extracted.get("form_type") or ""))
            hardened = run_default_extraction(
                text=summary.raw_text or "",
                form_type=hardened_form,
            )
            if hardened.get("status") in {"ok", "needs_review", "low_confidence"}:
                extracted = _apply_hardened_overlay(extracted, hardened)
                final_score = hardened.get("final_score")
                if isinstance(final_score, (float, int)):
                    overall_confidence = _clamp01(float(final_score) / 100.0)
        except Exception as exc:
            logger.warning("ACORD[hardened] re-extract overlay failed: %s", exc)

    await postgrest_patch(
        "acord_extraction_runs",
        f"id=eq.{quote(run_id, safe='')}",
        {
            "status": "draft",
            "form_type_detected": extracted.get("form_type"),
            "overall_confidence": overall_confidence,
            "original_extracted_json": extracted,
            "extracted_json": extracted,
            **({"raw_text": summary.raw_text} if summary.raw_text else {}),
        },
    )

    return AcordExtractResponse(
        run_id=run_id,
        status="draft",
        overall_confidence=overall_confidence,
        extracted=extracted,
    )


@router.post("/runs/{run_id}/submit")
async def submit_acord_run(
    run_id: str,
    body: AcordSubmitRequest,
    authorization: str | None = Header(default=None),
):
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    run_rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=*&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not run_rows:
        raise HTTPException(status_code=404, detail="Run not found")
    run = run_rows[0]
    if run.get("created_by") != user_id:
        raise HTTPException(status_code=403, detail="Forbidden")

    # Persist feedback
    try:
        await postgrest_insert(
            "acord_extraction_feedback",
            {
                "run_id": run_id,
                "created_by": user_id,
                "actor_role": "user",
                "thumbs_up": body.thumbs_up,
                "notes": body.notes,
                "corrected_json": body.corrected_json,
            },
        )
    except HTTPException as exc:
        # If Supabase schema cache is missing the table, PostgREST throws PGRST205.
        # Provide a clearer error so the user can apply migrations.
        detail = str(exc.detail or "")
        if "acord_extraction_feedback" in detail and (
            "PGRST205" in detail or "schema cache" in detail or "relation" in detail
        ):
            raise HTTPException(
                status_code=500,
                detail="ACORD feedback table is missing in Supabase. Please run migrations (20260318120000_acord_extraction_workflow.sql and 20260318150000_acord_extraction_feedback_resync.sql).",
            ) from exc
        raise

    extracted = run.get("extracted_json") or {}
    # Treat explicitly provided corrected JSON as authoritative training truth,
    # even when it is an empty object.
    if body.corrected_json is not None:
        extracted = body.corrected_json

    overall_confidence = float(run.get("overall_confidence") or 0.0)
    next_status = "needs_admin_review"
    queue_reason = "low_confidence_or_user_disagreed"
    if body.thumbs_up and overall_confidence >= _confidence_threshold() and not body.require_admin_approval_for_training:
        next_status = "approved"
        queue_reason = None

    await postgrest_patch(
        "acord_extraction_runs",
        f"id=eq.{quote(run_id, safe='')}",
        {
            "status": "submitted" if next_status != "approved" else "approved",
            "extracted_json": extracted,
        },
    )

    if next_status == "approved":
        # If there was a queue row, mark approved.
        try:
            await postgrest_patch(
                "acord_admin_queue",
                f"run_id=eq.{quote(run_id, safe='')}",
                {"state": "approved"},
            )
        except Exception:
            pass
        return {"status": "approved"}

    # Upsert into admin queue via REST (merge on conflict)
    # Using PostgREST upsert semantics requires Prefer header; easiest is insert and ignore conflict.
    try:
        await postgrest_insert(
            "acord_admin_queue",
            {
                "run_id": run_id,
                "priority": 0,
                "reason": queue_reason,
                "state": "open",
            },
        )
    except Exception:
        # Likely already exists; patch it.
        await postgrest_patch(
            "acord_admin_queue",
            f"run_id=eq.{quote(run_id, safe='')}",
            {"reason": queue_reason, "state": "open"},
        )

    await postgrest_patch(
        "acord_extraction_runs",
        f"id=eq.{quote(run_id, safe='')}",
        {"status": "needs_admin_review"},
    )

    return {"status": "needs_admin_review"}


async def _require_admin(authorization: str | None) -> dict:
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")
    roles = await postgrest_get("user_roles", f"select=role&user_id=eq.{quote(user_id, safe='')}&limit=1")
    role = roles[0].get("role") if roles else None
    if role not in {"admin", "global_admin"}:
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


async def _queue_or_restart_training_job(
    *,
    run_id: str,
    admin_id: Optional[str],
    background_tasks: BackgroundTasks,
) -> None:
    """
    Ensure a training job is queued for a run.

    Behavior:
    - If no job exists: create + spawn.
    - If a job is queued/running: keep as-is (avoid duplicate workers).
    - If latest job is completed/failed: create a NEW job row and spawn.
      This preserves full per-run history in acord_training_jobs.
    """
    existing_jobs = await postgrest_get(
        "acord_training_jobs",
        f"select=id,status&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=1",
    )
    if not existing_jobs:
        job = await create_job_row(run_id=run_id, created_by=admin_id)
        background_tasks.add_task(spawn_job_runner, job_id=str(job["id"]), run_id=run_id)
        return

    existing = existing_jobs[0]
    job_id = str(existing.get("id"))
    status = str(existing.get("status") or "queued").lower()

    if status in {"queued", "running"}:
        logger.info(
            "ACORD[approve] training job already %s for run_id=%s (id=%s) — skipping respawn.",
            status, run_id, job_id,
        )
        return

    # Keep history immutable: create a new job row for each restart.
    job = await create_job_row(run_id=run_id, created_by=admin_id)
    new_job_id = str(job["id"])
    logger.info(
        "ACORD[approve] creating new training job for run_id=%s (new_id=%s, prev_id=%s, prev_status=%s).",
        run_id, new_job_id, job_id, status,
    )
    background_tasks.add_task(spawn_job_runner, job_id=new_job_id, run_id=run_id)


@router.get("/admin/queue")
async def list_admin_queue(
    authorization: str | None = Header(default=None),
    states: Optional[str] = Query(default="open,in_progress", description="Comma-separated states, e.g. 'open,in_progress,approved,rejected'"),
    form_type: Optional[str] = Query(default=None, description="Filter by ACORD form type, e.g. '25' or '125'"),
    conf_min: Optional[float] = Query(default=None, ge=0.0, le=1.0, description="Minimum overall_confidence (0–1)"),
    conf_max: Optional[float] = Query(default=None, ge=0.0, le=1.0, description="Maximum overall_confidence (0–1)"),
    order_by: str = Query(default="priority", description="priority|created_at|updated_at"),
    order_dir: str = Query(default="desc", description="asc|desc"),
    page: int = Query(default=1, ge=1),
    limit: int = Query(default=25, ge=1, le=100),
):
    """List admin queue items with optional filtering and pagination."""
    await _require_admin(authorization)

    # Sanitize ordering inputs
    _valid_order = {"priority", "created_at", "updated_at"}
    _valid_dir = {"asc", "desc"}
    ob = order_by if order_by in _valid_order else "priority"
    od = order_dir if order_dir in _valid_dir else "desc"

    offset = (page - 1) * limit

    # Build states filter
    state_list = [s.strip() for s in (states or "open,in_progress").split(",") if s.strip()]
    state_fragment = f"&state=in.({','.join(quote(s, safe='') for s in state_list)})"

    # Build form-type filter (on the embedded resource)
    form_type_fragment = ""
    if form_type:
        form_type_fragment = f"&acord_extraction_runs.form_type_detected=eq.{quote(form_type.strip(), safe='')}"

    qs = (
        f"select=*,acord_extraction_runs(*)"
        f"{state_fragment}"
        f"{form_type_fragment}"
        f"&order={ob}.{od},created_at.asc"
        f"&limit={limit}&offset={offset}"
    )
    rows = await postgrest_get("acord_admin_queue", qs)

    # Client-side confidence filter (overall_confidence lives on the joined run)
    if conf_min is not None or conf_max is not None:
        def _conf_ok(row: dict) -> bool:
            conf = float((row.get("acord_extraction_runs") or {}).get("overall_confidence") or 0)
            if conf_min is not None and conf < conf_min:
                return False
            if conf_max is not None and conf > conf_max:
                return False
            return True
        rows = [r for r in rows if _conf_ok(r)]

    return {"queue": rows, "page": page, "limit": limit}


@router.post("/admin/{run_id}/review")
async def admin_review_run(
    run_id: str,
    body: AcordAdminReviewRequest,
    background_tasks: BackgroundTasks,
    authorization: str | None = Header(default=None),
):
    user = await _require_admin(authorization)
    admin_id = user.get("id")

    decision = (body.decision or "").strip().lower()
    if decision not in {"approve", "rework", "reject"}:
        raise HTTPException(status_code=400, detail="Invalid decision")

    run_rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=*&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not run_rows:
        raise HTTPException(status_code=404, detail="Run not found")
    run = run_rows[0]

    # Final training truth resolution priority:
    # 1) Admin-provided corrected_json in this approval request
    # 2) Most recent corrected_json from feedback history (usually user edits)
    # 3) Current run.extracted_json
    extracted = run.get("extracted_json") or {}
    if body.corrected_json is not None:
        extracted = body.corrected_json
    else:
        latest_hist = await _latest_corrected_json_for_run(run_id)
        if latest_hist is not None:
            extracted = latest_hist

    await postgrest_insert(
        "acord_extraction_feedback",
        {
            "run_id": run_id,
            "created_by": admin_id,
            "actor_role": "admin",
            "thumbs_up": decision == "approve",
            "notes": body.notes,
            "corrected_json": body.corrected_json,
        },
    )

    if decision == "approve":
        await postgrest_patch(
            "acord_extraction_runs",
            f"id=eq.{quote(run_id, safe='')}",
            {"status": "approved", "extracted_json": extracted},
        )
        await postgrest_patch(
            "acord_admin_queue",
            f"run_id=eq.{quote(run_id, safe='')}",
            {"state": "approved", "assigned_to": body.assigned_to},
        )

        await _queue_or_restart_training_job(
            run_id=run_id,
            admin_id=admin_id,
            background_tasks=background_tasks,
        )

        return {"status": "approved"}

    if decision == "reject":
        await postgrest_patch(
            "acord_extraction_runs",
            f"id=eq.{quote(run_id, safe='')}",
            {"status": "rejected", "extracted_json": extracted},
        )
        await postgrest_patch(
            "acord_admin_queue",
            f"run_id=eq.{quote(run_id, safe='')}",
            {"state": "rejected", "assigned_to": body.assigned_to},
        )
        return {"status": "rejected"}

    # rework
    await postgrest_patch(
        "acord_extraction_runs",
        f"id=eq.{quote(run_id, safe='')}",
        {"status": "needs_admin_review", "extracted_json": extracted},
    )
    await postgrest_patch(
        "acord_admin_queue",
        f"run_id=eq.{quote(run_id, safe='')}",
        {"state": "rework", "assigned_to": body.assigned_to},
    )
    return {"status": "needs_admin_review"}


# -- Batch review --------------------------------------------------------------

@router.post("/admin/batch-review")
async def batch_review_runs(
    body: AcordBatchReviewRequest,
    background_tasks: BackgroundTasks,
    authorization: str | None = Header(default=None),
):
    """Approve or reject multiple runs in one request (max 50). No per-run correction."""
    user = await _require_admin(authorization)
    admin_id = user.get("id")

    decision = (body.decision or "").strip().lower()
    if decision not in {"approve", "reject"}:
        raise HTTPException(status_code=400, detail="Batch decision must be 'approve' or 'reject'")

    run_status = "approved" if decision == "approve" else "rejected"
    queue_state = run_status

    results: list[dict] = []
    for run_id in body.run_ids:
        try:
            # Feedback row
            await postgrest_insert(
                "acord_extraction_feedback",
                {
                    "run_id": run_id,
                    "created_by": admin_id,
                    "actor_role": "admin",
                    "thumbs_up": decision == "approve",
                    "notes": body.notes,
                },
            )
            # Update run status
            await postgrest_patch(
                "acord_extraction_runs",
                f"id=eq.{quote(run_id, safe='')}",
                {"status": run_status},
            )
            # Update queue state
            await postgrest_patch(
                "acord_admin_queue",
                f"run_id=eq.{quote(run_id, safe='')}",
                {"state": queue_state},
            )
            # Spawn training job for approvals (deduplication guard included)
            if decision == "approve":
                await _queue_or_restart_training_job(
                    run_id=run_id,
                    admin_id=admin_id,
                    background_tasks=background_tasks,
                )

            results.append({"run_id": run_id, "ok": True})
        except Exception as exc:
            logger.warning("Batch review failed for run_id=%s: %s", run_id, exc)
            results.append({"run_id": run_id, "ok": False, "error": str(exc)})

    succeeded = sum(1 for r in results if r["ok"])
    return {"decision": decision, "succeeded": succeeded, "total": len(results), "results": results}


# -- User run history ----------------------------------------------------------

@router.get("/runs")
async def list_user_runs(
    authorization: str | None = Header(default=None),
    page: int = Query(default=1, ge=1),
    limit: int = Query(default=20, ge=1, le=100),
    status: Optional[str] = Query(default=None, description="Filter by status"),
):
    """List the authenticated user's own extraction runs, newest first."""
    user = await verify_user(authorization)
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    offset = (page - 1) * limit
    qs = (
        f"select=id,created_at,updated_at,source_filename,form_type_detected,"
        f"overall_confidence,status"
        f"&created_by=eq.{quote(user_id, safe='')}"
        f"&order=created_at.desc"
        f"&limit={limit}&offset={offset}"
    )
    if status:
        qs += f"&status=eq.{quote(status, safe='')}"

    rows = await postgrest_get("acord_extraction_runs", qs)
    return {"runs": rows, "page": page, "limit": limit}


# -- Admin queue stats & detail ------------------------------------------------

@router.get("/admin/queue/stats")
async def admin_queue_stats(authorization: str | None = Header(default=None)):
    """Return open/in_progress/total counts for the admin queue badge."""
    await _require_admin(authorization)
    all_rows = await postgrest_get("acord_admin_queue", "select=state")
    counts: dict[str, int] = {}
    for r in all_rows:
        s = r.get("state") or "unknown"
        counts[s] = counts.get(s, 0) + 1
    return {
        "open": counts.get("open", 0),
        "in_progress": counts.get("in_progress", 0),
        "rework": counts.get("rework", 0),
        "approved": counts.get("approved", 0),
        "rejected": counts.get("rejected", 0),
        "total": len(all_rows),
    }


@router.get("/admin/queue/{run_id}/detail")
async def get_admin_queue_item(run_id: str, authorization: str | None = Header(default=None)):
    """Return queue row metadata + nested run for a single item."""
    await _require_admin(authorization)
    rows = await postgrest_get(
        "acord_admin_queue",
        f"select=*,acord_extraction_runs(*)&run_id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Queue item not found")
    return {"item": rows[0]}


@router.patch("/admin/queue/{run_id}/detail")
async def patch_admin_queue_item(
    run_id: str,
    body: dict,
    authorization: str | None = Header(default=None),
):
    """Update priority, assigned_to, or state on a queue item."""
    await _require_admin(authorization)
    allowed = {"priority", "assigned_to", "state"}
    patch = {k: v for k, v in body.items() if k in allowed}
    if not patch:
        raise HTTPException(status_code=400, detail=f"No valid fields to patch. Allowed: {allowed}")
    await postgrest_patch("acord_admin_queue", f"run_id=eq.{quote(run_id, safe='')}", patch)
    return {"updated": patch}


# -- Training job status -------------------------------------------------------

@router.get("/admin/jobs")
async def list_training_jobs(
    authorization: str | None = Header(default=None),
    status: Optional[str] = Query(default=None),
    page: int = Query(default=1, ge=1),
    limit: int = Query(default=20, ge=1, le=100),
):
    """List fine-tuning jobs (admin only), newest first."""
    await _require_admin(authorization)
    offset = (page - 1) * limit
    qs = (
        f"select=id,created_at,updated_at,run_id,status,started_at,finished_at,error,log_path"
        f"&order=created_at.desc&limit={limit}&offset={offset}"
    )
    if status:
        qs += f"&status=eq.{quote(status, safe='')}"
    rows = await postgrest_get("acord_training_jobs", qs)
    return {"jobs": rows, "page": page, "limit": limit}


@router.get("/admin/jobs/by-run/{run_id}")
async def get_job_by_run_id(run_id: str, authorization: str | None = Header(default=None)):
    """Return the training job for a given run_id (if any)."""
    await _require_admin(authorization)
    rows = await postgrest_get(
        "acord_training_jobs",
        f"select=*&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=1",
    )
    if not rows:
        return {"job": None}
    return {"job": rows[0]}


@router.get("/admin/jobs/by-run/{run_id}/history")
async def get_job_history_by_run_id(
    run_id: str,
    authorization: str | None = Header(default=None),
    limit: int = Query(default=25, ge=1, le=200),
):
    """Return fine-tuning job history for a run_id, newest first."""
    await _require_admin(authorization)
    rows = await postgrest_get(
        "acord_training_jobs",
        (
            "select=id,created_at,updated_at,run_id,status,started_at,finished_at,"
            "error,log_path,dataset_path,output_dir"
            f"&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit={limit}"
        ),
    )
    return {"jobs": rows}


@router.get("/admin/jobs/{job_id}")
async def get_training_job(job_id: str, authorization: str | None = Header(default=None)):
    """Return full detail of a single training job."""
    await _require_admin(authorization)
    rows = await postgrest_get(
        "acord_training_jobs",
        f"select=*&id=eq.{quote(job_id, safe='')}&limit=1",
    )
    if not rows:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"job": rows[0]}


@router.get("/admin/jobs/{job_id}/eval")
async def get_job_eval_results(job_id: str, authorization: str | None = Header(default=None)):
    """Return evaluation metrics for a training job (seen/paraphrased/oos)."""
    await _require_admin(authorization)
    rows = await postgrest_get(
        "acord_eval_results",
        f"select=*&job_id=eq.{quote(job_id, safe='')}&order=eval_set.asc",
    )
    return {"eval_results": rows}


@router.get("/admin/runs/{run_id}/health-card")
async def get_run_health_card(run_id: str, authorization: str | None = Header(default=None)):
    """
    Single endpoint for operational health:
    - extraction confidence + feedback-derived confidence signal
    - latest training job status
    - latest eval metrics and gate pass/fail snapshot
    """
    await _require_admin(authorization)

    run_rows = await postgrest_get(
        "acord_extraction_runs",
        f"select=id,created_at,updated_at,status,form_type_detected,overall_confidence,extracted_json,original_extracted_json"
        f"&id=eq.{quote(run_id, safe='')}&limit=1",
    )
    if not run_rows:
        raise HTTPException(status_code=404, detail="Run not found")
    run = run_rows[0]

    feedback_rows = await postgrest_get(
        "acord_extraction_feedback",
        f"select=corrected_json,thumbs_up,actor_role,created_at&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=50",
    )
    original_json = run.get("original_extracted_json")
    if original_json is None:
        original_json = run.get("extracted_json") or {}
    edited_json = run.get("extracted_json") or {}
    confidence_eval = _evaluate_confidence_and_feedback(
        base_confidence=float(run.get("overall_confidence") or 0.0),
        run_status=str(run.get("status") or ""),
        original_json=original_json if isinstance(original_json, dict) else {},
        edited_json=edited_json if isinstance(edited_json, dict) else {},
        feedback_rows=feedback_rows if isinstance(feedback_rows, list) else [],
    )

    job_rows = await postgrest_get(
        "acord_training_jobs",
        f"select=*&run_id=eq.{quote(run_id, safe='')}&order=created_at.desc&limit=1",
    )
    latest_job = job_rows[0] if job_rows else None
    eval_rows: list[dict] = []
    gate_snapshot: Optional[dict] = None
    if latest_job:
        eval_rows = await postgrest_get(
            "acord_eval_results",
            f"select=*&job_id=eq.{quote(str(latest_job.get('id')), safe='')}&order=eval_set.asc",
        )
        gate_snapshot = _quality_gate_snapshot_from_eval_rows(eval_rows)

    return {
        "run": run,
        "confidence_evaluation": confidence_eval,
        "latest_training_job": latest_job,
        "latest_eval_results": eval_rows,
        "quality_gate_snapshot": gate_snapshot,
    }


def _read_tail_text(path: Path, max_lines: int) -> str:
    """
    Read last N lines without loading the entire file.

    Note: This is a best-effort utility for log tails, not a strict performance-optimized implementation.
    """
    try:
        stat = path.stat()
    except Exception:
        return ""

    if stat.st_size <= 2_000_000:  # 2MB: small enough to read fully
        try:
            return "\n".join(path.read_text(encoding="utf-8", errors="replace").splitlines()[-max_lines:])
        except Exception:
            return ""

    # For larger files, read backwards in blocks until we have enough lines.
    # We'll accumulate bytes blocks then decode once.
    lines: list[str] = []
    chunk_size = 16_384
    with path.open("r", encoding="utf-8", errors="replace") as f:
        f.seek(0, os.SEEK_END)
        end = f.tell()
        pos = end
        while pos > 0 and len(lines) <= max_lines:
            read_size = min(chunk_size, pos)
            pos -= read_size
            f.seek(pos)
            chunk = f.read(read_size)
            # prepend chunk lines to keep correct ordering
            lines = (chunk.splitlines() + lines)[- (max_lines + 1):]
        return "\n".join(lines[-max_lines:])


def _progress_from_log_tail(log_text: str, status: str) -> Optional[int]:
    """
    Best-effort progress extraction from training output.

    We look for patterns like: "23%|", " 23%|" etc.
    If not present, fall back to stage markers (start training / running evaluation / adapter saved).
    """
    if not log_text:
        return None

    if status in {"completed", "failed"}:
        return 100

    percent_matches = re.findall(r"(\d{1,3})%\s*\|", log_text)
    if not percent_matches:
        percent_matches = re.findall(r"(\d{1,3})%", log_text)

    last_pct: Optional[int] = None
    if percent_matches:
        try:
            last_pct = int(percent_matches[-1])
        except Exception:
            last_pct = None

    stage_pct: Optional[int] = None
    if "Starting QLoRA training" in log_text:
        stage_pct = 35
    if "Running evaluation" in log_text or "Running evaluation with fine-tuned adapter" in log_text:
        stage_pct = max(stage_pct or 0, 85)
    if "Adapter saved to" in log_text:
        stage_pct = max(stage_pct or 0, 98)

    # Don't trust generic percent bars unless we see real training progress markers.
    # Many progress bars (e.g. dataset split/map) may show 100% before actual training begins.
    has_training_progress = bool(
        re.search(r"Epoch\s*\d+\s*/\s*\d+", log_text) or "global_step" in log_text or "loss" in log_text
    )
    if not has_training_progress:
        last_pct = None

    if last_pct is None:
        return stage_pct
    if stage_pct is None:
        return max(0, min(100, last_pct))
    return max(0, min(100, max(last_pct, stage_pct)))


@router.get("/admin/jobs/{job_id}/log")
async def get_job_log_tail(
    job_id: str,
    authorization: str | None = Header(default=None),
    tail: int = Query(default=200, ge=1, le=2000, description="Tail lines to return"),
):
    """
    Return the last lines of the fine-tuning log file and a best-effort progress percent.
    """
    await _require_admin(authorization)

    rows = await postgrest_get(
        "acord_training_jobs",
        f"select=id,created_at,updated_at,status,log_path,error&id=eq.{quote(job_id, safe='')}&limit=1",
    )

    if not rows:
        raise HTTPException(status_code=404, detail="Job not found")

    job = rows[0]
    status = job.get("status") or "queued"
    log_path = job.get("log_path")

    tail_text = ""
    progress_percent: Optional[int] = None

    if log_path:
        try:
            resolved = Path(str(log_path)).resolve()
            backend_dir = Path(__file__).resolve().parents[2]
            allowed_base = (backend_dir / "fine_tuning" / "runs").resolve()

            # Only allow reading logs produced by our training runner.
            if not any(p == allowed_base for p in resolved.parents) and resolved != allowed_base:
                raise HTTPException(status_code=403, detail="Log path is not allowed.")

            if resolved.exists():
                tail_text = _read_tail_text(resolved, max_lines=tail)
                progress_percent = _progress_from_log_tail(tail_text, status=str(status))
            else:
                tail_text = f"[log missing] file not found at: {resolved}"
        except HTTPException:
            raise
        except Exception as exc:
            tail_text = f"[log tail unavailable] {exc}"
            progress_percent = _progress_from_log_tail(tail_text, status=str(status))
    else:
        tail_text = "[log missing] log_path is not set for this job row."

    # Provide useful fallback text even when file tail cannot be read.
    if not tail_text.strip():
        err = str(job.get("error") or "").strip()
        if err:
            tail_text = f"[no log tail] Using stored job error:\n{err}"
        else:
            tail_text = "[no log tail] Job log is empty or inaccessible."

    return {
        "job_id": job_id,
        "status": status,
        "updated_at": job.get("updated_at"),
        "progress_percent": progress_percent,
        "tail_text": tail_text,
        "error": job.get("error"),
    }
