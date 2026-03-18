"""
generate_docs.py  –  Activity Logs Full Documentation generator
Produces Activity_Logs_Full_Documentation_FINAL.docx with real embedded diagrams.

Requirements (install once):
    pip install python-docx matplotlib pillow
"""

import io
import os
import textwrap
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import matplotlib.patheffects as pe
import numpy as np
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
C_NAVY   = "#0D1B2A"
C_BLUE   = "#1565C0"
C_LBLUE  = "#42A5F5"
C_TEAL   = "#00796B"
C_GREEN  = "#2E7D32"
C_ORANGE = "#E65100"
C_PURPLE = "#6A1B9A"
C_RED    = "#C62828"
C_GREY   = "#546E7A"
C_BG     = "#F5F7FA"
C_WHITE  = "#FFFFFF"

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 1 – System Architecture
# ═══════════════════════════════════════════════════════════════
def draw_architecture() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 9), facecolor=C_BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 9); ax.axis("off")
    ax.set_facecolor(C_BG)

    def box(x, y, w, h, label, sub="", bg=C_BLUE, fg=C_WHITE, fs=10):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                               linewidth=1.5, edgecolor=bg,
                               facecolor=bg + "22" if fg == C_WHITE else bg)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.15 if sub else 0), label,
                ha="center", va="center", fontsize=fs, fontweight="bold",
                color=bg, wrap=True)
        if sub:
            ax.text(x + w/2, y + h/2 - 0.22, sub,
                    ha="center", va="center", fontsize=7.5, color=C_GREY)

    def arrow(x1, y1, x2, y2, color=C_GREY, label=""):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=1.5))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx+0.05, my, label, fontsize=7.5, color=color)

    # Layers ──────────────────────────────────────────────────
    # Client layer
    ax.text(0.3, 8.6, "CLIENT LAYER", fontsize=8, color=C_GREY, style="italic")
    box(0.5, 7.8, 2.8, 0.7, "Browser / Next.js 14", "Activity.tsx  (Auth + System tabs)", C_TEAL)
    box(3.5, 7.8, 2.8, 0.7, "Mobile / API Consumer", "Bearer JWT token", C_TEAL)

    # API Gateway
    ax.text(0.3, 7.3, "API GATEWAY  (FastAPI + Uvicorn)", fontsize=8, color=C_GREY, style="italic")
    box(0.5, 6.5, 5.8, 0.7, "/api/activity/system  (GET)", "RateLimiter · JWT verify · RBAC", C_BLUE)

    # Backend services
    ax.text(0.3, 6.05, "BACKEND SERVICES", fontsize=8, color=C_GREY, style="italic")
    box(0.5, 5.0, 1.7, 0.9, "supabase.py", "insert_audit_log()\nPII scrub · SHA-256", C_NAVY)
    box(2.4, 5.0, 1.7, 0.9, "logger/", "Pass-1 field scrub\nPass-2 Presidio NLP", C_NAVY)
    box(4.3, 5.0, 1.7, 0.9, "SHAP Engine", "generate_shap_\nreasoning()", C_PURPLE)

    # Database layer
    ax.text(0.3, 4.55, "DATABASE  (Supabase / PostgreSQL)", fontsize=8, color=C_GREY, style="italic")
    box(0.5, 3.4, 2.5, 1.0, "audit_logs", "sequence_num · chain_hash\nshap_values · model_id", C_GREEN)
    box(3.2, 3.4, 2.5, 1.0, "auth_audit", "action_code · outcome_code\nintegrity_hash", C_GREEN)
    box(5.9, 3.4, 2.5, 1.0, "DB Triggers", "prevent_modification()\ncompute_chain_hash()", C_ORANGE)

    # Compliance
    ax.text(0.3, 3.0, "COMPLIANCE STANDARDS", fontsize=8, color=C_GREY, style="italic")
    for i, (lbl, col) in enumerate([("EU AI Act\nArt.12-13", C_RED),
                                     ("SOC2\nCC7.2 / CC9.1", C_BLUE),
                                     ("NAIC\nAI Bulletin", C_TEAL),
                                     ("ATNA\nCodes", C_PURPLE)]):
        box(0.5 + i*2.1, 1.9, 1.8, 0.9, lbl, "", col, C_WHITE, 8)

    # CI/CD (right column)
    ax.text(9.5, 8.6, "CI / CD  (.github/workflows)", fontsize=8, color=C_GREY, style="italic")
    box(9.5, 7.7, 4.0, 0.7, "backend-ci.yml", "lint · pytest · sonar · docker · ACR push", C_ORANGE)
    box(9.5, 6.8, 4.0, 0.7, "frontend-ci.yml", "eslint · build · Vercel deploy", C_ORANGE)
    box(9.5, 5.9, 4.0, 0.7, "db-migrate.yml", "supabase db push · migrations", C_ORANGE)
    box(9.5, 5.0, 4.0, 0.7, "deploy-env.yml", "dev → staging → production", C_ORANGE)

    # Arrows (left column)
    arrow(1.9, 7.8, 1.9, 7.2, C_TEAL, "HTTPS")
    arrow(4.9, 7.8, 4.9, 7.2, C_TEAL, "JWT")
    arrow(3.4, 6.5, 1.35, 5.9, C_BLUE, "call")
    arrow(3.4, 6.5, 3.25, 5.9, C_BLUE)
    arrow(3.4, 6.5, 5.15, 5.9, C_BLUE)
    arrow(1.35, 5.0, 1.75, 4.4, C_NAVY)
    arrow(3.25, 5.0, 3.25, 4.4, C_NAVY)
    arrow(5.15, 5.0, 7.15, 4.4, C_PURPLE, "SHAP→reasoning")

    ax.set_title("NeuraSence Cloud – Activity Logs System Architecture",
                 fontsize=13, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 2 – Database Schema (ER-style)
# ═══════════════════════════════════════════════════════════════
def draw_db_schema() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 8), facecolor=C_BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis("off")

    def entity(ax, x, y, title, cols, color=C_BLUE, width=4.0):
        row_h = 0.36
        total_h = 0.5 + row_h * len(cols)
        # Header
        hdr = FancyBboxPatch((x, y + total_h - 0.5), width, 0.5,
                              boxstyle="round,pad=0.05", linewidth=1.5,
                              edgecolor=color, facecolor=color)
        ax.add_patch(hdr)
        ax.text(x + width/2, y + total_h - 0.25, title,
                ha="center", va="center", fontsize=9.5,
                fontweight="bold", color=C_WHITE)
        # Rows
        for i, (pk, cname, ctype, note) in enumerate(cols):
            ry = y + total_h - 0.5 - row_h * (i + 1)
            bg = "#E3F2FD" if pk else C_WHITE
            ax.add_patch(FancyBboxPatch((x, ry), width, row_h - 0.04,
                                         boxstyle="square,pad=0",
                                         linewidth=0.5,
                                         edgecolor="#BDBDBD", facecolor=bg))
            lbl = ("[PK] " if pk == "PK" else "[FK] " if pk == "FK" else "     ") + cname
            ax.text(x + 0.15, ry + row_h/2 - 0.02, lbl,
                    va="center", fontsize=7.5,
                    color=C_NAVY if not pk else color, fontweight="bold" if pk else "normal")
            ax.text(x + width - 0.1, ry + row_h/2 - 0.02, ctype,
                    ha="right", va="center", fontsize=7, color=C_GREY)
            if note:
                ax.text(x + width/2, ry + 0.05, note,
                        ha="center", va="bottom", fontsize=6.2, color=C_ORANGE)
        # Border
        ax.add_patch(FancyBboxPatch((x, y), width, total_h,
                                     boxstyle="round,pad=0.05",
                                     linewidth=1.5, edgecolor=color,
                                     facecolor="none"))
        return total_h

    # audit_logs table
    audit_cols = [
        ("PK", "id",             "uuid",        ""),
        ("FK", "user_id",        "uuid",        "→ auth.users(id)"),
        ("",   "action",         "text NOT NULL",""),
        ("",   "resource_type",  "text NOT NULL",""),
        ("",   "resource_id",    "text",        ""),
        ("",   "details",        "jsonb",       ""),
        ("",   "ip_address",     "text",        ""),
        ("",   "user_agent",     "text",        ""),
        ("",   "previous_value", "jsonb",       "PII-scrubbed"),
        ("",   "new_value",      "jsonb",       "PII-scrubbed"),
        ("",   "model_id",       "text",        "AI model identifier"),
        ("",   "prediction",     "jsonb",       "model output"),
        ("",   "shap_values",    "jsonb",       "{feature: float}"),
        ("",   "reasoning",      "text",        "auto-generated NL"),
        ("",   "integrity_hash", "text",        "SHA-256"),
        ("",   "sequence_num",   "bigint IDENTITY","monotonic"),
        ("",   "chain_hash",     "text",        "SHA-256(prev∥self)"),
        ("",   "created_at",     "timestamptz", "default now()"),
    ]
    h1 = entity(ax, 0.3, 0.2, "public.audit_logs", audit_cols, C_GREEN, 4.6)

    # auth_audit table
    auth_cols = [
        ("PK", "id",             "uuid",        ""),
        ("FK", "user_id",        "uuid NOT NULL","→ auth.users(id)"),
        ("",   "email",          "text",        ""),
        ("",   "role",           "text",        ""),
        ("",   "event",          "text",        "login/logout/…"),
        ("",   "action_code",    "text",        "C/R/U/D/E (ATNA)"),
        ("",   "outcome_code",   "integer",     "0/4/8/12 (ATNA)"),
        ("",   "resource_type",  "text",        ""),
        ("",   "resource_id",    "text",        ""),
        ("",   "integrity_hash", "text",        "SHA-256"),
        ("",   "created_at",     "timestamptz", "default now()"),
    ]
    entity(ax, 5.3, 1.8, "public.auth_audit", auth_cols, C_TEAL, 4.4)

    # auth.users (external)
    user_cols = [
        ("PK", "id",             "uuid",        ""),
        ("",   "email",          "text",        ""),
        ("",   "role",           "text",        ""),
        ("",   "created_at",     "timestamptz", ""),
    ]
    entity(ax, 9.9, 4.5, "auth.users  (Supabase)", user_cols, C_NAVY, 3.8)

    # DB Triggers
    ax.text(9.9, 4.2, "DB TRIGGERS", fontsize=9, fontweight="bold", color=C_ORANGE)
    for i, t in enumerate([
        "prevent_audit_modification()",
        "compute_audit_chain_hash()",
    ]):
        ax.text(9.9, 3.85 - i*0.35, f"▸ {t}", fontsize=8, color=C_ORANGE)

    # Relationships
    ax.annotate("", xy=(5.3, 4.0), xytext=(4.9, 4.0),
                arrowprops=dict(arrowstyle="-|>", color=C_GREEN, lw=1.5))
    ax.text(5.05, 4.1, "user_id FK", fontsize=7, color=C_GREEN)

    ax.annotate("", xy=(9.9, 5.5), xytext=(9.7, 5.5),
                arrowprops=dict(arrowstyle="-|>", color=C_TEAL, lw=1.5))
    ax.text(8.6, 5.6, "user_id FK", fontsize=7, color=C_TEAL)

    ax.set_title("Activity Logs – Database Schema (Entity Relationship)",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    fig.tight_layout(pad=0.5)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 3 – Audit Log Data Flow
# ═══════════════════════════════════════════════════════════════
def draw_data_flow() -> bytes:
    fig, ax = plt.subplots(figsize=(13, 7), facecolor=C_BG)
    ax.set_xlim(0, 13); ax.set_ylim(0, 7); ax.axis("off")

    STEPS = [
        ("Route Handler\n(FastAPI)", C_BLUE,    1.0),
        ("verify_user()\nverify_admin()", C_TEAL, 3.0),
        ("Perform\nBusiness Action", C_GREY,    5.0),
        ("insert_audit_log()\nsupabase.py", C_NAVY, 7.0),
        ("_scrub_audit_value()\nPII Removal", C_PURPLE, 9.0),
        ("SHA-256\nintegrity_hash", C_ORANGE,   11.0),
    ]

    y = 4.2
    for label, color, x in STEPS:
        circ = plt.Circle((x, y), 0.65, color=color + "33",
                           ec=color, lw=2, zorder=3)
        ax.add_patch(circ)
        ax.text(x, y, label, ha="center", va="center",
                fontsize=7.8, fontweight="bold", color=color, zorder=4)

    # Arrows between steps
    for i in range(len(STEPS) - 1):
        x1 = STEPS[i][2] + 0.65
        x2 = STEPS[i+1][2] - 0.65
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="-|>", color=C_GREY, lw=1.5))

    # DB step (below)
    db_steps = [
        ("PostgREST\nINSERT", C_GREEN, 1.5),
        ("DB Trigger\ncompute_chain_hash()", C_ORANGE, 4.5),
        ("Immutable Row\n+ chain_hash stored", C_GREEN, 7.5),
        ("verify_audit_ledger()\nintegrity check", C_TEAL, 10.5),
    ]
    y2 = 1.9
    for label, color, x in db_steps:
        rect = FancyBboxPatch((x-1.1, y2-0.55), 2.2, 1.1,
                               boxstyle="round,pad=0.1", linewidth=1.5,
                               edgecolor=color, facecolor=color + "22")
        ax.add_patch(rect)
        ax.text(x, y2, label, ha="center", va="center",
                fontsize=7.8, color=color, fontweight="bold")

    for i in range(len(db_steps) - 1):
        x1 = db_steps[i][2] + 1.1
        x2 = db_steps[i+1][2] - 1.1
        ax.annotate("", xy=(x2, y2), xytext=(x1, y2),
                    arrowprops=dict(arrowstyle="-|>", color=C_GREEN, lw=1.5))

    # Vertical connector from SHA step down to DB layer
    ax.annotate("", xy=(1.5, y2 + 0.55), xytext=(11.0, y - 0.65),
                arrowprops=dict(arrowstyle="-|>",
                                connectionstyle="arc3,rad=0.25",
                                color=C_ORANGE, lw=1.5))
    ax.text(7.5, 3.1, "INSERT with integrity_hash", fontsize=7.5,
            color=C_ORANGE, style="italic")

    # Labels
    ax.text(0.3, 4.95, "APPLICATION LAYER", fontsize=8, color=C_GREY, style="italic")
    ax.text(0.3, 2.65, "DATABASE LAYER", fontsize=8, color=C_GREY, style="italic")

    ax.set_title("Activity Log Data Flow – Request → Immutable Ledger",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 4 – Sequence Diagram (API call)
# ═══════════════════════════════════════════════════════════════
def draw_sequence() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 9), facecolor=C_BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 9); ax.axis("off")

    actors = [
        ("Browser\n(Activity.tsx)", 1.2,  C_TEAL),
        ("FastAPI\nRoute",          3.5,  C_BLUE),
        ("supabase.py\nCore",       6.0,  C_NAVY),
        ("logger/\nPII+SHAP",       8.5,  C_PURPLE),
        ("PostgreSQL\n+ Triggers",  11.5, C_GREEN),
    ]

    TOP = 8.5
    BOT = 0.3

    # Lifelines
    for name, x, col in actors:
        rect = FancyBboxPatch((x - 0.7, TOP - 0.05), 1.4, 0.65,
                               boxstyle="round,pad=0.07", linewidth=1.5,
                               edgecolor=col, facecolor=col)
        ax.add_patch(rect)
        ax.text(x, TOP + 0.27, name, ha="center", va="center",
                fontsize=8, fontweight="bold", color=C_WHITE)
        ax.plot([x, x], [TOP - 0.05, BOT], color=col, lw=1, ls="--", alpha=0.5)

    def msg(y, x1, x2, label, color=C_GREY, ret=False):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=color, lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        mx = (x1 + x2) / 2
        offset = 0.12 if x1 < x2 else -0.12
        ax.text(mx, y + 0.13, label, ha="center", fontsize=7.5, color=color)

    def act(x, y, h, col):
        ax.add_patch(FancyBboxPatch((x - 0.12, y - h), 0.24, h,
                                     boxstyle="square,pad=0",
                                     facecolor=col + "55", edgecolor=col, lw=1))

    # Sequence messages (y from top to bottom)
    msg(8.0, 1.2, 3.5, "GET /api/activity/system?page=0", C_TEAL)
    act(3.5, 8.0, 0.4, C_BLUE)

    msg(7.5, 3.5, 6.0, "verify_user(JWT)", C_BLUE)
    act(6.0, 7.5, 0.3, C_NAVY)
    msg(7.1, 6.0, 3.5, "UserContext {role, tenant}", C_NAVY, ret=True)

    msg(6.7, 3.5, 6.0, "postgrest_get(audit_logs, filters)", C_BLUE)
    act(6.0, 6.7, 0.3, C_NAVY)
    msg(6.3, 6.0, 11.5, "SELECT * … LIMIT 25", C_NAVY)
    act(11.5, 6.3, 0.25, C_GREEN)
    msg(5.9, 11.5, 6.0, "[{log rows}]", C_GREEN, ret=True)
    msg(5.5, 6.0, 3.5, "logs[]", C_NAVY, ret=True)

    msg(5.1, 3.5, 1.2, "200 OK {logs, page, has_more}", C_BLUE, ret=True)

    # --- INSERT flow ---
    ax.text(0.3, 4.75, "── INSERT FLOW (any route action) ──", fontsize=8,
            color=C_GREY, style="italic")

    msg(4.5, 3.5, 6.0, "insert_audit_log(action, resource, …)", C_BLUE)
    act(6.0, 4.5, 0.35, C_NAVY)

    msg(4.05, 6.0, 8.5, "_scrub_audit_value(details)", C_NAVY)
    act(8.5, 4.05, 0.25, C_PURPLE)
    msg(3.7, 8.5, 6.0, "sanitised_dict", C_PURPLE, ret=True)

    msg(3.35, 6.0, 8.5, "generate_shap_reasoning()", C_NAVY)
    act(8.5, 3.35, 0.25, C_PURPLE)
    msg(3.0, 8.5, 6.0, "reasoning_text", C_PURPLE, ret=True)

    msg(2.65, 6.0, 6.0, "SHA-256(action∥user∥ts∥…) → integrity_hash", C_ORANGE)
    ax.text(6.3, 2.75, "SHA-256 computed", fontsize=7, color=C_ORANGE, style="italic")

    msg(2.3, 6.0, 11.5, "INSERT INTO audit_logs (…, integrity_hash)", C_NAVY)
    act(11.5, 2.3, 0.35, C_GREEN)
    msg(1.85, 11.5, 11.5, "Trigger: compute_audit_chain_hash()", C_ORANGE)
    ax.annotate("", xy=(11.5, 1.55), xytext=(11.5, 1.85),
                arrowprops=dict(arrowstyle="-|>", color=C_ORANGE, lw=1))
    ax.text(11.65, 1.7, "chain_hash written", fontsize=7, color=C_ORANGE)
    msg(1.4, 11.5, 6.0, "row inserted (immutable)", C_GREEN, ret=True)

    ax.set_title("Activity Logs – API & Insert Sequence Diagram",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 5 – Role-Based Access Control
# ═══════════════════════════════════════════════════════════════
def draw_rbac() -> bytes:
    fig, ax = plt.subplots(figsize=(12, 6), facecolor=C_BG)
    ax.set_xlim(0, 12); ax.set_ylim(0, 6); ax.axis("off")

    roles = [
        ("global_admin", C_RED,    1.0),
        ("admin",        C_ORANGE, 3.0),
        ("user",         C_BLUE,   5.5),
        ("viewer",       C_TEAL,   7.5),
        ("guest",        C_GREY,   10.0),
    ]

    perms = {
        "global_admin": ["All logs (no filter)",   "Export CSV",   "All tenants",    "All resources"],
        "admin":        ["Non-admin logs (RLS)",    "Export CSV",   "Own tenant",     "All resources"],
        "user":         ["Own logs only",           "Export CSV",   "Own tenant",     "Own resources"],
        "viewer":       ["Own logs only",           "Read-only",    "Own tenant",     "Own resources"],
        "guest":        ["403 Forbidden",           "—",            "—",              "—"],
    }

    for role, col, x in roles:
        # Role box
        rect = FancyBboxPatch((x - 0.6, 4.6), 1.2, 0.7,
                               boxstyle="round,pad=0.08", linewidth=2,
                               edgecolor=col, facecolor=col)
        ax.add_patch(rect)
        ax.text(x, 4.95, role, ha="center", va="center",
                fontsize=8.5, fontweight="bold", color=C_WHITE)

        # Permission rows
        for i, perm in enumerate(perms[role]):
            icon = "✓" if perm not in ("—", "403 Forbidden") else ("✗" if perm == "403 Forbidden" else "–")
            colour = C_GREEN if icon == "✓" else (C_RED if icon == "✗" else C_GREY)
            ax.text(x, 4.0 - i * 0.65, f"{icon}  {perm}",
                    ha="center", va="center", fontsize=7.5, color=colour)

    # Column headers
    for i, h in enumerate(["Scope", "Export", "Tenant", "Resources"]):
        ax.text(11.4, 4.0 - i * 0.65, h, ha="right", va="center",
                fontsize=7.5, color=C_GREY, style="italic")

    # Horizontal divider
    ax.axhline(4.45, color=C_GREY, lw=0.5, ls="--", alpha=0.4)

    ax.set_title("Role-Based Access Control  –  Activity Logs Visibility",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 6 – Cryptographic Ledger Chain
# ═══════════════════════════════════════════════════════════════
def draw_ledger() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 5), facecolor=C_BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 5); ax.axis("off")

    blocks = [
        ("GENESIS\nblock",  "—",              "SHA-256\n('GENESIS'∥H₀)", C_ORANGE),
        ("Row #1\nseq=1",   "H₀=SHA256(…)",   "chain₁=SHA256\n(genesis∥H₁)", C_GREEN),
        ("Row #2\nseq=2",   "H₁=SHA256(…)",   "chain₂=SHA256\n(chain₁∥H₂)", C_GREEN),
        ("Row #3\nseq=3",   "H₂=SHA256(…)",   "chain₃=SHA256\n(chain₂∥H₃)", C_GREEN),
        ("Row N\nseq=N",    "H_{N-1}",        "chainN=SHA256\n(chain_{N-1}∥HN)", C_BLUE),
    ]

    bw = 2.2
    gap = 0.5
    y0 = 1.2
    bh = 2.2
    xs = [0.4 + i * (bw + gap) for i in range(len(blocks))]

    for i, (title, integ, chain, col) in enumerate(blocks):
        x = xs[i]
        # Block outer
        ax.add_patch(FancyBboxPatch((x, y0), bw, bh,
                                     boxstyle="round,pad=0.1", linewidth=2,
                                     edgecolor=col, facecolor=col + "18"))
        ax.text(x + bw/2, y0 + bh - 0.3, title,
                ha="center", va="center", fontsize=8.5,
                fontweight="bold", color=col)
        ax.text(x + bw/2, y0 + bh/2 + 0.0, f"integrity:\n{integ}",
                ha="center", va="center", fontsize=7, color=C_NAVY)
        ax.text(x + bw/2, y0 + 0.4, chain,
                ha="center", va="center", fontsize=7,
                color=col, fontweight="bold")

        if i < len(blocks) - 1:
            ax.annotate("", xy=(xs[i+1], y0 + bh/2),
                        xytext=(x + bw, y0 + bh/2),
                        arrowprops=dict(arrowstyle="-|>", color=col, lw=2))

    # "…" between last two
    ax.text((xs[-2] + xs[-1]) / 2, y0 + bh/2, "…",
            ha="center", va="center", fontsize=18, color=C_GREY)

    # Tamper note
    ax.text(7.0, 0.5,
            "WARNING: Removing, inserting, or reordering any row breaks all subsequent chain_hashes -> tampering is detectable",
            ha="center", fontsize=8.5, color=C_RED,
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#FFEBEE", edgecolor=C_RED))

    ax.set_title("Cryptographic Audit Ledger – Chain-Hash Linking",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 7 – PII Scrubbing Pipeline
# ═══════════════════════════════════════════════════════════════
def draw_pii() -> bytes:
    fig, ax = plt.subplots(figsize=(13, 5), facecolor=C_BG)
    ax.set_xlim(0, 13); ax.set_ylim(0, 5); ax.axis("off")

    stages = [
        ("Raw\nPayload",        C_RED,    0.8),
        ("Pass-1\nField-Name\nKeyword Match", C_ORANGE, 3.0),
        ("Pass-2\nPresidio NLP\nContent Scan", C_PURPLE, 5.8),
        ("Sanitised\nPayload",  C_GREEN,  8.6),
        ("audit_logs\nINSERT",  C_NAVY,   10.8),
    ]

    y = 2.6
    for label, col, x in stages:
        rect = FancyBboxPatch((x - 0.9, y - 0.75), 1.8, 1.5,
                               boxstyle="round,pad=0.1", linewidth=2,
                               edgecolor=col, facecolor=col + "22")
        ax.add_patch(rect)
        ax.text(x, y, label, ha="center", va="center",
                fontsize=8.5, color=col, fontweight="bold")

    for i in range(len(stages) - 1):
        x1 = stages[i][2] + 0.9
        x2 = stages[i+1][2] - 0.9
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="-|>", color=C_GREY, lw=1.5))

    # Redacted field list
    fields = ["password", "secret", "token", "api_key", "email",
              "ssn", "credit_card", "dob", "full_name", "phone"]
    ax.text(3.0, 4.6, "Redacted field names:", fontsize=7.5, color=C_ORANGE, fontweight="bold")
    ax.text(3.0, 4.25, "  ·  ".join(fields), fontsize=7.2, color=C_ORANGE)

    # Presidio entities
    ax.text(5.8, 4.6, "Presidio entities detected:", fontsize=7.5, color=C_PURPLE, fontweight="bold")
    ax.text(5.8, 4.25, "PERSON · EMAIL_ADDRESS · PHONE · CREDIT_CARD · US_SSN · IBAN · LOCATION",
            fontsize=7.2, color=C_PURPLE)

    ax.set_title("PII Scrubbing Pipeline – Two-Pass Protection",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 8 – SHAP AI Explainability
# ═══════════════════════════════════════════════════════════════
def draw_shap() -> bytes:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5), facecolor=C_BG)
    fig.patch.set_facecolor(C_BG)

    # Left: SHAP waterfall bar chart
    ax = axes[0]
    ax.set_facecolor(C_BG)
    features = ["transaction_amount", "account_age", "hour_of_day",
                "device_trust", "geo_velocity", "txn_freq"]
    vals     = [0.42, -0.31, 0.18, -0.14, 0.09, -0.06]
    colors   = [C_GREEN if v > 0 else C_RED for v in vals]
    bars = ax.barh(features, vals, color=colors, edgecolor="white", height=0.55)
    ax.axvline(0, color=C_GREY, lw=1)
    for bar, val in zip(bars, vals):
        ax.text(val + (0.01 if val > 0 else -0.01), bar.get_y() + bar.get_height()/2,
                f"{val:+.2f}", va="center",
                ha="left" if val > 0 else "right",
                fontsize=8, color=C_NAVY, fontweight="bold")
    ax.set_xlabel("SHAP Value (impact on prediction)", fontsize=8, color=C_GREY)
    ax.set_title("SHAP Feature Importance\n(model: fraud-v3, prediction: high_risk 87%)",
                 fontsize=9, color=C_NAVY)
    ax.tick_params(labelsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # Right: Generated reasoning text
    ax2 = axes[1]
    ax2.set_facecolor(C_WHITE)
    ax2.axis("off")
    reasoning = textwrap.dedent("""\
        Model 'fraud-v3' predicted 'high_risk'
        (confidence 87%).

        Top contributing factors:

        1. transaction_amount  +0.42  ↑  (pushed toward high_risk)
        2. account_age         -0.31  ↓  (pushed away from high_risk)
        3. hour_of_day         +0.18  ↑  (pushed toward high_risk)
        4. device_trust        -0.14  ↓  (pushed away from high_risk)
        5. geo_velocity        +0.09  ↑  (pushed toward high_risk)

        This reasoning is stored verbatim in
        audit_logs.reasoning for compliance audit.
    """)
    ax2.text(0.05, 0.95, "Auto-Generated Reasoning (stored in audit_logs)",
             transform=ax2.transAxes, fontsize=9, fontweight="bold", color=C_NAVY, va="top")
    ax2.text(0.05, 0.80, reasoning,
             transform=ax2.transAxes, fontsize=8.5, color=C_NAVY,
             va="top", family="monospace",
             bbox=dict(boxstyle="round,pad=0.4", facecolor="#E8F5E9", edgecolor=C_GREEN))

    fig.suptitle("AI Explainability – SHAP Values & Audit Reasoning (EU AI Act Art.13)",
                 fontsize=11, fontweight="bold", color=C_NAVY)
    fig.tight_layout(pad=1.5)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 9 – CI/CD Pipeline
# ═══════════════════════════════════════════════════════════════
def draw_cicd() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 6), facecolor=C_BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 6); ax.axis("off")

    # Three tracks
    tracks = [
        ("BACKEND  (backend-ci.yml)",   C_BLUE,   4.8,
         ["Code Push\n/ PR", "flake8\nLint", "pytest\nCoverage", "SonarQube\nScan",
          "Docker\nBuild", "ACR\nPush", "Deploy\nDev/Staging/Prod"]),
        ("FRONTEND (frontend-ci.yml)",  C_TEAL,   3.0,
         ["Code Push\n/ PR", "ESLint", "Next.js\nBuild", "Vercel\nPreview", "Vercel\nProd"]),
        ("DATABASE (db-migrate.yml)",   C_GREEN,  1.2,
         ["Migration\nFile Push", "supabase\ndb push", "Validate\nMigrations", "Dev→Staging\n→Prod"]),
    ]

    for (track_name, col, y, steps) in tracks:
        ax.text(0.2, y + 0.45, track_name, fontsize=8.5,
                fontweight="bold", color=col)
        n = len(steps)
        xs = [1.2 + i * (12.5 / (n - 1)) for i in range(n)]
        for j, (x, step) in enumerate(zip(xs, steps)):
            rect = FancyBboxPatch((x - 0.55, y - 0.35), 1.1, 0.7,
                                   boxstyle="round,pad=0.07", linewidth=1.5,
                                   edgecolor=col, facecolor=col + "22")
            ax.add_patch(rect)
            ax.text(x, y, step, ha="center", va="center",
                    fontsize=7.2, color=col, fontweight="bold")
            if j < n - 1:
                ax.annotate("", xy=(xs[j+1] - 0.55, y),
                            xytext=(x + 0.55, y),
                            arrowprops=dict(arrowstyle="-|>", color=col, lw=1.2))

    ax.set_title("CI/CD Pipeline  –  Backend · Frontend · Database Migration",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 10 – Frontend UI wireframe
# ═══════════════════════════════════════════════════════════════
def draw_frontend_ui() -> bytes:
    fig, ax = plt.subplots(figsize=(13, 7), facecolor=C_WHITE)
    ax.set_xlim(0, 13); ax.set_ylim(0, 7); ax.axis("off")

    # Browser chrome
    ax.add_patch(FancyBboxPatch((0.2, 0.1), 12.6, 6.7,
                                 boxstyle="round,pad=0.1", linewidth=2,
                                 edgecolor="#BDBDBD", facecolor="#FAFAFA"))

    # Top bar
    ax.add_patch(plt.Rectangle((0.2, 6.3), 12.6, 0.5, color="#1565C0"))
    ax.text(6.5, 6.55, "NeuraSence Cloud  –  Activity Logs",
            ha="center", va="center", fontsize=11, color=C_WHITE, fontweight="bold")

    # Tabs
    for i, (tab, active) in enumerate([("Auth Events", True), ("System Events", False)]):
        x = 0.5 + i * 2.8
        col = C_WHITE if active else "#E3F2FD"
        ax.add_patch(FancyBboxPatch((x, 5.9), 2.5, 0.35,
                                     boxstyle="round,pad=0.05", linewidth=1.5,
                                     edgecolor=C_BLUE, facecolor=col))
        ax.text(x + 1.25, 6.075, tab, ha="center", va="center",
                fontsize=9, color=C_BLUE, fontweight="bold" if active else "normal")

    # Filter bar
    ax.add_patch(FancyBboxPatch((0.4, 5.3), 12.2, 0.5,
                                 boxstyle="round,pad=0.05", linewidth=1,
                                 edgecolor="#BDBDBD", facecolor="#EFF3FB"))
    ax.text(0.6, 5.55, "Filters:", fontsize=8, color=C_GREY, fontweight="bold")
    for i, flt in enumerate(["Event / Action", "Date From", "Date To", "Resource Type"]):
        x = 1.5 + i * 2.7
        ax.add_patch(FancyBboxPatch((x, 5.36), 2.3, 0.28,
                                     boxstyle="round,pad=0.03", linewidth=1,
                                     edgecolor="#BDBDBD", facecolor=C_WHITE))
        ax.text(x + 0.1, 5.5, flt, fontsize=7.2, color=C_GREY)

    for lbl, x, col in [("Apply", 12.1, C_BLUE), ("Clear", 11.5, C_GREY)]:
        ax.add_patch(FancyBboxPatch((x - 0.4, 5.37), 0.78, 0.26,
                                     boxstyle="round,pad=0.04", linewidth=1,
                                     edgecolor=col, facecolor=col))
        ax.text(x + 0.0, 5.5, lbl, ha="center", fontsize=7.5,
                color=C_WHITE, fontweight="bold")

    # Table header
    cols_w = [1.6, 1.8, 1.5, 1.0, 1.6, 2.5, 1.2]
    cols_h = ["When", "Email / User", "Role", "Outcome", "Event", "Resource", "IP"]
    x = 0.4
    ax.add_patch(plt.Rectangle((0.4, 4.95), 12.2, 0.3, color="#1565C0"))
    for w, h in zip(cols_w, cols_h):
        ax.text(x + w/2, 5.10, h, ha="center", va="center",
                fontsize=7.5, color=C_WHITE, fontweight="bold")
        x += w

    # Table rows (mock data)
    rows = [
        ["2026-03-17 12:51", "alice@corp.com", "admin",     "✓ 0", "login",         "auth_session:a1b2", "192.168.1.5"],
        ["2026-03-17 12:49", "bob@corp.com",   "user",      "✓ 0", "model_predict", "model:fraud-v3",    "10.0.0.22"],
        ["2026-03-17 12:48", "carol@corp.com", "viewer",    "✗ 8", "login",         "auth_session:c3d4", "172.16.0.1"],
        ["2026-03-17 12:45", "dave@corp.com",  "global_adm","✓ 0", "device_create", "device:d5e6",       "192.168.2.1"],
    ]
    for ri, row in enumerate(rows):
        y = 4.6 - ri * 0.35
        bg = "#F5F7FA" if ri % 2 == 0 else C_WHITE
        ax.add_patch(plt.Rectangle((0.4, y - 0.15), 12.2, 0.34, color=bg))
        x = 0.4
        for ci, (val, w) in enumerate(zip(row, cols_w)):
            col = C_RED if "✗" in str(val) else (C_GREEN if "✓" in str(val) else C_NAVY)
            ax.text(x + w/2, y + 0.02, val, ha="center", va="center",
                    fontsize=7, color=col)
            x += w

    # Pagination + Export
    ax.text(0.6, 0.5, "← Previous", fontsize=8, color=C_BLUE)
    ax.text(6.5, 0.5, "Page 1 of 12", ha="center", fontsize=8, color=C_GREY)
    ax.text(11.2, 0.5, "Next →", fontsize=8, color=C_BLUE)
    ax.add_patch(FancyBboxPatch((12.2, 0.32), 0.9, 0.32,
                                 boxstyle="round,pad=0.04", linewidth=1,
                                 edgecolor=C_GREEN, facecolor=C_GREEN))
    ax.text(12.65, 0.48, "CSV", ha="center", fontsize=8, color=C_WHITE, fontweight="bold")

    ax.set_title("Activity Page UI – Auth Events Tab (Wireframe)",
                 fontsize=11, fontweight="bold", color=C_NAVY, pad=8)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 11 – Compliance Mapping
# ═══════════════════════════════════════════════════════════════
def draw_compliance() -> bytes:
    fig, ax = plt.subplots(figsize=(13, 6), facecolor=C_BG)
    ax.set_xlim(0, 13); ax.set_ylim(0, 6); ax.axis("off")

    standards = [
        ("EU AI Act\nArt.12-13", C_RED,    1.2,
         ["Immutable audit logs\n≥ 6 months retention",
          "AI decision\nexplainability (SHAP)",
          "Human oversight\nlog access",
          "High-risk AI\nrecord-keeping"]),
        ("SOC2\nCC7.2 / CC9.1", C_BLUE,   4.3,
         ["Audit log\ncollection",
          "Anomaly\ndetection support",
          "System monitoring\nlogs",
          "Change-tracking\n(prev/new value)"]),
        ("NAIC\nAI Bulletin", C_TEAL,     7.4,
         ["AI governance\ntrail",
          "Model ID +\nprediction stored",
          "Tenant isolation\n(RLS)",
          "Audit export\nfor regulators"]),
        ("ATNA\nStandard", C_PURPLE,      10.5,
         ["action_code:\nC/R/U/D/E",
          "outcome_code:\n0/4/8/12",
          "resource_type /\nresource_id",
          "user_id\n+ timestamp"]),
    ]

    for (std, col, x, items) in standards:
        ax.add_patch(FancyBboxPatch((x - 1.1, 4.0), 2.2, 1.0,
                                     boxstyle="round,pad=0.1", linewidth=2,
                                     edgecolor=col, facecolor=col))
        ax.text(x, 4.5, std, ha="center", va="center",
                fontsize=9, fontweight="bold", color=C_WHITE)
        for i, item in enumerate(items):
            ay = 3.4 - i * 0.8
            ax.add_patch(FancyBboxPatch((x - 0.95, ay - 0.3), 1.9, 0.65,
                                         boxstyle="round,pad=0.07", linewidth=1,
                                         edgecolor=col, facecolor=col + "18"))
            ax.text(x, ay + 0.02, item, ha="center", va="center",
                    fontsize=7.5, color=C_NAVY)
            ax.annotate("", xy=(x, ay + 0.35), xytext=(x, 4.0),
                        arrowprops=dict(arrowstyle="-", color=col, lw=1, ls="--"))

    ax.set_title("Regulatory Compliance Mapping  –  EU AI Act · SOC2 · NAIC · ATNA",
                 fontsize=12, fontweight="bold", color=C_NAVY, pad=10)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# Word Document Builder
# ═══════════════════════════════════════════════════════════════

def hex_to_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            r, g, b = hex_to_rgb(color)
            run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_para(doc, text, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        r, g, b = hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)
    if size:
        run.font.size = Pt(size)
    return p


def add_image(doc, img_bytes: bytes, caption: str, width=Inches(6.2)):
    buf = io.BytesIO(img_bytes)
    doc.add_picture(buf, width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        r, g, b = hex_to_rgb(C_GREY)
        run.font.color.rgb = RGBColor(r, g, b)


def add_table(doc, headers, rows, header_bg=C_BLUE):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
    return table


def build_docx(diagrams: dict) -> str:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ═══════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ACTIVITY LOGS")
    run.bold = True; run.font.size = Pt(36)
    r, g, b = hex_to_rgb(C_NAVY); run.font.color.rgb = RGBColor(r, g, b)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Full Technical Documentation")
    run2.font.size = Pt(22)
    r, g, b = hex_to_rgb(C_BLUE); run2.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"NeuraSence Cloud  ·  Next.js 14 + FastAPI  ·  Supabase / PostgreSQL\n"
                 f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n"
                 f"Revision: v3.0  ·  Compliance: EU AI Act · SOC2 · NAIC · ATNA"
                 ).font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", 1, C_NAVY)
    toc_items = [
        ("1.", "Executive Summary",                   "3"),
        ("2.", "System Architecture",                 "4"),
        ("3.", "Database Schema",                     "5"),
        ("4.", "Data Flow",                           "6"),
        ("5.", "API Reference",                       "7"),
        ("6.", "Sequence Diagram",                    "8"),
        ("7.", "Role-Based Access Control (RBAC)",    "9"),
        ("8.", "Cryptographic Audit Ledger",          "10"),
        ("9.", "PII Scrubbing Pipeline",              "11"),
        ("10.","AI Explainability & SHAP Values",     "12"),
        ("11.","Frontend UI (Activity Page)",         "13"),
        ("12.","CI/CD Pipeline",                      "14"),
        ("13.","Compliance Mapping",                  "15"),
        ("14.","Migration Timeline",                  "16"),
        ("15.","Test Coverage",                       "17"),
    ]
    for num, name, pg in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"  {num}  {name}").font.size = Pt(11)
        p.add_run(f"{'.' * (55 - len(num) - len(name))}{pg}").font.size = Pt(11)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "1. Executive Summary", 1, C_NAVY)
    add_para(doc, (
        "The Activity Logs system is a compliance-grade, append-only audit ledger built into "
        "the NeuraSence Cloud platform. It captures every user action, system event, and AI "
        "decision in two immutable PostgreSQL tables (audit_logs and auth_audit), protected by "
        "SHA-256 integrity hashing and a cryptographic chain-hash ledger that detects any "
        "post-write tampering.\n\n"
        "The system satisfies three regulatory frameworks simultaneously:"
    ))
    for bullet in [
        "EU AI Act Articles 12-13  –  immutable AI decision records with SHAP explainability",
        "SOC2 CC7.2 / CC9.1  –  comprehensive audit log collection and system monitoring",
        "NAIC AI Bulletin  –  AI governance trail with model ID and tenant isolation",
        "ATNA Standard  –  action_code (C/R/U/D/E) and outcome_code (0/4/8/12) fields",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet).font.size = Pt(10)

    add_para(doc, "\nKey technical highlights:", bold=True)
    stats = [
        ("Two audit tables",      "auth_audit (auth events) + audit_logs (system events)"),
        ("Immutability",          "Database triggers block all UPDATE / DELETE operations"),
        ("Integrity",             "SHA-256 per row + chain-hash linking across rows"),
        ("AI Explainability",     "SHAP values + auto-generated natural-language reasoning"),
        ("PII Protection",        "Two-pass scrubbing: field-name keywords + Presidio NLP"),
        ("RBAC",                  "5-tier role system enforced at RLS and API level"),
        ("Frontend",              "Two-tab Activity page with filtering, pagination, CSV export"),
        ("CI/CD",                 "Automated lint, test, SonarQube, Docker, ACR, deploy pipeline"),
    ]
    add_table(doc, ["Feature", "Description"], stats, C_NAVY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "2. System Architecture", 1, C_NAVY)
    add_para(doc, (
        "The diagram below shows all major components and their relationships. "
        "The client layer (Next.js 14) communicates with FastAPI over HTTPS using Bearer JWT tokens. "
        "The backend validates tokens, enforces RBAC, then reads from or writes to Supabase / PostgreSQL "
        "through PostgREST. CI/CD pipelines automate testing and deployment across three environments."
    ))
    add_image(doc, diagrams["architecture"],
              "Figure 1 – NeuraSence Cloud Activity Logs System Architecture", Inches(6.4))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 3. DATABASE SCHEMA
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "3. Database Schema", 1, C_NAVY)
    add_para(doc, (
        "Two tables store audit data. audit_logs holds all system and AI events; "
        "auth_audit holds authentication events. Both use FK constraints to auth.users "
        "with ON DELETE CASCADE. All tables are append-only (see Section 8)."
    ))
    add_image(doc, diagrams["db_schema"],
              "Figure 2 – Entity Relationship Diagram (audit_logs + auth_audit + auth.users)", Inches(6.4))

    add_heading(doc, "3.1  audit_logs columns", 2, C_GREEN)
    al_cols = [
        ["id",             "uuid PK",          "Unique row identifier"],
        ["user_id",        "uuid FK",          "→ auth.users(id) ON DELETE CASCADE"],
        ["action",         "text NOT NULL",    "Verb describing the operation"],
        ["resource_type",  "text NOT NULL",    "Type of resource acted upon"],
        ["resource_id",    "text",             "ID of the specific resource"],
        ["details",        "jsonb",            "Freeform context data"],
        ["ip_address",     "text",             "Client IP captured from request"],
        ["user_agent",     "text",             "Browser / SDK identifier"],
        ["previous_value", "jsonb",            "Before state (PII-scrubbed)"],
        ["new_value",      "jsonb",            "After state (PII-scrubbed)"],
        ["model_id",       "text",             "AI model identifier (EU AI Act)"],
        ["prediction",     "jsonb",            "Model output / decision"],
        ["shap_values",    "jsonb",            "{feature_name: float} SHAP contributions"],
        ["reasoning",      "text",             "Auto-generated NL explanation"],
        ["integrity_hash", "text",             "SHA-256(action∥user_id∥ts∥…)"],
        ["sequence_num",   "bigint IDENTITY",  "Monotonically increasing insert order"],
        ["chain_hash",     "text",             "SHA-256(prev_chain∥integrity_hash)"],
        ["created_at",     "timestamptz",      "Default now(), set by DB"],
    ]
    add_table(doc, ["Column", "Type", "Purpose"], al_cols, C_GREEN)

    add_heading(doc, "3.2  auth_audit columns", 2, C_TEAL)
    aa_cols = [
        ["id",             "uuid PK",    "Unique row identifier"],
        ["user_id",        "uuid FK",    "→ auth.users(id) ON DELETE CASCADE"],
        ["email",          "text",       "User email at event time"],
        ["role",           "text",       "User role at event time"],
        ["event",          "text",       "login / logout / token_refresh / …"],
        ["action_code",    "text",       "ATNA: C=Create, R=Read, U=Update, D=Delete, E=Execute"],
        ["outcome_code",   "integer",    "ATNA: 0=Success, 4=Minor, 8=Serious, 12=Major"],
        ["resource_type",  "text",       "Resource type"],
        ["resource_id",    "text",       "Resource ID"],
        ["integrity_hash", "text",       "SHA-256 tamper-evidence hash"],
        ["created_at",     "timestamptz","Default now()"],
    ]
    add_table(doc, ["Column", "Type", "Purpose"], aa_cols, C_TEAL)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 4. DATA FLOW
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "4. Data Flow", 1, C_NAVY)
    add_para(doc, (
        "Every audit record follows a deterministic pipeline from the route handler through "
        "PII scrubbing, SHA-256 hashing, PostgREST insertion, and database trigger chaining."
    ))
    add_image(doc, diagrams["data_flow"],
              "Figure 3 – Activity Log Data Flow: Request to Immutable Ledger", Inches(6.2))

    flow_steps = [
        ["1", "Route Handler", "FastAPI endpoint receives request with Authorization header"],
        ["2", "verify_user()", "JWT validated; UserContext (role, tenant) extracted"],
        ["3", "Business Action", "Core logic executes (create device, run model, etc.)"],
        ["4", "insert_audit_log()", "Called with action, resource_type, details, shap_values"],
        ["5", "_scrub_audit_value()", "PII field names redacted from details/prev/new values"],
        ["6", "SHA-256 hash", "integrity_hash = SHA-256(action∥user_id∥resource∥ts∥…)"],
        ["7", "PostgREST INSERT", "Row written to audit_logs table"],
        ["8", "DB Trigger", "compute_audit_chain_hash() chains to previous row"],
        ["9", "Immutable Row", "prevent_audit_modification() blocks any future UPDATE/DELETE"],
    ]
    add_table(doc, ["Step", "Component", "Description"], flow_steps, C_BLUE)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 5. API REFERENCE
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "5. API Reference", 1, C_NAVY)
    add_heading(doc, "GET /api/activity/system", 2, C_BLUE)
    add_para(doc, "Returns a paginated list of system audit log entries. Requires Bearer JWT.")

    params = [
        ["page",          "integer", "No",  "0",    "Page index (0-based)"],
        ["action",        "string",  "No",  "—",    "Partial match on action field (ilike)"],
        ["resource_type", "string",  "No",  "—",    "Exact match on resource_type"],
        ["date_from",     "ISO-8601","No",  "—",    "Filter: created_at ≥ date_from"],
        ["date_to",       "ISO-8601","No",  "—",    "Filter: created_at ≤ date_to"],
    ]
    add_table(doc, ["Parameter", "Type", "Required", "Default", "Description"], params, C_NAVY)

    add_heading(doc, "Response Schema", 3, C_BLUE)
    resp = [
        ["logs",      "array",   "Array of SystemLogRow objects (max 25)"],
        ["page",      "integer", "Current page number"],
        ["page_size", "integer", "25 (fixed)"],
        ["has_more",  "boolean", "True if another page exists"],
    ]
    add_table(doc, ["Field", "Type", "Description"], resp, C_TEAL)

    add_heading(doc, "HTTP Status Codes", 3, C_BLUE)
    statuses = [
        ["200", "OK",           "Logs returned successfully"],
        ["401", "Unauthorized", "Missing or invalid JWT token"],
        ["403", "Forbidden",    "Guest role not permitted"],
        ["500", "Server Error", "Unexpected backend failure"],
    ]
    add_table(doc, ["Code", "Status", "Meaning"], statuses, C_GREY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 6. SEQUENCE DIAGRAM
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "6. Sequence Diagram", 1, C_NAVY)
    add_para(doc, (
        "The sequence diagram illustrates two flows: the read path (GET /api/activity/system) "
        "and the write path (insert_audit_log called after any business action)."
    ))
    add_image(doc, diagrams["sequence"],
              "Figure 4 – API & Insert Sequence Diagram", Inches(6.4))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 7. RBAC
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "7. Role-Based Access Control (RBAC)", 1, C_NAVY)
    add_para(doc, (
        "Access to audit logs is controlled at two layers: the FastAPI route checks the "
        "decoded JWT role, and Supabase RLS policies enforce row-level restrictions in the database."
    ))
    add_image(doc, diagrams["rbac"],
              "Figure 5 – RBAC Matrix: Role vs. Visibility Scope", Inches(6.2))

    rbac_data = [
        ["global_admin", "All logs (all tenants, all roles)", "Yes", "All tenants", "All"],
        ["admin",        "All logs for own tenant (RLS)",     "Yes", "Own tenant",  "All resources"],
        ["user",         "Own logs only",                     "Yes", "Own tenant",  "Own resources"],
        ["viewer",       "Own logs only (read-only)",         "Yes", "Own tenant",  "Own resources"],
        ["guest",        "403 Forbidden",                     "No",  "—",           "—"],
    ]
    add_table(doc, ["Role", "Visibility", "Export", "Tenant Scope", "Resources"],
              rbac_data, C_NAVY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 8. CRYPTOGRAPHIC LEDGER
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "8. Cryptographic Audit Ledger", 1, C_NAVY)
    add_para(doc, (
        "Each row in audit_logs cryptographically binds to its predecessor via chain_hash. "
        "Inserting, removing, or reordering any row invalidates all subsequent chain_hashes, "
        "making tampering immediately detectable by verify_audit_ledger()."
    ))
    add_image(doc, diagrams["ledger"],
              "Figure 6 – Cryptographic Chain-Hash Ledger", Inches(6.4))

    add_heading(doc, "Chain-Hash Algorithm", 2, C_GREEN)
    add_para(doc, "Each row's chain_hash is computed by the DB trigger before INSERT:", bold=True)
    p = doc.add_paragraph()
    p.add_run(
        "  chain_hash(N) = SHA-256( chain_hash(N-1) ∥ integrity_hash(N) )\n"
        "  chain_hash(1) = SHA-256( 'GENESIS'       ∥ integrity_hash(1) )"
    ).font.name = "Courier New"

    add_heading(doc, "Verification", 2, C_GREEN)
    p = doc.add_paragraph()
    p.add_run("  SELECT * FROM public.verify_audit_ledger() WHERE NOT is_valid;").font.name = "Courier New"
    add_para(doc, "Returns rows where stored_chain ≠ recomputed_chain → indicates tampering or data loss.")

    add_heading(doc, "Immutability Triggers", 2, C_ORANGE)
    imm = [
        ["prevent_audit_modification()", "audit_logs",  "BEFORE UPDATE OR DELETE", "Raises exception unconditionally"],
        ["prevent_audit_modification()", "auth_audit",  "BEFORE UPDATE OR DELETE", "Raises exception unconditionally"],
        ["compute_audit_chain_hash()",   "audit_logs",  "BEFORE INSERT",           "Computes and writes chain_hash"],
    ]
    add_table(doc, ["Function", "Table", "Trigger Event", "Effect"], imm, C_ORANGE)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 9. PII SCRUBBING
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "9. PII Scrubbing Pipeline", 1, C_NAVY)
    add_para(doc, (
        "Before any data is written to the audit tables, all field values pass through a "
        "two-pass PII scrubbing system to comply with GDPR and HIPAA requirements."
    ))
    add_image(doc, diagrams["pii"],
              "Figure 7 – Two-Pass PII Scrubbing Pipeline", Inches(6.2))

    add_heading(doc, "Pass-1: Field-Name Keyword Match (Synchronous)", 2, C_ORANGE)
    add_para(doc, "Redacted field names (case-insensitive):")
    fields_text = (
        "password · passwd · secret · token · api_key · authorization · email · ssn · "
        "social_security · credit_card · card_number · cvv · dob · date_of_birth · "
        "birth_date · full_name · first_name · last_name · mobile · phone · phone_number · address"
    )
    p = doc.add_paragraph()
    p.add_run(fields_text).font.size = Pt(9)

    add_heading(doc, "Pass-2: Presidio NLP Content Scan (Async)", 2, C_PURPLE)
    add_para(doc, "Detected entity types (Microsoft Presidio):")
    entities = "PERSON · EMAIL_ADDRESS · PHONE_NUMBER · CREDIT_CARD · US_SSN · IBAN_CODE · LOCATION"
    p = doc.add_paragraph()
    p.add_run(entities).font.size = Pt(9)
    add_para(doc, "Pass-2 runs in a dedicated ThreadPoolExecutor (non-blocking). "
                  "If Presidio is unavailable, the system degrades gracefully to Pass-1 only.")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 10. AI EXPLAINABILITY
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "10. AI Explainability & SHAP Values", 1, C_NAVY)
    add_para(doc, (
        "The EU AI Act (Art.13) requires that high-risk AI systems provide meaningful "
        "explanations of automated decisions. The Activity Logs system stores SHAP feature "
        "contributions alongside each AI prediction and auto-generates a natural-language "
        "reasoning string for every model inference event."
    ))
    add_image(doc, diagrams["shap"],
              "Figure 8 – SHAP Feature Importance & Auto-Generated Reasoning (EU AI Act Art.13)",
              Inches(6.4))

    add_heading(doc, "Fields Stored per AI Decision", 2, C_PURPLE)
    ai_fields = [
        ["model_id",    "text",  "Identifier of the model that made the prediction"],
        ["prediction",  "jsonb", "Full model output object {class, confidence, …}"],
        ["shap_values", "jsonb", "{feature_name: shap_float, …} for top-N features"],
        ["reasoning",   "text",  "Auto-generated: 'Model X predicted Y. Top factors: …'"],
    ]
    add_table(doc, ["Field", "Type", "Content"], ai_fields, C_PURPLE)

    add_heading(doc, "generate_shap_reasoning() Algorithm", 2, C_PURPLE)
    alg = [
        ["1", "Sort features by |shap_value| descending"],
        ["2", "Take top_n features (default: 5)"],
        ["3", "Format each as: feature_name (±value, ↑/↓)"],
        ["4", "Prefix with model prediction and confidence level"],
        ["5", "Return full string → stored in audit_logs.reasoning"],
    ]
    add_table(doc, ["Step", "Action"], alg, C_NAVY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 11. FRONTEND UI
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "11. Frontend UI (Activity Page)", 1, C_NAVY)
    add_para(doc, (
        "The Activity page (frontend/src/app-pages/Activity.tsx) provides a two-tab interface "
        "for viewing and exporting audit records."
    ))
    add_image(doc, diagrams["frontend_ui"],
              "Figure 9 – Activity Page Wireframe (Auth Events Tab)", Inches(6.4))

    add_heading(doc, "Tab 1 – Auth Events", 2, C_TEAL)
    add_para(doc, "Data source: Supabase RLS query on auth_audit table.")
    auth_cols2 = [
        ["When",     "Formatted created_at timestamp"],
        ["Email",    "User email"],
        ["Role",     "Role badge (highlights admin roles)"],
        ["Event",    "Event name (login, logout, token_refresh, …)"],
        ["Action",   "ATNA action_code (C/R/U/D/E)"],
        ["Resource", "resource_type:resource_id"],
        ["Outcome",  "OutcomeBadge: Success/Minor/Serious/Major"],
    ]
    add_table(doc, ["Column", "Content"], auth_cols2, C_TEAL)

    add_heading(doc, "Tab 2 – System Events", 2, C_BLUE)
    add_para(doc, "Data source: Backend API GET /api/activity/system.")
    sys_cols = [
        ["When",        "Formatted created_at timestamp"],
        ["Action",      "Action string (e.g. device_create, model_predict)"],
        ["Resource",    "resource_type:truncated_resource_id"],
        ["Model",       "ModelBadge with model_id (monospace)"],
        ["AI Reasoning","Inline ShapFactors (top 3 features with ↑/↓)"],
        ["Details",     "Collapsible JSON details"],
        ["IP",          "Client IP address"],
    ]
    add_table(doc, ["Column", "Content"], sys_cols, C_BLUE)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 12. CI/CD PIPELINE
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "12. CI/CD Pipeline", 1, C_NAVY)
    add_para(doc, (
        "Three GitHub Actions workflows automate quality gates and deployments."
    ))
    add_image(doc, diagrams["cicd"],
              "Figure 10 – CI/CD Pipeline: Backend · Frontend · DB Migrations", Inches(6.4))

    cicd_data = [
        ["backend-ci.yml",  "Push/PR to v*-dev/staging/v1", "flake8 · pytest · SonarQube · Docker · ACR · Deploy"],
        ["frontend-ci.yml", "Push/PR to v*-dev/staging/v1", "ESLint · next build · Vercel preview · Vercel prod"],
        ["db-migrate.yml",  "Migration file push",           "supabase db push · validate · dev→staging→prod"],
        ["deploy-env.yml",  "Push to v1-dev/staging/v1",    "Azure Container Apps deployment per environment"],
    ]
    add_table(doc, ["Workflow", "Trigger", "Steps"], cicd_data, C_ORANGE)

    add_heading(doc, "Backend CI Jobs", 2, C_ORANGE)
    jobs = [
        ["lint",            "flake8 with max-line-length=120; fails on any violation"],
        ["test",            "pytest with coverage; skips if no test files present"],
        ["sonarqube-scan",  "On PR to v1-dev and push to v1-dev only"],
        ["docker-build",    "Builds container image with cache"],
        ["acr-push",        "Pushes image to Azure Container Registry"],
        ["deploy-dev",      "Triggered on push to v1-dev branch"],
        ["deploy-staging",  "Triggered on push to v1-staging branch"],
        ["deploy-production","Triggered on push to v1 branch"],
    ]
    add_table(doc, ["Job", "Description"], jobs, C_NAVY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 13. COMPLIANCE MAPPING
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "13. Compliance Mapping", 1, C_NAVY)
    add_para(doc, "The system satisfies multiple regulatory frameworks through its design.")
    add_image(doc, diagrams["compliance"],
              "Figure 11 – Regulatory Compliance Mapping", Inches(6.4))

    comp_data = [
        ["EU AI Act Art.12", "Immutable logs",         "DB triggers prevent UPDATE/DELETE; chain-hash detects tampering"],
        ["EU AI Act Art.13", "Transparency / explain", "SHAP values + reasoning stored per AI decision"],
        ["EU AI Act Art.12", "Retention ≥ 6 months",  "pg_cron archival to cold storage recommended"],
        ["SOC2 CC7.2",       "Audit log collection",   "All user and system events captured"],
        ["SOC2 CC9.1",       "System monitoring",      "Structured logs + request middleware"],
        ["NAIC AI Bulletin",  "AI governance trail",   "model_id + prediction + shap_values per inference"],
        ["NAIC",             "Tenant isolation",       "RLS policies per role and tenant_id"],
        ["ATNA",             "Standard audit codes",   "action_code (C/R/U/D/E) + outcome_code (0/4/8/12)"],
    ]
    add_table(doc, ["Standard", "Requirement", "Implementation"], comp_data, C_RED)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 14. MIGRATION TIMELINE
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "14. Migration Timeline", 1, C_NAVY)
    migrations = [
        ["20260313091500", "auth_audit_logs.sql",             "Create auth_audit table + RLS policies"],
        ["20260316000000", "auth_audit_user_id_fk.sql",       "Add FK constraint to auth.users"],
        ["20260316000001", "audit_immutability.sql",          "Append-only triggers on both tables"],
        ["20260316000002", "audit_logs_change_tracking.sql",  "Add previous_value / new_value columns"],
        ["20260317000002", "audit_ledger_shap.sql",           "Cryptographic ledger + SHAP columns"],
        ["20260317000003", "verify_audit_ledger_fix.sql",     "Fix NULL chain_hash for pre-migration rows"],
    ]
    add_table(doc, ["Timestamp", "File", "Change"], migrations, C_GREEN)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 15. TEST COVERAGE
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "15. Test Coverage", 1, C_NAVY)
    add_para(doc, (
        "insert_audit_log() is mocked in all production integration tests to prevent "
        "test runs from polluting the audit ledger."
    ))
    tests = [
        ["test_device_v1_production.py",     "Device API",    "Mocks insert_audit_log; tests device lifecycle"],
        ["test_tenants_production.py",       "Tenants API",   "Mocks insert_audit_log; tests tenant CRUD"],
        ["test_tenant_provisioning_sla.py",  "Provisioning",  "Mocks insert_audit_log; tests SLA workflows"],
    ]
    add_table(doc, ["Test File", "Scope", "Audit Mock Strategy"], tests, C_NAVY)

    add_heading(doc, "Mock Pattern", 2, C_BLUE)
    p = doc.add_paragraph()
    p.add_run(
        "  async def fake_insert_audit_log(**_kwargs):\n"
        "      pass\n\n"
        "  monkeypatch.setattr(device_routes, 'insert_audit_log', fake_insert_audit_log)"
    ).font.name = "Courier New"

    add_para(doc, (
        "\nThis pattern ensures business logic is tested in isolation while maintaining "
        "full confidence that the real insert_audit_log() is exercised in staging/production "
        "environments via integration smoke tests."
    ))

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph("─" * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph(
        f"Generated by generate_docs.py  ·  NeuraSence Cloud  ·  "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n"
        "This document is auto-generated from source code and migration files."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        run.font.size = Pt(8)
        r, g, b = hex_to_rgb(C_GREY)
        run.font.color.rgb = RGBColor(r, g, b)

    # Save
    out_path = os.path.join(OUT_DIR, "Activity_Logs_Full_Documentation_FINAL.docx")
    doc.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Rendering diagrams ...")
    diagrams = {
        "architecture": draw_architecture(),
        "db_schema":    draw_db_schema(),
        "data_flow":    draw_data_flow(),
        "sequence":     draw_sequence(),
        "rbac":         draw_rbac(),
        "ledger":       draw_ledger(),
        "pii":          draw_pii(),
        "shap":         draw_shap(),
        "frontend_ui":  draw_frontend_ui(),
        "cicd":         draw_cicd(),
        "compliance":   draw_compliance(),
    }
    print(f"  OK  {len(diagrams)} diagrams generated")
    print("Building Word document ...")
    path = build_docx(diagrams)
    print(f"  OK  Saved -> {path}")
