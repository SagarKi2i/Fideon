"""
fix_and_finalize.py
────────────────────────────────────────────────────────────────
1. Opens the original BACKUP document (untouched source of truth)
2. Fixes every factual error found by cross-checking real source code
3. Replaces all Mermaid/code-block paragraphs with real PNG diagrams
4. Saves as  Activity_Logs_Full_Documentation_20260317_125220.docx

Run:  python fix_and_finalize.py
────────────────────────────────────────────────────────────────
"""

import io, os, shutil, sys
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ──────────────────────────────────────────────────────
DIR    = os.path.dirname(os.path.abspath(__file__))
BACKUP = os.path.join(DIR, "Activity_Logs_Full_Documentation_20260317_125220_BACKUP.docx")
OUT    = os.path.join(DIR, "Activity_Logs_Full_Documentation_20260317_125220.docx")

# ── palette ────────────────────────────────────────────────────
C_NAVY="#0D1B2A"; C_BLUE="#1565C0"; C_TEAL="#00796B"; C_GREEN="#2E7D32"
C_ORANGE="#E65100"; C_PURPLE="#6A1B9A"; C_RED="#C62828"; C_GREY="#546E7A"
C_BG="#F5F7FA"; C_WHITE="#FFFFFF"


# ══════════════════════════════════════════════════════════════
# HELPER: fig → BytesIO stream
# ══════════════════════════════════════════════════════════════
def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
# HELPER: replace paragraph with image (proper docx API)
# ══════════════════════════════════════════════════════════════
def replace_para_with_picture(doc, para_idx, stream, width=Inches(5.8), caption=""):
    target = doc.paragraphs[para_idx]._element
    # add at end via proper API (handles rId correctly)
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(stream, width=width)
    img_e = img_p._element
    img_e.getparent().remove(img_e)
    target.addprevious(img_e)
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap_p.add_run(caption)
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x54,0x6E,0x7A)
        cap_e = cap_p._element
        cap_e.getparent().remove(cap_e)
        target.addprevious(cap_e)
    target.getparent().remove(target)


# ══════════════════════════════════════════════════════════════
# HELPER: fix paragraph text (find by partial match, replace content)
# ══════════════════════════════════════════════════════════════
def fix_para(doc, search, new_text, bold=False, color_hex=None, size=None):
    """Find the first paragraph whose text contains `search` and replace it."""
    for para in doc.paragraphs:
        if search in para.text:
            # wipe all runs
            for r in para.runs:
                r.text = ""
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run()
            run.text = new_text
            run.bold = bold
            if size:
                run.font.size = Pt(size)
            if color_hex:
                r2,g2,b2 = int(color_hex[0:2],16),int(color_hex[2:4],16),int(color_hex[4:6],16)
                run.font.color.rgb = RGBColor(r2,g2,b2)
            return True
    return False


def insert_paras_after(doc, search, new_paragraphs):
    """
    Insert new paragraphs immediately after the paragraph whose text contains `search`.
    new_paragraphs: list of (text, bold, italic, color_hex, size, style_name)
    """
    for i, para in enumerate(doc.paragraphs):
        if search in para.text:
            anchor = para._element
            for (txt, bold, italic, col, sz, sty) in reversed(new_paragraphs):
                new_p = doc.add_paragraph()
                if sty:
                    try: new_p.style = doc.styles[sty]
                    except: pass
                r = new_p.add_run(txt)
                r.bold = bold; r.italic = italic
                if sz: r.font.size = Pt(sz)
                if col:
                    r2,g2,b2 = int(col[0:2],16),int(col[2:4],16),int(col[4:6],16)
                    r.font.color.rgb = RGBColor(r2,g2,b2)
                e = new_p._element
                e.getparent().remove(e)
                anchor.addnext(e)
            return True
    return False


# ══════════════════════════════════════════════════════════════
# ─────────────────  DIAGRAM RENDERERS  ───────────────────────
# ══════════════════════════════════════════════════════════════

def _box(ax,x,y,w,h,lbl,sub="",bg=C_BLUE,fs=10):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.1",
                                 linewidth=1.5,edgecolor=bg,facecolor=bg+"22"))
    ax.text(x+w/2,y+h/2+(0.13 if sub else 0),lbl,ha="center",va="center",
            fontsize=fs,fontweight="bold",color=bg)
    if sub: ax.text(x+w/2,y+h/2-0.2,sub,ha="center",va="center",fontsize=7.5,color=C_GREY)

def _arr(ax,x1,y1,x2,y2,col=C_GREY,lbl=""):
    ax.annotate("",xy=(x2,y2),xytext=(x1,y1),
                arrowprops=dict(arrowstyle="-|>",color=col,lw=1.4))
    if lbl: ax.text((x1+x2)/2+0.05,(y1+y2)/2+0.1,lbl,fontsize=7.5,color=col)

# ── 1. Role Hierarchy ──────────────────────────────────────────
def diag_role_hierarchy():
    fig,ax=plt.subplots(figsize=(10,6),facecolor=C_BG)
    ax.set_xlim(0,10);ax.set_ylim(0,6);ax.axis("off")
    nodes=[("global_admin\nFull system access",5.0,5.0,C_RED),
           ("admin\nTenant management",2.5,3.5,C_ORANGE),
           ("user\nStandard access",5.0,3.5,C_BLUE),
           ("viewer\nRead-only",7.5,3.5,C_TEAL),
           ("guest\nPublic / limited",5.0,1.8,C_GREY)]
    for lbl,x,y,col in nodes:
        ax.add_patch(FancyBboxPatch((x-1.1,y-0.4),2.2,0.8,boxstyle="round,pad=0.1",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,y,lbl,ha="center",va="center",fontsize=8.5,fontweight="bold",color=C_WHITE)
    for x1,y1,x2,y2 in [(5.0,4.6,2.5,3.9),(5.0,4.6,5.0,3.9),
                          (5.0,4.6,7.5,3.9),(5.0,3.1,5.0,2.2)]:
        _arr(ax,x1,y1,x2,y2,C_GREY)
    ax.text(5.0,0.6,"Higher role inherits all capabilities of roles below it",
            ha="center",fontsize=8,color=C_GREY,style="italic")
    ax.set_title("Role Hierarchy Diagram",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 2. ER Diagram ──────────────────────────────────────────────
def diag_er():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    def entity(x,y,title,fields,col,w=2.8):
        h=0.4+0.32*len(fields)
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=C_WHITE))
        ax.add_patch(FancyBboxPatch((x,y+h-0.38),w,0.38,boxstyle="square,pad=0",
                                     linewidth=0,facecolor=col))
        ax.text(x+w/2,y+h-0.19,title,ha="center",va="center",
                fontsize=9,fontweight="bold",color=C_WHITE)
        for i,(f,t) in enumerate(fields):
            fy=y+h-0.38-0.32*(i+1)+0.1
            ax.text(x+0.15,fy,f,fontsize=7.5,color=C_NAVY)
            ax.text(x+w-0.1,fy,t,ha="right",fontsize=7,color=C_GREY)
        return x+w/2,y+h/2
    cx1,cy1=entity(0.3,3.5,"auth.users",[("id PK","uuid"),("email","text"),("role","text")],C_NAVY)
    entity(4.2,4.0,"auth_audit",[("id PK","uuid"),("user_id FK","uuid"),("event","text"),
                                  ("action_code","text"),("outcome_code","int"),
                                  ("integrity_hash","text")],C_TEAL,3.0)
    entity(4.2,1.5,"audit_logs",[("id PK","uuid"),("user_id FK","uuid"),("action","text"),
                                  ("shap_values","jsonb"),("chain_hash","text"),
                                  ("sequence_num","bigint")],C_GREEN,3.0)
    entity(8.8,4.2,"user_roles",[("id PK","uuid"),("user_id FK","uuid"),("role_id FK","uuid")],C_PURPLE)
    entity(8.8,2.2,"roles",[("id PK","uuid"),("name","text"),("description","text")],C_ORANGE)
    entity(8.8,0.2,"device_sync_logs",[("id PK","uuid"),("device_id FK","uuid"),
                                        ("synced_at","timestamp")],C_BLUE)
    for x1,y1,x2,y2,lbl,col in [
        (cx1+1.4,cy1,5.7,4.7,"1:N",C_TEAL),(cx1+1.4,cy1,5.7,2.8,"1:N",C_GREEN),
        (cx1+1.4,cy1,9.3,4.7,"1:N",C_NAVY)]:
        _arr(ax,x1,y1,x2,y2,col,lbl)
    ax.set_title("Entity-Relationship Diagram",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 3. Schema Column Detail ────────────────────────────────────
def diag_schema_detail():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    def cls_box(x,y,name,attrs,col,w=5.5):
        h=0.55+0.3*len(attrs)
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=C_WHITE))
        ax.add_patch(FancyBboxPatch((x,y+h-0.45),w,0.45,boxstyle="square,pad=0",
                                     linewidth=0,facecolor=col))
        ax.text(x+w/2,y+h-0.22,name,ha="center",va="center",
                fontsize=10,fontweight="bold",color=C_WHITE)
        for i,row in enumerate(attrs):
            ry=y+h-0.45-0.3*(i+1)+0.08
            ax.text(x+0.15,ry,row,fontsize=7.5,color=C_NAVY,fontfamily="monospace")
    cls_box(0.3,3.5,"auth_audit",["+UUID id  [PK]","+UUID user_id  [FK]","+text email",
        "+text role","+text event","+text action_code  // C/R/U/D/E",
        "+int  outcome_code  // 0/4/8/12","+text resource_type","+text resource_id",
        "+text integrity_hash  // SHA-256","+timestamptz created_at"],C_TEAL)
    cls_box(7.0,0.4,"audit_logs  (Cryptographic Ledger)",["+UUID id  [PK]",
        "+UUID user_id  [FK]","+text action","+text resource_type",
        "+jsonb details  // NOT in integrity_hash",
        "+jsonb previous_value  // PII-scrubbed (top-level only)",
        "+jsonb new_value  // PII-scrubbed (top-level only)",
        "+text model_id","+jsonb prediction",
        "+jsonb shap_values  // {feature: float}",
        "+text reasoning  // auto-generated, format: +0.4200",
        "+text integrity_hash  // SHA-256 (excl. details)",
        "+bigint sequence_num  // IDENTITY",
        "+text chain_hash  // SHA-256 ledger","+timestamptz created_at"],C_GREEN)
    ax.set_title("Schema Diagram - Supabase Column-Level Detail (corrected)",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 4. RLS Policy ──────────────────────────────────────────────
def diag_rls():
    fig,ax=plt.subplots(figsize=(11,8),facecolor=C_BG)
    ax.set_xlim(0,11);ax.set_ylim(0,8);ax.axis("off")
    def diamond(x,y,w,h,lbl,col):
        pts=[(x+w/2,y+h),(x+w,y+h/2),(x+w/2,y),(x,y+h/2)]
        ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",edgecolor=col,lw=2))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    def rect(x,y,w,h,lbl,col):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8.5,fontweight="bold",color=C_WHITE)
    rect(4.0,7.2,3.0,0.55,"SELECT on auth_audit",C_BLUE)
    diamond(3.5,5.9,4.0,1.1,"Who is the\ncurrent user?",C_NAVY)
    rect(0.2,4.5,2.5,0.55,"global_admin",C_RED)
    rect(3.0,4.5,2.5,0.55,"admin",C_ORANGE)
    rect(5.7,4.5,2.5,0.55,"user / viewer",C_BLUE)
    rect(8.4,4.5,2.5,0.55,"guest",C_GREY)
    rect(0.2,3.1,2.5,0.8,"ALL rows\n(no filter)",C_RED)
    rect(3.0,3.1,2.5,0.8,"Non-admin rows\n(own tenant)",C_ORANGE)
    rect(5.7,3.1,2.5,0.8,"Own rows only\n(user_id=auth.uid())",C_BLUE)
    rect(8.4,3.1,2.5,0.8,"403 Forbidden\n(RLS blocks)",C_GREY)
    rect(3.5,1.5,4.0,0.7,"Row returned to client",C_GREEN)
    for x1,y1,x2,y2 in [(5.5,7.2,5.5,7.0),(5.5,5.9,1.45,5.05),(5.5,5.9,4.25,5.05),
                          (5.5,5.9,6.95,5.05),(5.5,5.9,9.65,5.05)]:
        _arr(ax,x1,y1,x2,y2,C_GREY)
    for x in [1.45,4.25,6.95]: _arr(ax,x,3.1,5.5,2.2,C_GREEN)
    _arr(ax,9.65,3.1,8.9,2.2,C_RED,"denied")
    ax.set_title("Row-Level Security (RLS) Policy - auth_audit SELECT",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 5. Data Flow ───────────────────────────────────────────────
def diag_data_flow():
    fig,ax=plt.subplots(figsize=(13,6),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,6);ax.axis("off")
    steps=[("Browser\nAuth.tsx / Activity.tsx",C_TEAL,1.0),
           ("Supabase\nAuth Event",C_NAVY,3.0),
           ("FastAPI\nRoute Handler",C_BLUE,5.0),
           ("insert_audit_log()\nPII Scrub + SHA-256",C_ORANGE,7.2),
           ("PostgREST\nINSERT",C_GREEN,9.4),
           ("DB Trigger\nchain_hash",C_PURPLE,11.4)]
    y=3.2
    for lbl,col,x in steps:
        ax.add_patch(FancyBboxPatch((x-0.85,y-0.5),1.7,1.0,boxstyle="round,pad=0.1",
                                     linewidth=2,edgecolor=col,facecolor=col+"22"))
        ax.text(x,y,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    for i in range(len(steps)-1):
        _arr(ax,steps[i][2]+0.85,y,steps[i+1][2]-0.85,y,C_GREY)
    ax.add_patch(FancyBboxPatch((8.5,0.7),3.5,0.9,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_GREEN,facecolor=C_GREEN+"22"))
    ax.text(10.25,1.15,"audit_logs\n(Immutable Ledger)",ha="center",
            fontsize=9,color=C_GREEN,fontweight="bold")
    _arr(ax,10.25,2.7,10.25,1.6,C_GREEN,"stored")
    ax.set_title("Data Flow Diagram - How Audit Logs Are Created",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 6. Sequence: Login ─────────────────────────────────────────
def diag_seq_login():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    actors=[("User",1.2,C_TEAL),("Auth.tsx",3.2,C_BLUE),
            ("Supabase\nAuth",5.8,C_NAVY),("insert_\naudit_log",8.4,C_ORANGE),
            ("audit_logs\nDB",11.0,C_GREEN)]
    TOP=6.5;BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.7,TOP-0.05),1.4,0.5,boxstyle="round,pad=0.07",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[(6.0,1.2,3.2,"login(email,password)",C_TEAL,False),
          (5.5,3.2,5.8,"signInWithPassword()",C_BLUE,False),
          (5.0,5.8,3.2,"session token",C_NAVY,True),
          (4.5,3.2,8.4,"insert_audit_log('login',user_id,...)",C_BLUE,False),
          (4.0,8.4,11.0,"INSERT INTO auth_audit",C_ORANGE,False),
          (3.5,11.0,8.4,"row written",C_GREEN,True),
          (3.0,3.2,1.2,"redirect to dashboard",C_BLUE,True)]
    for y,x1,x2,lbl,col,ret in msgs:
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram - Login Audit Event",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 7. Sequence: Admin Approves Pod ───────────────────────────
def diag_seq_pod():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    actors=[("Admin",1.2,C_ORANGE),("PodActivation\nRequests.tsx",3.2,C_BLUE),
            ("FastAPI\n/pod_activation",5.8,C_NAVY),
            ("insert_\naudit_log",8.4,C_ORANGE),("audit_logs",11.0,C_GREEN)]
    TOP=6.5;BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.7,TOP-0.05),1.4,0.5,boxstyle="round,pad=0.07",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[(6.0,1.2,3.2,"click Approve Pod",C_ORANGE,False),
          (5.5,3.2,5.8,"PATCH /pod_activation/{id}",C_BLUE,False),
          (5.0,5.8,5.8,"verify_admin(JWT)",C_NAVY,False),
          (4.5,5.8,8.4,"insert_audit_log('pod_approve',...)",C_NAVY,False),
          (4.0,8.4,11.0,"INSERT INTO audit_logs",C_ORANGE,False),
          (3.5,11.0,8.4,"chain_hash written",C_GREEN,True),
          (3.0,5.8,3.2,"200 OK",C_NAVY,True),
          (2.5,3.2,1.2,"show success toast",C_BLUE,True)]
    for y,x1,x2,lbl,col,ret in msgs:
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram - Admin Approves Pod",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 8. Sequence: Viewing Logs ──────────────────────────────────
def diag_seq_view():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    actors=[("User/Admin",1.3,C_BLUE),("Activity.tsx",3.5,C_TEAL),
            ("Supabase\nauth_audit",6.0,C_NAVY),("FastAPI\n/api/activity",8.5,C_ORANGE),
            ("audit_logs",11.2,C_GREEN)]
    TOP=6.5;BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.8,TOP-0.05),1.6,0.5,boxstyle="round,pad=0.07",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=7.5,fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[(6.0,1.3,3.5,"navigate to /activity",C_BLUE,False),
          (5.5,3.5,6.0,"SELECT * FROM auth_audit (RLS)",C_TEAL,False),
          (5.0,6.0,3.5,"auth events[]",C_NAVY,True),
          (4.5,3.5,8.5,"GET /api/activity/system?page=0",C_TEAL,False),
          (4.0,8.5,11.2,"SELECT * FROM audit_logs (RBAC)",C_ORANGE,False),
          (3.5,11.2,8.5,"system logs[]",C_GREEN,True),
          (3.0,8.5,3.5,"200 OK {logs, has_more}",C_ORANGE,True),
          (2.5,3.5,1.3,"render two-tab Activity page",C_TEAL,True)]
    for y,x1,x2,lbl,col,ret in msgs:
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram - Viewing Activity Logs",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 9. Component Architecture ──────────────────────────────────
def diag_component_arch():
    fig,ax=plt.subplots(figsize=(13,8),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,8);ax.axis("off")
    ax.add_patch(FancyBboxPatch((0.3,6.2),12.4,1.4,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_BLUE,facecolor=C_BLUE+"11"))
    ax.text(0.7,7.45,"App Routes (App.tsx)",fontsize=8,color=C_BLUE,style="italic")
    for r,x in [("/activity",2.0),("/activity (admin)",4.5),("/chat",6.8),("/admin",8.9),("/workflow-ai",11.0)]:
        ax.add_patch(FancyBboxPatch((x-0.9,6.35),1.8,0.65,boxstyle="round,pad=0.07",
                                     linewidth=1.5,edgecolor=C_BLUE,facecolor=C_BLUE))
        ax.text(x,6.67,r,ha="center",va="center",fontsize=7.5,fontweight="bold",color=C_WHITE)
    comps=[("Activity.tsx\n(AuthEventsTab+\nSystemEventsTab)",2.0,4.5,C_TEAL,1.3),
           ("AdminActivity.tsx",4.5,5.0,C_ORANGE,1.1),
           ("Chat.tsx",6.8,5.2,C_BLUE,1.0),("WorkflowAI.tsx",10.0,5.0,C_PURPLE,1.2)]
    for nm,x,y,col,h in comps:
        ax.add_patch(FancyBboxPatch((x-0.9,y-0.4),1.8,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col+"22"))
        ax.text(x,y+h/2-0.4,nm,ha="center",va="center",fontsize=7.5,color=col,fontweight="bold")
    for nm,x,y,col in [("Supabase RLS\nauth_audit",2.0,2.8,C_NAVY),
                        ("FastAPI\n/api/activity/system",5.0,2.8,C_ORANGE),
                        ("FastAPI\n/api/workflow-ai",9.5,2.8,C_PURPLE)]:
        ax.add_patch(FancyBboxPatch((x-1.1,y-0.4),2.2,0.9,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col+"22"))
        ax.text(x,y+0.05,nm,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    for nm,x,col in [("auth_audit\ntable",2.0,C_TEAL),("audit_logs\ntable",5.0,C_GREEN)]:
        ax.add_patch(FancyBboxPatch((x-1.0,0.5),2.0,0.9,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,0.95,nm,ha="center",va="center",fontsize=8.5,fontweight="bold",color=C_WHITE)
    for x1,y1,x2,y2,col in [(2.0,6.35,2.0,5.5,C_TEAL),(4.5,6.35,4.5,5.5,C_ORANGE),
                              (2.0,4.5,2.0,3.25,C_TEAL),(2.0,4.5,4.5,3.25,C_TEAL),
                              (2.0,2.4,2.0,1.4,C_NAVY),(4.5,2.4,4.5,1.4,C_GREEN)]:
        _arr(ax,x1,y1,x2,y2,col)
    ax.set_title("Component Architecture Diagram",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 10. Hash Flow: auth_audit ──────────────────────────────────
def diag_hash_auth():
    fig,ax=plt.subplots(figsize=(11,5),facecolor=C_BG)
    ax.set_xlim(0,11);ax.set_ylim(0,5);ax.axis("off")
    inputs=[("user_id",0.7,3.8),("role",0.7,3.1),("event",0.7,2.4),
            ("action_code",0.7,1.7),("outcome_code",0.7,1.0),("resource_type",0.7,0.3)]
    for lbl,x,y in inputs:
        ax.add_patch(FancyBboxPatch((x-0.6,y-0.22),1.7,0.44,boxstyle="round,pad=0.05",
                                     linewidth=1.5,edgecolor=C_BLUE,facecolor=C_BLUE+"22"))
        ax.text(x+0.25,y,lbl,ha="center",va="center",fontsize=8.5,color=C_BLUE)
        _arr(ax,x+0.85,y,3.5,2.0,C_GREY)
    ax.add_patch(FancyBboxPatch((3.3,1.3),2.0,1.4,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_ORANGE,facecolor=C_ORANGE+"22"))
    ax.text(4.3,2.0,"SHA-256\nhash()",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_ORANGE)
    _arr(ax,5.3,2.0,7.0,2.0,C_ORANGE)
    ax.add_patch(FancyBboxPatch((7.0,1.5),3.5,1.0,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_GREEN,facecolor=C_GREEN+"22"))
    ax.text(8.75,2.0,"integrity_hash\n(stored in DB)",ha="center",va="center",
            fontsize=9,fontweight="bold",color=C_GREEN)
    ax.text(5.5,4.5,"NOTE: email is EXCLUDED from hash (PII) - masked as u***@domain in logger",
            ha="center",fontsize=8.5,color=C_RED,
            bbox=dict(boxstyle="round,pad=0.3",facecolor="#FFEBEE",edgecolor=C_RED))
    ax.set_title("auth_audit Integrity Hash Flow (Frontend SHA-256)",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 11. Hash Flow: audit_logs (CORRECTED - details excluded) ───
def diag_hash_audit():
    fig,ax=plt.subplots(figsize=(13,7),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,7);ax.axis("off")
    # Fields IN hash
    in_hash=["user_id","action","resource_type","resource_id",
             "previous_value*","new_value*","model_id","prediction",
             "shap_values","reasoning","created_at"]
    for i,f in enumerate(in_hash):
        y=6.5-i*0.52
        ax.add_patch(FancyBboxPatch((0.2,y-0.2),2.5,0.42,boxstyle="round,pad=0.04",
                                     linewidth=1,edgecolor=C_GREEN,facecolor=C_GREEN+"18"))
        ax.text(1.45,y,f,ha="center",va="center",fontsize=7.8,color=C_GREEN)
        _arr(ax,2.7,y,4.3,3.5,C_GREY)
    # details NOT in hash
    ax.add_patch(FancyBboxPatch((0.2,0.3),2.5,0.5,boxstyle="round,pad=0.04",
                                 linewidth=2,edgecolor=C_RED,facecolor=C_RED+"18"))
    ax.text(1.45,0.55,"details  [EXCLUDED]",ha="center",va="center",
            fontsize=8,color=C_RED,fontweight="bold")
    ax.text(1.45,0.1,"details is stored but NOT hashed",ha="center",
            fontsize=7,color=C_RED,style="italic")
    ax.add_patch(FancyBboxPatch((4.3,2.5),2.2,2.0,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_ORANGE,facecolor=C_ORANGE+"22"))
    ax.text(5.4,3.5,"SHA-256\nintegrity\nhash",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_ORANGE)
    _arr(ax,6.5,3.5,8.0,3.5,C_ORANGE,"integrity_hash")
    ax.add_patch(FancyBboxPatch((8.0,2.7),2.5,1.6,boxstyle="round,pad=0.08",
                                 linewidth=2,edgecolor=C_PURPLE,facecolor=C_PURPLE+"22"))
    ax.text(9.25,3.5,"DB Trigger\ncompute_\nchain_hash()",ha="center",va="center",
            fontsize=9,fontweight="bold",color=C_PURPLE)
    _arr(ax,10.5,3.5,11.5,3.5,C_PURPLE,"chain_hash")
    ax.add_patch(FancyBboxPatch((11.5,3.0),1.3,1.0,boxstyle="round,pad=0.08",
                                 linewidth=2,edgecolor=C_GREEN,facecolor=C_GREEN+"22"))
    ax.text(12.15,3.5,"stored",ha="center",va="center",fontsize=9,color=C_GREEN,fontweight="bold")
    ax.text(3.0,0.3,"* previous_value and new_value: scrubbed TOP-LEVEL keys only (_scrub_audit_value)",
            fontsize=7.5,color=C_ORANGE,style="italic")
    ax.set_title("audit_logs Integrity Hash Flow (CORRECTED: details excluded, top-level PII scrub)",
                 fontsize=11,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 12. Access Control Matrix ──────────────────────────────────
def diag_access_matrix():
    fig,ax=plt.subplots(figsize=(10,6),facecolor=C_BG)
    ax.set_xlim(-1,11);ax.set_ylim(-1,7);ax.axis("off")
    ax.axvline(5,color=C_GREY,lw=1,ls="--",alpha=0.4)
    ax.axhline(3,color=C_GREY,lw=1,ls="--",alpha=0.4)
    ax.text(2.5,6.3,"Own Records Only",ha="center",fontsize=9,color=C_GREY)
    ax.text(7.5,6.3,"All Records",ha="center",fontsize=9,color=C_GREY)
    ax.text(-0.7,4.5,"High\nPrivilege",ha="center",fontsize=9,color=C_GREY,rotation=90)
    ax.text(-0.7,1.5,"Low\nPrivilege",ha="center",fontsize=9,color=C_GREY,rotation=90)
    roles=[(1.5,5.2,"global_admin\nAll rows, all tenants",C_RED),
           (3.0,4.0,"admin\nTenant rows",C_ORANGE),
           (2.0,1.8,"user\nOwn rows",C_BLUE),
           (3.5,1.5,"viewer\nOwn rows\n(read-only)",C_TEAL),
           (1.2,0.5,"guest\n403 Forbidden",C_GREY)]
    for x,y,lbl,col in roles:
        ax.add_patch(plt.Circle((x,y),0.55,color=col+"44",ec=col,lw=2))
        ax.text(x,y,lbl,ha="center",va="center",fontsize=7.5,fontweight="bold",color=col)
    ax.set_title("Access Control Matrix - Audit Row Visibility by Role",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 13. Route Guard ────────────────────────────────────────────
def diag_route_guard():
    fig,ax=plt.subplots(figsize=(10,8),facecolor=C_BG)
    ax.set_xlim(0,10);ax.set_ylim(0,8);ax.axis("off")
    def diam(x,y,w,h,lbl,col):
        pts=[(x+w/2,y+h),(x+w,y+h/2),(x+w/2,y),(x,y+h/2)]
        ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",edgecolor=col,lw=2))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    def rct(x,y,w,h,lbl,col):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8.5,fontweight="bold",color=C_WHITE)
    rct(3.5,7.2,3.0,0.55,"User navigates to URL",C_BLUE)
    diam(3.3,5.9,3.4,1.1,"Is user\nauthenticated?",C_NAVY)
    rct(0.5,4.5,3.0,0.55,"Redirect to /login",C_RED)
    diam(3.3,4.1,3.4,1.1,"Has required\nrole?",C_NAVY)
    rct(7.0,4.5,2.5,0.55,"403 Forbidden",C_ORANGE)
    rct(3.5,2.8,3.0,0.7,"Render protected\npage",C_GREEN)
    rct(3.5,1.6,3.0,0.6,"Log route access\nin audit trail",C_TEAL)
    _arr(ax,5.0,7.2,5.0,7.0,C_BLUE)
    _arr(ax,5.0,5.9,2.0,5.05,C_RED,"No")
    _arr(ax,5.0,5.9,5.0,5.2,C_NAVY,"Yes")
    _arr(ax,5.0,4.1,8.25,4.5,C_ORANGE,"No")
    _arr(ax,5.0,4.1,5.0,3.5,C_GREEN,"Yes")
    _arr(ax,5.0,2.8,5.0,2.2,C_TEAL)
    ax.set_title("Navigation & Route Guard Diagram",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 14. Backend Logging Architecture (CORRECTED) ───────────────
def diag_backend_logging():
    fig,ax=plt.subplots(figsize=(14,7),facecolor=C_BG)
    ax.set_xlim(0,14);ax.set_ylim(0,7);ax.axis("off")
    # logger module box
    ax.add_patch(FancyBboxPatch((0.2,1.0),5.5,5.5,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_PURPLE,facecolor=C_PURPLE+"09"))
    ax.text(0.5,6.35,"backend/app/logger/__init__.py",fontsize=8,color=C_PURPLE,style="italic",fontweight="bold")
    items=[("configure_logging()\n- structlog JSON\n- RotatingFile 10MB x5\n- removes uvicorn handlers",1.5,5.0,C_NAVY),
           ("RequestLogging\nMiddleware\nrequest_id, method, path,\nclient_ip, status, duration_ms",1.5,3.3,C_TEAL),
           ("REDACTED_FIELDS\n20 fields, NO email\n(email -> u***@domain)",3.5,5.0,C_ORANGE),
           ("pii_scrubber_\nprocessor()\nPass-1 ONLY in structlog\n(not automatic Pass-2)",3.5,3.3,C_ORANGE),
           ("scrub_text()\nasync Pass-2 Presidio\nMUST be called explicitly\nby caller",3.5,1.5,C_RED),
           ("setup_logging(app)\nget_logger(name)",1.5,1.5,C_NAVY)]
    for nm,x,y,col in items:
        ax.add_patch(FancyBboxPatch((x-0.7,y-0.5),2.1,1.15,boxstyle="round,pad=0.08",
                                     linewidth=1.5,edgecolor=col,facecolor=col+"22"))
        ax.text(x+0.35,y,nm,ha="center",va="center",fontsize=7,color=col,fontweight="bold")
    # SHAP box
    ax.add_patch(FancyBboxPatch((0.2,0.2),5.5,0.7,boxstyle="round,pad=0.05",
                                 linewidth=2,edgecolor=C_GREEN,facecolor=C_GREEN+"18"))
    ax.text(2.95,0.55,"generate_shap_reasoning()  +  log_ai_decision_audit()  [also emits ai_decision_audited structlog line]",
            ha="center",va="center",fontsize=7.5,color=C_GREEN,fontweight="bold")
    # Outputs
    for nm,x,y,col in [("structlog\nJSON stdout",8.5,5.5,C_TEAL),
                        ("Rotating\napp.log\n(10MB x5)",8.5,4.0,C_NAVY),
                        ("audit_logs\n(PostgreSQL)",8.5,2.5,C_GREEN),
                        ("Presidio\nThreadPool\nmax_workers=2",11.5,4.0,C_PURPLE)]:
        ax.add_patch(FancyBboxPatch((x-0.8,y-0.5),2.6,1.1,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+0.5,y,nm,ha="center",va="center",fontsize=8,fontweight="bold",color=C_WHITE)
    _arr(ax,5.7,5.5,7.7,5.8,C_TEAL)
    _arr(ax,5.7,4.0,7.7,4.3,C_NAVY)
    _arr(ax,5.7,2.5,7.7,2.8,C_GREEN)
    _arr(ax,5.7,2.0,10.7,4.0,C_PURPLE,"Pass-2 async")
    ax.set_title("Backend Logging Architecture (logger/__init__.py) - CORRECTED",
                 fontsize=11,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 15. Full System Interaction ────────────────────────────────
def diag_full_system():
    fig,ax=plt.subplots(figsize=(14,9),facecolor=C_BG)
    ax.set_xlim(0,14);ax.set_ylim(0,9);ax.axis("off")
    users=[("global_admin",1.0,8.2,C_RED),("admin",3.5,8.2,C_ORANGE),
           ("user",6.0,8.2,C_BLUE),("viewer",8.5,8.2,C_TEAL),("guest",11.0,8.2,C_GREY)]
    for nm,x,y,col in users:
        ax.add_patch(plt.Circle((x,y),0.45,color=col,ec=col,lw=2))
        ax.text(x,y,nm[:6],ha="center",va="center",fontsize=7,fontweight="bold",color=C_WHITE)
        _arr(ax,x,y-0.45,7.0,7.0,col)
    ax.add_patch(FancyBboxPatch((5.0,6.4),4.0,0.7,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_NAVY,facecolor=C_NAVY))
    ax.text(7.0,6.75,"Next.js 14 Frontend",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_WHITE)
    _arr(ax,7.0,6.4,7.0,5.8,C_NAVY,"HTTPS/JWT")
    ax.add_patch(FancyBboxPatch((5.0,5.1),4.0,0.7,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_BLUE,facecolor=C_BLUE))
    ax.text(7.0,5.45,"FastAPI + Uvicorn",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_WHITE)
    for nm,x,y,col in [("verify_user\nverify_admin",2.5,4.0,C_TEAL),
                        ("insert_audit_log\n+ PII scrub (top-level)\n+ SHA-256 (excl. details)",7.0,4.0,C_ORANGE),
                        ("SHAP reasoning\ngenerator\n+0.4200 format",11.5,4.0,C_PURPLE)]:
        ax.add_patch(FancyBboxPatch((x-1.2,y-0.5),2.4,1.1,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col+"22"))
        ax.text(x,y,nm,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    for x1,y1,x2,y2,col in [(7.0,5.1,2.5,4.6,C_TEAL),(7.0,5.1,7.0,4.6,C_ORANGE),
                              (7.0,5.1,11.5,4.6,C_PURPLE)]:
        _arr(ax,x1,y1,x2,y2,col)
    for nm,x,y,col in [("auth_audit\ntable",3.0,2.3,C_TEAL),
                        ("audit_logs\ntable",7.0,2.3,C_GREEN),
                        ("auth.users\ntable",11.0,2.3,C_NAVY)]:
        ax.add_patch(FancyBboxPatch((x-1.1,y-0.45),2.2,0.9,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,y,nm,ha="center",va="center",fontsize=9,fontweight="bold",color=C_WHITE)
    _arr(ax,2.5,3.5,3.0,2.75,C_TEAL)
    _arr(ax,7.0,3.5,7.0,2.75,C_GREEN)
    ax.add_patch(FancyBboxPatch((5.0,0.5),4.0,0.7,boxstyle="round,pad=0.1",
                                 linewidth=2,edgecolor=C_ORANGE,facecolor=C_ORANGE))
    ax.text(7.0,0.85,"DB Triggers (immutability + chain_hash)",
            ha="center",va="center",fontsize=9,fontweight="bold",color=C_WHITE)
    _arr(ax,7.0,2.3-0.45,7.0,1.2,C_ORANGE)
    ax.set_title("Full System Interaction Diagram (CORRECTED)",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 16. Ledger Trigger Sequence ────────────────────────────────
def diag_ledger_trigger():
    fig,ax=plt.subplots(figsize=(12,6),facecolor=C_BG)
    ax.set_xlim(0,12);ax.set_ylim(0,6);ax.axis("off")
    actors=[("Application\ninsert_audit_log",1.5,C_BLUE),
            ("PostgreSQL\nBEFORE INSERT trigger",4.5,C_NAVY),
            ("pg_advisory\nxact_lock",7.5,C_ORANGE),
            ("audit_logs\ntable",10.5,C_GREEN)]
    TOP=5.5;BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.8,TOP-0.05),1.6,0.5,boxstyle="round,pad=0.07",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[(5.0,1.5,4.5,"INSERT row (integrity_hash set, details NOT in hash)",C_BLUE,False),
          (4.5,4.5,7.5,"LOCK (serialize inserts)",C_NAVY,False),
          (4.0,7.5,4.5,"lock acquired",C_ORANGE,True),
          (3.5,4.5,4.5,"fetch prev chain_hash",C_NAVY,False),
          (3.0,4.5,4.5,"SHA-256(prev||integrity_hash) -> chain_hash",C_NAVY,False),
          (2.5,4.5,10.5,"write row + chain_hash",C_NAVY,False),
          (2.0,10.5,4.5,"row written",C_GREEN,True),
          (1.5,4.5,7.5,"UNLOCK",C_NAVY,False)]
    for y,x1,x2,lbl,col,ret in msgs:
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Cryptographic Ledger Trigger - compute_audit_chain_hash()",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 17. SHAP Flow (CORRECTED - details excluded note) ──────────
def diag_shap_flow():
    fig,ax=plt.subplots(figsize=(13,5),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,5);ax.axis("off")
    steps=[("Route handler\ncalls log_ai_\ndecision_audit()",C_BLUE,1.2),
           ("reasoning\nprovided?",C_NAVY,3.4),
           ("generate_shap_\nreasoning()\nformat: +0.4200 (4dp)",C_PURPLE,5.8),
           ("insert_audit_log(\n...reasoning,\nshap_values)",C_ORANGE,8.4),
           ("ai_decision_\naudited structlog\n+ audit_logs row",C_GREEN,11.0)]
    y=2.5
    for lbl,col,x in steps:
        if "?" in lbl:
            pts=[(x,y+0.55),(x+0.85,y+1.1),(x+1.7,y+0.55),(x+0.85,y)]
            ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",edgecolor=col,lw=2))
            ax.text(x+0.85,y+0.55,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
        else:
            ax.add_patch(FancyBboxPatch((x-0.85,y-0.6),1.7,1.2,boxstyle="round,pad=0.1",
                                         linewidth=2,edgecolor=col,facecolor=col+"22"))
            ax.text(x,y,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    for i in range(len(steps)-1):
        if "?" in steps[i][0]:
            _arr(ax,steps[i][2]+1.7,y,steps[i+1][2]-0.85,y,C_GREY,"No")
        else:
            _arr(ax,steps[i][2]+0.85,y,steps[i+1][2]-0.85,y,C_GREY)
    _arr(ax,3.4+0.85,y+1.1,steps[3][2]-0.85,y+0.5,C_GREEN,"Yes (skip gen)")
    ax.text(6.5,0.5,"CORRECTED: log_ai_decision_audit() emits BOTH an audit_logs row AND a structlog 'ai_decision_audited' line",
            ha="center",fontsize=8,color=C_RED,
            bbox=dict(boxstyle="round,pad=0.3",facecolor="#FFEBEE",edgecolor=C_RED))
    ax.set_title("SHAP AI Decision Audit Flow (CORRECTED)",fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 18. SHAP Sequence ──────────────────────────────────────────
def diag_shap_seq():
    fig,ax=plt.subplots(figsize=(13,6),facecolor=C_BG)
    ax.set_xlim(0,13);ax.set_ylim(0,6);ax.axis("off")
    actors=[("AI Model\n/ Route",1.5,C_PURPLE),("logger/\n__init__.py",4.0,C_NAVY),
            ("generate_shap_\nreasoning()",6.5,C_TEAL),
            ("insert_\naudit_log()",9.0,C_ORANGE),("audit_logs\nDB",11.5,C_GREEN)]
    TOP=5.5;BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.75,TOP-0.05),1.5,0.5,boxstyle="round,pad=0.07",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=7.5,fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[(5.0,1.5,4.0,"log_ai_decision_audit(shap_values, prediction)",C_PURPLE,False),
          (4.4,4.0,6.5,"generate_shap_reasoning()",C_NAVY,False),
          (3.8,6.5,4.0,"'Model X predicted Y. Top: feat (+0.4200, up)...'",C_TEAL,True),
          (3.2,4.0,9.0,"insert_audit_log(...shap_values, reasoning)",C_NAVY,False),
          (2.6,9.0,11.5,"INSERT INTO audit_logs",C_ORANGE,False),
          (2.0,11.5,9.0,"row + chain_hash",C_GREEN,True)]
    for y,x1,x2,lbl,col,ret in msgs:
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.text(4.0,0.5,"Also: structlog.get_logger('ai_decision').info('ai_decision_audited', ...)",
            fontsize=8,color=C_PURPLE,style="italic")
    ax.set_title("SHAP Data Flow Sequence (CORRECTED: +0.4200 format, structlog also emitted)",
                 fontsize=11,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 19. Ledger Verification ────────────────────────────────────
def diag_verify():
    fig,ax=plt.subplots(figsize=(10,8),facecolor=C_BG)
    ax.set_xlim(0,10);ax.set_ylim(0,8);ax.axis("off")
    def diam(x,y,w,h,lbl,col):
        pts=[(x+w/2,y+h),(x+w,y+h/2),(x+w/2,y),(x,y+h/2)]
        ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",edgecolor=col,lw=2))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8,color=col,fontweight="bold")
    def rct(x,y,w,h,lbl,col):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",fontsize=8.5,fontweight="bold",color=C_WHITE)
    rct(3.0,7.2,4.0,0.6,"Compliance Audit Triggered",C_BLUE)
    rct(3.0,6.2,4.0,0.6,"verify_audit_ledger()",C_NAVY)
    rct(3.0,5.2,4.0,0.7,"Iterate rows by\nsequence_num",C_TEAL)
    diam(2.8,3.8,4.4,1.2,"chain_hash=NULL?\n(pre-migration row)",C_ORANGE)
    rct(0.3,3.0,2.5,0.6,"is_valid=NULL\n(sentinel)",C_ORANGE)
    rct(3.0,2.6,4.0,0.7,"Recompute:\nSHA-256(prev||integrity)",C_PURPLE)
    diam(2.8,1.3,4.4,1.1,"stored==\ncomputed?",C_NAVY)
    rct(0.3,0.5,2.5,0.6,"TAMPERED\nis_valid=false",C_RED)
    rct(7.2,0.5,2.5,0.6,"VALID\nis_valid=true",C_GREEN)
    for x1,y1,x2,y2,col,lbl in [
        (5.0,7.2,5.0,6.8,C_BLUE,""),(5.0,6.2,5.0,5.9,C_NAVY,""),
        (5.0,5.2,5.0,5.0,C_TEAL,""),(5.0,3.8,5.0,3.3,C_ORANGE,"No"),
        (2.8,4.4,1.55,3.6,C_ORANGE,"Yes"),(5.0,2.6,5.0,2.4,C_PURPLE,""),
        (5.0,1.3,5.0,1.1,C_NAVY,""),(2.8,1.85,1.55,1.1,C_RED,"No"),
        (7.2,1.85,8.45,1.1,C_GREEN,"Yes")]:
        _arr(ax,x1,y1,x2,y2,col,lbl)
    ax.set_title("Audit Ledger Verification - verify_audit_ledger()",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 20. Migration Timeline ─────────────────────────────────────
def diag_migration_timeline():
    fig,ax=plt.subplots(figsize=(14,5),facecolor=C_BG)
    ax.set_xlim(0,14);ax.set_ylim(0,5);ax.axis("off")
    events=[
        ("2026-03-13","20260313091500",
         "auth_audit table\n+RLS policies\n+integrity_hash",C_TEAL),
        ("2026-03-16\n(000000)","FK Constraint",
         "user_id FK\n-> auth.users",C_BLUE),
        ("2026-03-16\n(000001)","Immutability",
         "prevent_audit_\nmodification()",C_RED),
        ("2026-03-16\n(000002)","Change Tracking",
         "previous_value +\nnew_value columns\n(top-level scrub)",C_ORANGE),
        ("2026-03-17\n(000002)","Ledger + SHAP",
         "chain_hash\nsequence_num\nshap_values\nmodel_id\nprediction\nreasoning",C_PURPLE),
        ("2026-03-17\n(000003)","Ledger Fix",
         "NULL chain_hash\nfor pre-migration\nrows",C_GREEN),
    ]
    ax.plot([0.8,13.2],[2.5,2.5],color=C_GREY,lw=2)
    xs=np.linspace(1.2,12.8,len(events))
    for i,(dt,title,desc,col) in enumerate(events):
        x=xs[i]
        ax.add_patch(plt.Circle((x,2.5),0.2,color=col,zorder=5))
        if i%2==0:
            ax.text(x,3.2,title,ha="center",fontsize=8.5,fontweight="bold",color=col)
            ax.text(x,4.2,desc,ha="center",fontsize=7.5,color=C_NAVY)
            ax.plot([x,x],[2.7,3.0],color=col,lw=1,ls="--")
        else:
            ax.text(x,1.8,title,ha="center",fontsize=8.5,fontweight="bold",color=col)
            ax.text(x,0.5,desc,ha="center",fontsize=7.5,color=C_NAVY)
            ax.plot([x,x],[1.9,2.3],color=col,lw=1,ls="--")
        ax.text(x,2.5+(0.4 if i%2==0 else -0.65),dt,ha="center",fontsize=7,color=C_GREY)
    ax.set_title("Migration History & Schema Evolution Timeline",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)

# ── 21. PII Two-Pass Pipeline (CORRECTED) ─────────────────────
def diag_pii_corrected():
    fig,axes=plt.subplots(1,2,figsize=(14,6),facecolor=C_BG)
    fig.patch.set_facecolor(C_BG)

    # LEFT: Pass-1 (logger.REDACTED_FIELDS) vs _AUDIT_PII_FIELDS
    ax=axes[0]; ax.set_facecolor(C_BG); ax.axis("off")
    ax.text(0.5,1.0,"PASS-1: Field-Name Keyword Match (Synchronous, in structlog pipeline)",
            ha="center",fontsize=9,fontweight="bold",color=C_ORANGE,transform=ax.transAxes)
    logger_fields="password, passwd, secret, token, api_key, authorization,\n"\
                  "ssn, social_security, credit_card, card_number, cvv,\n"\
                  "dob, date_of_birth, birth_date, full_name, first_name,\n"\
                  "last_name, mobile, phone, phone_number, address\n"\
                  "(20 fields)"
    ax.text(0.02,0.82,"logger.REDACTED_FIELDS (20 fields - NO email):\n"+logger_fields,
            fontsize=8,color=C_NAVY,transform=ax.transAxes,va="top",
            bbox=dict(boxstyle="round,pad=0.4",facecolor="#FFF3E0",edgecolor=C_ORANGE))
    ax.text(0.02,0.38,"email handling:\n"
                      "NOT in REDACTED_FIELDS\n"
                      "Detected via '@' check\n"
                      "Masked as: u***@domain.com\n"
                      "(NOT [REDACTED])",
            fontsize=8.5,color=C_RED,transform=ax.transAxes,va="top",fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.4",facecolor="#FFEBEE",edgecolor=C_RED))
    ax.text(0.02,0.10,"supabase._AUDIT_PII_FIELDS (21 fields) DOES include email -> [REDACTED]\n"
                      "Only applies to previous_value / new_value (top-level keys only, no nesting)",
            fontsize=7.5,color=C_PURPLE,transform=ax.transAxes,va="top",style="italic")

    # RIGHT: Pass-2 Presidio
    ax2=axes[1]; ax2.set_facecolor(C_BG); ax2.axis("off")
    ax2.text(0.5,1.0,"PASS-2: Presidio NLP Content Scan (Async - NOT automatic)",
             ha="center",fontsize=9,fontweight="bold",color=C_PURPLE,transform=ax2.transAxes)
    ax2.text(0.02,0.82,
             "IMPORTANT: Pass-2 does NOT run automatically.\n"
             "The structlog pii_scrubber_processor uses Pass-1 ONLY.\n\n"
             "To use Pass-2, callers must explicitly call:\n"
             "   safe_text = await scrub_text(user_text)\n\n"
             "Entities detected by Presidio:\n"
             "  PERSON, EMAIL_ADDRESS, PHONE_NUMBER,\n"
             "  CREDIT_CARD, US_SSN, IBAN_CODE, LOCATION\n\n"
             "Replaced with: <PERSON>, <EMAIL_ADDRESS>, etc.\n"
             "(NOT [REDACTED] - angle-bracket format)\n\n"
             "Runs in ThreadPoolExecutor(max_workers=2)\n"
             "Graceful degradation if Presidio unavailable\n"
             "IP_ADDRESS intentionally excluded (forensics)",
             fontsize=8.5,color=C_NAVY,transform=ax2.transAxes,va="top",
             bbox=dict(boxstyle="round,pad=0.4",facecolor="#F3E5F5",edgecolor=C_PURPLE))

    fig.suptitle("PII Scrubbing Pipeline - CORRECTED (logger/__init__.py + supabase.py)",
                 fontsize=11,fontweight="bold",color=C_NAVY)
    fig.tight_layout(pad=1.5)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════
# INJECTION MAP  (paragraph index -> renderer)
# Indices from the BACKUP document (255 paragraphs)
# ══════════════════════════════════════════════════════════════
INJECTION_MAP = {
    40:  (diag_role_hierarchy,    "Figure 2.1 - Role Hierarchy Diagram"),
    78:  (diag_er,                "Figure 5 - ER Diagram: Entity-Relationship"),
    81:  (diag_schema_detail,     "Figure 6 - Schema Diagram: Column-Level Detail (corrected: details not hashed, top-level PII scrub only)"),
    84:  (diag_rls,               "Figure 7 - Row-Level Security (RLS) Policy Diagram"),
    89:  (diag_data_flow,         "Figure 8 - Data Flow Diagram: How Logs Are Created"),
    92:  (diag_seq_login,         "Figure 9 - Sequence Diagram: Login Audit Event"),
    95:  (diag_seq_pod,           "Figure 10 - Sequence Diagram: Admin Approves Pod"),
    98:  (diag_seq_view,          "Figure 11 - Sequence Diagram: Viewing Activity Logs"),
    101: (diag_component_arch,    "Figure 12 - Component Architecture Diagram"),
    105: (diag_hash_auth,         "Figure 13.1 - auth_audit Integrity Hash Flow (corrected: email masked as u***@domain, not [REDACTED])"),
    107: (diag_hash_audit,        "Figure 13.2 - audit_logs Hash + Chain Flow (corrected: details excluded from hash, top-level PII scrub only)"),
    126: (diag_access_matrix,     "Figure 16 - Access Control Matrix: Audit Row Visibility by Role"),
    136: (diag_route_guard,       "Figure 17 - Navigation & Route Guard Diagram"),
    139: (diag_backend_logging,   "Figure 18 - Backend Logging Architecture (corrected: Pass-2 not automatic, structlog line from log_ai_decision_audit)"),
    144: (diag_full_system,       "Figure 19 - Full System Interaction Diagram (corrected)"),
    181: (diag_ledger_trigger,    "Figure 21 - Cryptographic Ledger Trigger Sequence"),
    213: (diag_shap_flow,         "Figure 22 - SHAP AI Decision Audit Flow (corrected: +0.4200 format, structlog also emitted)"),
    223: (diag_shap_seq,          "Figure 22b - SHAP Data Flow Sequence (corrected)"),
    242: (diag_verify,            "Figure 24 - Audit Ledger Verification Flow"),
    249: (diag_migration_timeline,"Figure 25 - Migration History & Schema Evolution Timeline"),
}

# ══════════════════════════════════════════════════════════════
# TEXT CORRECTIONS  (search string -> corrected full paragraph)
# ══════════════════════════════════════════════════════════════
TEXT_FIXES = [
    # FIX 1: PII fields list - remove 'email' from list, correct count
    ("Covered fields (20 total): password, passwd, secret, token, api_key, authorization,",
     "logger.REDACTED_FIELDS - 20 fields (email NOT included here): "
     "password, passwd, secret, token, api_key, authorization, ssn, social_security, "
     "credit_card, card_number, cvv, dob, date_of_birth, birth_date, full_name, first_name, "
     "last_name, mobile, phone, phone_number, address. "
     "Email is handled separately: detected via '@' character and masked as u***@domain (not [REDACTED]). "
     "NOTE: supabase._AUDIT_PII_FIELDS (used in _scrub_audit_value) has 21 fields and DOES include email -> [REDACTED], "
     "but this only applies to previous_value/new_value top-level keys."
    ),

    # FIX 2: Pass-2 scrubbing description - clarify NOT automatic
    ("After field-name scrubbing, all remaining string values are scanned for PII content",
     "Pass-2 (Presidio) is NOT automatic. The structlog pii_scrubber_processor runs Pass-1 ONLY. "
     "Pass-2 must be explicitly called by the route handler or service using: "
     "safe_text = await scrub_text(user_supplied_text). "
     "When called, it offloads NLP analysis to a dedicated ThreadPoolExecutor(max_workers=2) "
     "so the asyncio event loop is never blocked. "
     "Presidio replaces entities with angle-bracket tags: <PERSON>, <EMAIL_ADDRESS>, etc. "
     "(not [REDACTED]). Graceful degradation to Pass-1 only if Presidio is unavailable."
    ),

    # FIX 3: SHAP format - 4 decimal places
    ("Top factors: transaction_amount (+0.42, up), account_age (-0.31, down)",
     "Top factors: transaction_amount (+0.4200, up), account_age (-0.3100, down), "
     "hour_of_day (+0.1800, up). "
     "NOTE: actual format is :+.4f (4 decimal places), e.g. +0.4200 not +0.42."
    ),

    # FIX 4: integrity_hash fields - clarify details excluded
    ("integrity_hash: SHA-256 over all non-PII application-layer fields including SHAP",
     "integrity_hash: SHA-256 over these fields: user_id, action, resource_type, resource_id, "
     "previous_value (json, scrubbed), new_value (json, scrubbed), model_id, prediction (json), "
     "shap_values (json), reasoning, created_at. "
     "IMPORTANT: 'details' field is stored in the row but is NOT included in the integrity_hash."
    ),

    # FIX 5: _scrub_audit_value depth
    ("Only top-level keys are scrubbed",
     "Only top-level keys are scrubbed (no recursion into nested dicts). "
     "This is intentional - previous_value/new_value must contain only shallow non-PII fields "
     "(role names, status strings, UUIDs). "
     "Contrast with logger._scrub_value() which recurses up to _MAX_SCRUB_DEPTH=10."
    ),
]

# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    if not os.path.exists(BACKUP):
        print(f"ERROR: Backup not found: {BACKUP}")
        sys.exit(1)

    print(f"Loading backup: {BACKUP}")
    doc = Document(BACKUP)
    print(f"  {len(doc.paragraphs)} paragraphs loaded")

    # ── Step 1: Apply text corrections ──────────────────────────
    print("\nApplying text corrections ...")
    for search, replacement in TEXT_FIXES:
        found = fix_para(doc, search[:50], replacement)
        status = "OK" if found else "NOT FOUND"
        print(f"  [{status}] fix: '{search[:55]}...'")

    # ── Step 2: Insert PII corrected diagram into section 20 ────
    # Replace the Mermaid code block at para 147 (Pass 1 flowchart)
    # We'll inject our corrected PII diagram after the "Pass 1" heading
    # (by searching for that heading's text)
    print("\nInjecting extra corrected PII diagram ...")
    pii_inserted = False
    for i, para in enumerate(doc.paragraphs):
        if "Pass 1" in para.text and "Field-name Keyword" in para.text and para.style.name.startswith("Heading"):
            # Find the next Normal paragraph (the code/flowchart block) after this heading
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                txt = doc.paragraphs[j].text.strip()
                if txt.startswith(("flowchart","graph","sequenceDiagram","erDiagram","classDiagram")):
                    stream = diag_pii_corrected()
                    replace_para_with_picture(doc, j, stream,
                        caption="Figure 20 - PII Scrubbing Pipeline (CORRECTED: email masked not [REDACTED], Pass-2 not automatic)")
                    pii_inserted = True
                    print("  OK  PII corrected diagram injected")
                    break
            break
    if not pii_inserted:
        print("  SKIP PII diagram (heading/code not found, will be in main injection pass)")

    # ── Step 3: Replace all remaining Mermaid blocks with images ─
    print("\nReplacing Mermaid code blocks with PNG diagrams ...")
    # Process in REVERSE order so removals don't shift earlier indices
    for orig_idx in sorted(INJECTION_MAP.keys(), reverse=True):
        fn, caption = INJECTION_MAP[orig_idx]
        if orig_idx >= len(doc.paragraphs):
            print(f"  SKIP  para {orig_idx} (out of range)")
            continue
        # Only replace if the paragraph looks like a diagram code block
        txt = doc.paragraphs[orig_idx].text.strip()
        is_code = any(txt.startswith(k) for k in
                      ("flowchart","graph","sequenceDiagram","erDiagram",
                       "classDiagram","quadrantChart","timeline","sequenceDiag"))
        if not is_code:
            print(f"  SKIP  para {orig_idx} (not a code block: '{txt[:40]}')")
            continue
        print(f"  OK    para {orig_idx}: {caption[:55]}...")
        stream = fn()
        replace_para_with_picture(doc, orig_idx, stream,
                                   width=Inches(5.8), caption=caption)

    # ── Step 4: Save ─────────────────────────────────────────────
    doc.save(OUT)
    print(f"\nSaved -> {OUT}")

    # Verify
    import zipfile
    with zipfile.ZipFile(OUT) as z:
        imgs = [n for n in z.namelist() if "media" in n]
    print(f"Embedded images: {len(imgs)}")
    print("Done.")
