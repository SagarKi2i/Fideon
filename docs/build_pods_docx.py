from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

NAVY  = RGBColor(0x1F, 0x39, 0x64)
BLUE  = RGBColor(0x2E, 0x74, 0xB5)
TEAL  = RGBColor(0x1F, 0x7A, 0x8C)
GRAY  = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def h(level, text):
    p = doc.add_heading(text, level=level)
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt({1:18,2:14,3:12,4:11}[level])
    run.font.color.rgb = {1:NAVY,2:BLUE,3:TEAL,4:GRAY}[level]
    p.style = doc.styles[f'Heading {level}']

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def sp():
    doc.add_paragraph('')

def shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)

def code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, hdr in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = ''
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, '2E74B5')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 0:
                shade_cell(cell, 'EBF3FB')
    if widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(widths):
                    cell.width = Inches(widths[ci])
    return t

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('NeuraBOX Insurance Pods')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('Reference Documentation')
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = BLUE

sp()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run('FNOL  •  Quote Generation  •  Policy Comparison  •  ACORD Form Extraction')
r3.font.size = Pt(11); r3.font.color.rgb = GRAY

tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run('Last updated: 2026-04-15')
r4.font.size = Pt(9); r4.font.color.rgb = GRAY; r4.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. POD ARCHITECTURE OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
h(1, '1. Pod Architecture Overview')
body('All four pods share the same generic pod framework. Each pod is identified by a pod_id string '
     'and plugs into the same extraction → feedback → admin-review → fine-tuning pipeline.')
sp()
h(3, 'Extraction Lifecycle')
for s in [
    'User uploads document / provides input',
    'POST /api/pods/{pod_id}/extract',
    'Pod-specific extractor (LLM + heuristics + OCR)',
    'Run persisted to pod_extraction_runs (status = draft)',
    'User reviews extracted JSON and submits feedback (thumbs_up/down + optional corrections)',
    'Confidence evaluation: score >= threshold → auto-approved; below threshold → admin queue',
    'Admin reviews: approve / rework / reject',
    'Training job spawned on approval (if AUTO_FINE_TUNE_ON_POD_APPROVAL = true)',
    'Quality gates evaluated → model update if gates pass',
]:
    bullet(s)

sp()
h(3, 'Pack Membership')
tbl(['Pack','Pods Included'],[
    ['Underwriting Pack','quote-generation, policy-comparison, acord-parser (+ 12 others)'],
    ['Claims Pack','claims-fnol (+ 4 others)'],
    ['Distribution Pack','policy-comparison (+ 6 others)'],
],[2.0,4.5])

# ════════════════════════════════════════════════════════════════════════════
# 2. SHARED DATABASE SCHEMA
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '2. Shared Database Schema')
body('Migration: supabase/migrations/20260320120000_pod_workflow_shared_tables.sql\n\n'
     'All four pods write to these five shared tables, identified by the pod_id column. '
     'The ACORD pod additionally has its own dedicated tables (Section 3).')

sp()
h(2, '2.1  pod_extraction_runs')
body('Core record created for every extraction attempt.')
tbl(['Column','Type','Description'],[
    ['id',                      'UUID PK',               'Auto-generated run identifier'],
    ['created_at',              'TIMESTAMPTZ NOT NULL',  'Row creation timestamp'],
    ['updated_at',              'TIMESTAMPTZ NOT NULL',  'Auto-updated on every change (trigger)'],
    ['created_by',              'UUID FK → auth.users',  'User who triggered the extraction'],
    ['pod_id',                  'TEXT NOT NULL',         'e.g. claims-fnol, quote-generation, policy-comparison, acord_form_understanding'],
    ['source_filename',         'TEXT',                  'Original uploaded filename'],
    ['source_mime',             'TEXT',                  'MIME type of uploaded file'],
    ['raw_text',                'TEXT',                  'Raw text extracted from document'],
    ['extracted_json',          'JSONB NOT NULL',        'Pod-specific structured output (see per-pod schema). Default: {}.'],
    ['original_extracted_json', 'JSONB NOT NULL',        'Immutable copy of the first extraction result. Set once on creation; never overwritten. Default: {}. (Added: migration 20260323120000)'],
    ['overall_confidence',      'DOUBLE PRECISION NOT NULL', 'Model confidence score 0.0 – 1.0. Default: 0.'],
    ['status',                  'TEXT NOT NULL',         "Lifecycle state. CHECK: draft | submitted | needs_admin_review | approved | rejected. Default: 'draft'."],
],[1.9,1.9,2.8]
)
body('Indexes: created_by, pod_id, status, created_at DESC.  '
     'RLS: users read/write own rows; admins manage all rows.  '
     'Note: UNIQUE(run_id) on pod_training_jobs was dropped in migration 20260323120000 to allow retrain history (multiple jobs per run).')

sp()
h(2, '2.2  pod_extraction_feedback')
body('User or admin correction attached to a run.')
tbl(['Column','Type','Description'],[
    ['id',             'UUID PK',                       ''],
    ['created_at',     'TIMESTAMPTZ NOT NULL',          ''],
    ['created_by',     'UUID FK → auth.users NOT NULL', ''],
    ['pod_id',         'TEXT NOT NULL',                 'Mirrors the parent run\'s pod_id'],
    ['run_id',         'UUID FK → pod_extraction_runs', 'CASCADE delete'],
    ['actor_role',     'TEXT NOT NULL',                 "CHECK: 'user' | 'admin'"],
    ['thumbs_up',      'BOOLEAN',                       'Positive / negative validation signal'],
    ['notes',          'TEXT',                          'Free-text annotation'],
    ['corrected_json', 'JSONB',                         'Full or partial corrected output JSON'],
],[1.8,2.0,2.8])

sp()
h(2, '2.3  pod_admin_queue')
body('One row per run requiring admin review. Created automatically when confidence falls below threshold.')
tbl(['Column','Type','Description'],[
    ['run_id',      'UUID PK FK → pod_extraction_runs', 'One queue entry per run. CASCADE delete.'],
    ['pod_id',      'TEXT NOT NULL',                    ''],
    ['created_at',  'TIMESTAMPTZ NOT NULL',             ''],
    ['updated_at',  'TIMESTAMPTZ NOT NULL',             ''],
    ['priority',    'INTEGER NOT NULL',                 'Higher value = reviewed first. Default: 0.'],
    ['reason',      'TEXT',                             'Why this run was queued'],
    ['assigned_to', 'UUID FK → auth.users',             'Admin assignee. SET NULL on user delete.'],
    ['state',       'TEXT NOT NULL',                    "CHECK: open | in_progress | approved | rework | rejected. Default: 'open'."],
],[2.2,1.8,2.6])

sp()
h(2, '2.4  pod_training_jobs')
body('Fine-tuning jobs triggered by admin approval. Multiple jobs per run are allowed (UNIQUE constraint removed in migration 20260323120000 to support retrain history).')
tbl(['Column','Type','Description'],[
    ['id',                     'UUID PK',                ''],
    ['created_at',             'TIMESTAMPTZ NOT NULL',   ''],
    ['updated_at',             'TIMESTAMPTZ NOT NULL',   ''],
    ['pod_id',                 'TEXT NOT NULL',          ''],
    ['run_id',                 'UUID FK → pod_extraction_runs', 'CASCADE delete. Multiple jobs per run allowed.'],
    ['created_by',             'UUID FK → auth.users',   'SET NULL on user delete.'],
    ['status',                 'TEXT NOT NULL',          "CHECK: queued | running | completed | failed. Default: 'queued'."],
    ['dataset_path',           'TEXT',                   'Path to prepared training dataset'],
    ['output_dir',             'TEXT',                   'Where fine-tuned weights are saved'],
    ['log_path',               'TEXT',                   'Training log file path'],
    ['error',                  'TEXT',                   'Error message if status = failed'],
    ['started_at',             'TIMESTAMPTZ',            'Job start time'],
    ['finished_at',            'TIMESTAMPTZ',            'Job completion/failure time'],
],[1.8,1.8,3.0])

sp()
h(2, '2.5  pod_eval_results')
body('Evaluation metrics recorded after each training job. UNIQUE on (job_id, eval_set).')
tbl(['Column','Type','Description'],[
    ['id',                 'UUID PK',                         ''],
    ['created_at',         'TIMESTAMPTZ NOT NULL',            ''],
    ['pod_id',             'TEXT NOT NULL',                   ''],
    ['job_id',             'UUID FK → pod_training_jobs',     'SET NULL on job delete.'],
    ['eval_set',           'TEXT NOT NULL',                   "CHECK: seen | paraphrased | oos | combined"],
    ['exact_match',        'DOUBLE PRECISION',                'Fraction of fields with exact match'],
    ['soft_accuracy',      'DOUBLE PRECISION',                'Near-match accuracy'],
    ['semantic_sim',       'DOUBLE PRECISION',                'Embedding cosine similarity'],
    ['hallucination_rate', 'DOUBLE PRECISION',                'Fraction of hallucinated fields'],
    ['refusal_rate',       'DOUBLE PRECISION',                'Fraction of refused outputs'],
    ['metrics_json',       'JSONB NOT NULL',                  'Full metrics payload. Default: {}.'],
    ['notes',              'TEXT',                            'Reviewer notes'],
],[1.9,1.8,3.0])
body('Unique constraint: pod_eval_results_job_evalset_unique on (job_id, eval_set).')

# ════════════════════════════════════════════════════════════════════════════
# 3. ACORD-SPECIFIC DATABASE TABLES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '3. ACORD-Specific Database Tables')
body('The ACORD pod has its own dedicated set of tables in addition to using the shared pod tables. '
     'These mirror the shared schema but are ACORD-specific and do not carry a pod_id column.')

sp()
h(2, '3.1  acord_extraction_runs')
body('Source migration: 20260318120000_acord_extraction_workflow.sql')
tbl(['Column','Type','Description'],[
    ['id',                 'UUID PK',               'Auto-generated run identifier'],
    ['created_at',         'TIMESTAMPTZ NOT NULL',  ''],
    ['updated_at',         'TIMESTAMPTZ NOT NULL',  'Auto-updated via trigger'],
    ['created_by',         'UUID FK → auth.users NOT NULL', 'CASCADE delete'],
    ['source_filename',    'TEXT',                  'Uploaded filename'],
    ['source_mime',        'TEXT',                  'MIME type'],
    ['form_type_detected', 'TEXT',                  'ACORD form type detected (e.g. "ACORD 25", "ACORD 125"). ACORD-only field.'],
    ['raw_text',           'TEXT',                  'Raw extracted text'],
    ['extracted_json',     'JSONB NOT NULL',         'Structured AcordFormSummary output. Default: {}.'],
    ['overall_confidence', 'DOUBLE PRECISION NOT NULL', '0.0 – 1.0. Default: 0.'],
    ['status',             'TEXT NOT NULL',          "CHECK: draft | submitted | needs_admin_review | approved | rejected. Default: 'draft'."],
],[1.9,1.9,2.8])
body('Indexes: created_by, status, created_at DESC.')
body('RLS policies:\n'
     '  • Users can insert own rows (created_by = auth.uid())\n'
     '  • Users can read/update own rows (draft/submitted only)\n'
     '  • Admins can manage all rows')

sp()
h(2, '3.2  acord_extraction_feedback')
body('Source migration: 20260318120000_acord_extraction_workflow.sql  '
     '(re-synced in 20260318150000_acord_extraction_feedback_resync.sql)')
tbl(['Column','Type','Description'],[
    ['id',             'UUID PK',                       ''],
    ['created_at',     'TIMESTAMPTZ NOT NULL',          ''],
    ['created_by',     'UUID FK → auth.users NOT NULL', 'CASCADE delete'],
    ['run_id',         'UUID FK → acord_extraction_runs NOT NULL', 'CASCADE delete'],
    ['actor_role',     'TEXT NOT NULL',                 "CHECK: 'user' | 'admin'"],
    ['thumbs_up',      'BOOLEAN',                       'Validation signal'],
    ['notes',          'TEXT',                          'Annotation text'],
    ['corrected_json', 'JSONB',                         'User or admin corrections'],
],[1.8,2.1,2.7])
body('Indexes: run_id, created_at DESC.')

sp()
h(2, '3.3  acord_admin_queue')
body('Source migration: 20260318120000_acord_extraction_workflow.sql')
tbl(['Column','Type','Description'],[
    ['run_id',      'UUID PK FK → acord_extraction_runs', 'CASCADE delete. One entry per run.'],
    ['created_at',  'TIMESTAMPTZ NOT NULL',               ''],
    ['updated_at',  'TIMESTAMPTZ NOT NULL',               'Auto-updated via trigger'],
    ['priority',    'INTEGER NOT NULL',                   'Higher = reviewed first. Default: 0.'],
    ['reason',      'TEXT',                               'Why queued'],
    ['assigned_to', 'UUID FK → auth.users',               'SET NULL on user delete'],
    ['state',       'TEXT NOT NULL',                      "CHECK: open | in_progress | approved | rework. Default: 'open'. (Note: no 'rejected' state unlike shared pod_admin_queue)"],
],[2.1,1.8,2.7])
body('Indexes: state, priority DESC.')

sp()
h(2, '3.4  acord_training_jobs')
body('Source migration: 20260318123000_acord_training_jobs.sql')
tbl(['Column','Type','Description'],[
    ['id',           'UUID PK',               ''],
    ['created_at',   'TIMESTAMPTZ NOT NULL',  ''],
    ['updated_at',   'TIMESTAMPTZ NOT NULL',  'Auto-updated via trigger'],
    ['run_id',       'UUID FK → acord_extraction_runs NOT NULL', 'CASCADE delete'],
    ['created_by',   'UUID FK → auth.users',  'SET NULL on user delete'],
    ['status',       'TEXT NOT NULL',         "CHECK: queued | running | completed | failed. Default: 'queued'."],
    ['dataset_path', 'TEXT',                  'Path to training dataset'],
    ['output_dir',   'TEXT',                  'Fine-tuned weights output directory'],
    ['log_path',     'TEXT',                  'Training log path'],
    ['error',        'TEXT',                  'Error message if failed'],
    ['started_at',   'TIMESTAMPTZ',           ''],
    ['finished_at',  'TIMESTAMPTZ',           ''],
],[1.7,1.8,3.1])
body('Indexes: run_id, status, created_at DESC.')

sp()
h(2, '3.5  acord_eval_results')
body('Source migration: 20260319110000_acord_eval_results_unique_job_evalset.sql  '
     '(separate from shared pod_eval_results)')
tbl(['Column','Type','Description'],[
    ['id',                 'UUID PK',   ''],
    ['created_at',         'TIMESTAMPTZ NOT NULL', ''],
    ['job_id',             'UUID FK → acord_training_jobs', 'SET NULL on job delete'],
    ['eval_set',           'TEXT NOT NULL', "CHECK: seen | paraphrased | oos | combined"],
    ['exact_match',        'DOUBLE PRECISION', ''],
    ['soft_accuracy',      'DOUBLE PRECISION', ''],
    ['semantic_sim',       'DOUBLE PRECISION', ''],
    ['hallucination_rate', 'DOUBLE PRECISION', ''],
    ['refusal_rate',       'DOUBLE PRECISION', ''],
    ['metrics_json',       'JSONB NOT NULL',   'Default: {}'],
    ['notes',              'TEXT',             ''],
],[1.9,1.8,3.0])
body('Unique constraint: (job_id, eval_set).')

sp()
h(2, '3.6  acord_extract_jobs  (Async Job Tracking)')
body('Source migrations: 20260331120000_acord_extract_jobs.sql, '
     '20260331133000_acord_extract_jobs_user_id_normalization.sql, '
     '20260331143000_acord_extract_jobs_phase.sql\n\n'
     'Persists async extraction job state across process restarts so the client can poll '
     'GET /api/acord/extract/status/{job_id} for progress. '
     'This is separate from acord_training_jobs — those are for fine-tuning; '
     'this table tracks live extraction jobs.')
tbl(['Column','Type','Description'],[
    ['job_id',     'UUID PK',              'Client-provided or server-generated job identifier'],
    ['created_at', 'TIMESTAMPTZ NOT NULL', ''],
    ['updated_at', 'TIMESTAMPTZ NOT NULL', 'Auto-updated via trigger'],
    ['user_id',    'UUID FK → auth.users NOT NULL', 'CASCADE delete'],
    ['status',     'TEXT NOT NULL',        "CHECK: queued | running | succeeded | failed. Default: 'queued'."],
    ['phase',      'TEXT',                 'Granular progress marker: queued | warming_model | generate_extracting | completed | failed. (Added: migration 20260331143000)'],
    ['error',      'TEXT',                 'Error message if status = failed'],
    ['result',     'JSONB',                'Full AcordExtractResponse payload once status = succeeded'],
],[1.6,2.0,3.0])
body('Indexes: user_id, status, created_at DESC.\n'
     'RLS: users can select/insert/update own rows; admins manage all rows.')

# ════════════════════════════════════════════════════════════════════════════
# 4. SHARED API SCHEMAS (Pydantic)
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '4. Shared API Schemas  (backend/app/schemas/pod_workflow.py)')

h(2, '4.1  PodExtractResponse')
body('Returned by POST /{pod_id}/extract for all generic pods.')
tbl(['Field','Type','Description'],[
    ['run_id',             'str',   'UUID of the created pod_extraction_runs row'],
    ['status',             'str',   "Lifecycle state, e.g. 'draft'"],
    ['overall_confidence', 'float', 'Model confidence score 0.0 – 1.0. Default: 0.0.'],
    ['extracted',          'dict',  'Pod-specific structured output (see per-pod schema). Default: {}.'],
],[2.0,1.0,3.6])

sp()
h(2, '4.2  PodSubmitRequest')
body('Sent by user to POST /{pod_id}/runs/{run_id}/submit.')
tbl(['Field','Type','Required','Description'],[
    ['thumbs_up',                          'bool',       'Yes', 'User validation signal (true = good extraction)'],
    ['require_admin_approval_for_training','bool',       'No',  'If true, route to admin queue even for high-confidence thumbs_up. Default: false.'],
    ['notes',                              'str | None', 'No',  'Free-text annotation'],
    ['corrected_json',                     'dict | None','No',  'User-corrected extracted fields (full or partial JSON)'],
],[2.4,1.0,0.8,2.4])

sp()
h(2, '4.3  PodAdminReviewRequest')
body('Sent by admin to POST /{pod_id}/admin/{run_id}/review.')
tbl(['Field','Type','Required','Description'],[
    ['decision',      'str',        'Yes', "approve | rework | reject"],
    ['notes',         'str | None', 'No',  'Review notes'],
    ['corrected_json','dict | None','No',  'Admin-corrected fields'],
    ['assigned_to',   'str | None', 'No',  'Reassign to another admin UUID'],
],[2.0,1.0,0.8,2.8])

sp()
h(2, '4.4  PodBatchReviewRequest')
tbl(['Field','Type','Required','Description'],[
    ['run_ids', 'list[str]',  'Yes', '1 – 50 run UUIDs to batch-review'],
    ['decision','str',        'Yes', "approve | reject"],
    ['notes',   'str | None', 'No',  'Applied to all reviewed runs'],
],[2.0,1.0,0.8,2.8])

sp()
h(2, '4.5  PodReExtractRequest')
tbl(['Field','Type','Required','Description'],[
    ['extraction_hint','str | None','No','Optional hint to guide re-extraction (pod-specific free text)'],
],[2.0,1.0,0.8,2.8])

# ════════════════════════════════════════════════════════════════════════════
# 5. ACORD-SPECIFIC API SCHEMAS (Pydantic)
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '5. ACORD-Specific API Schemas  (backend/app/schemas/acord_workflow.py)')
body('The ACORD pod uses its own schema classes rather than the generic pod schemas. '
     'Key differences: AcordExtractResponse adds partial/persist_error/warning fields; '
     'ReExtractRequest uses form_type_hint instead of extraction_hint; '
     'AcordExtractJobStatusResponse exposes async phase.')

sp()
h(2, '5.1  AcordExtractResponse')
body('Returned by POST /api/acord/extract (synchronous) and embedded in AcordExtractJobStatusResponse.')
tbl(['Field','Type','Description'],[
    ['run_id',             'str',        "UUID of the created acord_extraction_runs row. Default: ''."],
    ['status',             'str',        "Lifecycle state. Default: 'draft'."],
    ['overall_confidence', 'float',      'Confidence score 0.0 – 1.0. Default: 0.0.'],
    ['extracted',          'dict',       'Full AcordFormSummary as dict. Default: {}.'],
    ['partial',            'bool',       'True when extraction succeeded but DB persist failed. Client still receives extracted data. Default: false.'],
    ['persist_error',      'str | None', 'Error message from DB persist failure (when partial=true)'],
    ['warning',            'str | None', 'Human-readable note (e.g. "DB unavailable; submit may be disabled until saved")'],
],[2.0,1.2,3.4])

sp()
h(2, '5.2  AcordSubmitRequest')
body('Sent by user to POST /api/acord/runs/{run_id}/submit. Identical structure to PodSubmitRequest.')
tbl(['Field','Type','Required','Description'],[
    ['thumbs_up',                          'bool',       'Yes','User validation signal'],
    ['require_admin_approval_for_training','bool',       'No', 'Force admin queue even for high-confidence submissions. Default: false.'],
    ['notes',                              'str | None', 'No', 'Free-text annotation'],
    ['corrected_json',                     'dict | None','No', 'User-corrected extracted fields'],
],[2.4,1.0,0.8,2.4])

sp()
h(2, '5.3  AcordAdminReviewRequest')
tbl(['Field','Type','Required','Description'],[
    ['decision',      'str',        'Yes',"approve | rework | reject"],
    ['notes',         'str | None', 'No', 'Review notes'],
    ['corrected_json','dict | None','No', 'Admin-corrected fields'],
    ['assigned_to',   'str | None', 'No', 'Reassign to another admin UUID'],
],[2.0,1.0,0.8,2.8])

sp()
h(2, '5.4  AcordBatchReviewRequest')
tbl(['Field','Type','Required','Description'],[
    ['run_ids', 'list[str]',  'Yes','1 – 50 ACORD run UUIDs to batch-review'],
    ['decision','str',        'Yes',"approve | reject"],
    ['notes',   'str | None', 'No', 'Applied to all reviewed runs'],
],[2.0,1.0,0.8,2.8])

sp()
h(2, '5.5  ReExtractRequest')
body('Sent to POST /api/acord/runs/{run_id}/re-extract.  '
     'Note: uses form_type_hint (not the generic extraction_hint).')
tbl(['Field','Type','Required','Description'],[
    ['form_type_hint','str | None','No',"Override the form type for re-extraction. e.g. '25' or '125'. Overrides original hint."],
],[2.2,1.0,0.8,2.6])

sp()
h(2, '5.6  AcordExtractStartResponse')
body('Returned by POST /api/acord/extract/start (async job submission).')
tbl(['Field','Type','Description'],[
    ['job_id','str',"UUID of the created acord_extract_jobs row"],
    ['status','str',"Initial job status. Default: 'queued'."],
],[2.0,1.0,3.6])

sp()
h(2, '5.7  AcordExtractJobStatusResponse')
body('Returned by GET /api/acord/extract/status/{job_id}. '
     'Client polls this until status = succeeded or failed.')
tbl(['Field','Type','Description'],[
    ['job_id', 'str',                         'Job UUID'],
    ['status', 'str',                         "CHECK: queued | running | succeeded | failed"],
    ['phase',  'str | None',                  'Granular progress: queued | warming_model | generate_extracting | completed | failed'],
    ['result', 'AcordExtractResponse | None', 'Full extraction result. Populated when status = succeeded.'],
    ['error',  'str | None',                  'Error message. Populated when status = failed.'],
],[2.0,1.2,3.4])

sp()
h(2, '5.8  PreviewSftTrainingRecordBody')
body('Used by the admin preview endpoint to inspect what a training JSONL record would look like for a given run.')
tbl(['Field','Type','Description'],[
    ['extracted_json',  'dict',       'Extracted fields to format as training record. Default: {}.'],
    ['raw_text',        'str',        "Source text that produced the extraction. Default: ''."],
    ['source_filename', 'str | None', 'Original filename (informational)'],
],[2.0,1.2,3.4])

# ════════════════════════════════════════════════════════════════════════════
# 6. SHARED API REFERENCE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '6. Shared API Reference')
body('Base prefix: /api/pods/{pod_id}     Routes file: backend/app/routes/pods.py')

h(2, '6.1  User Endpoints')
tbl(['Method','Path','Description'],[
    ['POST','/{pod_id}/extract',                  'Submit file/text for extraction. Returns PodExtractResponse.'],
    ['GET', '/{pod_id}/runs/{run_id}',            'Retrieve a single extraction result.'],
    ['POST','/{pod_id}/runs/{run_id}/re-extract', 'Re-run extraction with optional hint. Returns PodExtractResponse.'],
    ['POST','/{pod_id}/runs/{run_id}/submit',     'Submit user feedback. Body: PodSubmitRequest.'],
    ['GET', '/{pod_id}/runs',                     'List paginated runs for the authenticated user.'],
],[0.8,2.8,3.0])

sp()
h(2, '6.2  Admin Endpoints')
tbl(['Method','Path','Description'],[
    ['GET',   '/{pod_id}/admin/queue/stats',             'Queue size and state breakdown'],
    ['GET',   '/{pod_id}/admin/queue',                   'Paginated queue (filterable by state/priority)'],
    ['POST',  '/{pod_id}/admin/{run_id}/review',         'Review a single run. Body: PodAdminReviewRequest.'],
    ['POST',  '/{pod_id}/admin/batch-review',            'Review 1–50 runs. Body: PodBatchReviewRequest.'],
    ['GET',   '/{pod_id}/admin/queue/{run_id}/detail',   'Full queue item detail'],
    ['PATCH', '/{pod_id}/admin/queue/{run_id}/detail',   'Update queue item (e.g. reassign)'],
    ['GET',   '/{pod_id}/admin/jobs',                    'List training jobs'],
    ['GET',   '/{pod_id}/admin/jobs/by-run/{run_id}',    'Get most recent job for a specific run'],
    ['GET',   '/{pod_id}/admin/jobs/by-run/{run_id}/history', 'All job history for a run'],
    ['GET',   '/{pod_id}/admin/jobs/{job_id}',           'Job details'],
    ['GET',   '/{pod_id}/admin/jobs/{job_id}/eval',      'Evaluation metrics (pod_eval_results rows)'],
    ['GET',   '/{pod_id}/admin/jobs/{job_id}/log',       'Raw training log'],
    ['GET',   '/{pod_id}/admin/runs/{run_id}/health-card','Confidence + feedback health metrics'],
],[0.8,2.8,3.0])

# ════════════════════════════════════════════════════════════════════════════
# 7. FNOL
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '7. Pod 1 — FNOL (Claims Intelligence)')

h(2, '7.1  Identity')
tbl(['Field','Value'],[
    ['Pod ID','claims-fnol'],
    ['Category','Claims'],
    ['Segment','Broker'],
    ['Pack','Claims Pack'],
    ['Frontend UI','frontend/src/components/playground/ClaimsFNOLUI.tsx'],
    ['Database','Shared pod tables (pod_id = claims-fnol)'],
    ['API prefix','/api/pods/claims-fnol'],
],[1.8,4.8])

sp()
h(2, '7.2  Description')
body('The FNOL pod is the first point of contact when a loss event occurs. It accepts a natural-language '
     'description of an incident plus optional supporting documents and produces a structured analysis report '
     'covering incident classification, coverage applicability, recommended next steps, and documentation '
     'requirements. Designed to accelerate the claims intake process for brokers and adjusters.')

sp()
h(2, '7.3  Required Input')
tbl(['Field','Type','Required','Details'],[
    ['description','string (textarea)','Yes',
     'Free-text incident description. Should include: date, location, parties involved, '
     'nature of loss, known damages. Minimum ~50 characters.'],
    ['file','File (PDF / DOCX / JPG / PNG)','No',
     'Supporting document: police report, photos, repair estimate, witness statement, etc.'],
],[1.2,1.5,0.8,3.1])

sp()
h(3, 'Example payload')
code('{\n  "type": "claims-fnol",\n  "description": "On March 15 2026 at 2pm, insured vehicle rear-ended a truck at 5th Ave\n'
     '               and Main St. $8,000 estimated damage. No injuries. Police report #2026-4821.",\n'
     '  "file": "police_report.pdf"\n}')

sp()
h(2, '7.4  Expected Output')
body('The pod returns a markdown-rendered FNOL Analysis Report. '
     'The full structured result is stored in pod_extraction_runs.extracted_json.')
tbl(['Output Section','Description'],[
    ['Incident Summary','Structured recap: date, location, parties, damages'],
    ['Loss Classification','Line of business, peril type, coverage trigger'],
    ['Parties Involved','Insured, third parties, witnesses'],
    ['Initial Coverage Assessment','Whether the loss appears covered under standard policy lines'],
    ['Documentation Checklist','Required supporting documents still needed'],
    ['Recommended Next Steps','Immediate actions for the broker/adjuster'],
    ['Subrogation Flag','Whether recovery from a third party may be possible'],
    ['Fraud Indicators','Patterns that warrant additional review'],
],[2.2,4.4])

sp()
h(3, 'extracted_json schema (stored in pod_extraction_runs)')
code('{\n  "incident_date":           "2026-03-15",\n'
     '  "incident_location":        "5th Ave & Main St",\n'
     '  "loss_type":                "collision",\n'
     '  "line_of_business":         "auto",\n'
     '  "estimated_damage":         8000,\n'
     '  "injuries_reported":        false,\n'
     '  "police_report_number":     "2026-4821",\n'
     '  "parties": [\n'
     '    { "role": "insured",      "name": "..." },\n'
     '    { "role": "third_party",  "name": "..." }\n'
     '  ],\n'
     '  "coverage_applicable":      true,\n'
     '  "subrogation_potential":    true,\n'
     '  "documentation_required":   ["repair_estimate", "police_report", "photos"],\n'
     '  "recommended_actions":      ["..."],\n'
     '  "confidence_notes":         "..."\n}')
body('Confidence: 0.0–1.0. Auto-approve threshold: POD_CONFIDENCE_THRESHOLD (default 0.85).')

# ════════════════════════════════════════════════════════════════════════════
# 8. QUOTE GENERATION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '8. Pod 2 — Quote Generation')

h(2, '8.1  Identity')
tbl(['Field','Value'],[
    ['Pod ID','quote-generation'],
    ['Category','Automation'],
    ['Segment','Broker'],
    ['Pack','Underwriting Pack'],
    ['Frontend UI','frontend/src/components/playground/QuoteGenerationUI.tsx'],
    ['Database','Shared pod tables (pod_id = quote-generation)'],
    ['API prefix','/api/pods/quote-generation'],
],[1.8,4.8])

sp()
h(2, '8.2  Description')
body('The Quote Generation pod navigates carrier websites or carrier APIs, submits coverage applications '
     'on behalf of the insured, collects quotes, and produces a side-by-side comparison proposal. '
     'Supports 18 major US carriers across 6 lines of business. '
     'Outputs a downloadable PDF proposal and email-ready summary.')

sp()
h(2, '8.3  Required Input')
h(3, 'Step 1 — Coverage Setup')
tbl(['Field','Type','Required','Options'],[
    ['insuranceType',    'string (select)',        'Yes','auto, home, commercial, general-liability, workers-comp, professional-liability'],
    ['selectedCarriers', 'string[] (multi-select)','Yes (>=1)','See carrier list below'],
],[1.5,1.5,1.0,2.6])

sp()
h(3, 'Supported Carriers (18)')
tbl(['ID','Display Name'],[
    ['progressive',    'Progressive'],   ['geico',          'GEICO'],
    ['state-farm',     'State Farm'],    ['allstate',        'Allstate'],
    ['liberty-mutual', 'Liberty Mutual'],['travelers',       'Travelers'],
    ['nationwide',     'Nationwide'],    ['farmers',         'Farmers Insurance'],
    ['usaa',           'USAA'],          ['american-family', 'American Family'],
    ['hartford',       'The Hartford'],  ['chubb',           'Chubb'],
    ['aig',            'AIG'],           ['zurich',          'Zurich'],
    ['hanover',        'The Hanover'],   ['cincinnati',      'Cincinnati Insurance'],
    ['erie',           'Erie Insurance'],['auto-owners',     'Auto-Owners'],
],[2.2,4.4])

sp()
h(3, 'Step 2 — Applicant Information')
tbl(['Field','Type','Required','Notes'],[
    ['name',           'string','Yes','Individual insured full name'],
    ['businessName',   'string','No', 'Required for commercial lines'],
    ['email',          'string','Yes','Proposal delivery address'],
    ['phone',          'string','No', 'Contact number'],
    ['address',        'string','Yes','Risk location address'],
    ['coverageAmount', 'number','Yes','Desired coverage limit in USD'],
],[1.5,1.0,1.0,3.1])

sp()
h(2, '8.4  Expected Output')
h(3, 'CarrierQuote object (TypeScript — one per selected carrier)')
code('interface CarrierQuote {\n'
     '  carrier:           string;                                   // e.g. "Progressive"\n'
     '  logo:              string;                                   // carrier logo URL\n'
     '  premium:           number;                                   // annual premium USD\n'
     '  coverage:          string;                                   // coverage description\n'
     '  deductible:        number;                                   // deductible USD\n'
     '  status:            "pending" | "fetching" | "complete" | "error";\n'
     '  features:          string[];                                 // coverage highlights\n'
     '  rating?:           number;                                   // customer satisfaction 1-5\n'
     '  claimsScore?:      number;                                   // claims satisfaction score\n'
     '  financialStrength?:string;                                   // AM Best / S&P rating\n}')

sp()
h(3, 'extracted_json schema (stored in pod_extraction_runs)')
code('{\n'
     '  "insurance_type": "auto",\n'
     '  "applicant": {\n'
     '    "name": "...", "business_name": "...", "email": "...",\n'
     '    "phone": "...", "address": "...", "coverage_amount": 500000\n'
     '  },\n'
     '  "quotes": [\n'
     '    {\n'
     '      "carrier": "Progressive", "premium": 1240, "coverage": "Full Coverage",\n'
     '      "deductible": 500, "features": ["Accident Forgiveness"],\n'
     '      "status": "complete", "rating": 4.2, "financial_strength": "A+"\n'
     '    }\n'
     '  ],\n'
     '  "recommended_carrier":     "Progressive",\n'
     '  "recommendation_rationale":"Lowest premium with comparable coverage and strong claims score.",\n'
     '  "total_savings_potential": 380\n}')

sp()
h(3, 'UI Output Sections')
tbl(['Section','Description'],[
    ['Quote Comparison Table','Side-by-side premium, deductible, coverage, features for all carriers'],
    ['Premium Analysis','Lowest / highest / average premium, total savings potential'],
    ['Coverage Details','Per-carrier policy terms, limits, and exclusions'],
    ['AI Recommendation','Ranked carrier recommendation with rationale bullets'],
    ['Proposal PDF','Downloadable client-ready PDF proposal (jsPDF)'],
    ['Email Preview','Draft email with proposal attachment ready to send'],
],[2.2,4.4])

# ════════════════════════════════════════════════════════════════════════════
# 9. POLICY COMPARISON
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '9. Pod 3 — Policy Comparison')

h(2, '9.1  Identity')
tbl(['Field','Value'],[
    ['Pod ID','policy-comparison'],
    ['Category','Analysis'],
    ['Segment','Broker'],
    ['Pack','Underwriting Pack, Distribution Pack'],
    ['Frontend UI','frontend/src/components/playground/PolicyComparisonUI.tsx'],
    ['Prompt builder','frontend/src/lib/policyComparisonPrompt.ts'],
    ['Database','Shared pod tables (pod_id = policy-comparison)'],
    ['API prefix','/api/pods/policy-comparison'],
],[1.8,4.8])

sp()
h(2, '9.2  Description')
body('The Policy Comparison pod accepts two insurance policy documents (current/expiring vs. proposed/renewal) '
     'and performs a deep structural analysis. It extracts key coverage fields from both, computes a '
     'clause-level diff (redline), calculates a deviation percentage, and produces a materiality-weighted '
     'recommendation. If premiums exceed a configurable threshold, it triggers a Quote Generation suggestion.')

sp()
h(2, '9.3  Required Input')
tbl(['Field','Type','Required','Notes'],[
    ['policyAFile',               'File (PDF/DOCX)','Yes', 'Current / expiring policy'],
    ['policyBFile',               'File (PDF/DOCX)','Yes', 'Proposed / renewal policy'],
    ['policyAName',               'string',         'Auto','Derived from filename'],
    ['policyBName',               'string',         'Auto','Derived from filename'],
    ['deviationThresholdPercent', 'number',         'No',  'Default 10. Triggers recommendation when exceeded.'],
],[1.8,1.4,0.9,2.5])
body('Document constraint: each document capped at 40,000 characters before being sent to the model.')

sp()
h(2, '9.4  LLM System Prompt Rules')
for r in [
    'Output is strict JSON only — no markdown or prose outside the JSON object.',
    'Compute deviation_percent as 0–100 (% of materially changed coverage/clauses).',
    'Set deviation_exceeds_threshold = deviation_percent > threshold.',
    'If deviation_exceeds_threshold, include recommendation with recommended_policy (A | B | NEITHER) and rationale list.',
    'Always include clause_diff.clauses[] with status = added | removed | changed, plus before/after text.',
    'Always include extracted_fields.policyA and .policyB with: carrier, premiums, limits, deductibles, effective dates, exclusions, endorsements.',
    'Include taxonomy: domain, doc types, lines of business.',
    'If data is missing, keep keys but set values to null and add a warning entry.',
]:
    bullet(r)

sp()
h(2, '9.5  Expected Output — PolicyComparisonStructured (TypeScript)')
code('type PolicyComparisonStructured = {\n'
     '  taxonomy: {\n'
     '    domain:            "insurance";\n'
     '    doc_type_a:        string;       // e.g. "Commercial General Liability Policy"\n'
     '    doc_type_b:        string;       // e.g. "Commercial General Liability Renewal"\n'
     '    lines_of_business: string[];     // e.g. ["GL", "PROPERTY"]\n'
     '  };\n'
     '  extracted_fields: {\n'
     '    policyA: {\n'
     '      carrier?:           string;\n'
     '      premium?:           number;\n'
     '      general_liability?: string;\n'
     '      deductible?:        number;\n'
     '      cyber_coverage?:    boolean;\n'
     '      epl_coverage?:      boolean;\n'
     '      water_damage?:      boolean;\n'
     '      effective_date?:    string;\n'
     '      expiration_date?:   string;\n'
     '      exclusions?:        string[];\n'
     '      endorsements?:      string[];\n'
     '      [key: string]:      any;\n'
     '    };\n'
     '    policyB: { /* same structure */ };\n'
     '  };\n'
     '  clause_diff: {\n'
     '    clauses: Array<{\n'
     '      id:       string;\n'
     '      title?:   string;\n'
     '      status:   "added" | "removed" | "changed";\n'
     '      before?:  string;   // policy A text\n'
     '      after?:   string;   // policy B text\n'
     '      path?:    string;   // document section path\n'
     '    }>;\n'
     '    meta?: Record<string, unknown>;\n'
     '  };\n'
     '  deviation_percent:           number;    // 0–100\n'
     '  deviation_exceeds_threshold: boolean;\n'
     '  recommendation?: {\n'
     '    recommended_policy: "A" | "B" | "NEITHER";\n'
     '    rationale:          string[];\n'
     '  };\n'
     '  warnings: string[];\n};')

sp()
h(3, 'UI Output Views')
tbl(['View','Description'],[
    ['Coverage View (default)',    'Side-by-side field table with premium diff and coverage gap highlights'],
    ['Clause Redline View',        'Clause-level diff: additions (green), removals (red), changes (yellow)'],
    ['AI Recommendation Banner',   'Policy A vs B vs Review recommendation with rationale'],
    ['Smart Quote Recommendation', 'Shown when premium > policyComparisonPremiumThreshold and smart recommendations enabled'],
],[2.2,4.4])

# ════════════════════════════════════════════════════════════════════════════
# 10. ACORD FORM EXTRACTION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '10. Pod 4 — ACORD Form Extraction')

h(2, '10.1  Identity')
tbl(['Field','Value'],[
    ['Pod ID','acord_form_understanding'],
    ['Category','Document Processing'],
    ['Segment','Broker'],
    ['Pack','Underwriting Pack'],
    ['Frontend UI','frontend/src/components/playground/ACORDParserUI.tsx'],
    ['Pydantic schemas','backend/Models/acord_form_understanding/schemas.py'],
    ['API schemas','backend/app/schemas/acord_workflow.py'],
    ['Database','Dedicated ACORD tables (Section 3) + shared pod tables'],
    ['API prefix','/api/acord  (dedicated router)'],
],[1.8,4.8])

sp()
h(2, '10.2  Description')
body('The ACORD Form Extraction pod parses and extracts structured data from ACORD standard insurance forms. '
     'Supports 8 form types via a multi-engine pipeline: fillable AcroForm widget extraction, '
     'multi-engine PDF-to-text, OCR (Tesseract / PaddleOCR), and LLM-based extraction '
     '(RunPod fine-tuned model with OpenAI fallback).')

sp()
h(2, '10.3  Supported Form Types')
tbl(['Form','Name','Use Case'],[
    ['ACORD 25','Certificate of Insurance','Proof of coverage for third parties / landlords'],
    ['ACORD 27','Evidence of Property Insurance','Mortgage / lender property proof'],
    ['ACORD 80','Garage Coverage Summary','Auto dealer / garage operations'],
    ['ACORD 85','General Liability Application','GL coverage application submission'],
    ['ACORD 90','Automobile Application','Commercial auto application'],
    ['ACORD 125','Commercial Insurance Application','Full commercial lines application (master form)'],
    ['ACORD 126','Commercial General Liability','CGL supplement / schedule'],
    ['ACORD 140','Property Loss Notice','Property claim first notice'],
],[1.0,2.2,3.4])

sp()
h(2, '10.4  Required Input')
tbl(['Field','Type','Required','Notes'],[
    ['formType',       'string (select)','Yes','ACORD 25 / 27 / 80 / 85 / 90 / 125 / 126 / 140'],
    ['file',           'File (PDF/DOCX)', 'Yes','Supports fillable and scanned/flattened PDFs'],
    ['form_type_hint', 'string',          'No', 'Override for re-extraction (sent via ReExtractRequest.form_type_hint). e.g. "25" or "125".'],
],[1.5,1.4,0.9,2.8])

sp()
h(2, '10.5  Extraction Pipeline Steps')
for i, s in enumerate([
    'PDF classification — detect fillable (AcroForm) vs flattened (scanned/print)',
    'Text extraction — try engines: BytesScout → pdfplumber → PyMuPDF → PyPDF2',
    'AcroForm extraction — extract fillable widget values if AcroForm layer present',
    'OCR — Tesseract or PaddleOCR for raster/scanned documents',
    'LLM structured extraction — RunPod fine-tuned model (fallback: OpenAI)',
    'Form-type-specific fallbacks — acord25_fallback.py, acord125_fallback.py',
    'Azure Document Intelligence — optional high-accuracy OCR overlay',
    'Confidence scoring — base score from engine + overlay adjustments',
    'Post-processing — normalization, validation, null handling',
], 1):
    bullet(f'{i}. {s}')

sp()
h(2, '10.6  Output Schema — AcordFormSummary (Pydantic)')
body('Defined in backend/Models/acord_form_understanding/schemas.py. '
     'This is what is stored in acord_extraction_runs.extracted_json.')

h(3, 'Top-level fields')
tbl(['Field','Python Type','Description'],[
    ['form_type',                   'Optional[str]',                   'e.g. "ACORD 25", "ACORD 125", "ACORD 140"'],
    ['form_version',                'Optional[str]',                   'Form version string'],
    ['certificate_number',          'Optional[str]',                   ''],
    ['revision_date',               'Optional[str]',                   ''],
    ['date',                        'Optional[str]',                   'Form date MM/DD/YYYY'],
    ['producer',                    'Optional[AcordProducer]',         'See AcordProducer schema below'],
    ['insured',                     'Optional[AcordInsured]',          'Primary named insured'],
    ['other_named_insureds',        'List[AcordOtherNamedInsured]',    'Additional named insureds (ACORD 125 p.1)'],
    ['holder',                      'Optional[AcordHolder]',           'Certificate holder / additional insured (ACORD 25)'],
    ['policy_info',                 'Optional[AcordPolicyInfo]',       'Policy-level info (ACORD 125)'],
    ['lines_of_business_indicated', 'List[str]',                       'LOBs checked on the form. e.g. ["BUSINESS OWNERS", "GL"]'],
    ['coverages',                   'List[AcordPolicyCoverage]',       'One entry per line of business'],
    ['premises',                    'List[AcordPremises]',             'Business locations/premises'],
    ['prior_carriers',              'List[AcordPriorCarrier]',         'Prior insurance history'],
    ['loss_history',                'List[AcordLossHistory]',          'Claims/loss history (typically 5 years)'],
    ['additional_interests',        'List[AcordAdditionalInterest]',   'Loss payees, mortgagees, etc.'],
    ['description_of_operations',   'Optional[str]',                   'Description of operations / locations / vehicles'],
    ['nature_of_business',          'Optional[str]',                   ''],
    ['cancellation_notice_days',    'Optional[int]',                   ''],
    ['additional_remarks',          'Optional[str]',                   ''],
    ['extra_fields',                'Optional[Dict[str, Any]]',        'Overflow: any LLM-extracted fields not in standard schema'],
    ['overall_confidence',          'float',                           'Final confidence score 0.0–1.0'],
    ['raw_text',                    'str',                             'Original extracted text from the document'],
],[2.2,2.0,2.5])

sp()
h(3, 'AcordProducer')
tbl(['Field','Python Type','Description'],[
    ['name',                    'Optional[str]',   'Agency / producer name'],
    ['name_confidence',         'Optional[float]', 'Confidence for name field. ge=0, le=1.'],
    ['contact_name',            'Optional[str]',   'Agent contact full name'],
    ['address',                 'Optional[str]',   ''],
    ['city',                    'Optional[str]',   ''],
    ['state',                   'Optional[str]',   ''],
    ['postal_code',             'Optional[str]',   ''],
    ['phone',                   'Optional[str]',   ''],
    ['fax',                     'Optional[str]',   ''],
    ['email',                   'Optional[str]',   ''],
    ['agency_customer_id',      'Optional[str]',   ''],
    ['subcode',                 'Optional[str]',   'Sub-producer code'],
    ['producer_license_no',     'Optional[str]',   'State producer license number'],
    ['national_producer_number','Optional[str]',   'NPN from NIPR'],
],[2.4,1.6,2.5])

sp()
h(3, 'AcordInsured (Primary Named Insured)')
tbl(['Field','Python Type','Description'],[
    ['name',            'Optional[str]',   'Named insured'],
    ['name_confidence', 'Optional[float]', 'Confidence score ge=0, le=1'],
    ['contact_name',    'Optional[str]',   ''],
    ['mailing_address', 'Optional[str]',   ''],
    ['city',            'Optional[str]',   ''],
    ['state',           'Optional[str]',   ''],
    ['postal_code',     'Optional[str]',   ''],
    ['phone',           'Optional[str]',   ''],
    ['fax',             'Optional[str]',   ''],
    ['email',           'Optional[str]',   ''],
    ['website',         'Optional[str]',   ''],
    ['entity_type',     'Optional[str]',   'Corporation / LLC / Partnership / Individual / Trust / etc.'],
    ['gl_code',         'Optional[str]',   ''],
    ['sic',             'Optional[str]',   'Standard Industrial Classification'],
    ['naics',           'Optional[str]',   'North American Industry Classification'],
    ['fein',            'Optional[str]',   'Federal Employer Identification Number'],
],[2.2,1.6,2.7])

sp()
h(3, 'AcordHolder (Certificate Holder / Additional Insured)')
tbl(['Field','Python Type','Description'],[
    ['name',                  'Optional[str]',  ''],
    ['address',               'Optional[str]',  ''],
    ['city',                  'Optional[str]',  ''],
    ['state',                 'Optional[str]',  ''],
    ['postal_code',           'Optional[str]',  ''],
    ['is_additional_insured', 'Optional[bool]', 'Whether holder is also an additional insured'],
    ['is_subrogation_waived', 'Optional[bool]', 'Whether subrogation is waived for this holder'],
],[2.4,1.6,2.5])

sp()
h(3, 'AcordPolicyInfo')
tbl(['Field','Python Type','Description'],[
    ['carrier',            'Optional[AcordCarrier]','Carrier name + NAIC number'],
    ['program_name',       'Optional[str]',         ''],
    ['program_code',       'Optional[str]',         ''],
    ['policy_number',      'Optional[str]',         ''],
    ['proposed_eff_date',  'Optional[str]',         'Policy effective date'],
    ['proposed_exp_date',  'Optional[str]',         'Policy expiration date'],
    ['billing_plan',       'Optional[str]',         'Direct / Agency'],
    ['payment_plan',       'Optional[str]',         'Annual / Monthly / etc.'],
    ['method_of_payment',  'Optional[str]',         'Cash / EFT / etc.'],
    ['deposit',            'Optional[str]',         ''],
    ['minimum_premium',    'Optional[str]',         ''],
    ['policy_premium',     'Optional[str]',         ''],
    ['transaction_type',   'Optional[str]',         'Quote / Issue Policy / Renew / Change / Cancel'],
    ['transaction_date',   'Optional[str]',         ''],
    ['underwriter',        'Optional[str]',         ''],
    ['underwriter_office', 'Optional[str]',         ''],
],[2.2,1.6,2.7])

sp()
h(3, 'AcordPolicyCoverage  (per line of business — includes per-field confidence scores)')
tbl(['Field','Python Type','Description'],[
    ['line_of_business',           'Optional[str]',          'GL | AUTO | WC | UMB | PROPERTY | CRIME | etc.'],
    ['block_confidence',           'Optional[float]',         'Confidence for this entire coverage block. ge=0, le=1.'],
    ['policy_number',              'Optional[str]',           ''],
    ['policy_number_confidence',   'Optional[float]',         'ge=0, le=1'],
    ['effective_date',             'Optional[date]',          ''],
    ['effective_date_confidence',  'Optional[float]',         'ge=0, le=1'],
    ['expiration_date',            'Optional[date]',          ''],
    ['expiration_date_confidence', 'Optional[float]',         'ge=0, le=1'],
    ['claims_made',                'Optional[bool]',          'Claims-made trigger basis'],
    ['occurrence_type',            'Optional[bool]',          'Occurrence trigger basis'],
    ['additional_insured',         'Optional[bool]',          ''],
    ['waiver_of_subrogation',      'Optional[bool]',          ''],
    ['--- GL Limits ---',          '---',                     '---'],
    ['each_occurrence',            'Optional[str]',           ''],
    ['damage_to_rented_premises',  'Optional[str]',           ''],
    ['medical_expense',            'Optional[str]',           ''],
    ['personal_advertising_injury','Optional[str]',           ''],
    ['general_aggregate',          'Optional[str]',           ''],
    ['products_comp_ops_aggregate','Optional[str]',           ''],
    ['--- Auto Limits ---',        '---',                     '---'],
    ['combined_single_limit',      'Optional[str]',           ''],
    ['bodily_injury_per_person',   'Optional[str]',           ''],
    ['bodily_injury_per_accident', 'Optional[str]',           ''],
    ['property_damage',            'Optional[str]',           ''],
    ['--- Umbrella/Excess ---',    '---',                     '---'],
    ['occurrence_limit',           'Optional[str]',           ''],
    ['aggregate_limit',            'Optional[str]',           ''],
    ['deductible',                 'Optional[str]',           ''],
    ['retention',                  'Optional[str]',           ''],
    ['retroactive_date',           'Optional[str]',           ''],
    ['--- Workers Comp ---',       '---',                     '---'],
    ['wc_statutory_limits',        'Optional[bool]',          ''],
    ['employer_liability_each_accident','Optional[str]',      ''],
    ['employer_liability_each_employee','Optional[str]',      ''],
    ['employer_liability_policy_limit', 'Optional[str]',      ''],
    ['insurers',                   'List[AcordCarrier]',      'Carriers for this coverage block. Default: [].'],
],[2.6,1.6,2.4])

sp()
h(3, 'AcordPremises')
tbl(['Field','Python Type','Description'],[
    ['location_number',          'Optional[str]',''],
    ['street',                   'Optional[str]',''],
    ['city',                     'Optional[str]',''],
    ['state',                    'Optional[str]',''],
    ['county',                   'Optional[str]',''],
    ['zip',                      'Optional[str]',''],
    ['interest',                 'Optional[str]','Owner Occupied / Tenant / etc.'],
    ['full_time_employees',      'Optional[str]',''],
    ['part_time_employees',      'Optional[str]',''],
    ['annual_revenues',          'Optional[str]',''],
    ['total_building_area_sqft', 'Optional[str]',''],
    ['description_of_operations','Optional[str]',''],
    ['area_leased_to_others',    'Optional[str]',''],
],[2.4,1.6,2.6])

sp()
h(3, 'AcordPriorCarrier')
tbl(['Field','Python Type','Description'],[
    ['year',            'Optional[str]',''],
    ['category',        'Optional[str]','General Liability / Automobile / Property / Other'],
    ['carrier',         'Optional[str]',''],
    ['policy_number',   'Optional[str]',''],
    ['premium',         'Optional[str]',''],
    ['effective_date',  'Optional[str]',''],
    ['expiration_date', 'Optional[str]',''],
],[2.4,1.6,2.6])

sp()
h(3, 'AcordLossHistory')
tbl(['Field','Python Type','Description'],[
    ['date_of_occurrence','Optional[str]', ''],
    ['line_type',         'Optional[str]', 'Line of business for this claim'],
    ['description',       'Optional[str]', ''],
    ['date_of_claim',     'Optional[str]', ''],
    ['amount_paid',       'Optional[str]', ''],
    ['amount_reserved',   'Optional[str]', ''],
    ['subrogation',       'Optional[bool]','Whether subrogation was pursued'],
    ['claim_open',        'Optional[bool]','Whether the claim is still open'],
],[2.4,1.6,2.6])

sp()
h(3, 'AcordAdditionalInterest')
tbl(['Field','Python Type','Description'],[
    ['interest_type',  'Optional[str]','Additional Insured / Loss Payee / Mortgagee / etc.'],
    ['name',           'Optional[str]',''],
    ['address',        'Optional[str]',''],
    ['location',       'Optional[str]',''],
    ['building',       'Optional[str]',''],
    ['loan_reference', 'Optional[str]',''],
],[2.4,1.6,2.6])

sp()
h(3, 'ExtractionMeta')
tbl(['Field','Python Type','Description'],[
    ['form_type_detected',        'Optional[str]',  'Form type the engine classified'],
    ['blank_in_document',         'List[str]',       'Fields present but left blank in document'],
    ['not_applicable_to_form_type','List[str]',      'Fields not part of this form type'],
    ['all_checked_items',         'List[str]',       'Every checkbox/item marked in document'],
    ['remarks',                   'List[str]',       'LLM notes about ambiguous/low-confidence values'],
    ['extraction_engine',         'Optional[str]',   'bytescout | pdfplumber | pymupdf | pypdf2 | ocr | txt | legacy'],
    ['base_confidence',           'Optional[float]', '0.0–1.0 from engine quality. ge=0, le=1.'],
    ['structured_response_source','Optional[str]',   'LLM RunPod | LLM OpenAI | Fallback'],
    ['pdf_form_classification',   'Optional[str]',   'fillable (AcroForm) | flattened (scanned)'],
    ['ocr_text_engine',           'Optional[str]',   'tesseract | paddle'],
],[2.6,1.4,2.6])

sp()
h(2, '10.7  ACORD API Endpoints  (/api/acord)')
h(3, 'User Endpoints')
tbl(['Method','Path','Description'],[
    ['POST','/api/acord/extract',                       'Synchronous extraction. Returns AcordExtractResponse.'],
    ['POST','/api/acord/extract/start',                 'Async job start. Returns AcordExtractStartResponse (job_id + status=queued).'],
    ['GET', '/api/acord/extract/status/{job_id}',       'Poll async job. Returns AcordExtractJobStatusResponse (status + phase + result).'],
    ['GET', '/api/acord/runs',                          'List past extractions (paginated).'],
    ['GET', '/api/acord/runs/{run_id}',                 'Get extraction details + metadata.'],
    ['POST','/api/acord/runs/{run_id}/submit',          'Submit run for review. Body: AcordSubmitRequest.'],
    ['POST','/api/acord/runs/{run_id}/re-extract',      'Re-extract. Body: ReExtractRequest (form_type_hint).'],
],[0.8,2.9,2.9])

sp()
h(3, 'Admin Endpoints')
tbl(['Method','Path','Description'],[
    ['GET',   '/api/acord/admin/queue',                    'Admin review queue'],
    ['POST',  '/api/acord/admin/{run_id}/review',          'Single review. Body: AcordAdminReviewRequest.'],
    ['POST',  '/api/acord/admin/batch-review',             'Batch review. Body: AcordBatchReviewRequest.'],
    ['GET',   '/api/acord/admin/queue/stats',              'Queue statistics'],
    ['GET',   '/api/acord/admin/queue/{run_id}/detail',    'Full queue item detail'],
    ['PATCH', '/api/acord/admin/queue/{run_id}/detail',    'Update queue item'],
    ['GET',   '/api/acord/admin/jobs',                     'List ACORD training jobs'],
    ['GET',   '/api/acord/admin/jobs/{job_id}',            'Training job details'],
    ['GET',   '/api/acord/admin/jobs/{job_id}/eval',       'Evaluation results (acord_eval_results)'],
    ['GET',   '/api/acord/admin/jobs/{job_id}/log',        'Training logs'],
    ['GET',   '/api/acord/admin/runs/{run_id}/health-card','Confidence + feedback health metrics'],
],[0.8,2.9,2.9])

sp()
h(2, '10.8  Frontend Output Tabs (ACORDParserUI)')
tbl(['Tab','Description'],[
    ['JSON View',    'Full normalized AcordFormSummary as formatted JSON'],
    ['Fields View',  'Flattened key-value table for quick review'],
    ['Edit Mode',    'Inline JSON editor for user corrections'],
    ['Changes View', 'Diff view before/after user edits'],
    ['Split View',   'Side-by-side original document vs extracted fields'],
],[2.0,4.6])

sp()
h(2, '10.9  Fine-Tuning Quality Gates')
tbl(['Environment Variable','Default','Meaning'],[
    ['FT_ACORD_QG_MIN_JSON_VALID_RATE',        '0.90','>=90% of outputs must be valid JSON'],
    ['FT_ACORD_QG_MIN_JSON_EXACT_MATCH',       '0.70','>=70% exact field match rate'],
    ['FT_ACORD_QG_MIN_JSON_FIELD_RECALL',      '0.80','>=80% field recall'],
    ['FT_ACORD_QG_MAX_JSON_EXTRA_FIELD_RATE',  '0.10','<=10% extra/hallucinated fields'],
    ['FT_ACORD_QG_MAX_OOS_HALLUCINATION_RATE', '0.25','<=25% hallucination on out-of-sample set'],
],[3.0,0.8,2.8])

# ════════════════════════════════════════════════════════════════════════════
# 11. CONFIDENCE & TRAINING WORKFLOW
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '11. Confidence & Training Workflow')

h(2, '11.1  Confidence Score Adjustments')
body('Defined in backend/app/services/pod_extraction.py.')
tbl(['Signal','Adjustment'],[
    ['Latest feedback is thumbs_up',        '+0.05'],
    ['Latest feedback is thumbs_down',      '-0.08'],
    ['JSON was manually corrected by user', '-0.10'],
    ['Admin approved the run',              '+0.04'],
    ['Admin rejected or requested rework',  '-0.06'],
],[4.0,2.6])
body('Final score clamped to [0.0, 1.0].')

sp()
h(2, '11.2  Auto-Approve Logic')
code('if thumbs_up == true AND adjusted_confidence >= POD_CONFIDENCE_THRESHOLD:\n'
     '    status = "approved"\n'
     '    trigger_training_job()  # if AUTO_FINE_TUNE_ON_POD_APPROVAL = true\n'
     'else:\n'
     '    status = "needs_admin_review"\n'
     '    create pod_admin_queue entry')

sp()
h(2, '11.3  Evaluation Sets')
tbl(['Set','Description'],[
    ['seen',        'Training examples — sanity check'],
    ['paraphrased', 'Rephrased versions of training examples (generalisation check)'],
    ['oos',         'Out-of-sample documents not seen during training'],
    ['combined',    'Weighted combination of all three (primary gate metric)'],
],[1.5,5.1])

# ════════════════════════════════════════════════════════════════════════════
# 12. ENVIRONMENT VARIABLES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '12. Environment Variables')
tbl(['Variable','Default','Applies To','Description'],[
    ['POD_CONFIDENCE_THRESHOLD',              '0.85',  'All pods',       'Minimum confidence for auto-approval'],
    ['ACORD_CONFIDENCE_THRESHOLD',            '0.85',  'ACORD only',     'ACORD-specific threshold override'],
    ['AUTO_FINE_TUNE_ON_POD_APPROVAL',        'true',  'All pods',       'Trigger training job on admin approval'],
    ['ACORD_POD_ID',                          'acord_form_understanding','ACORD','Pod identifier string'],
    ['RUNPOD_POD_ID',                         '-',     'ACORD',          'RunPod ML service instance ID'],
    ['FT_ACORD_QG_MIN_JSON_VALID_RATE',       '0.90',  'ACORD training', 'Quality gate: valid JSON rate'],
    ['FT_ACORD_QG_MIN_JSON_EXACT_MATCH',      '0.70',  'ACORD training', 'Quality gate: exact match rate'],
    ['FT_ACORD_QG_MIN_JSON_FIELD_RECALL',     '0.80',  'ACORD training', 'Quality gate: field recall'],
    ['FT_ACORD_QG_MAX_JSON_EXTRA_FIELD_RATE', '0.10',  'ACORD training', 'Quality gate: max extra field rate'],
    ['FT_ACORD_QG_MAX_OOS_HALLUCINATION_RATE','0.25',  'ACORD training', 'Quality gate: max hallucination rate'],
],[2.5,0.7,1.3,2.1])

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'c:\Users\samar\Downloads\neura-box-cloud-main\docs\pods_reference.docx'
doc.save(out)
print('Saved:', out)
