"""
Update script for pods_reference_main.docx
Changes applied:
1. ACORD Pod ID Consistency — remove "acord-parser" from P374 note
2. Clarify Extraction vs Inference LLM in sections 10.7 (P424) and 10.8 (P434)
3. Convert all input_schema / output_schema code blocks from examples to proper JSON Schema definitions
4. Policy Comparison extraction input — mark "role" as REQUIRED and enforced
5. Quote Generation — snake_case normalization note is already in P301; reinforced in inference input schema title/description
"""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document('docs/pods_reference_main.docx')


# ─── helpers ────────────────────────────────────────────────────────────────

def find_para_by_text(substring):
    """Return (index, paragraph) for the first paragraph whose text contains substring."""
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i, p
    return None, None


def replace_body_para(para, new_text):
    """Replace all runs of a body paragraph with new_text in the first run."""
    if not para.runs:
        run = para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ''


def replace_code_block(para, new_text):
    """Replace all runs of a code block paragraph with new_text in the first run."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ''


# ─── JSON Schema definitions ────────────────────────────────────────────────

SCHEMAS = {}

# 7.5 — FNOL Extraction Input
SCHEMAS['fnol_ext_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "FNOLExtractionInput",
  "description": "Files submitted to the FNOL Extraction LLM — 1 to 10 documents per request",
  "type": "object",
  "required": ["files"],
  "properties": {
    "files": {
      "type": "array",
      "minItems": 1,
      "maxItems": 10,
      "description": "Ordered array of documents to extract from",
      "items": {
        "type": "object",
        "required": ["index", "filename", "mime_type", "content"],
        "properties": {
          "index":     { "type": "integer", "minimum": 0,  "description": "0-based file position" },
          "filename":  { "type": "string",                 "description": "Original filename, e.g. police_report.pdf" },
          "mime_type": { "type": "string",                 "description": "e.g. application/pdf, image/jpeg" },
          "content":   { "type": "string",                 "description": "Raw file bytes (binary or base64)" }
        }
      }
    }
  }
}"""

# 7.5 — FNOL Extraction Output (per file)
SCHEMAS['fnol_ext_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "FNOLExtractionOutputPerFile",
  "description": "Extraction result for one FNOL document. Aggregated via merge (union) before Inference.",
  "type": "object",
  "required": ["file_index"],
  "properties": {
    "file_index":           { "type": "integer",           "description": "Matches the input files[].index" },
    "document_type":        { "type": "string",            "description": "e.g. police_report, repair_estimate, photos, witness_statement" },
    "incident_date":        { "type": ["string", "null"],  "description": "ISO 8601 incident date" },
    "incident_location":    { "type": ["string", "null"],  "description": "Physical location of the loss event" },
    "parties_mentioned":    { "type": "array",  "items": { "type": "string" }, "description": "All party names found in document" },
    "damage_estimate":      { "type": ["number", "null"],  "description": "Estimated damage in USD" },
    "police_report_number": { "type": ["string", "null"],  "description": "Police report reference" },
    "raw_text_excerpt":     { "type": "string",            "description": "Salient text extract used for field population" }
  }
}"""

# 7.6 — FNOL Inference Input
SCHEMAS['fnol_inf_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "FNOLInferenceInput",
  "description": "Input to the FNOL Inference LLM — incident description plus aggregated extraction results. extraction is null when no files were uploaded (files=[]).",
  "type": "object",
  "required": ["description"],
  "properties": {
    "description": {
      "type": "string",
      "description": "Free-text incident description from the user (min ~50 characters)"
    },
    "extraction": {
      "type": ["object", "null"],
      "description": "Aggregated extraction payload from Phase 1. Null when files=[].",
      "required": ["files_processed", "merged_fields"],
      "properties": {
        "files_processed": { "type": "integer", "description": "Number of files processed by extraction phase" },
        "merged_fields": {
          "type": "object",
          "description": "Union of all per-file extracted fields (arrays concatenated, scalars deduplicated)",
          "properties": {
            "incident_date":        { "type": ["string", "null"] },
            "incident_location":    { "type": ["string", "null"] },
            "parties_mentioned":    { "type": "array", "items": { "type": "string" } },
            "damage_estimate":      { "type": ["number", "null"] },
            "police_report_number": { "type": ["string", "null"] }
          }
        }
      }
    }
  }
}"""

# 7.6 — FNOL Inference Output
SCHEMAS['fnol_inf_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "FNOLInferenceOutput",
  "description": "FNOL Inference output — stored as extracted_json in pod_extraction_runs",
  "type": "object",
  "required": ["loss_type", "line_of_business", "coverage_applicable", "confidence_notes"],
  "properties": {
    "incident_date":          { "type": ["string", "null"], "description": "ISO 8601 date of incident" },
    "incident_location":      { "type": ["string", "null"] },
    "loss_type":              { "type": "string",  "description": "e.g. collision, fire, theft, slip_and_fall, water_damage" },
    "line_of_business":       { "type": "string",  "description": "e.g. auto, property, gl, wc" },
    "estimated_damage":       { "type": ["number", "null"], "description": "Estimated loss in USD" },
    "injuries_reported":      { "type": "boolean" },
    "police_report_number":   { "type": ["string", "null"] },
    "parties": {
      "type": "array",
      "items": {
        "type": "object",
        "required": ["role"],
        "properties": {
          "role": { "type": "string", "enum": ["insured", "third_party", "witness"] },
          "name": { "type": "string" }
        }
      }
    },
    "coverage_applicable":    { "type": "boolean" },
    "subrogation_potential":  { "type": "boolean" },
    "documentation_required": { "type": "array", "items": { "type": "string" } },
    "recommended_actions":    { "type": "array", "items": { "type": "string" } },
    "confidence_notes":       { "type": "string" }
  }
}"""

# 8.6 — Quote Generation Inference Input
SCHEMAS['quote_inf_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "QuoteGenerationInferenceInput",
  "description": "Input to the Quote Generation Inference LLM — snake_case fields. Frontend sends camelCase (insuranceType, selectedCarriers, businessName); backend normalises to snake_case before populating this schema.",
  "type": "object",
  "required": ["insurance_type", "selected_carriers", "applicant"],
  "properties": {
    "insurance_type": {
      "type": "string",
      "enum": ["auto", "home", "commercial", "general-liability", "workers-comp", "professional-liability"]
    },
    "selected_carriers": {
      "type": "array",
      "minItems": 1,
      "items": { "type": "string", "description": "Carrier ID, e.g. progressive, geico, state-farm" }
    },
    "applicant": {
      "type": "object",
      "required": ["name", "email", "address", "coverage_amount"],
      "properties": {
        "name":            { "type": "string" },
        "email":           { "type": "string",           "description": "Proposal delivery address" },
        "address":         { "type": "string",           "description": "Risk location address" },
        "coverage_amount": { "type": "number",           "description": "Desired coverage limit in USD" },
        "business_name":   { "type": ["string", "null"], "description": "Required for commercial lines" },
        "phone":           { "type": ["string", "null"] }
      }
    }
  }
}"""

# 8.6 — Quote Generation Inference Output
SCHEMAS['quote_inf_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "QuoteGenerationInferenceOutput",
  "description": "Quote Generation Inference output — stored as extracted_json in pod_extraction_runs",
  "type": "object",
  "required": ["insurance_type", "applicant", "quotes"],
  "properties": {
    "insurance_type": { "type": "string" },
    "applicant":      { "type": "object", "description": "Snake_case applicant fields (mirrors inference input)" },
    "quotes": {
      "type": "array",
      "description": "One entry per selected carrier",
      "items": {
        "type": "object",
        "required": ["carrier", "premium", "coverage", "deductible", "status"],
        "properties": {
          "carrier":            { "type": "string" },
          "premium":            { "type": "number",           "description": "Annual premium in USD" },
          "coverage":           { "type": "string" },
          "deductible":         { "type": "number",           "description": "USD" },
          "status":             { "type": "string",           "enum": ["pending", "fetching", "complete", "error"] },
          "features":           { "type": "array",            "items": { "type": "string" } },
          "rating":             { "type": ["number", "null"], "description": "Customer satisfaction 1-5" },
          "claims_score":       { "type": ["number", "null"] },
          "financial_strength": { "type": ["string", "null"], "description": "AM Best / S&P rating" }
        }
      }
    },
    "recommended_carrier":      { "type": ["string", "null"] },
    "recommendation_rationale": { "type": ["string", "null"] },
    "total_savings_potential":  { "type": ["number", "null"], "description": "USD savings vs highest quote" }
  }
}"""

# 9.6 — Policy Comparison Extraction Input
SCHEMAS['polcomp_ext_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "PolicyComparisonExtractionInput",
  "description": "Two policy documents for independent extraction. The 'role' field is REQUIRED and enforced by the backend — files[0].role must be 'policyA' (current/expiring) and files[1].role must be 'policyB' (proposed/renewal).",
  "type": "object",
  "required": ["files"],
  "properties": {
    "files": {
      "type": "array",
      "minItems": 2,
      "maxItems": 2,
      "description": "Exactly 2 files — one per policy",
      "items": {
        "type": "object",
        "required": ["index", "role", "filename", "mime_type", "content"],
        "properties": {
          "index":     { "type": "integer", "minimum": 0, "maximum": 1 },
          "role":      { "type": "string",  "enum": ["policyA", "policyB"],
                         "description": "REQUIRED and enforced. policyA = current/expiring; policyB = proposed/renewal" },
          "filename":  { "type": "string" },
          "mime_type": { "type": "string" },
          "content":   { "type": "string", "description": "Document text capped at 40,000 characters" }
        }
      }
    }
  }
}"""

# 9.6 — Policy Comparison Extraction Output
SCHEMAS['polcomp_ext_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "PolicyComparisonExtractionOutput",
  "description": "Paired extraction results — policyA and policyB extracted independently and kept as a matched pair (aggregation_strategy: pair) for Inference.",
  "type": "object",
  "required": ["policyA", "policyB"],
  "properties": {
    "policyA": {
      "type": "object",
      "description": "Extracted coverage fields from the current/expiring policy",
      "properties": {
        "carrier":           { "type": ["string", "null"] },
        "premium":           { "type": ["number", "null"],  "description": "Annual premium USD" },
        "effective_date":    { "type": ["string", "null"],  "description": "ISO 8601" },
        "expiration_date":   { "type": ["string", "null"] },
        "general_liability": { "type": ["string", "null"] },
        "deductible":        { "type": ["number", "null"] },
        "exclusions":        { "type": "array", "items": { "type": "string" } },
        "endorsements":      { "type": "array", "items": { "type": "string" } }
      }
    },
    "policyB": {
      "type": "object",
      "description": "Extracted coverage fields from the proposed/renewal policy (same structure as policyA)"
    }
  }
}"""

# 9.7 — Policy Comparison Inference Input
SCHEMAS['polcomp_inf_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "PolicyComparisonInferenceInput",
  "description": "Paired extraction results plus deviation threshold — passed to the Policy Comparison Inference LLM",
  "type": "object",
  "required": ["extraction"],
  "properties": {
    "deviation_threshold_percent": {
      "type": "number",
      "minimum": 0,
      "maximum": 100,
      "default": 10,
      "description": "% deviation above which a recommendation is generated"
    },
    "extraction": {
      "type": "object",
      "required": ["policyA", "policyB"],
      "description": "Paired extraction output from Phase 1 (see Section 9.6)",
      "properties": {
        "policyA": { "type": "object", "description": "Extracted fields from current/expiring policy" },
        "policyB": { "type": "object", "description": "Extracted fields from proposed/renewal policy" }
      }
    }
  }
}"""

# 9.7 — Policy Comparison Inference Output
SCHEMAS['polcomp_inf_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "PolicyComparisonInferenceOutput",
  "description": "Policy Comparison Inference output — stored as extracted_json in pod_extraction_runs",
  "type": "object",
  "required": ["taxonomy", "extracted_fields", "clause_diff", "deviation_percent", "deviation_exceeds_threshold", "warnings"],
  "properties": {
    "taxonomy": {
      "type": "object",
      "required": ["domain", "doc_type_a", "doc_type_b", "lines_of_business"],
      "properties": {
        "domain":            { "type": "string", "const": "insurance" },
        "doc_type_a":        { "type": "string", "description": "e.g. Commercial General Liability Policy" },
        "doc_type_b":        { "type": "string" },
        "lines_of_business": { "type": "array",  "items": { "type": "string" } }
      }
    },
    "extracted_fields": {
      "type": "object",
      "required": ["policyA", "policyB"],
      "properties": {
        "policyA": {
          "type": "object",
          "properties": {
            "carrier":           { "type": ["string", "null"] },
            "premium":           { "type": ["number", "null"] },
            "general_liability": { "type": ["string", "null"] },
            "deductible":        { "type": ["number", "null"] },
            "effective_date":    { "type": ["string", "null"] },
            "expiration_date":   { "type": ["string", "null"] },
            "exclusions":        { "type": "array", "items": { "type": "string" } },
            "endorsements":      { "type": "array", "items": { "type": "string" } }
          }
        },
        "policyB": { "type": "object", "description": "Same structure as policyA" }
      }
    },
    "clause_diff": {
      "type": "object",
      "required": ["clauses"],
      "properties": {
        "clauses": {
          "type": "array",
          "items": {
            "type": "object",
            "required": ["id", "status"],
            "properties": {
              "id":                { "type": "string" },
              "title":             { "type": ["string", "null"] },
              "status":            { "type": "string", "enum": ["added", "removed", "changed"] },
              "before":            { "type": ["string", "null"], "description": "Policy A text" },
              "after":             { "type": ["string", "null"], "description": "Policy B text" },
              "path":              { "type": ["string", "null"] },
              "materiality_score": { "type": "number",  "minimum": 0, "maximum": 1,
                                     "description": "1.0 = critical coverage change; 0.0 = administrative only" }
            }
          }
        }
      }
    },
    "deviation_percent":           { "type": "number",  "minimum": 0, "maximum": 100 },
    "deviation_exceeds_threshold": { "type": "boolean" },
    "recommendation": {
      "type": ["object", "null"],
      "description": "Present when deviation_exceeds_threshold = true",
      "properties": {
        "recommended_policy": { "type": "string", "enum": ["A", "B", "NEITHER"] },
        "rationale":          { "type": "array",  "items": { "type": "string" } }
      }
    },
    "warnings": { "type": "array", "items": { "type": "string" } }
  }
}"""

# 10.7 — ACORD Extraction Input
SCHEMAS['acord_ext_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "ACORDExtractionInput",
  "description": "Files submitted to the ACORD multi-engine Extraction pipeline (PDF classify -> text extract -> AcroForm -> OCR -> general LLM structuring). NOT the fine-tuned RunPod model — that is the Inference LLM in Section 10.8.",
  "type": "object",
  "required": ["files"],
  "properties": {
    "files": {
      "type": "array",
      "minItems": 1,
      "maxItems": 10,
      "items": {
        "type": "object",
        "required": ["index", "filename", "mime_type", "content"],
        "properties": {
          "index":          { "type": "integer", "minimum": 0 },
          "filename":       { "type": "string",  "description": "e.g. acord_125_widgets_inc.pdf" },
          "mime_type":      { "type": "string",  "description": "e.g. application/pdf" },
          "content":        { "type": "string",  "description": "Raw file bytes" },
          "form_type_hint": { "type": "string",  "description": "Optional form type override, e.g. '25' or '125'. Omit for auto-detection." }
        }
      }
    }
  }
}"""

# 10.7 — ACORD Extraction Output
SCHEMAS['acord_ext_out'] = """\
[
  {
    "$schema": "http://json-schema.org/draft-07/schema#",
    "title": "ACORDExtractionOutputPerFile",
    "description": "AcordFormSummary for one file. Aggregation strategy: list. Passed as raw_text + acroform_fields to the Inference LLM (Section 10.8).",
    "type": "object",
    "required": ["form_type", "overall_confidence"],
    "properties": {
      "form_type":          { "type": ["string", "null"], "description": "e.g. ACORD 25, ACORD 125, ACORD 140" },
      "form_version":       { "type": ["string", "null"] },
      "producer":           { "type": ["object", "null"], "description": "AcordProducer — see Section 10.6" },
      "insured":            { "type": ["object", "null"], "description": "AcordInsured — see Section 10.6" },
      "coverages":          { "type": "array",            "description": "List[AcordPolicyCoverage] — see Section 10.6" },
      "overall_confidence": { "type": "number", "minimum": 0, "maximum": 1 },
      "extraction_meta": {
        "type": ["object", "null"],
        "properties": {
          "extraction_engine":          { "type": "string", "enum": ["bytescout", "pdfplumber", "pymupdf", "pypdf2", "ocr", "txt", "legacy"] },
          "pdf_form_classification":    { "type": "string", "enum": ["fillable", "flattened"] },
          "structured_response_source": { "type": "string", "description": "LLM RunPod | LLM OpenAI | Fallback" },
          "base_confidence":            { "type": "number", "minimum": 0, "maximum": 1 }
        }
      }
    }
  }
]"""

# 10.8 — ACORD Inference Input
SCHEMAS['acord_inf_in'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "ACORDInferenceInput",
  "description": "Input to the ACORD fine-tuned RunPod Inference LLM (OpenAI fallback). Receives raw_text and AcroForm widget values produced by the Extraction pipeline. Files are NOT passed here — the Extraction phase handles all file-to-text conversion.",
  "type": "object",
  "required": ["raw_text"],
  "properties": {
    "raw_text": {
      "type": "string",
      "description": "Plain text produced by the Extraction pipeline engines (BytesScout / pdfplumber / OCR)"
    },
    "form_type_hint": {
      "type": ["string", "null"],
      "description": "Detected or user-specified form type, e.g. '125' or '25'"
    },
    "acroform_fields": {
      "type": ["object", "null"],
      "description": "Key-value map of fillable PDF widget values, keyed by field name. Present only when the PDF has an AcroForm layer.",
      "additionalProperties": { "type": "string" }
    },
    "extraction_engine": {
      "type": ["string", "null"],
      "description": "Which text extraction engine produced raw_text",
      "enum": ["bytescout", "pdfplumber", "pymupdf", "pypdf2", "ocr", "txt", "legacy", null]
    }
  }
}"""

# 10.8 — ACORD Inference Output
SCHEMAS['acord_inf_out'] = """\
{
  "$schema": "http://json-schema.org/draft-07/schema#",
  "title": "ACORDInferenceOutput",
  "description": "AcordFormSummary produced by the RunPod fine-tuned Inference LLM — stored as extracted_json in acord_extraction_runs. See Section 10.6 for the full Pydantic field-by-field schema.",
  "type": "object",
  "required": ["overall_confidence"],
  "properties": {
    "form_type":              { "type": ["string", "null"], "description": "e.g. ACORD 25, ACORD 125" },
    "form_version":           { "type": ["string", "null"] },
    "certificate_number":     { "type": ["string", "null"] },
    "producer":               { "type": ["object", "null"], "description": "AcordProducer — see Section 10.6" },
    "insured":                { "type": ["object", "null"], "description": "AcordInsured — see Section 10.6" },
    "holder":                 { "type": ["object", "null"], "description": "AcordHolder — see Section 10.6" },
    "coverages":              { "type": "array",            "description": "List[AcordPolicyCoverage] — see Section 10.6" },
    "premises":               { "type": "array" },
    "loss_history":           { "type": "array" },
    "additional_interests":   { "type": "array" },
    "description_of_operations": { "type": ["string", "null"] },
    "extra_fields":           { "type": ["object", "null"], "description": "LLM-extracted fields not in standard schema" },
    "overall_confidence":     { "type": "number", "minimum": 0, "maximum": 1 },
    "raw_text":               { "type": "string", "description": "Original extracted text from the document" }
  }
}"""


# ─── Updated body-text paragraphs ────────────────────────────────────────────

# 1. ACORD Pod ID note (remove "acord-parser" reference)
BODY_UPDATES = {
    'Earlier versions of documentation used': (
        'Pod ID standardization note: the canonical pod_id for this pod is acord_form_understanding '
        '— set in the ACORD_POD_ID env var, used in all acord_* DB migrations, and in the shared '
        'pod_extraction_runs.pod_id column. All API routes, training jobs, and admin queue entries '
        'use this identifier.'
    ),
    # 2. 10.7 description — clarify extraction vs inference
    'The ACORD Extraction phase is a multi-engine pipeline (not a single LLM call)': (
        'The ACORD Extraction phase is a multi-engine pipeline — NOT a single LLM call. '
        'Pipeline steps: PDF classification → multi-engine text extraction (BytesScout / pdfplumber / PyMuPDF / PyPDF2) '
        '→ AcroForm widget parsing → OCR (Tesseract / PaddleOCR) '
        '→ general LLM call for field structuring → form-type fallbacks '
        '→ Azure Document Intelligence overlay → confidence scoring → post-processing. '
        'See Section 10.5 for the full 9-step sequence.\n\n'
        'IMPORTANT DISTINCTION — Extraction LLM vs Inference LLM:\n'
        'Extraction Phase (this section): The LLM call at step 5 is a GENERAL-PURPOSE structuring call '
        'that converts raw extracted text into an initial AcordFormSummary draft. '
        'This uses a standard model (RunPod serverless or OpenAI). It is NOT the fine-tuned model.\n'
        'Inference Phase (Section 10.8): A SEPARATELY TRAINED, fine-tuned RunPod model that receives '
        'raw_text + acroform_fields and produces the final AcordFormSummary stored in '
        'acord_extraction_runs.extracted_json. This is the model that enters the training pipeline.'
    ),
    # 3. 10.8 description — clarify inference is fine-tuned model
    'The ACORD Inference LLM receives the raw extracted text plus AcroForm field values': (
        'The ACORD Inference LLM receives the raw extracted text plus AcroForm field values produced '
        'by the Extraction pipeline (Section 10.7) and outputs the final AcordFormSummary.\n\n'
        'IMPORTANT DISTINCTION — This is the RunPod fine-tuned model (with OpenAI as fallback): '
        'a DIFFERENT model from the general LLM used inside the Extraction pipeline’s step 5. '
        'The fine-tuned Inference model is what enters the ACORD training pipeline: when a run is '
        'approved (auto-approve or admin review), a training job is created for this model '
        '(NOT the Extraction LLM). '
        'Training data format: inference input_schema fields form the prompt; '
        'output_schema (AcordFormSummary) fields form the completion. '
        'For inference configuration, only input_schema and output_schema are defined here '
        '— file-to-text conversion is handled entirely by the Extraction phase.'
    ),
}

# ─── Schema paragraph locators ───────────────────────────────────────────────
# Maps a unique text substring found in the paragraph to the schema key to apply.

CODE_BLOCK_MAP = [
    # (substring_to_find_paragraph, schema_key)
    # 7.5 FNOL Extraction Input — the files array with "police_report.pdf"
    ('"index":     0,\n      "filename":  "police_report.pdf"', 'fnol_ext_in'),
    # 7.5 FNOL Extraction Output per file
    ('"file_index":           0,\n  "document_type":        "police_report"', 'fnol_ext_out'),
    # 7.6 FNOL Inference Input
    ('"description": "On March 15 2026 at 2pm, insured vehicle rear-ended a truck...',  'fnol_inf_in'),
    # 7.6 FNOL Inference Output
    ('"incident_date":           "2026-03-15",\n  "incident_location":       "5th Ave & Main St",\n  "lo', 'fnol_inf_out'),
    # 8.6 Quote Generation Inference Input
    ('"insurance_type":    "auto",\n  "selected_carriers":', 'quote_inf_in'),
    # 8.6 Quote Generation Inference Output
    ('"insurance_type": "auto",\n  "applicant": { "name": "Jane Smith"', 'quote_inf_out'),
    # 9.6 Policy Comparison Extraction Input (role array)
    ('"role": "policyA", "filename": "current_policy.pdf"', 'polcomp_ext_in'),
    # 9.6 Policy Comparison Extraction Output
    ('"policyA": {\n    "carrier":          "XYZ Insurance"', 'polcomp_ext_out'),
    # 9.7 Policy Comparison Inference Input
    ('"deviation_threshold_percent": 10,\n  "extraction": {', 'polcomp_inf_in'),
    # 9.7 Policy Comparison Inference Output
    ('"taxonomy": {\n    "domain": "insurance",\n    "doc_type_a": "Commercial General Liability Policy"', 'polcomp_inf_out'),
    # 10.7 ACORD Extraction Input
    ('"filename":       "acord_125_widgets_inc.pdf"', 'acord_ext_in'),
    # 10.7 ACORD Extraction Output
    ('"form_type":         "ACORD 125",\n    "form_version":      "2016/03",\n    "producer":', 'acord_ext_out'),
    # 10.8 ACORD Inference Input
    ('"raw_text":          "<text produced by the extraction pipeline engines>"', 'acord_inf_in'),
    # 10.8 ACORD Inference Output
    ('"form_type":              "ACORD 125",\n  "form_version":           "2016/03",\n  "producer":', 'acord_inf_out'),
]


# ─── Apply body-text changes ──────────────────────────────────────────────────
for substring, new_text in BODY_UPDATES.items():
    idx, para = find_para_by_text(substring)
    if para is None:
        print(f'WARNING: body paragraph not found for substring: {substring[:60]}')
        continue
    replace_body_para(para, new_text)
    print(f'Updated body paragraph P{idx}: {substring[:50]}...')


# ─── Apply code block changes ─────────────────────────────────────────────────
for substring, schema_key in CODE_BLOCK_MAP:
    idx, para = find_para_by_text(substring)
    if para is None:
        print(f'WARNING: code block not found for substring: {substring[:60]}')
        continue
    replace_code_block(para, SCHEMAS[schema_key])
    print(f'Replaced code block P{idx} with schema "{schema_key}": {substring[:50]}...')


# ─── Save ─────────────────────────────────────────────────────────────────────
doc.save('docs/pods_reference_main.docx')
print('\nDone. Saved to docs/pods_reference_main.docx')