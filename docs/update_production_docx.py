"""
Production-ready DOCX update script for pods_reference_main.docx.

Changes applied:
  1. Pack membership table  → new 4-pack segment structure (Brokers, MGA, Carriers, Others)
  2. FNOL identity table    → Pack: Brokers Pack
  3. Quote Gen identity     → Pack: Brokers Pack
  4. Policy Comp identity   → Pack: Brokers Pack
  5. ACORD identity         → Pack: Brokers Pack
  6. Pod Definition JSON    → "pack": "Brokers Pack"
  7. Fideon-Only Extensions → pack field description updated to new IDs
  8. Section 12 env vars    → adds body text paragraph after heading
"""

import io
import sys
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SRC = Path("pods_reference_main.docx")
DST = Path("pods_reference_production.docx")
shutil.copy(SRC, DST)

doc = Document(DST)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def get_body_elements(doc):
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        yield tag, child


def set_cell_text(cell, text):
    """Replace cell content preserving first paragraph's formatting."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for run in p.runs[1:]:
        run._element.getparent().remove(run._element)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def find_table_by_header(doc, *header_cells):
    """Return the first table whose first row starts with the given cell texts."""
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        first_row = [c.text.strip() for c in table.rows[0].cells]
        if all(first_row[i] == header_cells[i] for i in range(len(header_cells))):
            return table
    return None


def find_para_containing(doc, *substrings):
    """Return first paragraph whose text contains ALL given substrings."""
    for p in doc.paragraphs:
        if all(s in p.text for s in substrings):
            return p
    return None


def insert_paragraph_after(ref_para, text):
    """Insert a new paragraph immediately after ref_para using XML copy approach."""
    from docx.oxml import OxmlElement
    import copy
    # Clone the reference paragraph element to inherit its style
    new_p = copy.deepcopy(ref_para._element)
    # Clear all runs from the clone
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # Add a single run with the new text
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_el.append(t_el)
    new_p.append(r_el)
    ref_para._element.addnext(new_p)
    return new_p


# ─────────────────────────────────────────────────────────────────────────────
# 1. Pack membership table
# ─────────────────────────────────────────────────────────────────────────────
pack_table = find_table_by_header(doc, "Pack", "Pods Included")
if pack_table:
    # Remove all data rows
    for row in pack_table.rows[1:]:
        tr = row._tr
        tr.getparent().remove(tr)

    new_rows = [
        ("Brokers Pack",  "claims-fnol, quote-generation, policy-comparison, acord_form_understanding (+ 12 more broker pods)"),
        ("MGA Pack",      "mga-binding-authority, mga-program-underwriting, mga-bordereaux-generator, mga-capacity-matching, mga-producer-management, mga-treaty-compliance"),
        ("Carriers Pack", "carrier-submission-intake, carrier-submission-triage, carrier-risk-scoring, carrier-pricing-engine, carrier-claims-intake, carrier-claims-adjudication, carrier-fraud-detection, carrier-subrogation, carrier-policy-issuance, carrier-reinsurance"),
        ("Others Pack",   "(reserved — pods TBD)"),
    ]
    for pack_name, pods in new_rows:
        row = pack_table.add_row()
        set_cell_text(row.cells[0], pack_name)
        set_cell_text(row.cells[1], pods)
    print("✓ Pack membership table updated")
else:
    print("✗ Pack membership table NOT found")


# ─────────────────────────────────────────────────────────────────────────────
# 2–5. Pod identity tables — update Pack row
# ─────────────────────────────────────────────────────────────────────────────
pod_pack_updates = [
    ("claims-fnol",              "Brokers Pack"),
    ("quote-generation",         "Brokers Pack"),
    ("policy-comparison",        "Brokers Pack"),
    ("acord_form_understanding", "Brokers Pack"),
]

for pod_id, new_pack in pod_pack_updates:
    found = False
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        # Look for identity table: has "Pod ID" in first column and pod_id value in second
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells[0] == "Pod ID" and cells[1] == pod_id:
                # Found the identity table — now find the Pack row
                for r in table.rows:
                    rc = [c.text.strip() for c in r.cells]
                    if rc[0] == "Pack":
                        set_cell_text(r.cells[1], new_pack)
                        found = True
                        break
                break
        if found:
            break
    print(f"{'✓' if found else '✗'} Identity table Pack updated for {pod_id}")


# ─────────────────────────────────────────────────────────────────────────────
# 6. Pod Definition JSON example — update "pack" value
# ─────────────────────────────────────────────────────────────────────────────
for p in doc.paragraphs:
    if '"pack":' in p.text and "Claims Pack" in p.text:
        for run in p.runs:
            if "Claims Pack" in run.text:
                run.text = run.text.replace('"Claims Pack"', '"Brokers Pack"')
                print('✓ Pod Definition JSON "pack" value updated')
                break


# ─────────────────────────────────────────────────────────────────────────────
# 7. Fideon-Only Extensions table — update pack field description
# ─────────────────────────────────────────────────────────────────────────────
ext_table = find_table_by_header(doc, "Field", "Type", "Description")
if ext_table:
    for row in ext_table.rows:
        cells = [c.text.strip() for c in row.cells]
        if cells[0] == "pack":
            old_desc = row.cells[2].text
            if "Claims Pack" in old_desc or "Underwriting" in old_desc:
                set_cell_text(
                    row.cells[2],
                    'Segment pack membership. One of: "brokers" | "mga" | "carriers" | "others"',
                )
                print("✓ Fideon-Only Extensions pack description updated")
            break


# ─────────────────────────────────────────────────────────────────────────────
# 8. Section 12 — Environment Variables body text
# ─────────────────────────────────────────────────────────────────────────────
env_heading = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and "12" in p.text and "Environment" in p.text:
        env_heading = p
        break

if env_heading:
    env_body = (
        "The following environment variables control runtime behaviour across all four pods. "
        "Variables marked (required) have no default and must be set before starting the backend. "
        "All threshold variables accept float values in the range 0.0 – 1.0. "
        "ACORD-specific variables are only required when the ACORD Form Extraction pod is active."
    )
    insert_paragraph_after(env_heading, env_body)
    print("✓ Section 12 body text added")
else:
    print("✗ Section 12 heading not found")


# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"\n✓ Saved: {DST}")