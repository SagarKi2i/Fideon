"""
Generate pipeline architecture DOCX:
  Fine-Tune → Quantize → Azure Blob / SeaweedFS → Electron App
"""

import io, textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────
C = {
    "bg":       "#080C14",
    "purple":   "#AFA9EC",
    "purple_d": "#3C3489",
    "teal":     "#5DCAA5",
    "teal_d":   "#085041",
    "blue":     "#85B7EB",
    "blue_d":   "#0C447C",
    "amber":    "#EF9F27",
    "amber_d":  "#633806",
    "green":    "#4ADE80",
    "ink":      "#EDF2FF",
    "ink2":     "#8DA4C4",
    "red":      "#EF4444",
}

# ─────────────────────────────────────────────
# HELPER: rounded box
# ─────────────────────────────────────────────
def rbox(ax, x, y, w, h, fc, ec, text, fontsize=7.5, text_color="white",
         bold=False, pad=0.015, alpha=0.88, subtext=None, subtext_color=None):
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle=f"round,pad={pad}",
                         linewidth=1.2, edgecolor=ec,
                         facecolor=fc, alpha=alpha, zorder=3)
    ax.add_patch(box)
    ty = y + h / 2 + (0.012 if subtext else 0)
    ax.text(x + w/2, ty, text,
            ha="center", va="center",
            fontsize=fontsize, color=text_color,
            fontweight="bold" if bold else "normal",
            fontfamily="DejaVu Sans", zorder=4, wrap=False)
    if subtext:
        ax.text(x + w/2, y + h/2 - 0.018, subtext,
                ha="center", va="center",
                fontsize=6.2, color=subtext_color or C["ink2"],
                fontfamily="DejaVu Sans", zorder=4)

def arrow(ax, x1, y1, x2, y2, color="#475569", lw=1.4, label=None):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color,
                                lw=lw, mutation_scale=10),
                zorder=5)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.01, my, label, fontsize=5.8, color=C["ink2"],
                va="center", fontfamily="DejaVu Sans", zorder=6)

def stage_bg(ax, x, y, w, h, fc, ec, title, title_color):
    bg = FancyBboxPatch((x, y), w, h,
                        boxstyle="round,pad=0.008",
                        linewidth=1, edgecolor=ec,
                        facecolor=fc, alpha=0.28, zorder=1)
    ax.add_patch(bg)
    ax.text(x + w/2, y + h - 0.022, title,
            ha="center", va="top",
            fontsize=7.8, color=title_color,
            fontweight="bold", fontfamily="DejaVu Sans", zorder=2)


# ─────────────────────────────────────────────
# FIGURE 1 — Full pipeline overview
# ─────────────────────────────────────────────
def make_overview_diagram():
    fig, ax = plt.subplots(figsize=(14, 8.5))
    fig.patch.set_facecolor(C["bg"])
    ax.set_facecolor(C["bg"])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    # ── title ──
    ax.text(0.5, 0.975, "LLM Fine-Tune → Quantize → Azure Blob / SeaweedFS → Electron",
            ha="center", va="top", fontsize=10, fontweight="bold",
            color=C["ink"], fontfamily="DejaVu Sans")
    ax.text(0.5, 0.955, "End-to-end automated pipeline for insurance document AI",
            ha="center", va="top", fontsize=7.5, color=C["ink2"],
            fontfamily="DejaVu Sans")

    # ── Stage 1: Vast.ai ──
    stage_bg(ax, 0.02, 0.73, 0.96, 0.20, C["purple_d"], C["purple"],
             "Stage 1 — Vast.ai GPU Node  (GPU billing active)", C["purple"])
    rbox(ax, 0.04, 0.755, 0.17, 0.08, C["purple_d"], C["purple"],
         "ACORD Extraction", bold=True, text_color=C["purple"],
         subtext="Qwen2.5-VL-7B\nPOST /extract", subtext_color=C["ink2"])
    arrow(ax, 0.21, 0.795, 0.255, 0.795, C["purple"])
    rbox(ax, 0.255, 0.755, 0.17, 0.08, C["purple_d"], C["purple"],
         "QLoRA Fine-Tune", bold=True, text_color=C["purple"],
         subtext="Unsloth · LLaMA-Factory\nrank=16 · ~45 min", subtext_color=C["ink2"])
    arrow(ax, 0.425, 0.795, 0.47, 0.795, C["purple"])
    rbox(ax, 0.47, 0.755, 0.17, 0.08, C["purple_d"], C["purple"],
         "merge_and_unload()", bold=True, text_color=C["purple"],
         subtext="LoRA adapter → BF16\nfull merged weights", subtext_color=C["ink2"])
    arrow(ax, 0.64, 0.795, 0.685, 0.795, C["purple"])
    rbox(ax, 0.685, 0.755, 0.185, 0.08, C["purple_d"], C["purple"],
         "llama.cpp Quantize", bold=True, text_color=C["purple"],
         subtext="Q5_K_M · Q4_K_M · FP16\nsha256 · gpg signed", subtext_color=C["ink2"])
    rbox(ax, 0.885, 0.755, 0.09, 0.08, "#2A0A0A", C["red"],
         "POST /stop\nGPU ends", bold=False, text_color=C["red"],
         fontsize=6.5)

    # ── Arrow Stage 1 → Stage 2 ──
    arrow(ax, 0.5, 0.73, 0.5, 0.695, C["ink2"], label="q5km.gguf · manifest.json · *.sig")

    # ── Stage 2: Storage ──
    stage_bg(ax, 0.02, 0.515, 0.96, 0.175, C["blue_d"], C["blue"],
             "Stage 2 — Upload & Register  (CPU · on-premise)", C["blue"])
    rbox(ax, 0.04, 0.545, 0.19, 0.07, C["blue_d"], C["blue"],
         "Sign & Verify", bold=True, text_color=C["blue"],
         subtext="sha256sum · gpg --sign\nmanifest.json", subtext_color=C["ink2"])
    arrow(ax, 0.23, 0.58, 0.27, 0.58, C["blue"])
    rbox(ax, 0.27, 0.545, 0.19, 0.07, C["blue_d"], C["blue"],
         "Azure Blob Upload", bold=True, text_color=C["blue"],
         subtext="azure-storage-blob SDK\npresigned URL · tenant-scoped", subtext_color=C["ink2"])
    arrow(ax, 0.46, 0.58, 0.50, 0.58, C["blue"])
    rbox(ax, 0.50, 0.545, 0.19, 0.07, C["blue_d"], C["blue"],
         "SeaweedFS Mirror", bold=True, text_color=C["blue"],
         subtext="S3-compat · boto3\nair-gap / on-premise", subtext_color=C["ink2"])
    arrow(ax, 0.69, 0.58, 0.73, 0.58, C["blue"])
    rbox(ax, 0.73, 0.545, 0.235, 0.07, C["blue_d"], C["blue"],
         "Model Registry (Supabase)", bold=True, text_color=C["blue"],
         subtext="blob_url · sha256 · quant_level\nrollback_ptr · version", subtext_color=C["ink2"])

    # ── Arrow Stage 2 → Stage 3 ──
    arrow(ax, 0.5, 0.515, 0.5, 0.48, C["ink2"], label="heartbeat row: is_downloaded=false")

    # ── Stage 3: Electron ──
    stage_bg(ax, 0.02, 0.30, 0.96, 0.175, C["teal_d"], C["teal"],
             "Stage 3 — Electron Desktop App  (customer on-premise)", C["teal"])
    rbox(ax, 0.04, 0.33, 0.19, 0.07, C["teal_d"], C["teal"],
         "Device Heartbeat", bold=True, text_color=C["teal"],
         subtext="Supabase local_models\nnew row detected", subtext_color=C["ink2"])
    arrow(ax, 0.23, 0.365, 0.27, 0.365, C["teal"])
    rbox(ax, 0.27, 0.33, 0.19, 0.07, C["teal_d"], C["teal"],
         "\"Update Available\" UI", bold=True, text_color=C["teal"],
         subtext="badge · version info\nuser clicks Download", subtext_color=C["ink2"])
    arrow(ax, 0.46, 0.365, 0.50, 0.365, C["teal"])
    rbox(ax, 0.50, 0.33, 0.21, 0.07, C["teal_d"], C["teal"],
         "Stream Download", bold=True, text_color=C["teal"],
         subtext="presigned URL → temp file\nsha256 verify on-device", subtext_color=C["ink2"])
    arrow(ax, 0.71, 0.365, 0.75, 0.365, C["teal"])
    rbox(ax, 0.75, 0.33, 0.215, 0.07, C["teal_d"], C["teal"],
         "ollama create", bold=True, text_color=C["teal"],
         subtext="acord/model:v1.2-q5km\nModelfile → local serving", subtext_color=C["ink2"])

    # ── Arrow Stage 3 → Stage 4 ──
    arrow(ax, 0.5, 0.30, 0.5, 0.265, C["ink2"], label="OpenAI-compat · port 11434")

    # ── Stage 4: Serving ──
    stage_bg(ax, 0.02, 0.09, 0.96, 0.165, C["amber_d"], C["amber"],
             "Stage 4 — Local Inference  (zero data egress)", C["amber"])
    rbox(ax, 0.04, 0.12, 0.19, 0.065, C["amber_d"], C["amber"],
         "Ollama Runtime", bold=True, text_color=C["amber"],
         subtext="GGUF loaded into RAM\nCPU or GPU offload", subtext_color=C["ink2"])
    arrow(ax, 0.23, 0.1525, 0.27, 0.1525, C["amber"])
    rbox(ax, 0.27, 0.12, 0.19, 0.065, C["amber_d"], C["amber"],
         "OpenAI-compat API", bold=True, text_color=C["amber"],
         subtext="port 11434\nACORD extraction calls", subtext_color=C["ink2"])
    arrow(ax, 0.46, 0.1525, 0.50, 0.1525, C["amber"])
    rbox(ax, 0.50, 0.12, 0.19, 0.065, C["amber_d"], C["amber"],
         "Langfuse Tracing", bold=True, text_color=C["amber"],
         subtext="latency · accuracy\nHITL feedback loop", subtext_color=C["ink2"])
    arrow(ax, 0.69, 0.1525, 0.73, 0.1525, C["amber"])
    rbox(ax, 0.73, 0.12, 0.215, 0.065, C["amber_d"], C["amber"],
         "Rollback (1-click)", bold=True, text_color=C["amber"],
         subtext="ollama run model:v1.1-q5km\nsingle tag change · <5s", subtext_color=C["ink2"])

    # ── tools footer ──
    ax.text(0.5, 0.045,
            "Tools: Unsloth · LLaMA-Factory · llama.cpp Q5_K_M/Q4_K_M · azure-storage-blob · boto3 · SeaweedFS · sha256sum · gpg · Ollama · Langfuse",
            ha="center", va="top", fontsize=6.2, color=C["ink2"],
            fontfamily="DejaVu Sans")
    ax.text(0.5, 0.022,
            "~52 min GPU total (45 min train + 7 min quant) · GPU billing ends after Stage 1 · all subsequent stages run CPU-only",
            ha="center", va="top", fontsize=6.0, color="#4C6280",
            fontfamily="DejaVu Sans")

    plt.tight_layout(pad=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C["bg"], edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# FIGURE 2 — Existing codebase integration map
# ─────────────────────────────────────────────
def make_integration_diagram():
    fig, ax = plt.subplots(figsize=(14, 6.5))
    fig.patch.set_facecolor(C["bg"])
    ax.set_facecolor(C["bg"])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.5, 0.975, "Integration Map — New Code vs Existing Codebase",
            ha="center", va="top", fontsize=10, fontweight="bold",
            color=C["ink"], fontfamily="DejaVu Sans")
    ax.text(0.5, 0.952, "Green = already exists · Blue = extend · Orange = new file",
            ha="center", va="top", fontsize=7.5, color=C["ink2"],
            fontfamily="DejaVu Sans")

    # ── legend ──
    for i, (label, color) in enumerate([
        ("Already exists", C["teal"]),
        ("Extend existing", C["blue"]),
        ("New file / code", C["amber"]),
    ]):
        lx = 0.04 + i * 0.25
        rbox(ax, lx, 0.895, 0.18, 0.038, "#0D1422", color,
             label, fontsize=7, text_color=color, bold=True, pad=0.008)

    # ── Backend column ──
    ax.text(0.18, 0.845, "Backend (FastAPI · Python)", ha="center",
            fontsize=8, color=C["ink"], fontweight="bold", fontfamily="DejaVu Sans")
    items_be = [
        ("fine_tuning/job_runner.py", C["blue"],   "Extend: add quantize\n+ upload call after train"),
        ("fine_tuning/train.py",      C["teal"],   "Already runs QLoRA;\nno changes needed"),
        ("fine_tuning/quantize.py",   C["amber"],  "NEW: merge_and_unload()\n+ llama-quantize subprocess"),
        ("app/services/model_storage.py", C["amber"], "NEW: Azure Blob + SeaweedFS\nupload · presigned URLs"),
        ("app/routes/model_registry.py", C["blue"], "Extend: add blob_url,\nsha256, quant_level columns"),
        ("app/services/webhook_engine.py", C["teal"], "Already fires events;\nno changes needed"),
    ]
    for i, (name, color, note) in enumerate(items_be):
        y = 0.77 - i * 0.105
        rbox(ax, 0.04, y, 0.275, 0.085, "#0D1422", color,
             name, fontsize=6.8, text_color=color, bold=True,
             subtext=note, subtext_color=C["ink2"], pad=0.01)

    # ── DB column ──
    ax.text(0.5, 0.845, "Database (Supabase · PostgreSQL)", ha="center",
            fontsize=8, color=C["ink"], fontweight="bold", fontfamily="DejaVu Sans")
    items_db = [
        ("public.local_models",      C["blue"],   "Extend: heartbeat already syncs;\nadd new model row trigger"),
        ("public.model_registry",    C["blue"],   "Extend: ALTER TABLE → add\nblob_url, sha256, quant_level"),
        ("public.acord_training_jobs", C["teal"], "Already tracks job state;\nno changes needed"),
        ("20260409_quantize_cols.sql", C["amber"], "NEW migration: add columns\nto model_registry table"),
    ]
    for i, (name, color, note) in enumerate(items_db):
        y = 0.77 - i * 0.105
        rbox(ax, 0.36, y, 0.275, 0.085, "#0D1422", color,
             name, fontsize=6.8, text_color=color, bold=True,
             subtext=note, subtext_color=C["ink2"], pad=0.01)

    # ── Electron column ──
    ax.text(0.82, 0.845, "Electron App (TypeScript)", ha="center",
            fontsize=8, color=C["ink"], fontweight="bold", fontfamily="DejaVu Sans")
    items_el = [
        ("electron/ollama-backend.ts", C["blue"],  "Extend: add model:download-\nupdate IPC handler"),
        ("electron/device-client.ts",  C["teal"],  "Already handles heartbeat;\nno changes needed"),
        ("electron/main.ts",           C["blue"],  "Extend: show notification\nbadge on new model row"),
        ("electron/model-updater.ts",  C["amber"], "NEW: stream download,\nsha256 verify, ollama create"),
    ]
    for i, (name, color, note) in enumerate(items_el):
        y = 0.77 - i * 0.105
        rbox(ax, 0.68, y, 0.295, 0.085, "#0D1422", color,
             name, fontsize=6.8, text_color=color, bold=True,
             subtext=note, subtext_color=C["ink2"], pad=0.01)

    # ── arrows between columns ──
    for y_frac in [0.7, 0.595, 0.49]:
        arrow(ax, 0.315, y_frac + 0.042, 0.36, y_frac + 0.042, C["ink2"])
    for y_frac in [0.7, 0.595]:
        arrow(ax, 0.635, y_frac + 0.042, 0.68, y_frac + 0.042, C["ink2"])

    ax.text(0.5, 0.06,
            "Existing patterns reused: RunPod orchestrator → Vast.ai orchestrator (same poll/teardown pattern)  ·  "
            "litellm fallback chain unchanged  ·  Supabase RLS unchanged",
            ha="center", va="top", fontsize=6.4, color="#4C6280",
            fontfamily="DejaVu Sans")

    plt.tight_layout(pad=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C["bg"], edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def add_heading(doc, text, level=1, color_hex="2B2B6B"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_para(doc, text, bold=False, italic=False, size=10, color_hex=None,
             space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_bullet(doc, text, level=0, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_code(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x2D, 0xD4, 0xBF)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def stage_table(doc, rows_data, header_color, header_text_color="FFFFFF"):
    """rows_data: list of (col1, col2) strings. First row = header."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(2.2), Inches(4.4)]
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, (cell_text, width) in enumerate(zip(row_data, widths)):
            cell = row.cells[j]
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(header_text_color)
                set_cell_bg(cell, header_color)
            else:
                run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────
# BUILD DOCX
# ─────────────────────────────────────────────
def build_docx():
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.4)
        section.right_margin  = Cm(2.4)

    # ══════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("LLM Fine-Tune · Quantize · Azure Blob / SeaweedFS")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x2B, 0x2B, 0x6B)

    add_para(doc, "Distribution Pipeline — Architecture & Integration Guide",
             bold=True, size=14, color_hex="0C447C",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)

    add_para(doc, "Neura-Box Cloud  ·  v1.0  ·  April 2026",
             size=10, color_hex="8DA4C4",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)

    add_para(doc,
             "This document describes the end-to-end automated pipeline for taking a QLoRA fine-tuned LLM "
             "(trained on Vast.ai) through GPU-accelerated quantization, uploading signed GGUF artifacts to "
             "Azure Blob Storage (with SeaweedFS as an on-premise mirror), and delivering them to the Electron "
             "desktop app as a one-click model update — with zero data egress and full integrity verification.",
             size=10, color_hex="4C6280",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=20)

    add_divider(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1 — OVERVIEW DIAGRAM
    # ══════════════════════════════════════════
    add_heading(doc, "1. Pipeline Architecture Overview", level=1, color_hex="2B2B6B")
    add_para(doc,
             "The pipeline has four stages. GPU billing is active only during Stage 1 (~52 min total). "
             "All subsequent stages run on CPU-only infrastructure with zero data egress to the cloud.",
             size=10)

    overview_buf = make_overview_diagram()
    doc.add_picture(overview_buf, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 1 — Full pipeline: Vast.ai → Azure Blob / SeaweedFS → Electron Ollama",
             size=8, color_hex="8DA4C4", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=16)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 2 — STAGE BY STAGE
    # ══════════════════════════════════════════
    add_heading(doc, "2. Stage-by-Stage Breakdown", level=1, color_hex="2B2B6B")

    # ── Stage 1 ──
    add_heading(doc, "Stage 1 — Vast.ai GPU Node  (GPU billing active)", level=2, color_hex="3C3489")
    add_para(doc,
             "The Vast.ai node is provisioned via POST /start. The base model is loaded into VRAM. "
             "QLoRA fine-tuning runs in a background thread (~45 min). Once training completes, the LoRA "
             "adapter is merged into the base BF16 weights using merge_and_unload(), then llama.cpp "
             "GGML CUDA kernels quantize the merged model in ~7 min. The node is terminated immediately "
             "after — GPU billing ends here.", size=10)

    stage_table(doc, [
        ("Step", "Detail"),
        ("ACORD Extraction",     "Qwen2.5-VL-7B · POST /extract · page images → JSON fields"),
        ("HITL Corrections",     "Next.js review app · verified_text.json · JSONL builder · PII redacted"),
        ("QLoRA Fine-Tuning",    "Unsloth + LLaMA-Factory · rank=16 · α=32 · BF16 · 3 epochs · ~45 min"),
        ("merge_and_unload()",   "Merges LoRA adapter into base weights: W' = W + (α/r)·B·A"),
        ("llama.cpp Quantize",   "Q5_K_M (~10.7 GB, production)  ·  Q4_K_M (~8.0 GB, edge/air-gap)  ·  FP16 (CI reference)"),
        ("Sign Artifacts",       "sha256sum per file · gpg --detach-sign · manifest.json generated"),
        ("POST /stop",           "Graceful shutdown · VRAM released · GPU billing ends · ~52 min total billed"),
    ], header_color="3C3489")

    # ── Stage 2 ──
    add_heading(doc, "Stage 2 — Upload to Azure Blob & SeaweedFS  (CPU)", level=2, color_hex="0C447C")
    add_para(doc,
             "Signed GGUF artifacts are uploaded from the Vast.ai node (or the backend orchestrator) to "
             "Azure Blob Storage as the primary store. SeaweedFS acts as an on-premise S3-compatible mirror "
             "for air-gapped customers. Both use the same boto3/azure-sdk code path with different endpoint configs. "
             "On successful upload the model registry row is written to Supabase.", size=10)

    stage_table(doc, [
        ("Step", "Detail"),
        ("Azure Blob Upload",    "azure-storage-blob SDK · models/{domain}/v{ver}/{quant}/ · presigned URL · 24hr expiry · tenant-scoped"),
        ("SeaweedFS Mirror",     "boto3 S3 PutObject · same path schema · different endpoint URL in config · air-gap support"),
        ("SHA-256 Verify",       "Server-side checksum against manifest.json · mismatch → reject + alert webhook"),
        ("Registry Insert",      "Supabase model_registry: blob_url · sha256 · quant_level · docker_sha · rollback_ptr · version"),
        ("Heartbeat Trigger",    "New row inserted into public.local_models with is_downloaded=false for all active devices"),
    ], header_color="0C447C")

    # ── Stage 3 ──
    add_heading(doc, "Stage 3 — Electron Desktop App — Detection & Download  (on-premise)", level=2, color_hex="085041")
    add_para(doc,
             "The Electron app's existing device heartbeat (device-client.ts) polls Supabase for model "
             "sync state. When it detects a new local_models row with is_downloaded=false it surfaces an "
             "'Update Available' notification. The user clicks Download — a presigned Azure Blob URL is "
             "fetched from the backend, the GGUF is streamed to a temp file, sha256 is verified on-device, "
             "then ollama create registers the model.", size=10)

    stage_table(doc, [
        ("Step", "Detail"),
        ("Heartbeat Poll",        "device-client.ts polls /device/sync · new local_models row detected · is_downloaded=false"),
        ("UI Notification",       "main.ts: badge on tray icon + in-app banner · version info · quant level · file size"),
        ("Presigned URL Fetch",   "IPC: model:download-update → backend GET /models/download-url → returns Azure Blob presigned URL"),
        ("Stream Download",       "electron/model-updater.ts · httpx/fetch stream → temp file · progress bar via IPC"),
        ("SHA-256 Verify",        "sha256 computed on downloaded file · compared against manifest.json · mismatch → delete + alert"),
        ("ollama create",         "ollama create acord/qwen25-14b:v1.2-q5km -f Modelfile (points to .gguf path)"),
        ("Registry Update",       "PATCH /device/models → is_downloaded=true · updated_at · version confirmed"),
    ], header_color="085041")

    # ── Stage 4 ──
    add_heading(doc, "Stage 4 — Local Inference  (zero data egress)", level=2, color_hex="633806")
    add_para(doc,
             "Once registered in Ollama, the model serves ACORD extraction requests via the OpenAI-compatible "
             "API on port 11434. Langfuse tracing captures latency and accuracy. The Electron app can roll back "
             "to the previous version with a single tag change — no retraining required.", size=10)

    stage_table(doc, [
        ("Step", "Detail"),
        ("Ollama Runtime",        "GGUF loaded into RAM/VRAM · CPU or GPU offload · no PEFT library needed on serving node"),
        ("OpenAI-compat API",     "port 11434 · /api/chat · /api/generate · drop-in for existing ACORD extraction calls"),
        ("Langfuse Tracing",      "Latency · token throughput · extraction accuracy · HITL feedback loop"),
        ("Rollback",              "ollama run acord/model:v1.1-q5km — single tag change · <5 seconds · no retraining"),
        ("Air-gap Path",          "SeaweedFS presigned URL · or SFTP push to DMZ · or signed .tar.gz via USB courier"),
    ], header_color="633806", header_text_color="FFFFFF")

    add_divider(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 3 — INTEGRATION MAP
    # ══════════════════════════════════════════
    add_heading(doc, "3. Existing Codebase Integration", level=1, color_hex="2B2B6B")
    add_para(doc,
             "The diagram below maps each pipeline step to existing files in the neura-box-cloud repository. "
             "Green = no changes needed. Blue = extend. Orange = new file to create.",
             size=10)

    integ_buf = make_integration_diagram()
    doc.add_picture(integ_buf, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 2 — Integration map: new code vs existing files",
             size=8, color_hex="8DA4C4", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=16)

    # ── existing table ──
    add_heading(doc, "3.1  Files That Already Exist (No Changes)", level=2, color_hex="085041")
    stage_table(doc, [
        ("File", "Existing Role → Used As-Is"),
        ("fine_tuning/train.py",            "QLoRA training · already runs on GPU node · no changes"),
        ("fine_tuning/evaluate.py",         "Eval metrics · already called by job_runner · no changes"),
        ("app/services/runpod_orchestrator.py", "Pod provision/teardown pattern → copy for Vast.ai orchestrator"),
        ("app/services/webhook_engine.py",  "Already fires events · used for 'new model ready' webhook"),
        ("app/routes/device.py",            "Heartbeat + model sync · already writes local_models rows"),
        ("electron/device-client.ts",       "Heartbeat poll · already reads local_models · detect new rows"),
        ("electron/ollama-backend.ts",      "ollama create/delete/list · already implemented"),
        ("electron/preload.ts",             "IPC bridge · just add new channel names"),
    ], header_color="085041")

    add_heading(doc, "3.2  Files to Extend", level=2, color_hex="0C447C")
    stage_table(doc, [
        ("File", "What to Add"),
        ("fine_tuning/job_runner.py",       "After train completes: call quantize.py → model_storage.py → registry insert"),
        ("app/routes/model_registry.py",    "Handle new columns: blob_url · sha256 · quant_level · rollback_ptr"),
        ("electron/main.ts",                "IPC handler: model:download-update · tray badge on new model notification"),
        ("electron/ollama-backend.ts",      "modelCreateFromFile(path, modelName) helper for local GGUF registration"),
    ], header_color="0C447C")

    add_heading(doc, "3.3  New Files to Create", level=2, color_hex="633806")
    stage_table(doc, [
        ("New File", "Purpose"),
        ("fine_tuning/quantize.py",          "merge_and_unload() + llama-quantize subprocess · outputs q5km.gguf · sha256 · manifest.json"),
        ("app/services/model_storage.py",    "Azure Blob upload (azure-storage-blob) + SeaweedFS mirror (boto3) · presigned URL generation"),
        ("electron/model-updater.ts",        "Stream download from presigned URL · sha256 verify · call ollama create · IPC progress events"),
        ("supabase/migrations/YYYYMMDD_quantize_cols.sql", "ALTER TABLE model_registry ADD COLUMN blob_url, sha256, quant_level, rollback_ptr"),
    ], header_color="633806", header_text_color="FFFFFF")

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 4 — DB MIGRATION
    # ══════════════════════════════════════════
    add_heading(doc, "4. Database Schema Changes", level=1, color_hex="2B2B6B")
    add_para(doc, "Add the following columns to the existing model_registry table:", size=10)

    add_code(doc, "-- supabase/migrations/20260410_quantize_distribution_cols.sql")
    add_code(doc, "ALTER TABLE public.model_registry")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS blob_url        text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS sha256          text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS quant_level     text,   -- 'q5km' | 'q4km' | 'fp16'")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS rollback_ptr    text,   -- model_registry.id of prev version")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS seaweedfs_url   text;   -- optional on-premise mirror URL")
    doc.add_paragraph()
    add_para(doc,
             "The existing local_models table already has is_downloaded (boolean) and device_id. "
             "A new row is inserted here for each active device when a new model version is registered — "
             "the heartbeat picks it up automatically without any schema changes.", size=10)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 5 — AZURE BLOB vs SEAWEEDFS
    # ══════════════════════════════════════════
    add_heading(doc, "5. Azure Blob vs SeaweedFS — How They Coexist", level=1, color_hex="2B2B6B")

    stage_table(doc, [
        ("Dimension",           "Azure Blob (Primary)",         "SeaweedFS (Mirror)"),
        ("Purpose",             "Tenant download via presigned URL",    "On-premise / air-gap mirror"),
        ("SDK",                 "azure-storage-blob Python SDK",        "boto3 S3-compat API"),
        ("Auth",                "SAS token · tenant-scoped · 24hr",     "VAST_API_SECRET · weed filer endpoint"),
        ("Path schema",         "models/{domain}/v{ver}/{quant}/",      "Same path schema"),
        ("Integrity",           "SHA-256 verified on upload",           "Same — verified against manifest.json"),
        ("When used",           "Connected Electron app (default)",     "Air-gapped nodes · SFTP path · USB bundle"),
        ("Code change",         "New model_storage.py",                 "Same file · different endpoint env var"),
    ] , header_color="0C447C")

    add_para(doc,
             "Both targets are written in a single upload_model() function. "
             "Setting SEAWEEDFS_ENDPOINT in the environment enables the mirror — "
             "if unset, only Azure Blob is written.", size=10, space_before=6)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 6 — BUILD ORDER
    # ══════════════════════════════════════════
    add_heading(doc, "6. Suggested Implementation Order", level=1, color_hex="2B2B6B")
    add_para(doc,
             "Each step is independently testable before the next is started.", size=10)

    steps = [
        ("1", "fine_tuning/quantize.py",
         "Implement merge_and_unload() + llama-quantize subprocess. "
         "Test locally: give it a dummy adapter, confirm .gguf output + sha256 file."),
        ("2", "app/services/model_storage.py",
         "Azure Blob upload + SeaweedFS upload + presigned URL generation. "
         "Test with a dummy 1 MB file against a dev Azure container."),
        ("3", "Supabase migration",
         "Run ALTER TABLE on model_registry. "
         "Confirm local_models insert trigger works with a test row."),
        ("4", "job_runner.py extension",
         "Wire quantize.py → model_storage.py → registry insert into the post-train hook. "
         "Test full flow end-to-end with a small model on Vast.ai."),
        ("5", "electron/model-updater.ts",
         "Implement IPC download handler: presigned URL fetch → stream → sha256 verify → ollama create. "
         "Test with a pre-uploaded GGUF in Azure Blob."),
        ("6", "electron/main.ts notification",
         "On heartbeat response, check for new local_models rows. "
         "Show tray badge + in-app banner. Wire to model-updater.ts."),
        ("7", "End-to-end test",
         "Run a full training cycle on Vast.ai. Verify GGUF appears in Azure Blob, "
         "notification fires in Electron, download completes, ollama serves correctly."),
    ]

    for num, title, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f"  Step {num} — {title}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
        add_para(doc, f"     {desc}", size=9.5, color_hex="4C6280",
                 space_before=0, space_after=6)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 7 — TOOLCHAIN
    # ══════════════════════════════════════════
    add_heading(doc, "7. Toolchain Summary", level=1, color_hex="2B2B6B")

    stage_table(doc, [
        ("Stage",       "Tools"),
        ("Stage 1 — Training",      "Unsloth · LLaMA-Factory · HuggingFace PEFT + TRL · bitsandbytes · BF16"),
        ("Stage 1 — Quantization",  "llama.cpp (GGML CUDA build) · convert_hf_to_gguf.py · llama-quantize · gpg · sha256sum"),
        ("Stage 2 — Storage",       "azure-storage-blob SDK · boto3 (SeaweedFS S3-compat) · weed filer · manifest.json"),
        ("Stage 2 — Registry",      "Supabase PostgreSQL · FastAPI route · webhook_engine.py"),
        ("Stage 3 — Electron",      "electron/model-updater.ts · Node.js fetch stream · electron-store · IPC"),
        ("Stage 4 — Serving",       "Ollama (GGUF native) · OpenAI-compat API · Langfuse tracing · rollback by tag"),
    ], header_color="2B2B6B")

    add_divider(doc)

    # ══════════════════════════════════════════
    # FOOTER NOTE
    # ══════════════════════════════════════════
    add_para(doc,
             "GPU cost estimate: ~52 min per fine-tuning cycle (45 min QLoRA + 7 min quantization). "
             "All stages after GPU termination run on CPU-only infrastructure. "
             "Rollback to any prior version: single ollama tag change, under 5 seconds, no retraining required.",
             size=9, color_hex="8DA4C4", space_before=10)

    # ── save ──
    out_path = os.path.join(OUT_DIR, "LLM_Pipeline_Architecture.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_docx()
