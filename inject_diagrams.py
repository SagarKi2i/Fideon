"""
inject_diagrams.py
Opens the EXISTING Activity_Logs_Full_Documentation_20260317_125220.docx,
renders a real matplotlib image for every Mermaid / diagram code block,
inserts each image immediately after its code paragraph,
and saves as Activity_Logs_Full_Documentation_20260317_125220.docx
(overwrites in-place – original is backed up as _BACKUP.docx).

Requirements:  pip install python-docx matplotlib pillow
"""

import copy
import io
import os
import shutil

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ──────────────────────────────────────────────
C_NAVY   = "#0D1B2A"
C_BLUE   = "#1565C0"
C_LBLUE  = "#42A5F5"
C_TEAL   = "#00796B"
C_GREEN  = "#2E7D32"
C_ORANGE = "#E65100"
C_PURPLE = "#6A1B9A"
C_RED    = "#C62828"
C_GREY   = "#546E7A"
C_BG     = "#F5F7FA"
C_WHITE  = "#FFFFFF"

SRC = os.path.join(os.path.dirname(__file__),
                   "Activity_Logs_Full_Documentation_20260317_125220.docx")
BACKUP = SRC.replace(".docx", "_BACKUP.docx")
OUT    = SRC


# ═══════════════════════════════════════════════════════════════
# helpers
# ═══════════════════════════════════════════════════════════════
def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf


def replace_para_with_picture(doc, para_idx: int, stream, width=Inches(5.8), caption=""):
    """
    Replace doc.paragraphs[para_idx] (the Mermaid code block) with a real image.
    Strategy: use python-docx's own add_picture() on the REAL doc (so relationships
    are handled correctly), then move the resulting paragraph to the right position,
    then delete the original code paragraph.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    target_elem = doc.paragraphs[para_idx]._element

    # 1. Add image at end of doc using the proper API (correct rId wiring)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(stream, width=width)

    # 2. Detach the new image paragraph from the end and place it before target
    img_elem = img_para._element
    img_elem.getparent().remove(img_elem)
    target_elem.addprevious(img_elem)

    # 3. Caption paragraph (also added at end, then moved)
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        from docx.shared import RGBColor
        cap_run.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)
        cap_elem = cap_para._element
        cap_elem.getparent().remove(cap_elem)
        target_elem.addprevious(cap_elem)

    # 4. Remove original code paragraph
    target_elem.getparent().remove(target_elem)


# ═══════════════════════════════════════════════════════════════
# ── DIAGRAM RENDERERS ──────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, label, sub="", bg=C_BLUE, fg=C_WHITE, fs=10):
    rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                           linewidth=1.5, edgecolor=bg,
                           facecolor=bg + "22")
    ax.add_patch(rect)
    ax.text(x+w/2, y+h/2+(0.13 if sub else 0), label,
            ha="center", va="center", fontsize=fs,
            fontweight="bold", color=bg)
    if sub:
        ax.text(x+w/2, y+h/2-0.2, sub, ha="center", va="center",
                fontsize=7.5, color=C_GREY)


def _arr(ax, x1, y1, x2, y2, col=C_GREY, lbl="", style="-|>"):
    ax.annotate("", xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle=style, color=col, lw=1.4))
    if lbl:
        ax.text((x1+x2)/2+0.05, (y1+y2)/2+0.08, lbl, fontsize=7.5, color=col)


# ── 1. Role Hierarchy (graph TD, para 40) ────────────────────
def diag_role_hierarchy():
    fig, ax = plt.subplots(figsize=(10, 6), facecolor=C_BG)
    ax.set_xlim(0,10); ax.set_ylim(0,6); ax.axis("off")
    nodes = [
        ("global_admin\nFull system access", 5.0, 5.0, C_RED),
        ("admin\nTenant management", 2.5, 3.5, C_ORANGE),
        ("user\nStandard access", 5.0, 3.5, C_BLUE),
        ("viewer\nRead-only", 7.5, 3.5, C_TEAL),
        ("guest\nPublic / limited", 5.0, 1.8, C_GREY),
    ]
    for label, x, y, col in nodes:
        rect = FancyBboxPatch((x-1.1, y-0.4), 2.2, 0.8,
                               boxstyle="round,pad=0.1", linewidth=2,
                               edgecolor=col, facecolor=col)
        ax.add_patch(rect)
        ax.text(x, y, label, ha="center", va="center",
                fontsize=8.5, fontweight="bold", color=C_WHITE)
    edges = [(5.0,4.6, 2.5,3.9), (5.0,4.6, 5.0,3.9),
             (5.0,4.6, 7.5,3.9), (5.0,3.1, 5.0,2.2)]
    for x1,y1,x2,y2 in edges:
        _arr(ax, x1,y1, x2,y2, C_GREY)
    ax.text(5.0, 0.6, "Higher role inherits all capabilities of roles below it",
            ha="center", fontsize=8, color=C_GREY, style="italic")
    ax.set_title("Role Hierarchy Diagram", fontsize=12, fontweight="bold", color=C_NAVY)
    return fig_to_stream(fig)


# ── 2. ER Diagram (para 78) ────────────────────────────────────
def diag_er():
    fig, ax = plt.subplots(figsize=(13, 7), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis("off")

    def entity(x, y, title, fields, col, w=2.8):
        h = 0.4 + 0.32*len(fields)
        ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.08",
                                     linewidth=2, edgecolor=col, facecolor=C_WHITE))
        ax.add_patch(FancyBboxPatch((x,y+h-0.38), w, 0.38,
                                     boxstyle="square,pad=0",
                                     linewidth=0, facecolor=col))
        ax.text(x+w/2, y+h-0.19, title, ha="center", va="center",
                fontsize=9, fontweight="bold", color=C_WHITE)
        for i, (f, t) in enumerate(fields):
            fy = y+h-0.38-0.32*(i+1)+0.1
            ax.text(x+0.15, fy, f, fontsize=7.5, color=C_NAVY)
            ax.text(x+w-0.1, fy, t, ha="right", fontsize=7, color=C_GREY)
        return x+w/2, y+h/2

    cx1,cy1 = entity(0.3,3.5, "auth.users",
                     [("id PK","uuid"),("email","text"),("role","text")], C_NAVY)
    cx2,cy2 = entity(4.2,4.0, "auth_audit",
                     [("id PK","uuid"),("user_id FK","uuid"),("event","text"),
                      ("action_code","text"),("outcome_code","int"),
                      ("integrity_hash","text")], C_TEAL, 3.0)
    cx3,cy3 = entity(4.2,1.5, "audit_logs",
                     [("id PK","uuid"),("user_id FK","uuid"),("action","text"),
                      ("shap_values","jsonb"),("chain_hash","text"),
                      ("sequence_num","bigint")], C_GREEN, 3.0)
    cx4,cy4 = entity(8.8,4.2, "user_roles",
                     [("id PK","uuid"),("user_id FK","uuid"),
                      ("role_id FK","uuid")], C_PURPLE)
    cx5,cy5 = entity(8.8,2.2, "roles",
                     [("id PK","uuid"),("name","text"),
                      ("description","text")], C_ORANGE)
    cx6,cy6 = entity(8.8,0.2, "device_sync_logs",
                     [("id PK","uuid"),("device_id FK","uuid"),
                      ("synced_at","timestamp")], C_BLUE)

    for (x1,y1,x2,y2,lbl,col) in [
        (cx1+1.4,cy1, cx2,cy2,"1:N",C_TEAL),
        (cx1+1.4,cy1, cx3,cy3-0.3,"1:N",C_GREEN),
        (cx4,cy4, cx5,cy5+0.3,"N:1",C_PURPLE),
        (cx1+1.4,cy1, cx4,cy4,"1:N",C_NAVY),
    ]:
        _arr(ax, x1,y1, x2,y2, col, lbl)

    ax.set_title("Entity-Relationship Diagram – Activity Logs Tables",
                 fontsize=12, fontweight="bold", color=C_NAVY)
    return fig_to_stream(fig)


# ── 3. Schema Column Detail (classDiagram, para 81) ───────────
def diag_schema_detail():
    fig, ax = plt.subplots(figsize=(13,7), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis("off")

    def cls_box(x, y, name, attrs, col, w=5.5):
        rows = attrs
        h = 0.55 + 0.3*len(rows)
        ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.08",
                                     linewidth=2, edgecolor=col, facecolor=C_WHITE))
        ax.add_patch(FancyBboxPatch((x,y+h-0.45), w, 0.45,
                                     boxstyle="square,pad=0", linewidth=0, facecolor=col))
        ax.text(x+w/2, y+h-0.22, name, ha="center", va="center",
                fontsize=10, fontweight="bold", color=C_WHITE)
        for i, row in enumerate(rows):
            ry = y+h-0.45-0.3*(i+1)+0.08
            ax.text(x+0.15, ry, row, fontsize=7.5, color=C_NAVY,
                    fontfamily="monospace")

    cls_box(0.3, 3.5, "auth_audit", [
        "+UUID id  [PK]",
        "+UUID user_id  [FK]",
        "+text email",
        "+text role",
        "+text event",
        "+text action_code  // C/R/U/D/E",
        "+int  outcome_code  // 0/4/8/12",
        "+text resource_type",
        "+text resource_id",
        "+text integrity_hash  // SHA-256",
        "+timestamptz created_at",
    ], C_TEAL)

    cls_box(7.0, 0.4, "audit_logs  (Cryptographic Ledger)", [
        "+UUID id  [PK]",
        "+UUID user_id  [FK]",
        "+text action",
        "+text resource_type",
        "+jsonb details",
        "+jsonb previous_value  // PII-scrubbed",
        "+jsonb new_value  // PII-scrubbed",
        "+text model_id",
        "+jsonb prediction",
        "+jsonb shap_values",
        "+text reasoning",
        "+text integrity_hash  // SHA-256",
        "+bigint sequence_num  // IDENTITY",
        "+text chain_hash  // SHA-256 ledger",
        "+timestamptz created_at",
    ], C_GREEN)

    ax.set_title("Schema Diagram – Supabase Column-Level Detail",
                 fontsize=12, fontweight="bold", color=C_NAVY)
    return fig_to_stream(fig)


# ── 4. RLS Policy Diagram (flowchart TD, para 84) ─────────────
def diag_rls():
    fig, ax = plt.subplots(figsize=(11, 8), facecolor=C_BG)
    ax.set_xlim(0,11); ax.set_ylim(0,8); ax.axis("off")

    def diamond(x, y, w, h, label, col):
        dx, dy = w/2, h/2
        pts = [(x,y+dy),(x+dx,y+h),(x+w,y+dy),(x+dx,y)]
        patch = plt.Polygon(pts, closed=True, facecolor=col+"22",
                             edgecolor=col, linewidth=2)
        ax.add_patch(patch)
        ax.text(x+w/2, y+dy, label, ha="center", va="center",
                fontsize=8, color=col, fontweight="bold")

    def rect(x, y, w, h, label, col):
        ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.08",
                                     linewidth=2, edgecolor=col, facecolor=col))
        ax.text(x+w/2, y+h/2, label, ha="center", va="center",
                fontsize=8.5, fontweight="bold", color=C_WHITE)

    rect(4.0, 7.2, 3.0, 0.55, "SELECT on auth_audit", C_BLUE)
    diamond(3.5, 5.9, 4.0, 1.1, "Who is the\ncurrent user?", C_NAVY)

    rect(0.2, 4.5, 2.5, 0.55, "global_admin", C_RED)
    rect(3.0, 4.5, 2.5, 0.55, "admin", C_ORANGE)
    rect(5.7, 4.5, 2.5, 0.55, "user / viewer", C_BLUE)
    rect(8.4, 4.5, 2.5, 0.55, "guest", C_GREY)

    rect(0.2, 3.1, 2.5, 0.8, "ALL rows\n(no filter)", C_RED)
    rect(3.0, 3.1, 2.5, 0.8, "Non-admin rows\n(own tenant)", C_ORANGE)
    rect(5.7, 3.1, 2.5, 0.8, "Own rows only\n(user_id = auth.uid())", C_BLUE)
    rect(8.4, 3.1, 2.5, 0.8, "403 Forbidden\n(RLS blocks)", C_GREY)

    rect(3.5, 1.5, 4.0, 0.7, "Row returned to client\n(via PostgREST / RLS)", C_GREEN)

    for x1,y1,x2,y2 in [(5.5,7.2, 5.5,7.0),(5.5,5.9, 1.45,5.05),
                          (5.5,5.9, 4.25,5.05),(5.5,5.9, 6.95,5.05),
                          (5.5,5.9, 9.65,5.05)]:
        _arr(ax, x1,y1, x2,y2, C_GREY)
    for x in [1.45,4.25,6.95]:
        _arr(ax, x,3.1, 5.5,2.2, C_GREEN)
    _arr(ax, 9.65,3.1, 8.9,2.2, C_RED, "denied")

    ax.set_title("Row-Level Security (RLS) Policy – auth_audit SELECT",
                 fontsize=12, fontweight="bold", color=C_NAVY)
    return fig_to_stream(fig)


# ── 5. Data Flow (flowchart LR, para 89) ──────────────────────
def diag_data_flow():
    fig, ax = plt.subplots(figsize=(13,6), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis("off")
    steps = [
        ("Browser\nAuth.tsx / Activity.tsx", C_TEAL,   1.0),
        ("Supabase\nAuth Event", C_NAVY,                3.0),
        ("FastAPI\nRoute Handler", C_BLUE,               5.0),
        ("insert_audit_log()\nPII Scrub + SHA-256", C_ORANGE, 7.2),
        ("PostgREST\nINSERT", C_GREEN,                   9.4),
        ("DB Trigger\nchain_hash", C_PURPLE,            11.4),
    ]
    y = 3.2
    for lbl, col, x in steps:
        ax.add_patch(FancyBboxPatch((x-0.85, y-0.5), 1.7, 1.0,
                                     boxstyle="round,pad=0.1", linewidth=2,
                                     edgecolor=col, facecolor=col+"22"))
        ax.text(x, y, lbl, ha="center", va="center",
                fontsize=8, color=col, fontweight="bold")
    for i in range(len(steps)-1):
        _arr(ax, steps[i][2]+0.85, y, steps[i+1][2]-0.85, y, C_GREY)
    # secondary: audit_logs table at bottom
    ax.add_patch(FancyBboxPatch((8.5,0.7),3.5,0.9, boxstyle="round,pad=0.1",
                                 linewidth=2, edgecolor=C_GREEN, facecolor=C_GREEN+"22"))
    ax.text(10.25,1.15,"audit_logs\n(Immutable Ledger)", ha="center",
            fontsize=9, color=C_GREEN, fontweight="bold")
    _arr(ax,10.25,2.7, 10.25,1.6, C_GREEN,"stored")
    ax.set_title("Data Flow Diagram – How Audit Logs Are Created",
                 fontsize=12, fontweight="bold", color=C_NAVY)
    return fig_to_stream(fig)


# ── 6. Sequence: Login Audit Event (para 92) ──────────────────
def diag_seq_login():
    fig, ax = plt.subplots(figsize=(13,7), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis("off")
    actors = [("User",1.2,C_TEAL),("Auth.tsx",3.2,C_BLUE),
              ("Supabase\nAuth",5.8,C_NAVY),("insert_\naudit_log",8.4,C_ORANGE),
              ("audit_logs\nDB",11.0,C_GREEN)]
    TOP=6.5; BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.7,TOP-0.05),1.4,0.5,
                                     boxstyle="round,pad=0.07",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,
                fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[
        (6.0,1.2,3.2,"login(email,password)",C_TEAL),
        (5.5,3.2,5.8,"signInWithPassword()",C_BLUE),
        (5.0,5.8,3.2,"session token",C_NAVY,True),
        (4.5,3.2,8.4,"insert_audit_log('login',user_id,...)",C_BLUE),
        (4.0,8.4,11.0,"INSERT INTO auth_audit",C_ORANGE),
        (3.5,11.0,8.4,"row written",C_GREEN,True),
        (3.0,3.2,1.2,"redirect to dashboard",C_BLUE,True),
    ]
    for item in msgs:
        y,x1,x2,lbl,col = item[0],item[1],item[2],item[3],item[4]
        ret = item[5] if len(item)>5 else False
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram – Login Audit Event",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 7. Sequence: Admin Approves Pod (para 95) ─────────────────
def diag_seq_pod():
    fig, ax = plt.subplots(figsize=(13,7), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis("off")
    actors=[("Admin",1.2,C_ORANGE),("PodActivation\nRequests.tsx",3.2,C_BLUE),
            ("FastAPI\n/pod_activation",5.8,C_NAVY),
            ("insert_\naudit_log",8.4,C_ORANGE),("audit_logs",11.0,C_GREEN)]
    TOP=6.5; BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.7,TOP-0.05),1.4,0.5,
                                     boxstyle="round,pad=0.07",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,
                fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[
        (6.0,1.2,3.2,"click Approve Pod",C_ORANGE),
        (5.5,3.2,5.8,"PATCH /pod_activation/{id}",C_BLUE),
        (5.0,5.8,5.8,"verify_admin(JWT)",C_NAVY),
        (4.5,5.8,5.8,"update pod status=approved",C_NAVY),
        (4.0,5.8,8.4,"insert_audit_log('pod_approve',...)",C_NAVY),
        (3.5,8.4,11.0,"INSERT INTO audit_logs",C_ORANGE),
        (3.0,11.0,8.4,"chain_hash written",C_GREEN,True),
        (2.5,5.8,3.2,"200 OK",C_NAVY,True),
        (2.0,3.2,1.2,"show success toast",C_BLUE,True),
    ]
    for item in msgs:
        y,x1,x2,lbl,col = item[0],item[1],item[2],item[3],item[4]
        ret = item[5] if len(item)>5 else False
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram – Admin Approves Pod",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 8. Sequence: Viewing Activity Logs (para 98) ──────────────
def diag_seq_view():
    fig, ax = plt.subplots(figsize=(13,7), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis("off")
    actors=[("User/Admin/\nGlobal Admin",1.3,C_BLUE),("Activity.tsx",3.5,C_TEAL),
            ("Supabase\nauth_audit",6.0,C_NAVY),("FastAPI\n/api/activity",8.5,C_ORANGE),
            ("audit_logs",11.2,C_GREEN)]
    TOP=6.5; BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.8,TOP-0.05),1.6,0.5,
                                     boxstyle="round,pad=0.07",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=7.5,
                fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[
        (6.0,1.3,3.5,"navigate to /activity",C_BLUE),
        (5.5,3.5,6.0,"SELECT * FROM auth_audit (RLS)",C_TEAL),
        (5.0,6.0,3.5,"auth events[]",C_NAVY,True),
        (4.5,3.5,8.5,"GET /api/activity/system?page=0",C_TEAL),
        (4.0,8.5,11.2,"SELECT * FROM audit_logs (RBAC)",C_ORANGE),
        (3.5,11.2,8.5,"system logs[]",C_GREEN,True),
        (3.0,8.5,3.5,"200 OK {logs, has_more}",C_ORANGE,True),
        (2.5,3.5,1.3,"render two-tab Activity page",C_TEAL,True),
    ]
    for item in msgs:
        y,x1,x2,lbl,col = item[0],item[1],item[2],item[3],item[4]
        ret = item[5] if len(item)>5 else False
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Sequence Diagram – Viewing Activity Logs",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 9. Component Architecture (graph TB, para 101) ────────────
def diag_component_arch():
    fig, ax = plt.subplots(figsize=(13,8), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,8); ax.axis("off")

    # App Routes box
    ax.add_patch(FancyBboxPatch((0.3,6.2),12.4,1.4,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_BLUE,facecolor=C_BLUE+"11"))
    ax.text(0.7,7.45,"App Routes (App.tsx)",fontsize=8,color=C_BLUE,style="italic")
    routes=[("/activity",2.0),("/activity (admin)",4.5),
            ("/chat",6.8),("/admin",8.9),("/workflow-ai",11.0)]
    for r,x in routes:
        ax.add_patch(FancyBboxPatch((x-0.9,6.35),1.8,0.65,
                                     boxstyle="round,pad=0.07",linewidth=1.5,
                                     edgecolor=C_BLUE,facecolor=C_BLUE))
        ax.text(x,6.67,r,ha="center",va="center",
                fontsize=7.5,fontweight="bold",color=C_WHITE)

    # Component layer
    comps=[("Activity.tsx\n(AuthEventsTab+\nSystemEventsTab)",2.0,4.5,C_TEAL,1.3),
           ("AdminActivity.tsx",4.5,5.0,C_ORANGE,1.1),
           ("Chat.tsx",6.8,5.2,C_BLUE,1.0),
           ("WorkflowAI.tsx",10.0,5.0,C_PURPLE,1.2)]
    for (nm,x,y,col,h) in comps:
        ax.add_patch(FancyBboxPatch((x-0.9,y-0.4),1.8,h,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col+"22"))
        ax.text(x,y+h/2-0.4,nm,ha="center",va="center",
                fontsize=7.5,color=col,fontweight="bold")

    # API layer
    apis=[("Supabase RLS\nauth_audit",2.0,2.8,C_NAVY),
          ("FastAPI\n/api/activity/system",5.0,2.8,C_ORANGE),
          ("FastAPI\n/api/workflow-ai",9.5,2.8,C_PURPLE)]
    for nm,x,y,col in apis:
        ax.add_patch(FancyBboxPatch((x-1.1,y-0.4),2.2,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col+"22"))
        ax.text(x,y+0.05,nm,ha="center",va="center",
                fontsize=8,color=col,fontweight="bold")

    # DB
    for nm,x,col in [("auth_audit\ntable",2.0,C_TEAL),
                      ("audit_logs\ntable",5.0,C_GREEN)]:
        ax.add_patch(FancyBboxPatch((x-1.0,0.5),2.0,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,0.95,nm,ha="center",va="center",
                fontsize=8.5,fontweight="bold",color=C_WHITE)

    # arrows
    for (x1,y1,x2,y2,col) in [
        (2.0,6.35,2.0,5.5,C_TEAL),(4.5,6.35,4.5,5.5,C_ORANGE),
        (2.0,4.5,2.0,3.25,C_TEAL),(2.0,4.5,4.5,3.25,C_TEAL),
        (2.0,2.4,2.0,1.4,C_NAVY),(4.5,2.4,4.5,1.4,C_GREEN),
    ]:
        _arr(ax,x1,y1,x2,y2,col)

    ax.set_title("Component Architecture Diagram",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 10. Audit Integrity Hash Flow – auth_audit (para 105) ─────
def diag_hash_auth():
    fig, ax = plt.subplots(figsize=(11,5), facecolor=C_BG)
    ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis("off")
    inputs=[("user_id",0.7,3.5),("role",0.7,2.8),("event",0.7,2.1),
            ("action_code",0.7,1.4),("outcome_code",0.7,0.7),
            ("resource_type",0.7,0.0)]
    for lbl,x,y in inputs:
        ax.add_patch(FancyBboxPatch((x-0.6,y-0.22),1.7,0.44,
                                     boxstyle="round,pad=0.05",linewidth=1.5,
                                     edgecolor=C_BLUE,facecolor=C_BLUE+"22"))
        ax.text(x+0.25,y,lbl,ha="center",va="center",fontsize=8.5,color=C_BLUE)
        _arr(ax,x+0.85,y, 3.5,2.0,C_GREY)

    ax.add_patch(FancyBboxPatch((3.3,1.3),2.0,1.4,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_ORANGE,facecolor=C_ORANGE+"22"))
    ax.text(4.3,2.0,"SHA-256\nhash()",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_ORANGE)

    _arr(ax,5.3,2.0, 7.0,2.0,C_ORANGE)
    ax.add_patch(FancyBboxPatch((7.0,1.5),3.5,1.0,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_GREEN,facecolor=C_GREEN+"22"))
    ax.text(8.75,2.0,"integrity_hash\n(stored in DB)",ha="center",va="center",
            fontsize=9,fontweight="bold",color=C_GREEN)

    ax.text(5.5,4.4,"Note: email is EXCLUDED from hash (PII protection)",
            ha="center",fontsize=8.5,color=C_RED,
            bbox=dict(boxstyle="round,pad=0.3",facecolor="#FFEBEE",edgecolor=C_RED))

    ax.set_title("Audit Integrity Hash Flow – auth_audit (Frontend SHA-256)",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 11. Audit Integrity Hash Flow – audit_logs (para 107) ─────
def diag_hash_audit():
    fig, ax = plt.subplots(figsize=(13,6), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis("off")

    app_fields=["user_id","action","resource_type","resource_id",
                "ip_address","model_id","shap_values","prediction","reasoning"]
    for i,f in enumerate(app_fields):
        y=5.3-i*0.52
        ax.add_patch(FancyBboxPatch((0.2,y-0.2),2.3,0.42,
                                     boxstyle="round,pad=0.04",linewidth=1,
                                     edgecolor=C_BLUE,facecolor=C_BLUE+"18"))
        ax.text(1.35,y,f,ha="center",va="center",fontsize=7.8,color=C_BLUE)
        _arr(ax,2.5,y, 4.0,3.0,C_GREY)

    ax.add_patch(FancyBboxPatch((4.0,2.2),2.2,1.6,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_ORANGE,facecolor=C_ORANGE+"22"))
    ax.text(5.1,3.0,"SHA-256\nintegrity\nhash",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_ORANGE)

    _arr(ax,6.2,3.0, 7.6,3.0,C_ORANGE,"integrity_hash")

    ax.add_patch(FancyBboxPatch((7.6,2.3),2.5,1.4,
                                 boxstyle="round,pad=0.08",linewidth=2,
                                 edgecolor=C_PURPLE,facecolor=C_PURPLE+"22"))
    ax.text(8.85,3.0,"DB Trigger\ncompute_\nchain_hash()",ha="center",va="center",
            fontsize=9,fontweight="bold",color=C_PURPLE)

    _arr(ax,10.1,3.0, 11.4,3.0,C_PURPLE,"chain_hash")
    ax.add_patch(FancyBboxPatch((11.4,2.4),1.4,1.2,
                                 boxstyle="round,pad=0.08",linewidth=2,
                                 edgecolor=C_GREEN,facecolor=C_GREEN+"22"))
    ax.text(12.1,3.0,"stored\nin DB",ha="center",va="center",
            fontsize=9,fontweight="bold",color=C_GREEN)

    ax.set_title("Audit Integrity Hash Flow – audit_logs (Backend SHA-256 + Ledger)",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 12. Access Control Matrix / quadrant (para 126) ───────────
def diag_access_matrix():
    fig, ax = plt.subplots(figsize=(10,6), facecolor=C_BG)
    ax.set_xlim(-1,11); ax.set_ylim(-1,7); ax.axis("off")
    ax.axvline(5,color=C_GREY,lw=1,ls="--",alpha=0.4)
    ax.axhline(3,color=C_GREY,lw=1,ls="--",alpha=0.4)
    ax.text(2.5,6.3,"Own Records Only",ha="center",fontsize=9,color=C_GREY)
    ax.text(7.5,6.3,"All Records",ha="center",fontsize=9,color=C_GREY)
    ax.text(-0.7,4.5,"High\nPrivilege",ha="center",fontsize=9,color=C_GREY,rotation=90)
    ax.text(-0.7,1.5,"Low\nPrivilege",ha="center",fontsize=9,color=C_GREY,rotation=90)

    roles=[(1.5,5.2,"global_admin\nAll rows, all tenants",C_RED),
           (3.0,4.0,"admin\nTenant rows",C_ORANGE),
           (2.0,1.8,"user\nOwn rows",C_BLUE),
           (3.5,1.5,"viewer\nOwn rows\n(read-only)",C_TEAL),
           (1.2,0.5,"guest\n403 Forbidden",C_GREY)]
    for x,y,lbl,col in roles:
        ax.add_patch(plt.Circle((x,y),0.55,color=col+"44",ec=col,lw=2))
        ax.text(x,y,lbl,ha="center",va="center",fontsize=7.5,
                fontweight="bold",color=col)

    ax.set_title("Access Control Matrix – Audit Row Visibility by Role",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 13. Navigation Route Guard (flowchart TD, para 136) ───────
def diag_route_guard():
    fig, ax = plt.subplots(figsize=(10,8), facecolor=C_BG)
    ax.set_xlim(0,10); ax.set_ylim(0,8); ax.axis("off")

    def diam(x,y,w,h,lbl,col):
        pts=[(x+w/2,y+h),(x+w,y+h/2),(x+w/2,y),(x,y+h/2)]
        ax.add_patch(plt.Polygon(pts,closed=True,
                                  facecolor=col+"22",edgecolor=col,lw=2))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",
                fontsize=8,color=col,fontweight="bold")

    def rct(x,y,w,h,lbl,col):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",
                fontsize=8.5,fontweight="bold",color=C_WHITE)

    rct(3.5,7.2,3.0,0.55,"User navigates to URL",C_BLUE)
    diam(3.3,5.9,3.4,1.1,"Is user\nauthenticated?",C_NAVY)
    rct(0.5,4.5,3.0,0.55,"Redirect to /login",C_RED)
    diam(3.3,4.1,3.4,1.1,"Has required\nrole?",C_NAVY)
    rct(7.0,4.5,2.5,0.55,"403 Forbidden",C_ORANGE)
    rct(3.5,2.8,3.0,0.7,"Render protected\npage (Activity, Admin…)",C_GREEN)
    rct(3.5,1.6,3.0,0.6,"Log route access\nin audit trail",C_TEAL)

    _arr(ax,5.0,7.2,5.0,7.0,C_BLUE)
    _arr(ax,5.0,5.9,2.0,5.05,C_RED,"No")
    _arr(ax,5.0,5.9,5.0,5.2,C_NAVY,"Yes")
    _arr(ax,5.0,4.1,8.25,4.5,C_ORANGE,"No")
    _arr(ax,5.0,4.1,5.0,3.5,C_GREEN,"Yes")
    _arr(ax,5.0,2.8,5.0,2.2,C_TEAL)

    ax.set_title("Navigation & Route Guard Diagram",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 14. Backend Logging Architecture (flowchart LR, para 139) ─
def diag_backend_logging():
    fig, ax = plt.subplots(figsize=(13,6), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis("off")
    # FastAPI box
    ax.add_patch(FancyBboxPatch((0.2,1.2),4.0,4.2,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_BLUE,facecolor=C_BLUE+"09"))
    ax.text(0.5,5.2,"FastAPI Application",fontsize=8,color=C_BLUE,style="italic")
    components=[("factory.py\n(app init)",1.0,4.0,C_NAVY),
                ("RequestLogging\nMiddleware",1.0,2.8,C_TEAL),
                ("Routes\n(activity.py etc.)",2.8,4.0,C_BLUE),
                ("insert_audit_log()\nsupabase.py",2.8,2.8,C_ORANGE)]
    for nm,x,y,col in components:
        ax.add_patch(FancyBboxPatch((x-0.6,y-0.4),2.0,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col+"22"))
        ax.text(x+0.4,y,nm,ha="center",va="center",fontsize=8,
                color=col,fontweight="bold")

    # Logger module
    ax.add_patch(FancyBboxPatch((5.2,2.0),2.4,2.4,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_PURPLE,facecolor=C_PURPLE+"09"))
    ax.text(5.4,4.25,"logger/__init__.py",fontsize=8,color=C_PURPLE,style="italic")
    ax.text(6.4,3.5,"Pass-1\nField Scrub",ha="center",fontsize=8,color=C_PURPLE)
    ax.text(6.4,2.5,"Pass-2\nPresidio NLP",ha="center",fontsize=8,color=C_PURPLE)

    # Outputs
    for nm,x,y,col in [("structlog\nJSON stdout",9.0,4.2,C_TEAL),
                        ("Rotating\napp.log",9.0,3.1,C_NAVY),
                        ("audit_logs\n(PostgreSQL)",9.0,2.0,C_GREEN)]:
        ax.add_patch(FancyBboxPatch((x-0.7,y-0.4),2.4,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x+0.5,y,nm,ha="center",va="center",fontsize=8,
                fontweight="bold",color=C_WHITE)

    # Arrows
    _arr(ax,4.2,3.5, 5.2,3.5,C_PURPLE,"PII scrub")
    _arr(ax,7.6,3.8, 8.3,4.45,C_TEAL)
    _arr(ax,7.6,3.2, 8.3,3.35,C_NAVY)
    _arr(ax,7.6,2.5, 8.3,2.25,C_GREEN)

    ax.set_title("Backend Logging Architecture",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 15. Full System Interaction (flowchart TB, para 144) ───────
def diag_full_system():
    fig, ax = plt.subplots(figsize=(14,9), facecolor=C_BG)
    ax.set_xlim(0,14); ax.set_ylim(0,9); ax.axis("off")
    users=[("global_admin",1.0,8.2,C_RED),("admin",3.5,8.2,C_ORANGE),
           ("user",6.0,8.2,C_BLUE),("viewer",8.5,8.2,C_TEAL),
           ("guest",11.0,8.2,C_GREY)]
    for nm,x,y,col in users:
        ax.add_patch(plt.Circle((x,y),0.45,color=col,ec=col,lw=2))
        ax.text(x,y,nm[:6],ha="center",va="center",fontsize=7,
                fontweight="bold",color=C_WHITE)
    for nm,x,y,col in users:
        _arr(ax,x,y-0.45, 7.0,7.0,col)

    ax.add_patch(FancyBboxPatch((5.0,6.4),4.0,0.7,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_NAVY,facecolor=C_NAVY))
    ax.text(7.0,6.75,"Next.js 14 Frontend",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_WHITE)

    _arr(ax,7.0,6.4, 7.0,5.8,C_NAVY,"HTTPS/JWT")
    ax.add_patch(FancyBboxPatch((5.0,5.1),4.0,0.7,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_BLUE,facecolor=C_BLUE))
    ax.text(7.0,5.45,"FastAPI + Uvicorn",ha="center",va="center",
            fontsize=10,fontweight="bold",color=C_WHITE)

    for nm,x,y,col in [("verify_user\nverify_admin",2.5,4.0,C_TEAL),
                        ("insert_audit_log\nPII scrub",7.0,4.0,C_ORANGE),
                        ("SHAP reasoning\ngenerator",11.5,4.0,C_PURPLE)]:
        ax.add_patch(FancyBboxPatch((x-1.2,y-0.4),2.4,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col+"22"))
        ax.text(x,y,nm,ha="center",va="center",fontsize=8,
                color=col,fontweight="bold")
    _arr(ax,7.0,5.1, 2.5,4.5,C_TEAL)
    _arr(ax,7.0,5.1, 7.0,4.5,C_ORANGE)
    _arr(ax,7.0,5.1, 11.5,4.5,C_PURPLE)

    for nm,x,y,col in [("auth_audit\ntable",3.0,2.3,C_TEAL),
                        ("audit_logs\ntable",7.0,2.3,C_GREEN),
                        ("auth.users\ntable",11.0,2.3,C_NAVY)]:
        ax.add_patch(FancyBboxPatch((x-1.1,y-0.45),2.2,0.9,
                                     boxstyle="round,pad=0.08",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,y,nm,ha="center",va="center",fontsize=9,
                fontweight="bold",color=C_WHITE)

    _arr(ax,2.5,3.6, 3.0,2.75,C_TEAL)
    _arr(ax,7.0,3.6, 7.0,2.75,C_GREEN)

    ax.add_patch(FancyBboxPatch((5.0,0.5),4.0,0.7,
                                 boxstyle="round,pad=0.1",linewidth=2,
                                 edgecolor=C_ORANGE,facecolor=C_ORANGE))
    ax.text(7.0,0.85,"DB Triggers (immutability + chain_hash)",
            ha="center",va="center",fontsize=9,fontweight="bold",color=C_WHITE)
    _arr(ax,7.0,2.3-0.45, 7.0,1.2,C_ORANGE)

    ax.set_title("Full System Interaction Diagram",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 16. Cryptographic Ledger Trigger Sequence (para 181) ───────
def diag_ledger_trigger():
    fig, ax = plt.subplots(figsize=(12,6), facecolor=C_BG)
    ax.set_xlim(0,12); ax.set_ylim(0,6); ax.axis("off")
    actors=[("Application\ninsert_audit_log",1.5,C_BLUE),
            ("PostgreSQL\nBEFORE INSERT trigger",4.5,C_NAVY),
            ("pg_advisory\nxact_lock",7.5,C_ORANGE),
            ("audit_logs\ntable",10.5,C_GREEN)]
    TOP=5.5; BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.8,TOP-0.05),1.6,0.5,
                                     boxstyle="round,pad=0.07",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=8,
                fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[
        (5.0,1.5,4.5,"INSERT row (integrity_hash set)",C_BLUE),
        (4.5,4.5,7.5,"LOCK (serialize inserts)",C_NAVY),
        (4.0,7.5,4.5,"lock acquired",C_ORANGE,True),
        (3.5,4.5,4.5,"fetch prev chain_hash",C_NAVY),
        (3.0,4.5,4.5,"SHA-256(prev||integrity_hash) -> chain_hash",C_NAVY),
        (2.5,4.5,10.5,"write row + chain_hash",C_NAVY),
        (2.0,10.5,4.5,"row written",C_GREEN,True),
        (1.5,4.5,7.5,"UNLOCK",C_NAVY),
    ]
    for item in msgs:
        y,x1,x2,lbl,col = item[0],item[1],item[2],item[3],item[4]
        ret = item[5] if len(item)>5 else False
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("Cryptographic Ledger Trigger – compute_audit_chain_hash()",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 17. SHAP Flow (flowchart LR, para 213) ────────────────────
def diag_shap_flow():
    fig, ax = plt.subplots(figsize=(12,4), facecolor=C_BG)
    ax.set_xlim(0,12); ax.set_ylim(0,4); ax.axis("off")
    steps=[
        ("Route handler\ncalls log_ai_\ndecision_audit()",C_BLUE,1.2),
        ("reasoning\nprovided?",C_NAVY,3.4),
        ("generate_shap_\nreasoning(shap_values,\nprediction, model_id)",C_PURPLE,5.8),
        ("insert_audit_log(\n...reasoning,\nshap_values)",C_ORANGE,8.4),
        ("audit_logs\nrow written",C_GREEN,11.0),
    ]
    y=2.2
    for lbl,col,x in steps:
        if "?" in lbl:
            pts=[(x,y+0.55),(x+0.85,y+1.1),(x+1.7,y+0.55),(x+0.85,y)]
            ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",
                                      edgecolor=col,lw=2))
            ax.text(x+0.85,y+0.55,lbl,ha="center",va="center",
                    fontsize=8,color=col,fontweight="bold")
        else:
            ax.add_patch(FancyBboxPatch((x-0.85,y-0.55),1.7,1.1,
                                         boxstyle="round,pad=0.1",linewidth=2,
                                         edgecolor=col,facecolor=col+"22"))
            ax.text(x,y,lbl,ha="center",va="center",fontsize=8,
                    color=col,fontweight="bold")
    for i in range(len(steps)-1):
        if "?" in steps[i][0]:
            _arr(ax,steps[i][2]+1.7,y, steps[i+1][2]-0.85,y,C_GREY,"No")
        else:
            _arr(ax,steps[i][2]+0.85,y, steps[i+1][2]-0.85,y,C_GREY)
    # Yes branch from diamond
    _arr(ax,3.4+0.85,y+1.1, steps[3][2]-0.85,y+0.6,C_GREEN,"Yes (skip)")
    ax.set_title("SHAP AI Decision Audit – log_ai_decision_audit() Flow",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 18. SHAP Data Flow Sequence (para 223) ────────────────────
def diag_shap_seq():
    fig, ax = plt.subplots(figsize=(13,6), facecolor=C_BG)
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis("off")
    actors=[("AI Model\n/ Route",1.5,C_PURPLE),
            ("logger/\n__init__.py",4.0,C_NAVY),
            ("generate_shap_\nreasoning()",6.5,C_TEAL),
            ("insert_\naudit_log()",9.0,C_ORANGE),
            ("audit_logs\nDB",11.5,C_GREEN)]
    TOP=5.5; BOT=0.3
    for nm,x,col in actors:
        ax.add_patch(FancyBboxPatch((x-0.75,TOP-0.05),1.5,0.5,
                                     boxstyle="round,pad=0.07",linewidth=2,
                                     edgecolor=col,facecolor=col))
        ax.text(x,TOP+0.2,nm,ha="center",va="center",fontsize=7.5,
                fontweight="bold",color=C_WHITE)
        ax.plot([x,x],[TOP-0.05,BOT],color=col,lw=1,ls="--",alpha=0.5)
    msgs=[
        (5.0,1.5,4.0,"log_ai_decision_audit(shap_values, prediction)",C_PURPLE),
        (4.4,4.0,6.5,"generate_shap_reasoning(shap_values, prediction)",C_NAVY),
        (3.8,6.5,4.0,"reasoning text",C_TEAL,True),
        (3.2,4.0,9.0,"insert_audit_log(..., shap_values, reasoning)",C_NAVY),
        (2.6,9.0,11.5,"INSERT INTO audit_logs",C_ORANGE),
        (2.0,11.5,9.0,"row + chain_hash",C_GREEN,True),
    ]
    for item in msgs:
        y,x1,x2,lbl,col=item[0],item[1],item[2],item[3],item[4]
        ret=item[5] if len(item)>5 else False
        ax.annotate("",xy=(x2,y),xytext=(x1,y),
                    arrowprops=dict(arrowstyle="<-" if ret else "-|>",
                                    color=col,lw=1.3,
                                    linestyle="dashed" if ret else "solid"))
        ax.text((x1+x2)/2,y+0.12,lbl,ha="center",fontsize=7.5,color=col)
    ax.set_title("SHAP Data Flow Sequence – AI Decision to Audit Ledger",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 19. Audit Ledger Verification Flow (flowchart TD, para 242)
def diag_verify():
    fig, ax = plt.subplots(figsize=(10,8), facecolor=C_BG)
    ax.set_xlim(0,10); ax.set_ylim(0,8); ax.axis("off")

    def diam(x,y,w,h,lbl,col):
        pts=[(x+w/2,y+h),(x+w,y+h/2),(x+w/2,y),(x,y+h/2)]
        ax.add_patch(plt.Polygon(pts,closed=True,facecolor=col+"22",
                                  edgecolor=col,lw=2))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",
                fontsize=8,color=col,fontweight="bold")

    def rct(x,y,w,h,lbl,col):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.08",
                                     linewidth=2,edgecolor=col,facecolor=col))
        ax.text(x+w/2,y+h/2,lbl,ha="center",va="center",
                fontsize=8.5,fontweight="bold",color=C_WHITE)

    rct(3.0,7.2,4.0,0.6,"Compliance Audit Triggered",C_BLUE)
    rct(3.0,6.2,4.0,0.6,"Call verify_audit_ledger()",C_NAVY)
    rct(3.0,5.2,4.0,0.7,"Iterate rows ordered\nby sequence_num",C_TEAL)
    diam(2.8,3.8,4.4,1.2,"chain_hash = NULL?\n(pre-migration row)",C_ORANGE)
    rct(0.3,3.0,2.5,0.6,"Mark is_valid=NULL\n(sentinel)",C_ORANGE)
    rct(3.0,2.6,4.0,0.7,"Recompute chain:\nSHA-256(prev||integrity)",C_PURPLE)
    diam(2.8,1.3,4.4,1.1,"stored_chain\n== computed?",C_NAVY)
    rct(0.3,0.5,2.5,0.6,"TAMPERED\nis_valid=false",C_RED)
    rct(7.2,0.5,2.5,0.6,"VALID\nis_valid=true",C_GREEN)

    for (x1,y1,x2,y2,col,lbl) in [
        (5.0,7.2,5.0,6.8,C_BLUE,""),(5.0,6.2,5.0,5.9,C_NAVY,""),
        (5.0,5.2,5.0,5.0,C_TEAL,""),(5.0,3.8,5.0,3.3,C_ORANGE,"No"),
        (2.8,4.4,1.55,3.6,C_ORANGE,"Yes"),(5.0,2.6,5.0,2.4,C_PURPLE,""),
        (5.0,1.3,5.0,1.1,C_NAVY,""),(2.8,1.85,1.55,1.1,C_RED,"No"),
        (7.2,1.85,8.45,1.1,C_GREEN,"Yes"),
    ]:
        _arr(ax,x1,y1,x2,y2,col,lbl)

    ax.set_title("Audit Ledger Verification Flow – verify_audit_ledger()",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ── 20. Migration Timeline (timeline, para 249) ───────────────
def diag_migration_timeline():
    fig, ax = plt.subplots(figsize=(14,5), facecolor=C_BG)
    ax.set_xlim(0,14); ax.set_ylim(0,5); ax.axis("off")
    events=[
        ("2026-03-13",  "20260313091500",
         "auth_audit table\n+ RLS policies\n+ integrity_hash",C_TEAL),
        ("2026-03-16\n(000000)", "FK Constraint",
         "auth_audit.user_id\nFK -> auth.users",C_BLUE),
        ("2026-03-16\n(000001)", "Immutability",
         "prevent_audit_\nmodification() triggers",C_RED),
        ("2026-03-16\n(000002)", "Change Tracking",
         "previous_value +\nnew_value columns",C_ORANGE),
        ("2026-03-17\n(000002)", "Ledger + SHAP",
         "chain_hash + sequence_num\nshap_values + model_id\nprediction + reasoning",C_PURPLE),
        ("2026-03-17\n(000003)", "Ledger Fix",
         "NULL chain_hash\nfor pre-migration\nrows handled",C_GREEN),
    ]
    # Draw timeline line
    ax.plot([0.8,13.2],[2.5,2.5],color=C_GREY,lw=2)
    xs=np.linspace(1.2,12.8,len(events))
    for i,(dt,title,desc,col) in enumerate(events):
        x=xs[i]
        ax.plot([x,x],[2.5,2.5],color=col,lw=2)
        ax.add_patch(plt.Circle((x,2.5),0.2,color=col,zorder=5))
        # Alternate above/below
        if i%2==0:
            ax.text(x,3.3,title,ha="center",fontsize=8.5,fontweight="bold",color=col)
            ax.text(x,4.3,desc,ha="center",fontsize=7.5,color=C_NAVY)
            ax.plot([x,x],[2.7,3.1],color=col,lw=1,ls="--")
        else:
            ax.text(x,1.7,title,ha="center",fontsize=8.5,fontweight="bold",color=col)
            ax.text(x,0.5,desc,ha="center",fontsize=7.5,color=C_NAVY)
            ax.plot([x,x],[1.9,2.3],color=col,lw=1,ls="--")
        ax.text(x,2.5+(0.4 if i%2==0 else -0.65),dt,
                ha="center",fontsize=7,color=C_GREY)

    ax.set_title("Migration History & Schema Evolution Timeline",
                 fontsize=12,fontweight="bold",color=C_NAVY)
    return fig_to_stream(fig)


# ═══════════════════════════════════════════════════════════════
# INJECTION MAP
# Maps paragraph index  ->  (diagram_function, caption)
# ═══════════════════════════════════════════════════════════════
INJECTION_MAP = {
    40:  (diag_role_hierarchy,    "Figure: Role Definitions & Hierarchy"),
    78:  (diag_er,                "Figure: ER Diagram – Entity-Relationship"),
    81:  (diag_schema_detail,     "Figure: Schema Diagram – Supabase Column-Level Detail"),
    84:  (diag_rls,               "Figure: Row-Level Security (RLS) Policy Diagram"),
    89:  (diag_data_flow,         "Figure: Data Flow Diagram – How Logs Are Created"),
    92:  (diag_seq_login,         "Figure: Sequence Diagram – Login Audit Event"),
    95:  (diag_seq_pod,           "Figure: Sequence Diagram – Admin Approves Pod"),
    98:  (diag_seq_view,          "Figure: Sequence Diagram – Viewing Activity Logs"),
    101: (diag_component_arch,    "Figure: Component Architecture Diagram"),
    105: (diag_hash_auth,         "Figure: auth_audit Integrity Hash Flow"),
    107: (diag_hash_audit,        "Figure: audit_logs Integrity Hash + Chain Flow"),
    126: (diag_access_matrix,     "Figure: Access Control Matrix – Audit Row Visibility"),
    136: (diag_route_guard,       "Figure: Navigation & Route Guard Diagram"),
    139: (diag_backend_logging,   "Figure: Backend Logging Architecture"),
    144: (diag_full_system,       "Figure: Full System Interaction Diagram"),
    181: (diag_ledger_trigger,    "Figure: Cryptographic Ledger Trigger Sequence"),
    213: (diag_shap_flow,         "Figure: SHAP AI Decision Audit Flow"),
    223: (diag_shap_seq,          "Figure: SHAP Data Flow Sequence"),
    242: (diag_verify,            "Figure: Audit Ledger Verification Flow"),
    249: (diag_migration_timeline,"Figure: Migration History & Schema Evolution"),
}


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    # Always read from the original backup so we start clean
    if not os.path.exists(BACKUP):
        shutil.copy2(SRC, BACKUP)
        print(f"Backup created: {BACKUP}")
    else:
        print(f"Reading from backup: {BACKUP}")

    doc = Document(BACKUP)  # always start from the untouched original
    total = len(doc.paragraphs)
    print(f"Document loaded: {total} paragraphs")

    # Sort injection indices descending so insertions don't shift earlier indices
    sorted_indices = sorted(INJECTION_MAP.keys(), reverse=True)

    for orig_idx in sorted_indices:
        fn, caption = INJECTION_MAP[orig_idx]
        # Check index still valid
        if orig_idx >= len(doc.paragraphs):
            print(f"  SKIP idx={orig_idx} (out of range after insertions)")
            continue
        print(f"  Rendering diagram for para {orig_idx}: {caption[:50]}...")
        stream = fn()
        replace_para_with_picture(doc, orig_idx, stream,
                                   width=Inches(5.8), caption=caption)

    doc.save(OUT)
    print(f"\nDone. Saved -> {OUT}")
    print(f"Original backed up -> {BACKUP}")
