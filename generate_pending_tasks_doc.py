"""
Generate pending tasks DOCX for Fideon OS backend/frontend work.
"""

import textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    else:
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return h


def para(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def status_badge(doc, text, color):
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # background highlight via shading on the run
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    fill = color.replace("#", "")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)
    return p


def draw_progress_chart() -> bytes:
    """Bar chart showing done vs pending for each item."""
    items   = ["C1\nRoute Migration", "C2\nVector Ingestion", "W2\nCircuit Breaker",
               "W3\nRoute Auth", "C3\nFrontend Guard"]
    done    = [20,  70, 60, 20, 0]
    pending = [80, 30, 40, 80, 100]

    x = range(len(items))
    fig, ax = plt.subplots(figsize=(11, 4.5))
    fig.patch.set_facecolor("white")

    bars_done    = ax.bar(x, done,    label="Done",    color="#059669", width=0.5)
    bars_pending = ax.bar(x, pending, bottom=done, label="Pending", color="#DC2626", width=0.5, alpha=0.85)

    for i, (d, p) in enumerate(zip(done, pending)):
        if d > 0:
            ax.text(i, d / 2, f"{d}%", ha="center", va="center",
                    fontsize=9, fontweight="bold", color="white")
        ax.text(i, d + p / 2, f"{p}%", ha="center", va="center",
                fontsize=9, fontweight="bold", color="white")

    ax.set_xticks(list(x))
    ax.set_xticklabels(items, fontsize=9)
    ax.set_ylabel("Completion %", fontsize=9)
    ax.set_ylim(0, 115)
    ax.set_yticks([0, 25, 50, 75, 100])
    ax.legend(fontsize=9, loc="upper right")
    ax.set_title("Pending Tasks — Completion Status  |  Fideon OS", fontsize=11,
                 fontweight="bold", color="#1E3A5F", pad=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_facecolor("#F8FAFC")

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def draw_effort_chart() -> bytes:
    """Pie chart of estimated effort per task."""
    labels  = ["C1 Route Migration\n(2 days)", "C2 Vector Ingestion\n(2 hrs)",
               "W2 Circuit Breaker\n(1 hr)", "W3 Route Auth\n(2 days)", "C3 Frontend Guard\n(30 min)"]
    sizes   = [480, 120, 60, 480, 30]   # minutes
    colors  = ["#DC2626", "#D97706", "#F59E0B", "#B91C1C", "#059669"]
    explode = (0.04, 0.04, 0.04, 0.04, 0.04)

    fig, ax = plt.subplots(figsize=(8, 6))
    fig.patch.set_facecolor("white")
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, explode=explode,
        autopct="%1.0f%%", startangle=140,
        textprops={"fontsize": 8.5},
        pctdistance=0.75,
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_fontweight("bold")
        at.set_color("white")
    ax.set_title("Effort Distribution — Pending Tasks", fontsize=11,
                 fontweight="bold", color="#1E3A5F", pad=10)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ── Document ──────────────────────────────────────────────────────────────────

def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width        = Cm(21)
    section.page_height       = Cm(29.7)
    section.left_margin       = section.right_margin  = Cm(2.0)
    section.top_margin        = section.bottom_margin = Cm(1.8)

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_heading("Fideon OS", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Pending Implementation Tasks — Backend & Frontend")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(13)
    s.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()
    m = doc.add_paragraph("Version: 1.0   |   Branch: v1-dev   |   April 2026")
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.runs[0].font.size = Pt(9)
    m.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_page_break()

    # ── TOC ───────────────────────────────────────────────────────────────────
    heading(doc, "Table of Contents", 1)
    toc = [
        "1.  Summary — What Is Done vs Pending",
        "2.  Progress & Effort Charts",
        "3.  C1  — Migrate Existing Routes to tenant_scoped helpers",
        "4.  C2  — Pass tenant_id Through Vector Ingestion",
        "5.  W2  — Wire Circuit Breaker record_success / record_failure",
        "6.  W3  — Migrate Existing Routes to CurrentTenant Dependency",
        "7.  C3  — Frontend apiBaseUrl.ts HTTP Guard",
        "8.  Recommended Implementation Order",
        "9.  Definition of Done Checklist",
    ]
    for item in toc:
        bullet(doc, item)
    doc.add_page_break()

    # ── Section 1: Summary table ───────────────────────────────────────────────
    heading(doc, "1. Summary — What Is Done vs Pending", 1)
    summary_rows = [
        ("C1", "Service role key — migrate routes to tenant_scoped",      "CRITICAL", "Pending", "26 route files, ~80+ call sites"),
        ("C2", "pgvector tenant_id — wire through ingestion pipeline",     "CRITICAL", "Partial", "vectorstore_ingestion.py + acord.py"),
        ("W2", "Circuit breaker — wire record_success/failure in llm.py",  "WARNING",  "Partial", "1 file, ~15 lines"),
        ("W3", "CurrentTenant dependency — migrate all routes",            "WARNING",  "Pending", "26 route files"),
        ("C3", "Frontend apiBaseUrl.ts — HTTP fallback guard",             "CRITICAL", "Pending", "1 file, 5 lines"),
    ]
    add_table(doc,
              ["ID", "Task", "Severity", "Status", "Scope"],
              summary_rows,
              col_widths=[0.4, 2.8, 0.85, 0.75, 1.8])
    doc.add_paragraph()
    para(doc,
         "C1 and W3 touch the same 26 route files — implement together to avoid doing two passes "
         "over the same code. Total combined effort: 2 days, not 4.",
         bold=True)
    doc.add_page_break()

    # ── Section 2: Charts ──────────────────────────────────────────────────────
    heading(doc, "2. Progress & Effort Charts", 1)

    heading(doc, "2.1 Completion Status per Task", 2)
    progress_png = draw_progress_chart()
    doc.add_picture(io.BytesIO(progress_png), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    heading(doc, "2.2 Effort Distribution", 2)
    effort_png = draw_effort_chart()
    doc.add_picture(io.BytesIO(effort_png), width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── Section 3: C1 ─────────────────────────────────────────────────────────
    heading(doc, "3. C1 — Migrate Routes to tenant_scoped Helpers", 1)

    p = doc.add_paragraph()
    r = p.add_run("Severity: CRITICAL   |   Status: PENDING   |   Effort: 2 days (combined with W3)")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    heading(doc, "3.1 What Was Built", 2)
    para(doc, "The following helpers were added to backend/app/core/supabase.py:")
    bullet(doc, "tenant_scoped_get(table, query, tenant_id)  — appends &tenant_id=eq.<id> to every GET")
    bullet(doc, "tenant_scoped_insert(table, data, tenant_id) — injects tenant_id into INSERT payload")
    bullet(doc, "tenant_scoped_patch(table, query, data, tenant_id) — scopes PATCH to tenant's rows only")
    bullet(doc, "MULTI_TENANT_TABLES frozenset — 24 tables registered as requiring tenant scoping")
    bullet(doc, "_assert_tenant_id() guard — raises ValueError immediately if tenant_id is missing")

    heading(doc, "3.2 What Is Still Missing", 2)
    para(doc, "26 route files still call postgrest_get() / postgrest_insert() / postgrest_patch() "
         "directly without any tenant_id filter. A bug in any of these can silently return "
         "another tenant's data.")

    heading(doc, "3.3 Affected Files", 2)
    affected = [
        ("backend/app/routes/activity.py",          "postgrest_get on audit_logs — no tenant filter"),
        ("backend/app/routes/admin.py",              "user listing, role changes — no tenant filter"),
        ("backend/app/routes/agents.py",             "agent_pipelines reads — no tenant filter"),
        ("backend/app/routes/chat.py",               "chat_conversations reads/writes — no tenant filter"),
        ("backend/app/routes/decision_reviews.py",   "decision_reviews reads — no tenant filter"),
        ("backend/app/routes/device.py",             "devices reads/writes — no tenant filter"),
        ("backend/app/routes/device_admin.py",       "devices admin reads — no tenant filter"),
        ("backend/app/routes/federated_learning.py", "federated_rounds reads — no tenant filter"),
        ("backend/app/routes/model_registry.py",     "activated_models reads — no tenant filter"),
        ("backend/app/routes/notifications.py",      "notifications reads — no tenant filter"),
        ("backend/app/routes/pods.py",               "pod extraction runs — no tenant filter"),
        ("backend/app/routes/acord.py",              "acord_extraction_runs — no tenant filter"),
        ("backend/app/routes/settings.py",           "app_users / tenants — no tenant filter"),
        ("backend/app/routes/tenants.py",            "tenants reads — no tenant filter"),
        ("backend/app/routes/webhooks.py",           "webhooks reads/writes — no tenant filter"),
        ("backend/app/routes/workflow_ai.py",        "workflows reads — no tenant filter"),
        ("backend/app/services/acord_training.py",   "training_jobs reads — no tenant filter"),
        ("backend/app/services/pod_training.py",     "training_jobs reads — no tenant filter"),
        ("backend/app/services/webhook_engine.py",   "webhooks service reads — no tenant filter"),
    ]
    add_table(doc, ["File", "Issue"], affected, col_widths=[3.2, 3.4])
    doc.add_paragraph()

    heading(doc, "3.4 How to Fix (Pattern)", 2)
    para(doc, "Before (current — unsafe):", bold=True)
    code_block(doc, textwrap.dedent("""\
        # Any route — current pattern
        user = await verify_user(authorization)
        user_id = user.get("id")
        rows = await postgrest_get("devices", f"user_id=eq.{user_id}")
        # Bug: if user_id is from another tenant this returns their devices
    """))
    para(doc, "After (correct — tenant-scoped):", bold=True)
    code_block(doc, textwrap.dedent("""\
        # Any route — safe pattern
        from app.core.deps import CurrentTenant
        from app.core.supabase import tenant_scoped_get

        async def list_devices(ctx: CurrentTenant):
            rows = await tenant_scoped_get(
                "devices",
                f"user_id=eq.{ctx['user_id']}",
                ctx["tenant_id"],           # always scoped — no cross-tenant leakage
            )
            return rows
    """))
    doc.add_page_break()

    # ── Section 4: C2 ─────────────────────────────────────────────────────────
    heading(doc, "4. C2 — Pass tenant_id Through Vector Ingestion", 1)

    p = doc.add_paragraph()
    r = p.add_run("Severity: CRITICAL   |   Status: PARTIAL   |   Effort: 2 hours")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    heading(doc, "4.1 What Was Built", 2)
    bullet(doc, "pgvector_store.py — upsert_chunks() now accepts tenant_id parameter")
    bullet(doc, "pgvector_store.py — query_similar() now accepts tenant_id parameter")
    bullet(doc, "query_similar() without tenant_id now restricts to tenant_id IS NULL rows only")
    bullet(doc, "Migration 20260421000000_vector_service_role.sql — vector_service role created")

    heading(doc, "4.2 What Is Still Missing", 2)
    para(doc, "The callers of upsert_chunks() and query_similar() have not been updated "
         "to pass tenant_id. Chunks are still being stored and queried with tenant_id=NULL.")

    heading(doc, "4.3 Files to Fix", 2)
    c2_files = [
        ("backend/app/services/vectorstore_ingestion.py",
         "ingest_text_into_vectorstore() calls upsert_chunks() without tenant_id. "
         "Add tenant_id parameter and pass it through."),
        ("backend/app/routes/acord.py",
         "_ingest_acord_into_vectorstore() calls vectorstore_ingestion without tenant_id. "
         "Extract tenant_id from request context and pass down."),
        ("backend/LLM/rag/generator.py",
         "query_similar() calls may lack tenant_id. "
         "Verify and add tenant_id to all RAG query calls."),
    ]
    add_table(doc, ["File", "Fix Required"], c2_files, col_widths=[3.0, 3.6])
    doc.add_paragraph()

    heading(doc, "4.4 How to Fix", 2)
    para(doc, "vectorstore_ingestion.py — add tenant_id parameter:", bold=True)
    code_block(doc, textwrap.dedent("""\
        # Before
        def ingest_text_into_vectorstore(*, collection_name, doc_id, text, ...):
            upsert_chunks(collection_name=collection_name, chunks=chunks)

        # After
        def ingest_text_into_vectorstore(*, collection_name, doc_id, text,
                                            tenant_id: str | None = None, ...):
            upsert_chunks(collection_name=collection_name, chunks=chunks,
                          tenant_id=tenant_id)
    """))
    para(doc, "acord.py — pass tenant_id from request context:", bold=True)
    code_block(doc, textwrap.dedent("""\
        # Before
        await _ingest_acord_into_vectorstore(text=extracted_text, pod_id=pod_id)

        # After
        ctx = await get_user_context(authorization)
        await _ingest_acord_into_vectorstore(
            text=extracted_text,
            pod_id=pod_id,
            tenant_id=ctx["tenant_id"],    # scopes chunks to this tenant only
        )
    """))
    doc.add_page_break()

    # ── Section 5: W2 ─────────────────────────────────────────────────────────
    heading(doc, "5. W2 — Wire Circuit Breaker record_success / record_failure", 1)

    p = doc.add_paragraph()
    r = p.add_run("Severity: WARNING   |   Status: PARTIAL   |   Effort: 1 hour")
    r.bold = True
    r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    heading(doc, "5.1 What Was Built", 2)
    bullet(doc, "resilience.py — SimpleCircuitBreaker class with CLOSED / OPEN / HALF states")
    bullet(doc, "resilience.py — runpod_circuit_breaker singleton instance")
    bullet(doc, "resilience.py — @runpod_retry decorator (tenacity, 3 retries, 2s-8s backoff)")
    bullet(doc, "llm.py — @runpod_retry applied to _runpod_generate_text()")
    bullet(doc, "llm.py — runpod_circuit_breaker.before_call() added to _offline_fallback_stream()")

    heading(doc, "5.2 What Is Still Missing", 2)
    para(doc,
         "record_success() and record_failure() are never called after RunPod requests complete. "
         "The circuit breaker tracks zero failures and never opens. "
         "It blocks calls when OPEN (before_call works) but it can never become OPEN "
         "because the failure count never increases.")

    heading(doc, "5.3 File to Fix", 2)
    para(doc, "File: backend/app/services/llm.py", bold=True)

    heading(doc, "5.4 How to Fix", 2)
    code_block(doc, textwrap.dedent("""\
        # backend/app/services/llm.py
        from app.core.resilience import runpod_circuit_breaker

        async def _runpod_generate_text(payload: dict, model_name: str) -> str:
            runpod_token = _clean_bearer_token(FIDEON_SECRET_KEY or RUNPOD_API_KEY)
            if not runpod_token:
                raise RuntimeError("RunPod token not configured")

            # Step 1 — check circuit before hitting RunPod
            runpod_circuit_breaker.before_call()

            try:
                result = await _do_runpod_call(payload, model_name)
                # Step 2 — record success so circuit stays CLOSED
                runpod_circuit_breaker.record_success()
                return result
            except Exception as exc:
                # Step 3 — record failure; after 5 failures circuit opens
                runpod_circuit_breaker.record_failure()
                raise
    """))
    para(doc,
         "Same pattern applies to _try_runpod_openai_compat_stream() if it exists in llm.py.")
    doc.add_page_break()

    # ── Section 6: W3 ─────────────────────────────────────────────────────────
    heading(doc, "6. W3 — Migrate Existing Routes to CurrentTenant Dependency", 1)

    p = doc.add_paragraph()
    r = p.add_run("Severity: WARNING   |   Status: PENDING   |   Effort: 2 days (combined with C1)")
    r.bold = True
    r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    heading(doc, "6.1 What Was Built", 2)
    bullet(doc, "deps.py — get_current_tenant() FastAPI dependency")
    bullet(doc, "deps.py — get_current_tenant_strict() — raises 403 if tenant_id is NULL")
    bullet(doc, "deps.py — get_current_admin() — enforces admin / global_admin role")
    bullet(doc, "deps.py — CurrentTenant, CurrentTenantStrict, CurrentAdmin Annotated shortcuts")
    bullet(doc, "deps.py — GetDB — Annotated[DBRepository, Depends(get_db)]")

    heading(doc, "6.2 What Is Still Missing", 2)
    para(doc,
         "All 26 route files still repeat the same 4-line auth pattern manually on every endpoint. "
         "This is error-prone — if any route forgets to call verify_user() it becomes unauthenticated. "
         "With CurrentTenant the dependency is declared in the function signature and FastAPI "
         "enforces it automatically.")

    heading(doc, "6.3 How to Fix (Pattern)", 2)
    para(doc, "Before (current — manual, repetitive):", bold=True)
    code_block(doc, textwrap.dedent("""\
        @router.get("/devices")
        async def list_devices(authorization: Optional[str] = Header(None)):
            user = await verify_user(authorization)            # repeated in every route
            user_id = user.get("id")
            profile = await postgrest_get("app_users",
                f"user_id=eq.{user_id}&select=tenant_id&limit=1")
            tenant_id = profile[0].get("tenant_id")           # repeated in every route
            rows = await postgrest_get("devices",
                f"user_id=eq.{user_id}")                      # no tenant scope
            return rows
    """))
    para(doc, "After (correct — dependency injection):", bold=True)
    code_block(doc, textwrap.dedent("""\
        from app.core.deps import CurrentTenant, GetDB
        from app.core.db import DBQuery

        @router.get("/devices")
        async def list_devices(ctx: CurrentTenant, db: GetDB):
            # ctx["user_id"], ctx["tenant_id"], ctx["role"] all pre-resolved
            rows = await db.tenant_get(
                "devices",
                DBQuery(filters={"user_id": ctx["user_id"]}, limit=100),
                tenant_id=ctx["tenant_id"],
            )
            return rows
    """))

    heading(doc, "6.4 Migration Strategy", 2)
    para(doc, "Recommended approach — migrate one route file at a time:")
    bullet(doc, "1. Pick one route file (start with the smallest, e.g. activity.py)")
    bullet(doc, "2. Replace verify_user() + manual tenant_id lookup with ctx: CurrentTenant")
    bullet(doc, "3. Replace postgrest_get/insert/patch with db.tenant_get/insert/update")
    bullet(doc, "4. Run existing tests to confirm no regression")
    bullet(doc, "5. Repeat for next file")
    para(doc, "Do NOT migrate all 26 files in one PR — it will be impossible to review.")
    doc.add_page_break()

    # ── Section 7: C3 ─────────────────────────────────────────────────────────
    heading(doc, "7. C3 — Frontend apiBaseUrl.ts HTTP Guard", 1)

    p = doc.add_paragraph()
    r = p.add_run("Severity: CRITICAL   |   Status: PENDING   |   Effort: 30 minutes")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    heading(doc, "7.1 Problem", 2)
    para(doc,
         "frontend/src/lib/apiBaseUrl.ts currently falls back to http://127.0.0.1:8080 "
         "if NEXT_PUBLIC_API_URL is not set. In a misconfigured Azure App Service deployment, "
         "all API calls — including JWT Bearer tokens — would be sent over plaintext HTTP.")

    heading(doc, "7.2 File to Change", 2)
    para(doc, "File: frontend/src/lib/apiBaseUrl.ts", bold=True)

    heading(doc, "7.3 Exact Code Change", 2)
    para(doc, "Current code (unsafe):", bold=True)
    code_block(doc, textwrap.dedent("""\
        export function getApiBaseUrl(): string {
          const url = process.env.NEXT_PUBLIC_API_URL;
          if (!url) {
            return 'http://127.0.0.1:8080';   // silently returns HTTP in production
          }
          return url;
        }
    """))
    para(doc, "Fixed code (safe — throws in production if env var is missing):", bold=True)
    code_block(doc, textwrap.dedent("""\
        export function getApiBaseUrl(): string {
          const url = process.env.NEXT_PUBLIC_API_URL;
          if (!url) {
            const isLocalhost =
              typeof window !== 'undefined' &&
              (window.location.hostname === 'localhost' ||
               window.location.hostname === '127.0.0.1');
            if (!isLocalhost) {
              throw new Error(
                '[Fideon OS] NEXT_PUBLIC_API_URL is not set. ' +
                'Set it to the HTTPS backend URL in your Azure App Service config.'
              );
            }
            return 'http://127.0.0.1:8080';   // local dev only
          }
          return url;
        }
    """))

    heading(doc, "7.4 Additional Required Steps", 2)
    bullet(doc, "Add NEXT_PUBLIC_API_URL=https://<backend>.azurewebsites.net to GitHub Actions "
               "secrets for staging and production environments")
    bullet(doc, "Add NEXT_PUBLIC_API_URL to the deploy-app-service.yml env var sync step")
    bullet(doc, "Never copy http://localhost:8000 from .env.local to App Service config")
    doc.add_page_break()

    # ── Section 8: Implementation Order ───────────────────────────────────────
    heading(doc, "8. Recommended Implementation Order", 1)
    para(doc,
         "Ordered by: highest security impact first, then effort efficiency "
         "(batching overlapping files together).")

    order_rows = [
        ("1", "C3  Frontend HTTP guard",
         "30 min",  "CRITICAL", "Single file, zero risk of regression. Do first."),
        ("2", "W2  Circuit breaker wiring",
         "1 hr",    "WARNING",  "Single file (llm.py), ~15 lines. Quick win."),
        ("3", "C2  Vector ingestion tenant_id",
         "2 hrs",   "CRITICAL", "2-3 files. Low regression risk."),
        ("4", "C1 + W3  Route migration",
         "2 days",  "CRITICAL", "26 files. Do together. One file per PR."),
    ]
    add_table(doc,
              ["Order", "Task", "Effort", "Severity", "Notes"],
              order_rows,
              col_widths=[0.5, 2.2, 0.7, 0.85, 2.4])
    doc.add_paragraph()
    para(doc,
         "Quick wins (C3 + W2 + C2) can be done in a single half-day sprint. "
         "Route migration (C1 + W3) should be a separate sprint with dedicated PR reviews.",
         bold=True)
    doc.add_page_break()

    # ── Section 9: Definition of Done ─────────────────────────────────────────
    heading(doc, "9. Definition of Done Checklist", 1)
    para(doc, "A task is complete ONLY when all items below are checked.")

    checklist = {
        "C3 — Frontend HTTP Guard": [
            "apiBaseUrl.ts throws on non-localhost when NEXT_PUBLIC_API_URL is unset",
            "NEXT_PUBLIC_API_URL added to GitHub Actions staging + production secrets",
            "deploy-app-service.yml syncs NEXT_PUBLIC_API_URL to App Service",
            "Tested: deploy without env var causes visible build/runtime error (not silent HTTP)",
        ],
        "W2 — Circuit Breaker Wiring": [
            "runpod_circuit_breaker.before_call() called before every RunPod HTTP request",
            "runpod_circuit_breaker.record_success() called after every successful response",
            "runpod_circuit_breaker.record_failure() called after every failed response",
            "Manual test: 5 consecutive RunPod failures cause circuit to OPEN",
            "Manual test: after 60s, circuit moves to HALF-OPEN and one probe is allowed",
        ],
        "C2 — Vector Ingestion tenant_id": [
            "vectorstore_ingestion.py ingest_text_into_vectorstore() accepts tenant_id param",
            "acord.py passes tenant_id from request context to ingestion function",
            "rag/generator.py query_similar() calls include tenant_id",
            "Verified: new chunks have tenant_id populated in rag_chunks table",
            "Verified: query_similar() without tenant_id returns zero rows (not all tenants)",
        ],
        "C1 + W3 — Route Migration": [
            "Every route that reads a multi-tenant table uses tenant_scoped_get() or db.tenant_get()",
            "Every route that writes a multi-tenant table uses tenant_scoped_insert() or db.tenant_insert()",
            "No route calls postgrest_get() on a MULTI_TENANT_TABLES table without tenant_id filter",
            "All routes use ctx: CurrentTenant instead of manual verify_user() + tenant lookup",
            "Cross-tenant test: JWT from tenant A cannot return data from tenant B on any endpoint",
            "All existing tests pass with no regression",
        ],
    }

    for task_name, items in checklist.items():
        p = doc.add_paragraph()
        r = p.add_run(task_name)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        for item in items:
            bullet(doc, f"[ ]  {item}", level=1)
        doc.add_paragraph()

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    out_path = "fideon_pending_tasks.docx"
    print("Building pending tasks document...")
    doc = build_doc()
    doc.save(out_path)
    print(f"Saved: {out_path}")
