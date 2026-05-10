"""Generate Fine-Tuning Pipeline Architecture document with proper diagrams."""
import os, tempfile
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── tmp image paths ──────────────────────────────────────────────────────────
_TMP = Path(tempfile.gettempdir())
def _p(name): return str(_TMP / f"fideon_{name}.png")

# ─── Colour palette ───────────────────────────────────────────────────────────
C = dict(
    blue='#1A56DB', blue_dark='#1E40AF', blue_light='#DBEAFE', blue_xlight='#EFF6FF',
    green='#059669', green_light='#DCFCE7',
    amber='#D97706', amber_light='#FEF3C7',
    purple='#7C3AED', purple_light='#F3E8FF',
    red='#DC2626', red_light='#FEF2F2',
    slate='#334155', gray='#64748B',
    yellow_light='#FEFCE8', yellow_edge='#CA8A04',
    white='#FFFFFF', bg='#F8FAFC',
)

# ─── Diagram primitives ───────────────────────────────────────────────────────
def _box(ax, cx, cy, w, h, text, fill, edge, fs=8.5, bold=False, tc='#1e293b', zorder=3):
    ax.add_patch(mpatches.FancyBboxPatch(
        (cx-w/2, cy-h/2), w, h, boxstyle="round,pad=0.08",
        facecolor=fill, edgecolor=edge, linewidth=1.5, zorder=zorder))
    ax.text(cx, cy, text, ha='center', va='center', fontsize=fs,
            fontweight='bold' if bold else 'normal', color=tc,
            multialignment='center', zorder=zorder+1)

def _diamond(ax, cx, cy, w, h, text, fill, edge, fs=8):
    pts = [(cx, cy+h/2), (cx+w/2, cy), (cx, cy-h/2), (cx-w/2, cy)]
    ax.add_patch(plt.Polygon(pts, closed=True, facecolor=fill, edgecolor=edge, linewidth=1.5, zorder=3))
    ax.text(cx, cy, text, ha='center', va='center', fontsize=fs, multialignment='center', zorder=4)

def _arr(ax, x1, y1, x2, y2, color='#64748b', lw=1.5):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw, mutation_scale=14), zorder=5)

def _band(ax, y_bot, y_top, label, hc, fc, x1=0.3, x2=12.7):
    w = x2 - x1
    ax.add_patch(mpatches.FancyBboxPatch(
        (x1, y_bot), w, y_top-y_bot, boxstyle="round,pad=0.05",
        facecolor=fc, edgecolor=hc, linewidth=2, zorder=1))
    ax.add_patch(mpatches.Rectangle((x1, y_top-0.48), w, 0.48, facecolor=hc, zorder=2))
    ax.text((x1+x2)/2, y_top-0.24, label, ha='center', va='center',
            fontsize=9.5, fontweight='bold', color='white', zorder=3)


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 1 — Full Pipeline Flowchart
# ══════════════════════════════════════════════════════════════════════════════
def gen_pipeline(path):
    fig = plt.figure(figsize=(13, 25), facecolor='white')
    ax  = fig.add_axes([0.01, 0.005, 0.98, 0.99])
    ax.set_xlim(0, 13); ax.set_ylim(0, 25); ax.axis('off')

    ax.text(6.5, 24.5, "FIDEON OS — Fine-Tuning Pipeline: Complete Flow",
            ha='center', va='center', fontsize=13, fontweight='bold', color=C['blue_dark'])

    _band(ax, 18.8, 24.1, "PHASE 1: DATA COLLECTION  (Electron App)",      C['blue'],   C['blue_xlight'])
    _band(ax, 13.4, 18.3, "PHASE 2: TRAINING  (RunPod GPU)",                C['green'],  C['green_light'])
    _band(ax,  9.2, 12.9, "PHASE 3: QUANTIZATION  (RunPod)",                C['amber'],  C['amber_light'])
    _band(ax,  6.0,  8.7, "PHASE 4: UPLOAD TO SEAWEEDFS",                   C['purple'], C['purple_light'])
    _band(ax,  0.3,  5.5, "PHASE 5: ELECTRON DOWNLOAD & DEPLOY",            C['red'],    C['red_light'])

    BW, CX, BH = 9.2, 6.5, 0.52

    # ── Phase 1 ──────────────────────────────────────────────────────────────
    ys = [23.35, 22.55, 21.75, 20.95, 20.15]
    texts_p1 = [
        "User uploads ACORD form (PDF)",
        "AI extracts fields → displays result to user  (displayed_output saved)",
        "User reviews, corrects mistakes, clicks  APPROVE",
        "Source PDF  →  uploaded to SeaweedFS  (source_file_url stored)",
        "Supabase stores:  source_file_url  +  original_extracted_json  +  corrected_json",
    ]
    for i, (y, t) in enumerate(zip(ys, texts_p1)):
        fill = C['blue_light'] if i != 4 else '#ECFDF5'
        edge = C['blue'] if i != 4 else C['green']
        _box(ax, CX, y, BW+0.4, BH, t, fill, edge, fs=8)
        if i < len(ys)-1: _arr(ax, CX, y-BH/2, CX, ys[i+1]+BH/2)

    ax.text(CX+0.3, 19.73, "records ≥ 19  →  training job queued",
            fontsize=7.5, color=C['gray'], style='italic', ha='center')
    _arr(ax, CX, 19.89, CX, 18.82)

    # ── Phase 2 ──────────────────────────────────────────────────────────────
    ys2 = [18.0, 17.2, 16.4, 15.6]
    texts_p2 = [
        "Check SeaweedFS for latest  model.safetensors",
        "Found → download to RunPod   |   Not found → use base model (qwen2.5-72b)",
        "Stratified sample:  60% recent / 30% hard examples / 10% OOS  (cap 1500 rows)",
        "QLoRA training  →  Merge: loaded_weights + new_LoRA  =  merged_model",
    ]
    for i, (y, t) in enumerate(zip(ys2, texts_p2)):
        _box(ax, CX, y, BW+0.6, BH, t, C['green_light'], C['green'], fs=8)
        if i < len(ys2)-1: _arr(ax, CX, y-BH/2, CX, ys2[i+1]+BH/2)

    _arr(ax, CX, 15.34, CX, 14.82)
    _diamond(ax, CX, 14.48, 5.0, 0.72, "Quality Gate\nPASS ?", C['yellow_light'], C['yellow_edge'])
    # FAIL
    _arr(ax, CX+2.5, 14.48, 10.6, 14.48, color=C['red'])
    ax.text(9.0, 14.6, "FAIL", fontsize=8, color=C['red'], fontweight='bold')
    _box(ax, 11.6, 14.48, 2.0, 0.6, "ABORT\nKeep prev", C['red_light'], C['red'], fs=8, tc=C['red'])
    # PASS
    _arr(ax, CX, 14.12, CX, 13.42, color=C['green'])
    ax.text(CX+0.3, 13.75, "PASS", fontsize=8, color=C['green'], fontweight='bold')

    # ── Phase 3 ──────────────────────────────────────────────────────────────
    ys3 = [12.6, 11.8, 11.0, 10.2]
    texts_p3 = [
        "Convert merged model → GGUF F16  (base for quantization — llama.cpp convert)",
        "Produce 3 variants:  model-f16.gguf (high)  |  model-q8_0.gguf (mid)  |  model-q4_k_m.gguf (low)",
        "GPG sign each GGUF  →  .sig files   |   Compute SHA-256 checksums per variant",
        "Build  manifest.json  (version, quant_variants, checksums, device_requirements)",
    ]
    for i, (y, t) in enumerate(zip(ys3, texts_p3)):
        _box(ax, CX, y, BW+0.6, BH, t, C['amber_light'], C['amber'], fs=8)
        if i < len(ys3)-1: _arr(ax, CX, y-BH/2, CX, ys3[i+1]+BH/2)
    _arr(ax, CX, 9.94, CX, 8.72)

    # ── Phase 4 ──────────────────────────────────────────────────────────────
    ys4 = [8.35, 7.5]
    _box(ax, CX, ys4[0], BW+1.2, BH,
         "Upload to SeaweedFS  broker/vN.N/:  model.safetensors  +  model-f16.gguf  +  model-q8_0.gguf  +  model-q4_k_m.gguf  +  *.sig  +  manifest.json",
         C['purple_light'], C['purple'], fs=7.5)
    _arr(ax, CX, ys4[0]-BH/2, CX, ys4[1]+BH/2)
    _box(ax, CX, ys4[1], BW+1.2, BH,
         "Register vN.N in  adapter_registry  (is_available=true,  canary_pct=10,  quant_variants=[f16,q8_0,q4_k_m])",
         C['purple_light'], C['purple'], fs=7.8)
    _arr(ax, CX, ys4[1]-BH/2, CX, 5.52)

    # ── Phase 5 ──────────────────────────────────────────────────────────────
    ys5 = [5.15, 4.32, 3.48, 2.65, 1.82]
    _box(ax, CX, ys5[0], BW, BH,
         "Electron My Models page  →  GET /api/v1/adapter/latest?domain=broker",
         C['red_light'], C['red'])
    _arr(ax, CX, ys5[0]-BH/2, CX, ys5[1]+0.36)

    _diamond(ax, CX, ys5[1], 4.8, 0.7, "Canary Gate\nDevice in cohort?", C['yellow_light'], C['yellow_edge'])
    _arr(ax, CX+2.4, ys5[1], 10.6, ys5[1], color=C['gray'])
    ax.text(9.1, ys5[1]+0.12, "NO", fontsize=8, color=C['gray'], fontweight='bold')
    _box(ax, 11.6, ys5[1], 2.0, 0.6, "available:\nfalse", C['bg'], C['gray'], fs=8)
    _arr(ax, CX, ys5[1]-0.35, CX, ys5[2]+BH/2, color=C['green'])
    ax.text(CX+0.3, ys5[1]-0.15, "YES", fontsize=8, color=C['green'], fontweight='bold')

    texts_p5 = [
        "ModelUpdateBanner: 'Update available vN.N — Download & Install'",
        "Read device config: RAM check → select variant  (≥32 GB=F16 | 16-32 GB=Q8 | <16 GB=Q4)",
        "GET /api/v1/adapter/download-url?quant=<selected>  →  Presigned SeaweedFS URL (1hr TTL)",
        "Download selected GGUF  →  Verify SHA-256 + GPG sig  →  ollama create  →  Model live offline",
    ]
    ys5_ext = [ys5[2], ys5[2]-0.78, ys5[2]-1.56, ys5[2]-2.34]
    for i in range(4):
        y = ys5_ext[i]
        fill = C['red_light'] if i != 1 else '#FFFBEB'
        edge = C['red'] if i != 1 else C['amber']
        _box(ax, CX, y, BW+0.4, BH, texts_p5[i], fill, edge, fs=8)
        if i < 3: _arr(ax, CX, y-BH/2, CX, ys5_ext[i+1]+BH/2)

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.08)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 2 — SeaweedFS Storage Layout
# ══════════════════════════════════════════════════════════════════════════════
def gen_seaweedfs(path):
    fig = plt.figure(figsize=(13, 8.5), facecolor='white')
    ax  = fig.add_axes([0.01, 0.01, 0.98, 0.98])
    ax.set_xlim(0, 13); ax.set_ylim(0, 8.5); ax.axis('off')

    ax.text(6.5, 8.1, "SeaweedFS  —  my-bucket/ Storage Layout",
            ha='center', va='center', fontsize=12, fontweight='bold', color=C['blue_dark'])

    def nd(cx, cy, w, h, text, fill, edge, fs=8.5, bold=False):
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx-w/2, cy-h/2), w, h, boxstyle="round,pad=0.08",
            facecolor=fill, edgecolor=edge, linewidth=1.5, zorder=3))
        ax.text(cx, cy, text, ha='center', va='center', fontsize=fs,
                fontweight='bold' if bold else 'normal', color='#1e293b',
                multialignment='center', zorder=4)

    def conn(x1, y1, x2, y2):
        ax.plot([x1, x2], [y1, y2], color='#94A3B8', linewidth=1.5, zorder=2)

    # Root
    nd(6.5, 7.6, 4.2, 0.55, "SeaweedFS  20.40.61.106:8333", C['blue'], C['blue_dark'], bold=True, fs=9.5)
    conn(6.5, 7.32, 6.5, 6.98)
    nd(6.5, 6.72, 2.5, 0.5, "my-bucket/", C['blue_light'], C['blue'], bold=True)

    # Domain folders
    for tx, lbl in [(2.2, "broker/"), (6.5, "claims/"), (10.8, "underwriter/")]:
        nd(tx, 5.85, 2.3, 0.48, lbl, C['purple_light'], C['purple'], bold=True)
        conn(6.5, 6.47, tx, 6.09)

    # Versions under broker
    for vx, vlbl, is_l in [(1.0,"v1.0.0/",False),(2.2,"v1.1.0/ ← latest",True),(3.4,"v1.2.0/",False)]:
        fill = C['green_light'] if is_l else C['bg']
        edge = C['green'] if is_l else '#CBD5E1'
        nd(vx, 4.9, 1.7, 0.46, vlbl, fill, edge, fs=7.8)
        conn(2.2, 5.61, vx, 5.13)

    # Files under v1.1.0
    files = [
        (0.7,  3.75, "model.safetensors\n(2–10 GB)",   C['green_light'],  C['green'],   "Re-training\nweights"),
        (2.1,  3.75, "model-f16.gguf\n(~28 GB)",       C['amber_light'],  C['amber'],   "High-end\n(≥32 GB RAM)"),
        (3.5,  3.75, "model-q8_0.gguf\n(~15 GB)",      C['amber_light'],  C['amber'],   "Mid-range\n(16–32 GB RAM)"),
        (4.9,  3.75, "model-q4_k_m.gguf\n(~8 GB)",     C['amber_light'],  C['amber'],   "Budget\n(<16 GB RAM)"),
        (6.3,  3.75, "*.gguf.sig\n(GPG sig)",           '#F0FDF4',         '#16A34A',    "Integrity\ncheck"),
        (7.7,  3.75, "manifest.json\n(checksums+\nvariants)", '#FFF7ED',   '#EA580C',    "Version\nmetadata"),
    ]
    for fx, fy, lbl, fill, edge, usage in files:
        nd(fx, fy, 1.5, 0.65, lbl, fill, edge, fs=7.5)
        conn(2.2, 4.67, fx, fy+0.32)
        ax.text(fx, fy-0.58, usage, ha='center', va='center', fontsize=6.5,
                color=C['gray'], multialignment='center')

    # Source PDF storage note
    nd(9.8, 4.9, 2.4, 0.48, "source_forms/\n<run_id>.pdf", '#FDF2FF', '#9333EA', fs=7.5, bold=False)
    conn(6.5, 6.47, 9.8, 6.09)
    ax.text(9.8, 4.3, "Original ACORD\nPDF per run", ha='center', fontsize=6.5,
            color='#9333EA', multialignment='center')

    # Legend
    ax.text(0.5, 2.15, "Legend:", fontsize=9, fontweight='bold', color=C['slate'])
    legend = [
        (C['green_light'],  C['green'],   "model.safetensors — re-training weights (server only, never sent to Electron)"),
        (C['amber_light'],  C['amber'],   "GGUF variants (F16/Q8/Q4) — Electron picks based on device RAM"),
        (C['purple_light'], C['purple'],  "Domain folder — strict isolation, no data crosses domains"),
        ('#FDF2FF',         '#9333EA',    "source_forms/ — original ACORD PDFs stored per run_id"),
    ]
    for i, (fill, edge, lbl) in enumerate(legend):
        ly = 1.72 - i*0.42
        ax.add_patch(mpatches.FancyBboxPatch((0.5, ly-0.14), 0.38, 0.3,
                     boxstyle="round,pad=0.02", facecolor=fill, edgecolor=edge, linewidth=1.2))
        ax.text(1.02, ly+0.01, lbl, fontsize=8, va='center', color='#1e293b')

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.1)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 3 — Progressive Fine-Tuning
# ══════════════════════════════════════════════════════════════════════════════
def gen_progressive(path):
    fig = plt.figure(figsize=(13, 5.5), facecolor='white')
    ax  = fig.add_axes([0.01, 0.06, 0.98, 0.88])
    ax.set_xlim(0, 13); ax.set_ylim(0, 5.5); ax.axis('off')

    ax.text(6.5, 5.15, "Progressive Fine-Tuning — Each Cycle Builds on Previous Weights",
            ha='center', va='center', fontsize=11.5, fontweight='bold', color=C['blue_dark'])

    def cycle(cx, version, source, batch, rfill):
        W, H = 3.0, 3.4
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx-W/2, 0.6), W, H, boxstyle="round,pad=0.1",
            facecolor=rfill, edgecolor=C['blue'], linewidth=2, zorder=2))
        # header
        ax.add_patch(mpatches.Rectangle((cx-W/2, 0.6+H-0.44), W, 0.44, facecolor=C['blue'], zorder=3))
        ax.text(cx, 0.6+H-0.22, f"CYCLE  {version}", ha='center', va='center',
                fontsize=9.5, fontweight='bold', color='white', zorder=4)
        # content
        for dy, txt, col, bld in [
            (0.8,  "INPUT:",          C['slate'], True),
            (0.45, source,            C['slate'], False),
            (0.0,  "+",               C['blue'],  False),
            (-0.32, batch,            C['slate'], False),
            (-0.72, "▼",              C['gray'],  False),
            (-1.0,  "OUTPUT:",        C['green'], True),
            (-1.3,  f"safetensors {version}\n+ GGUF {version}", C['green'], False),
        ]:
            ax.text(cx, 0.6+H-0.44+dy-0.5, txt, ha='center', va='center',
                    fontsize=7.5 if txt not in ("+","▼") else (13 if txt=="+" else 10),
                    fontweight='bold' if bld else 'normal', color=col,
                    multialignment='center', zorder=4)

    cycle(1.9,  "v1.0", "Base model\n(qwen2.5-72b)", "Batch 1 approvals", '#EFF6FF')
    cycle(6.5,  "v1.1", "safetensors v1.0\n(from SeaweedFS)", "Batch 2 approvals", '#ECFDF5')
    cycle(11.1, "v1.2", "safetensors v1.1\n(from SeaweedFS)", "Batch 3 approvals", '#F5F3FF')

    for ax_x, lbl in [(4.2, "safetensors\nv1.0"), (8.8, "safetensors\nv1.1")]:
        _arr(ax, ax_x-0.75, 2.3, ax_x+0.75, 2.3, color=C['green'], lw=2.5)
        ax.text(ax_x, 2.65, lbl, ha='center', va='center', fontsize=8,
                color=C['green'], fontweight='bold', multialignment='center')

    ax.text(6.5, 0.22,
            "Base model used ONLY in Cycle 1.  Every subsequent cycle starts from previous fine-tuned weights.  Knowledge accumulates — not reset.",
            ha='center', va='center', fontsize=8, color=C['slate'], style='italic',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFFBEB', edgecolor=C['amber'], linewidth=1.2))

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.1)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 4 — Stratified Sampling
# ══════════════════════════════════════════════════════════════════════════════
def gen_sampling(path):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5), facecolor='white')
    fig.patch.set_facecolor('white')
    fig.suptitle("Stratified Data Sampling — Training Cost Stays Fixed as Dataset Grows",
                 fontsize=11.5, fontweight='bold', color=C['blue_dark'], y=1.01)

    # Pie
    sizes  = [60, 30, 10]
    labels = ['Recent data\n(last 90 days)\n60%',
              'Hard examples\n(user corrections)\n30%',
              'OOS refusals\n(anti-hallucination)\n10%']
    colors_pie = [C['blue'], C['green'], C['amber']]
    wedges, _, autotexts = ax1.pie(
        sizes, labels=labels, colors=colors_pie, autopct='%1.0f%%',
        startangle=130, pctdistance=0.65, explode=(0.03,0.03,0.03),
        textprops={'fontsize': 8.5})
    for at in autotexts:
        at.set_fontsize(11); at.set_fontweight('bold'); at.set_color('white')
    ax1.set_title("1,500 training rows — composition", fontsize=9.5, color=C['slate'], pad=10)

    # Bar
    ax2.set_facecolor('white')
    db_sizes   = ['50', '500', '1,500', '5,000', '50,000']
    train_rows = [50,    500,   1500,    1500,     1500]
    bar_colors = [C['blue'], C['blue'], C['green'], C['green'], C['green']]
    bars = ax2.bar(db_sizes, train_rows, color=bar_colors, edgecolor='white', linewidth=0.5, alpha=0.85)
    ax2.axhline(y=1500, color=C['amber'], linewidth=2, linestyle='--', label='Cap = 1,500 rows')
    ax2.set_xlabel("Total approved records in DB", fontsize=9, color=C['slate'])
    ax2.set_ylabel("Training rows used", fontsize=9, color=C['slate'])
    ax2.set_title("Training time stays constant after cap", fontsize=9.5, color=C['slate'])
    ax2.tick_params(colors=C['slate'], labelsize=8)
    ax2.spines['top'].set_visible(False); ax2.spines['right'].set_visible(False)
    ax2.spines['left'].set_color('#CBD5E1'); ax2.spines['bottom'].set_color('#CBD5E1')
    ax2.legend(fontsize=8.5)
    for bar, val in zip(bars, train_rows):
        ax2.text(bar.get_x()+bar.get_width()/2, val+30, str(val),
                 ha='center', va='bottom', fontsize=8.5, fontweight='bold', color=C['slate'])
    ax2.set_ylim(0, 1900)

    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.12)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 5 — Canary Rollout
# ══════════════════════════════════════════════════════════════════════════════
def gen_canary(path):
    fig = plt.figure(figsize=(12, 5.5), facecolor='white')
    ax  = fig.add_axes([0.09, 0.14, 0.86, 0.76])
    ax.set_facecolor('#F8FAFC')
    fig.suptitle("Canary Rollout — Safe Progressive Model Updates",
                 fontsize=11.5, fontweight='bold', color=C['blue_dark'], y=0.98)

    stages = [
        (0, 0,   "Release\n(new version)",             C['gray']),
        (1, 10,  "Day 1\n10% devices",                 C['amber']),
        (2, 50,  "Day 3\n50% devices\n(if no issues)", C['blue']),
        (3, 100, "Day 7\n100% devices",                C['green']),
    ]
    for i, pct, lbl, color in stages:
        h = max(pct, 3)
        ax.bar(i, h, width=0.6, color=color, edgecolor='white', alpha=0.85, zorder=3)
        if pct: ax.text(i, pct+2.5, f"{pct}%", ha='center', va='bottom',
                        fontsize=12, fontweight='bold', color=color)
        ax.text(i, -5, lbl, ha='center', va='top', fontsize=8.2, color=C['slate'],
                multialignment='center')

    xs = [i for i,p,_,__ in stages if p > 0]
    ys = [p for _,p,__,___ in stages if p > 0]
    ax.plot(xs, ys, color=C['blue'], lw=2, linestyle='--', marker='o',
            markersize=8, markerfacecolor=C['blue'], zorder=4)
    ax.axhline(100, color=C['green'], lw=1, linestyle=':', alpha=0.5)

    ax.set_ylabel("% of devices receiving update", fontsize=9, color=C['slate'])
    ax.set_xticks([]); ax.set_yticks([0,10,50,100])
    ax.tick_params(colors=C['slate'])
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CBD5E1'); ax.spines['bottom'].set_color('#CBD5E1')
    ax.set_xlim(-0.6, 5.5); ax.set_ylim(-12, 115)

    ax.text(4.5, 75, "ROLLBACK:\nSet  blocked=true\n→ no device gets update\n→ prev version stays active",
            ha='center', va='center', fontsize=8.5, color=C['red'],
            bbox=dict(boxstyle='round,pad=0.4', facecolor=C['red_light'], edgecolor=C['red'], lw=1.5))
    ax.text(4.5, 32, "Gate:  SHA-256(device_id + version)\n% 100  <  canary_pct\n(deterministic — same device\nalways gets same answer)",
            ha='center', va='center', fontsize=8, color=C['slate'],
            bbox=dict(boxstyle='round,pad=0.4', facecolor='white', edgecolor='#CBD5E1', lw=1.2))

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.12)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 6 — Quality Gate
# ══════════════════════════════════════════════════════════════════════════════
def gen_quality_gate(path):
    fig = plt.figure(figsize=(13, 7), facecolor='white')
    ax  = fig.add_axes([0.01, 0.01, 0.98, 0.98])
    ax.set_xlim(0, 13); ax.set_ylim(0, 7); ax.axis('off')

    ax.text(6.5, 6.65, "Quality Gate — Model Evaluated Before Every Release",
            ha='center', va='center', fontsize=12, fontweight='bold', color=C['blue_dark'])

    # Input
    _box(ax, 6.5, 6.1, 4.5, 0.52, "Merged model  (post-training)", C['green_light'], C['green'], bold=True)
    _arr(ax, 6.5, 5.84, 6.5, 5.48)
    ax.text(6.85, 5.66, "evaluate on holdout split (10%)", fontsize=7.5, color=C['gray'], style='italic')

    # 5 check boxes
    checks = [
        (1.1,  "JSON Valid Rate\n≥ 80%",       "Malformed\nJSON output"),
        (3.25, "Field Recall\n≥ 30%",           "Missing\nrequired fields"),
        (5.4,  "Exact Match\n≥ 20%",            "Wrong\nfield values"),
        (7.55, "Extra Field Rate\n≤ 25%",       "Hallucinated\nfields"),
        (9.7,  "OOS Hallucination\n≤ 60%",      "Answering\nout-of-scope"),
    ]
    for cx, lbl, catches in checks:
        # fan arrow from input box
        ax.annotate('', xy=(cx, 4.78), xytext=(6.5, 5.46),
                    arrowprops=dict(arrowstyle='->', color='#94A3B8', lw=1.0, mutation_scale=11), zorder=2)
        _box(ax, cx, 4.42, 2.0, 0.68, lbl, C['blue_light'], C['blue'], fs=8, bold=True)
        ax.text(cx, 3.85, "catches:", fontsize=6.5, ha='center', color=C['gray'])
        ax.text(cx, 3.55, catches, fontsize=7, ha='center', color=C['slate'], multialignment='center')

    # Horizontal collector bar + single arrow to diamond
    ax.plot([1.1, 9.7], [3.1, 3.1], color='#94A3B8', lw=1.5, zorder=2)
    for cx, _, __ in checks:
        ax.plot([cx, cx], [3.38, 3.1], color='#94A3B8', lw=1.5, zorder=2)
    _arr(ax, 6.5, 3.1, 6.5, 2.72)

    # Decision diamond
    _diamond(ax, 6.5, 2.38, 5.2, 0.72, "ALL CHECKS\nPASS ?", C['yellow_light'], C['yellow_edge'], fs=8.5)

    # YES → right
    _arr(ax, 6.5+2.6, 2.38, 10.0, 2.38, color=C['green'], lw=2)
    ax.text(8.6, 2.55, "YES", fontsize=9, color=C['green'], fontweight='bold')
    _box(ax, 11.2, 2.38, 2.2, 0.68, "✅  Upload to\nSeaweedFS", C['green_light'], C['green'], bold=True)
    ax.text(11.2, 1.85, "New version enters\ncanary rollout", ha='center', fontsize=7.5,
            color=C['green'], multialignment='center')

    # NO → left
    _arr(ax, 6.5-2.6, 2.38, 3.0, 2.38, color=C['red'], lw=2)
    ax.text(4.3, 2.55, "NO", fontsize=9, color=C['red'], fontweight='bold')
    _box(ax, 1.8, 2.38, 2.2, 0.68, "❌  ABORT\nKeep prev version", C['red_light'], C['red'], bold=True, fs=8)
    ax.text(1.8, 1.85, "Alert ops team\nPrev model stays active", ha='center', fontsize=7.5,
            color=C['red'], multialignment='center')

    # Config note
    ax.add_patch(mpatches.FancyBboxPatch((0.3, 0.2), 12.4, 1.15,
                 boxstyle="round,pad=0.1", facecolor='#FFFBEB', edgecolor=C['amber'], lw=1.5, zorder=1))
    ax.text(6.5, 1.03,
            "Holdout is split BEFORE training — 10% of records never seen during training (FT_ACORD_HOLDOUT_RATIO=0.10)",
            ha='center', va='center', fontsize=8.5, color=C['slate'])
    ax.text(6.5, 0.62,
            "FT_QG_MIN_JSON_VALID_RATE=0.80  •  FT_QG_MIN_JSON_FIELD_RECALL=0.40  •  FT_QG_MAX_OOS_HALLUCINATION_RATE=0.60",
            ha='center', va='center', fontsize=8, color=C['gray'], style='italic')

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.1)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 7 — Domain Isolation
# ══════════════════════════════════════════════════════════════════════════════
def gen_domain_isolation(path):
    fig = plt.figure(figsize=(13, 6.5), facecolor='white')
    ax  = fig.add_axes([0.01, 0.01, 0.98, 0.98])
    ax.set_xlim(0, 13); ax.set_ylim(0, 6.5); ax.axis('off')

    ax.text(6.5, 6.15, "Domain Isolation — Each Domain Is Completely Separate",
            ha='center', va='center', fontsize=12, fontweight='bold', color=C['blue_dark'])

    domains = [
        (2.2,  "broker",       C['blue'],   C['blue_xlight'],   ["ACORD form data",    "v1.0, v1.1, v1.2"]),
        (6.5,  "claims",       C['green'],  C['green_light'],   ["Claims POD data",    "v1.0, v1.1"]),
        (10.8, "underwriter",  C['purple'], C['purple_light'],  ["Underwriting data",  "v1.0"]),
    ]
    for cx, name, edge, fill, details in domains:
        W, H = 3.8, 4.5
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx-W/2, 0.7), W, H, boxstyle="round,pad=0.1",
            facecolor=fill, edgecolor=edge, linewidth=2.5, zorder=2))
        ax.add_patch(mpatches.Rectangle((cx-W/2, 0.7+H-0.5), W, 0.5, facecolor=edge, zorder=3))
        ax.text(cx, 0.7+H-0.25, f"Domain: {name}", ha='center', va='center',
                fontsize=10, fontweight='bold', color='white', zorder=4)

        rows = [
            ("SeaweedFS path:", f"my-bucket/{name}/"),
            ("Training data:", details[0]),
            ("Versions:", details[1]),
            ("Device JWTs:", f"scoped to {name}"),
            ("Data leaks to\nother domains:", "NEVER"),
        ]
        for ri, (label, val) in enumerate(rows):
            y = 0.7+H-0.5-0.62*(ri+1)
            ax.text(cx-W/2+0.2, y, label, fontsize=7.8, va='center', color=edge, fontweight='bold')
            color_val = C['red'] if val == "NEVER" else '#1e293b'
            bold_val  = val == "NEVER"
            ax.text(cx+0.0, y, val, fontsize=7.8, va='center', color=color_val,
                    fontweight='bold' if bold_val else 'normal', multialignment='center', ha='center')

    # No-crossing arrows
    for xa, xb in [(4.1, 4.6), (8.4, 8.9)]:
        ax.annotate('', xy=(xb, 2.95), xytext=(xa, 2.95),
                    arrowprops=dict(arrowstyle='<->', color=C['red'], lw=2), zorder=5)
        ax.text((xa+xb)/2, 3.2, "NO DATA\nCROSSING", ha='center', fontsize=7.5,
                color=C['red'], fontweight='bold', multialignment='center')
        circ = plt.Circle(((xa+xb)/2, 2.95), 0.28, color=C['red'], fill=False, lw=2, zorder=5)
        ax.add_patch(circ)
        ax.plot([(xa+xb)/2-0.2, (xa+xb)/2+0.2], [2.75, 3.15], color=C['red'], lw=2, zorder=6)

    ax.text(6.5, 0.35,
            "Each device JWT is scoped to one domain.  Training jobs for domain 'broker' never touch 'claims' data.  SeaweedFS paths are prefixed by domain.",
            ha='center', va='center', fontsize=8.5, color=C['slate'], style='italic',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFFBEB', edgecolor=C['amber'], lw=1.2))

    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.1)
    plt.close()


# ══════════════════════════════════════════════════════════════════════════════
# Generate all diagram images
# ══════════════════════════════════════════════════════════════════════════════
print("Generating diagrams...", flush=True)
P_PIPELINE   = _p("pipeline")
P_SEAWEEDFS  = _p("seaweedfs")
P_PROGRESSIVE= _p("progressive")
P_SAMPLING   = _p("sampling")
P_CANARY     = _p("canary")
P_QGATE      = _p("qgate")
P_TENANT     = _p("tenant")

gen_pipeline(P_PIPELINE)    ; print("  [1/7] pipeline done")
gen_seaweedfs(P_SEAWEEDFS)  ; print("  [2/7] seaweedfs done")
gen_progressive(P_PROGRESSIVE);print("  [3/7] progressive done")
gen_sampling(P_SAMPLING)    ; print("  [4/7] sampling done")
gen_canary(P_CANARY)        ; print("  [5/7] canary done")
gen_quality_gate(P_QGATE)   ; print("  [6/7] quality gate done")
gen_domain_isolation(P_TENANT); print("  [7/7] domain isolation done")


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCX
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ─── helpers ─────────────────────────────────────────────────────────────────
def heading(text, level=1, color=RGBColor(0x1a,0x56,0xdb)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        run.font.size = Pt(16 if level==1 else 13 if level==2 else 11)
    return p

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3+level*0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text); run.font.size = Pt(10.5)
    return p

def add_img(path, w=Inches(6.4)):
    doc.add_picture(path, width=w)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph('─'*90)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xd1,0xd5,0xdb); run.font.size = Pt(8)

def colored_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]; cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1A56DB')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        fill = 'EFF6FF' if ri%2==0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]; cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ─── TITLE ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
for txt, sz, col in [
    ("FIDEON OS",                                        28, RGBColor(0x1a,0x56,0xdb)),
    ("Fine-Tuning Pipeline Architecture",                20, RGBColor(0x37,0x47,0x51)),
    ("User Approval → Training → Quantization → Electron Download", 12, RGBColor(0x6b,0x72,0x80)),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(sz); r.font.bold = sz >= 20
    r.font.color.rgb = col
doc.add_paragraph(); divider(); doc.add_paragraph()

# ─── SECTION 1: OVERVIEW ─────────────────────────────────────────────────────
heading("1. Overview", 1)
para("Fideon OS uses a continuous fine-tuning loop. Every time a user approves AI-extracted "
     "data in the Electron app, that approval becomes training data. Once enough approvals "
     "are collected, the system automatically trains a smarter model, quantizes it into multiple "
     "GGUF variants, stores everything in SeaweedFS, and makes the right variant available for "
     "the user's device to download — all without manual intervention.")
doc.add_paragraph()

heading("1a. What Is Stored at Each Approval", 2)
para("Three pieces of data are permanently stored together for every approved run:", size=11)
doc.add_paragraph()
colored_table(
    ["Data", "Where Stored", "Column / Path", "Used For"],
    [
        ["Source file (ACORD PDF)",    "SeaweedFS",       "source_forms/<run_id>.pdf",              "Audit trail, re-extraction if needed"],
        ["Displayed output (AI initial)", "Supabase",     "original_extracted_json",                "Baseline — what the AI showed before correction"],
        ["Corrected output (user truth)",  "Supabase",    "corrected_json  (feedback table)",       "Training label — the ground truth the model learns"],
    ],
    col_widths=[1.8, 1.4, 2.4, 2.3]
)
para("The training pair used for fine-tuning is:  INPUT = raw_text (from PDF)  →  OUTPUT = corrected_json. "
     "The original AI output is preserved separately so we can measure how much the user changed.", size=10)
doc.add_paragraph()

# ─── SECTION 2: FULL PIPELINE ────────────────────────────────────────────────
heading("2. Full Pipeline Flow", 1)
para("The diagram below shows all five phases of the pipeline, from user approval to offline "
     "model deployment on the Electron device.", size=11)
doc.add_paragraph()
add_img(P_PIPELINE, w=Inches(6.4))

# ─── SECTION 3: SEAWEEDFS ────────────────────────────────────────────────────
heading("3. SeaweedFS Storage Layout", 1)
para("Two file types are stored per version: PyTorch weights for re-training, and GGUF for inference.", size=11)
doc.add_paragraph()
add_img(P_SEAWEEDFS, w=Inches(6.4))
para("KEY RULES:", bold=True, color=RGBColor(0xdc,0x26,0x26))
bullet("GGUF = inference only. Electron downloads ONE variant based on device RAM and uses via Ollama.")
bullet("safetensors = training only. Never sent to Electron. Used as starting point for next fine-tune.")
bullet("3 GGUF variants stored per version: F16 (high-end) / Q8_0 (mid-range) / Q4_K_M (low-end).")
bullet("Source ACORD PDFs stored under source_forms/<run_id>.pdf — permanent audit trail.")
bullet("Each domain has its own folder in SeaweedFS. Data never crosses domains.")
doc.add_paragraph()

# ─── SECTION 4: WHY TWO FORMATS ──────────────────────────────────────────────
heading("4. Why Multiple File Formats?", 1)
doc.add_paragraph()
colored_table(
    ["Format", "File", "Size (14B model)", "Purpose", "Who Uses It"],
    [
        ["PyTorch safetensors", "model.safetensors",  "2–10 GB",  "Next fine-tune starting point",                    "Server/RunPod only"],
        ["GGUF F16",           "model-f16.gguf",      "~28 GB",   "Full precision inference — high-end devices",       "Electron (≥32 GB RAM)"],
        ["GGUF Q8_0",          "model-q8_0.gguf",     "~15 GB",   "8-bit quantized — mid-range devices",               "Electron (16–32 GB RAM)"],
        ["GGUF Q4_K_M",        "model-q4_k_m.gguf",   "~8 GB",    "4-bit quantized — budget/laptop devices",           "Electron (<16 GB RAM)"],
    ],
    col_widths=[1.5, 1.8, 1.4, 2.2, 1.5]
)
para("GGUF cannot be fine-tuned — it has no gradient support. safetensors preserves full "
     "precision weights needed for continued training. The Electron app reads available RAM "
     "from device configuration settings and downloads only the appropriate GGUF variant — "
     "not all three.", size=11)
doc.add_paragraph()

# ─── SECTION 5: PROGRESSIVE FINE-TUNING ──────────────────────────────────────
heading("5. Progressive Fine-Tuning — Each Cycle Builds on Previous", 1)
para("The system never goes back to the base model after the first training cycle. "
     "Each cycle downloads the previous safetensors from SeaweedFS and trains on top.", size=11)
doc.add_paragraph()
add_img(P_PROGRESSIVE, w=Inches(6.4))

# ─── SECTION 6: STRATIFIED SAMPLING ──────────────────────────────────────────
heading("6. Stratified Data Sampling — Handles Growing Dataset", 1)
para("After months of use, tens of thousands of approvals accumulate. Training on all data "
     "would be slow and expensive. Smart sampling keeps training time constant.", size=11)
doc.add_paragraph()
add_img(P_SAMPLING, w=Inches(6.3))
doc.add_paragraph()
colored_table(
    ["Bucket", "% of Training", "Source", "Why"],
    [
        ["Recent data",   "60%", "Approvals from last 90 days",            "Learn current patterns and new form types"],
        ["Hard examples", "30%", "Records where user corrected AI output",  "Never forget edge cases and mistakes"],
        ["OOS refusals",  "10%", "Out-of-scope examples (code-embedded)",   "Prevent hallucination on unknown inputs"],
    ],
    col_widths=[1.5, 1.2, 2.8, 2.4]
)

# ─── SECTION 7: DOMAIN ISOLATION ─────────────────────────────────────────────
heading("7. Domain Isolation", 1)
para("Every domain has completely separate data, models, and access tokens. "
     "No training data or model weights ever cross domain boundaries.", size=11)
doc.add_paragraph()
add_img(P_TENANT, w=Inches(6.4))

# ─── SECTION 8: CANARY ROLLOUT ────────────────────────────────────────────────
heading("8. Canary Rollout — Safe Model Updates", 1)
para("New model versions roll out gradually. This prevents a bad model from reaching all users "
     "at once. The gate is deterministic — the same device always gets the same answer.", size=11)
doc.add_paragraph()
add_img(P_CANARY, w=Inches(6.3))

# ─── SECTION 9: QUALITY GATE ─────────────────────────────────────────────────
heading("9. Quality Gate — Never Ship a Bad Model", 1)
para("After every training cycle, the model is evaluated against a holdout split before being "
     "allowed into SeaweedFS. A failure aborts the release and keeps the previous version active.", size=11)
doc.add_paragraph()
add_img(P_QGATE, w=Inches(6.4))
doc.add_paragraph()
colored_table(
    ["Check", "Threshold", "What it catches"],
    [
        ["JSON valid rate",        ">= 80%", "Model outputting malformed JSON"],
        ["Field recall",           ">= 30%", "Model missing required fields"],
        ["Exact match rate",       ">= 20%", "Model getting field values completely wrong"],
        ["Extra field rate",       "<= 25%", "Model hallucinating non-existent fields"],
        ["OOS hallucination rate", "<= 60%", "Model answering questions it should refuse"],
    ],
    col_widths=[2.2, 1.2, 3.5]
)

# ─── SECTION 10: CONFIG ──────────────────────────────────────────────────────
heading("10. Key Configuration Variables", 1)
doc.add_paragraph()
colored_table(
    ["Variable", "Default", "Purpose"],
    [
        ["ACORD_TRAINING_MIN_RECORDS",     "19",    "Min approvals before training triggers"],
        ["FT_MAX_TRAINING_ROWS",          "1500",  "Hard cap on training rows (sampling)"],
        ["FT_RECENT_DAYS",                "90",    "Defines 'recent' for 60% bucket"],
        ["FT_ACORD_HOLDOUT_RATIO",        "0.10",  "10% held out for quality gate evaluation"],
        ["FT_QG_MIN_JSON_VALID_RATE",     "0.80",  "Quality gate: JSON parse success rate"],
        ["FT_QG_MIN_JSON_FIELD_RECALL",   "0.40",  "Quality gate: required field coverage"],
        ["FT_QG_MAX_OOS_HALLUCINATION_RATE","0.60","Quality gate: out-of-scope refusal rate"],
        ["SEAWEEDFS_ENDPOINT",            "...",   "Storage server for models and weights"],
        ["SEAWEEDFS_BUCKET",              "my-bucket","Root bucket for all domain data"],
    ],
    col_widths=[3.2, 1.0, 3.0]
)

# ─── SECTION 11: BUILD STATUS ────────────────────────────────────────────────
heading("11. Build Status", 1)
doc.add_paragraph()
colored_table(
    ["Component", "Status", "Notes"],
    [
        ["User approval flow",                   "Built",   "Electron + backend routes"],
        ["Training data export (JSONL)",          "Built",   "export_approved_acord_dataset.py"],
        ["QLoRA fine-tuning",                    "Built",   "train_qlora_chat.py"],
        ["Quality gate",                         "Built",   "quality_gate.py"],
        ["GGUF quantization (single variant)",   "Built",   "Quantization/quantize.py"],
        ["SeaweedFS upload (GGUF)",               "Built",   "Quantization/upload.py"],
        ["adapter_registry in Supabase",          "Built",   "Migration applied"],
        ["Backend adapter API",                   "Built",   "adapter_registry.py routes"],
        ["Electron download + Ollama install",    "Built",   "model-updater.ts"],
        ["Canary rollout gate",                   "Built",   "canary_pct in adapter_registry"],
        ["Store source PDF to SeaweedFS",         "Needed",  "Upload ACORD PDF at upload time, save source_file_url"],
        ["Store displayed + corrected output",    "Needed",  "original_extracted_json + corrected_json linked per run"],
        ["3 GGUF variants (F16/Q8/Q4)",          "Needed",  "Run quantize 3x, upload all 3 to SeaweedFS"],
        ["Device RAM check + variant selection",  "Needed",  "Electron reads RAM, picks F16/Q8/Q4 before download"],
        ["safetensors upload to SeaweedFS",       "Needed",  "Save merged weights before quantize step"],
        ["Pull safetensors for next cycle",       "Needed",  "Download from SeaweedFS before training"],
        ["Stratified data sampling",              "Needed",  "Replace simple limit with 60/30/10 split"],
        ["Hard example flagging in DB",           "Needed",  "Mark corrected records as is_hard_example"],
        ["Auto canary progression",               "Needed",  "Bump canary_pct after N days no issues"],
    ],
    col_widths=[2.8, 1.1, 3.1]
)

# ─── SECTION 12: SIMPLE SUMMARY ──────────────────────────────────────────────
heading("12. Simple Summary", 1)
doc.add_paragraph()
steps = [
    ("1", "User uploads ACORD form, AI extracts fields",     "Source PDF saved to SeaweedFS. AI output (displayed_output) saved to Supabase."),
    ("2", "User corrects mistakes and clicks Approve",       "corrected_json saved. All 3 records linked: source PDF + displayed + corrected."),
    ("3", "Enough approvals collected (19+)",                "System starts training job on RunPod GPU"),
    ("4", "Training starts from PREVIOUS fine-tuned weights","Not from scratch — downloads safetensors from SeaweedFS"),
    ("5", "New model trained, quality checked, quantized",   "Produces safetensors + 3 GGUF variants: F16 / Q8_0 / Q4_K_M"),
    ("6", "All files uploaded to SeaweedFS",                 "New version registered in adapter_registry"),
    ("7", "Electron user opens My Models page",              "Banner: 'Model update available vN.N — Download & Install'"),
    ("8", "App reads device RAM, selects GGUF variant",      "≥32 GB=F16 | 16-32 GB=Q8_0 | <16 GB=Q4_K_M"),
    ("9", "User clicks Download",                            "Only the matching GGUF downloaded from SeaweedFS, installed into Ollama"),
]
t = doc.add_table(rows=len(steps), cols=3)
t.style = 'Table Grid'
for i, (num, action, result) in enumerate(steps):
    fill = '1A56DB' if i%2==0 else '1E429F'
    c0 = t.rows[i].cells[0]; c0.text = num
    c0.paragraphs[0].runs[0].font.bold = True
    c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
    c0.paragraphs[0].runs[0].font.size = Pt(14)
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = c0._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    tcPr.append(shd); c0.width = Inches(0.4)
    c1 = t.rows[i].cells[1]; c1.text = action
    c1.paragraphs[0].runs[0].font.bold = True; c1.paragraphs[0].runs[0].font.size = Pt(10)
    tc1 = c1._tc; tcPr1 = tc1.get_or_add_tcPr()
    shd1 = OxmlElement('w:shd'); shd1.set(qn('w:val'),'clear'); shd1.set(qn('w:color'),'auto'); shd1.set(qn('w:fill'),'EFF6FF')
    tcPr1.append(shd1)
    c2 = t.rows[i].cells[2]; c2.text = result
    c2.paragraphs[0].runs[0].font.size = Pt(10); c2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x37,0x47,0x51)
    tc2 = c2._tc; tcPr2 = tc2.get_or_add_tcPr()
    shd2 = OxmlElement('w:shd'); shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto')
    shd2.set(qn('w:fill'),'FFFFFF' if i%2==0 else 'F8FAFC'); tcPr2.append(shd2)

doc.add_paragraph(); divider()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = p.add_run("Fideon OS — Confidential — Fine-Tuning Pipeline Architecture v2.0")
fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x9c,0xa3,0xaf); fr.font.italic = True

out = "Fideon_FineTuning_Pipeline_Architecture.docx"
doc.save(out)
print(f"Saved: {out}")
