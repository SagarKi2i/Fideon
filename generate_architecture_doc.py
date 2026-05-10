"""
Generate architecture documentation DOCX for Fideon OS.
Reflects the existing codebase — excludes changes added in the current sprint.
"""

import io
import textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h


def para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


# ── Sequence Diagram ──────────────────────────────────────────────────────────

def draw_signup_sequence() -> bytes:
    ACTOR_LABELS = [
        "User\n(Browser)",
        "Next.js\nFrontend",
        "FastAPI\nBackend",
        "Azure\nAPIM",
        "Supabase\nGoTrue",
        "PostgreSQL\n+ Trigger",
    ]
    ACTOR_COLORS = ["#1D4ED8", "#047857", "#B45309", "#B91C1C", "#6D28D9", "#9D174D"]

    N        = len(ACTOR_LABELS)
    COL_GAP  = 3.0
    LEFT_PAD = 1.6
    xs       = [LEFT_PAD + i * COL_GAP for i in range(N)]
    FIG_W    = xs[-1] + LEFT_PAD + 0.2

    BOX_W    = 2.4
    BOX_H    = 0.70
    ROW_H    = 0.90
    NOTE_H   = 0.50

    STEPS = [
        (0, 1, "1  Fill signup form  (email · password · plan · device)",    False),
        (1, 2, "2  POST /api/v1/auth/signup   {email, password, metadata}",  False),
        (2, 2, "3  Validate input · rate-limit (5 req / min per IP)",        None),
        (2, 3, "4  Proxy request → APIM",                                    False),
        (3, 4, "5  POST /auth/v1/signup",                                    False),
        (4, 4, "6  Create  auth.users  row",                                 None),
        (4, 5, "7  Fire trigger:  on_auth_user_created()",                   False),
        (5, 5, "8  Insert → app_users · tenants · user_roles\n"
               "   activated_models · devices",                              None),
        (4, 3, "9  Return  {access_token,  refresh_token,  user}",           True),
        (3, 2, "10  Forward tokens",                                         True),
        (2, 2, "11  insert_auth_audit_row()   [SHA-256 · immutable]",        None),
        (2, 1, "12  Return  {access_token,  refresh_token,  user}",          True),
        (1, 1, "13  Store session  (Supabase JS · IndexedDB)",               None),
        (1, 2, "14  PATCH /api/v1/auth/profile/name   {full_name}",          False),
        (2, 3, "15  PATCH app_users.full_name  via PostgREST",               False),
        (3, 2, "16  200 OK  {success}",                                      True),
        (2, 1, "17  200 OK  {success}",                                      True),
        (1, 0, "18  Navigate  →  /dashboard",                                True),
    ]

    N_ROWS = len(STEPS)
    FIG_H  = BOX_H + 0.4 + N_ROWS * ROW_H + 0.5

    fig, ax = plt.subplots(figsize=(FIG_W * 0.82, FIG_H * 0.82))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, FIG_W)
    ax.set_ylim(FIG_H, 0)
    ax.axis("off")

    TOP_Y = 0.3
    for i, (label, color) in enumerate(zip(ACTOR_LABELS, ACTOR_COLORS)):
        x = xs[i]
        ax.add_patch(mpatches.FancyBboxPatch(
            (x - BOX_W / 2, TOP_Y), BOX_W, BOX_H,
            boxstyle="round,pad=0.08",
            linewidth=2, edgecolor=color, facecolor=color, zorder=3
        ))
        ax.text(x, TOP_Y + BOX_H / 2, label,
                ha="center", va="center", fontsize=9, fontweight="bold",
                color="white", linespacing=1.35, zorder=4)
        ax.plot([x, x], [TOP_Y + BOX_H, FIG_H - 0.2],
                color=color, linewidth=1.4, linestyle=(0, (6, 4)),
                alpha=0.5, zorder=1)

    for row_i, (fc, tc, label, is_return) in enumerate(STEPS):
        y = TOP_Y + BOX_H + 0.4 + (row_i + 0.5) * ROW_H
        fc_color = ACTOR_COLORS[fc]

        if is_return is None:
            x     = xs[fc]
            lines = label.split("\n")
            bh    = NOTE_H + 0.18 * (len(lines) - 1)
            bw    = min(BOX_W * 1.2, 4.0)
            fills = ["#EFF6FF","#ECFDF5","#FFFBEB","#FEF2F2","#F5F3FF","#FDF2F8"]
            ax.add_patch(mpatches.FancyBboxPatch(
                (x - bw / 2, y - bh / 2), bw, bh,
                boxstyle="round,pad=0.06",
                linewidth=1.5, edgecolor=fc_color, facecolor=fills[fc], zorder=3
            ))
            ax.text(x, y, label,
                    ha="center", va="center", fontsize=8.2,
                    color="#111827", linespacing=1.4, zorder=4)
        else:
            x1, x2 = xs[fc], xs[tc]
            linestyle = (0, (5, 3)) if is_return else "solid"
            ax.annotate(
                "", xy=(x2, y), xytext=(x1, y),
                arrowprops=dict(
                    arrowstyle="-|>",
                    color=fc_color, lw=1.8,
                    linestyle=linestyle,
                    mutation_scale=13,
                ), zorder=2
            )
            mid_x = (x1 + x2) / 2
            ax.text(mid_x, y - 0.17, label,
                    ha="center", va="bottom", fontsize=8.4,
                    color="#0F172A",
                    bbox=dict(boxstyle="round,pad=0.22", fc="white",
                              ec=fc_color, lw=1.0, alpha=0.96),
                    zorder=5)

    ax.set_title("User Signup — Sequence Diagram  ·  Fideon OS",
                 fontsize=13, fontweight="bold", color="#1E3A5F", pad=10)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ── Infrastructure Overview Diagram ───────────────────────────────────────────

def draw_infra_overview() -> bytes:
    fig, ax = plt.subplots(figsize=(22, 14))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 22)
    ax.set_ylim(0, 14)
    ax.axis("off")

    lanes = [
        (0.2,  3.9,  "CLIENT",          "#EFF6FF", "#93C5FD"),
        (4.3,  9.5,  "AZURE CLOUD",     "#F0FDF4", "#6EE7B7"),
        (14.1, 7.6,  "DATA & SERVICES", "#FFF7ED", "#FCD34D"),
    ]
    for lx, lw, ltitle, lbg, lborder in lanes:
        ax.add_patch(mpatches.FancyBboxPatch(
            (lx, 0.4), lw, 13.0,
            boxstyle="round,pad=0.15",
            linewidth=2, edgecolor=lborder, facecolor=lbg, zorder=0
        ))
        ax.text(lx + lw / 2, 13.25, ltitle,
                ha="center", va="bottom", fontsize=10, fontweight="bold",
                color="#374151")

    def box(cx, cy, w, h, title, subtitle, fill, edge):
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx, cy), w, h,
            boxstyle="round,pad=0.12",
            linewidth=2.2, edgecolor=edge, facecolor=fill, zorder=2
        ))
        lines = subtitle.split("\n")
        title_y = cy + h / 2 + (0.22 if subtitle else 0)
        ax.text(cx + w / 2, title_y, title,
                ha="center", va="center", fontsize=10, fontweight="bold",
                color=edge, zorder=3)
        if subtitle:
            for k, sl in enumerate(lines):
                ax.text(cx + w / 2,
                        cy + h / 2 - 0.18 - k * 0.27,
                        sl,
                        ha="center", va="center", fontsize=8.2,
                        color="#374151", zorder=3)

    BOXES = [
        # CLIENT
        (0.5,  6.0,  3.2, 1.4,
         "User",             "Browser / Desktop",          "#DBEAFE", "#1D4ED8"),
        # AZURE CLOUD
        (4.6,  10.0, 3.6, 1.4,
         "Next.js Frontend", "Azure App Service  ·  :443", "#D1FAE5", "#047857"),
        (4.6,   7.2, 3.6, 1.4,
         "FastAPI Backend",  "Azure App Service  ·  :8080","#FEF3C7", "#B45309"),
        (4.6,   1.8, 3.6, 1.4,
         "GitHub Actions",   "CI / CD Pipelines",          "#E2E8F0", "#475569"),
        (9.4,  10.0, 3.6, 1.4,
         "Azure ACR",        "neurapodacr.azurecr.io",     "#EDE9FE", "#5B21B6"),
        (9.4,   7.2, 3.6, 1.4,
         "Azure APIM",       "API Gateway\nRate Limiter",  "#FEE2E2", "#B91C1C"),
        (9.4,   3.8, 3.6, 1.4,
         "RunPod GPU",       "LLM Inference\nFine-tuning", "#FEF9C3", "#92400E"),
        # DATA & SERVICES
        (14.5, 11.5, 6.8, 1.4,
         "Supabase GoTrue",  "Auth  ·  JWT  ·  Sessions",  "#F3E8FF", "#6D28D9"),
        (14.5,  9.2, 6.8, 1.8,
         "PostgreSQL",
         "Supabase DB  ·  pgvector\nRow-Level Security (RLS)\nMulti-tenant schema",
         "#FCE7F3", "#9D174D"),
        (14.5,  7.5, 6.8, 1.2,
         "Supabase Storage", "PDFs  ·  ACORD forms  ·  POD files", "#D1FAE5", "#047857"),
        (14.5,  5.6, 6.8, 1.4,
         "SeaweedFS",
         "GGUF adapters  ·  fine-tuned models\n[DEV/TESTING — staging planned]",
         "#FEF9C3", "#92400E"),
        (14.5,  3.5, 6.8, 1.2,
         "DB Migrations",    "supabase/migrations/*.sql\nApplied via db-migrate.yml",
         "#FEF3C7", "#B45309"),
    ]
    for b in BOXES:
        box(*b)

    # DEV-only badge on SeaweedFS box
    ax.text(21.2, 6.55, "DEV",
            ha="center", va="center", fontsize=7, fontweight="bold",
            color="white",
            bbox=dict(boxstyle="round,pad=0.18", fc="#92400E", ec="none"),
            zorder=6)

    def arrow(x1, y1, x2, y2, label, color, bidir=False, rad=0.0):
        style = "<|-|>" if bidir else "-|>"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(
                        arrowstyle=style, color=color, lw=1.8,
                        mutation_scale=12,
                        connectionstyle=f"arc3,rad={rad}"
                    ), zorder=1)
        if label:
            mx = (x1 + x2) / 2
            my = (y1 + y2) / 2
            ax.text(mx, my, label,
                    ha="center", va="center", fontsize=8,
                    color="white", fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.25", fc=color,
                              ec="none", alpha=0.90),
                    zorder=4)

    arrow(3.7,  6.9,   4.6, 10.7,  "HTTPS",           "#1D4ED8", bidir=True)
    arrow(3.7,  6.6,   4.6,  7.9,  "HTTPS",           "#B45309", bidir=True)
    arrow(6.4, 10.0,   6.4,  8.6,  "API /api/v1/*",   "#047857", bidir=True)
    arrow(8.2,  7.9,   9.4,  7.9,  "Supabase proxy",  "#B45309")
    arrow(13.0, 8.2,  14.5, 12.2,  "Auth calls",      "#B91C1C")
    arrow(13.0, 7.9,  14.5, 10.1,  "PostgREST",       "#B91C1C")
    arrow(8.2,  7.5,   9.4,  4.5,  "LLM / fine-tune", "#92400E")
    arrow(8.2,  7.2,  14.5,  8.1,  "Upload docs",     "#047857")
    # RunPod → SeaweedFS (store fine-tuned adapters)
    arrow(13.0, 4.5,  14.5,  6.3,  "Save adapters",   "#92400E")
    arrow(6.4,  3.2,  11.2, 10.0,  "Push image",      "#475569")
    arrow(9.4, 10.7,   8.2, 11.4,  "Deploy",          "#5B21B6")
    arrow(9.4, 10.2,   8.2,  8.6,  "Deploy",          "#5B21B6")
    arrow(17.9, 4.7,  17.9,  9.2,  "Apply SQL",       "#B45309")

    ax.set_title("Infrastructure Overview — Fideon OS",
                 fontsize=15, fontweight="bold", color="#1E3A5F", pad=12)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ── Document Builder ─────────────────────────────────────────────────────────

def build_doc() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin  = section.bottom_margin = Cm(1.8)

    # ── Cover ──
    doc.add_paragraph()
    title = doc.add_heading("Fideon OS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Architecture & Developer Reference")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()
    meta = doc.add_paragraph("Version: 1.1   |   Branch: v1-dev   |   April 2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_page_break()

    # ── TOC ──
    heading(doc, "Table of Contents", 1)
    toc_items = [
        "1.  Technology Stack Overview",
        "2.  Backend: Database Configuration",
        "3.  Frontend: Backend Connection Configuration",
        "4.  Signup Flow — Sequence Diagram & Walkthrough",
        "5.  Infrastructure Overview Diagram",
        "6.  Backend Routes Reference",
        "7.  Frontend Pages Reference",
        "8.  CI/CD Pipeline Summary",
        "9.  Environment Variables Reference",
        "10. Security Review & Open Items",
    ]
    for item in toc_items:
        bullet(doc, item)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. TECH STACK
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "1. Technology Stack Overview", 1)
    para(doc, (
        "Fideon OS is a multi-tenant SaaS platform for AI-powered insurance form "
        "extraction, device management, and fine-tuning workflows. It is built on:"
    ))

    stack_rows = [
        ("Frontend",        "Next.js 14.2 + React + TypeScript",                 "SPA, routing, SSR support"),
        ("UI",              "Tailwind CSS + shadcn/ui",                           "Design system, component library"),

        ("Backend",         "FastAPI + Uvicorn/Gunicorn",                         "REST API, business logic, auth proxy layer"),
        ("Auth",            "Supabase GoTrue (JWT)",                             "Signup, login, token refresh, RLS"),
        ("Database",        "Supabase PostgreSQL 14+",                           "Primary data store, multi-tenant"),
        ("DB API",          "PostgREST (via Supabase)",                          "REST interface over PostgreSQL"),
        ("Vector DB",       "pgvector extension",                                "Embeddings, semantic search (backend/LLM/vectorstore/)"),
        ("Gateway",         "Azure API Management (APIM)",                       "Proxies all Supabase traffic, rate limiting"),
        ("Container",       "Docker + Azure App Service",                        "Deployment"),
        ("Registry",        "Azure Container Registry (ACR)",                    "Docker image storage (neurapodacr.azurecr.io)"),
        ("CI/CD",           "GitHub Actions",                                    "Build, lint, test, SonarQube, deploy"),
        ("LLM",             "Groq API (llama-3.3-70b-versatile)",                "Currently using Groq API key for help/assistance features as interim solution until Fideon OS own fine-tuned LLM model is production-ready. Full litellm fallback chain (RunPod, Gemini, OpenAI, Anthropic) is configured but Groq is primary."),
        ("GPU / Fine-tune", "RunPod + PyTorch + transformers + peft",            "Model training and inference"),

    ]
    add_table(doc, ["Layer", "Technology", "Notes"], stack_rows, col_widths=[1.3, 2.5, 2.8])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # 2. BACKEND DB CONFIG
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "2. Backend: Database Configuration", 1)

    heading(doc, "2.1 Overview", 2)
    para(doc, (
        "The backend does NOT use an ORM (e.g., SQLAlchemy). It communicates with PostgreSQL "
        "entirely through Supabase's PostgREST HTTP API and GoTrue auth service. All HTTP calls "
        "are asynchronous using httpx. The Supabase instance is accessed via the Azure APIM proxy."
    ))

    heading(doc, "2.2 Core Module: backend/app/core/supabase.py", 2)
    para(doc, "Provides all low-level DB interaction primitives:", bold=True)
    db_fns = [
        ("service_headers()",               "Returns {'Authorization': 'Bearer <SERVICE_ROLE_KEY>', 'apikey': ...}"),
        ("verify_user(token)",              "Calls GET /auth/v1/user — validates Bearer token, returns user dict"),
        ("get_user_context(token)",         "Calls verify_user(), looks up app_users + user_roles, returns dict with user_id, role, tenant_id"),
        ("postgrest_get(table, params)",    "GET /rest/v1/<table> with PostgREST query params (filters, selects)"),
        ("postgrest_insert(table, data)",   "POST /rest/v1/<table> — insert row(s)"),
        ("postgrest_patch(table, f, data)", "PATCH /rest/v1/<table>?<filter> — update rows"),
        ("postgrest_delete(table, filter)", "DELETE /rest/v1/<table>?<filter> — delete rows"),
        ("verify_admin(token)",             "Calls verify_user(), then checks user_roles for admin/global_admin. Raises 403 if not found. Used by admin-only routes."),
        ("insert_auth_audit_row(...)",      "Appends to auth_audit with SHA-256 integrity hash (immutable)"),
        ("insert_audit_log(...)",           "Appends to audit_logs with integrity hash (immutable)"),
    ]
    for fn, desc in db_fns:
        bullet(doc, f"{fn} — {desc}")

    heading(doc, "2.3 Environment Variables", 2)
    code_block(doc, textwrap.dedent("""\
        # Supabase connection (via APIM proxy)
        SUPABASE_URL=https://fideon-staging-apim.azure-api.net/supabase

        # JWT keys (from Supabase project settings)
        SUPABASE_ANON_KEY=<anon_jwt>            # public, used by frontend
        SUPABASE_SERVICE_ROLE_KEY=<service_jwt> # private, backend only (bypasses RLS)

        # Direct PostgreSQL connection (pgvector / psycopg3)
        PGVECTOR_DATABASE_URL=postgresql://user:pass@host:5432/postgres?sslmode=require
    """))

    heading(doc, "2.4 Key Database Tables", 2)
    table_rows = [
        ("auth.users",          "Supabase-managed, stores credentials + metadata"),
        ("app_users",           "User profiles: tenant_id, full_name, status, onboarding metadata"),
        ("user_roles",          "user_id → role (global_admin, admin, user, viewer, guest)"),
        ("tenants",             "Multi-tenant organisations (name, slug, plan, seat_limit)"),
        ("devices",             "Registered RPA devices (device_token, heartbeat, status)"),
        ("activated_models",    "Per-user active AI models from catalog"),
        ("model_catalog",       "Available AI models (domain, provider, version)"),
        ("audit_logs",          "Immutable audit trail (EU AI Act / SOC2 compliance)"),
        ("auth_audit",          "Immutable auth events (login, logout, signup)"),
        ("acord_training_jobs", "ACORD form fine-tuning job records"),
        ("pod_training_jobs",   "Custom form (POD) fine-tuning job records"),
        ("workflows",           "Agent workflow definitions"),
        ("documents",           "Uploaded PDFs/forms (references to Supabase Storage)"),
        ("rag_chunks",          "pgvector embeddings for semantic search"),
    ]
    add_table(doc, ["Table", "Purpose"], table_rows, col_widths=[2.0, 4.5])
    doc.add_paragraph()

    heading(doc, "2.5 Row-Level Security (RLS)", 2)
    para(doc, (
        "All tables have PostgreSQL Row-Level Security policies. The service role key bypasses "
        "RLS for backend operations. The anon key (frontend Supabase JS client) is restricted "
        "by RLS — users can only see their own tenant's data."
    ))

    heading(doc, "2.6 Audit Log Integrity", 2)
    para(doc, (
        "Every write to audit_logs and auth_audit includes a SHA-256 integrity hash computed "
        "over (user_id, action, resource_type, resource_id, created_at). Database triggers "
        "prevent UPDATE and DELETE on these tables. This satisfies EU AI Act Article 12, "
        "SOC2 CC7.2/CC9.1, and NAIC requirements."
    ))

    heading(doc, "2.7 Custom Access Token Hook (JWT Role Embedding)", 2)
    para(doc, "Migration: supabase/migrations/20260420100000_custom_access_token_hook.sql", bold=True)
    para(doc, (
        "A PostgreSQL function public.custom_access_token_hook(event jsonb) is registered as a "
        "Supabase Custom Access Token hook. Every JWT issued by GoTrue passes through this hook. "
        "The hook looks up the user's role in user_roles and embeds it into the token's app_metadata:"
    ))
    code_block(doc, textwrap.dedent("""\
        -- Embedded in every JWT after hook activation:
        {
          "app_metadata": {
            "role": "global_admin" | "admin" | "user" | "viewer" | "guest"
          }
        }
    """))
    bullet(doc, "Frontend reads role from token — zero extra HTTP calls on page load")
    bullet(doc, "Backward compatible — useUserRole HTTP fallback still works for pre-hook tokens")
    para(doc, "Activation steps (one-time, per environment):", bold=True)
    bullet(doc, "1. Run the migration (db-migrate.yml or Supabase CLI)")
    bullet(doc, "2. Supabase Dashboard → Authentication → Hooks")
    bullet(doc, "3. Enable 'Custom Access Token' hook, select: public.custom_access_token_hook")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. FRONTEND → BACKEND CONFIG
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "3. Frontend: Backend Connection Configuration", 1)

    heading(doc, "3.1 Environment Variables", 2)
    code_block(doc, textwrap.dedent("""\
        # frontend/.env.local  — LOCAL DEV ONLY
        NEXT_PUBLIC_SUPABASE_URL=https://fideon-staging-apim.azure-api.net/supabase
        NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY=<anon_jwt>
        NEXT_PUBLIC_API_URL=http://localhost:8000
        NEXT_PUBLIC_APP_BASE_URL=http://localhost:3000

        # Azure App Service (staging / production) — set as App Service env vars
        NEXT_PUBLIC_API_URL=https://<backend>.azurewebsites.net   # MUST be HTTPS
        NEXT_PUBLIC_APP_BASE_URL=https://<frontend>.azurewebsites.net
    """))

    heading(doc, "3.2 API Base URL Resolution (src/lib/apiBaseUrl.ts)", 2)
    para(doc, "Resolution order:", bold=True)
    bullet(doc, "1. NEXT_PUBLIC_API_URL environment variable (set at build/runtime)")
    bullet(doc, "2. Default fallback: http://127.0.0.1:8080  — local dev only")
    para(doc, "All API paths: apiUrl('/api/v1/...') → base + path")

    heading(doc, "3.3 Auth Token Management (src/lib/authHeader.ts)", 2)
    para(doc, "Called before every backend request:", bold=True)
    steps_fe = [
        "Get current Supabase session from IndexedDB via @supabase/supabase-js",
        "Decode JWT payload, check exp claim",
        "If expired or within 60s of expiry: POST /api/v1/auth/refresh (rate: 30/min)",
        "Update local Supabase session with new access_token + refresh_token",
        "Return header: { Authorization: 'Bearer <access_token>' }",
        "In-flight deduplication: concurrent refresh calls merged into one",
    ]
    for i, s in enumerate(steps_fe, 1):
        bullet(doc, f"{i}. {s}")

    heading(doc, "3.4 Backend API Wrapper (src/lib/backendApi.ts)", 2)
    para(doc, "All HTTP calls from the frontend go through this wrapper:", bold=True)
    code_block(doc, textwrap.dedent("""\
        backendFetch(path, options)       // raw fetch with 15s timeout + auth header
        backendGet<T>(path, options)      // GET + JSON parse
        backendPost<T>(path, body, ...)   // POST JSON body
        backendPatch<T>(path, body, ...)  // PATCH JSON body
        backendDelete<T>(path, ...)       // DELETE
    """))
    bullet(doc, "15-second default timeout (accommodates RunPod cold starts)")
    bullet(doc, "Auto-attaches Authorization: Bearer header via authHeader()")
    bullet(doc, "Combines external AbortSignal with internal timeout signal")
    bullet(doc, "No direct PostgREST calls — all DB access goes through FastAPI")

    heading(doc, "3.5 Supabase JS Client", 2)
    para(doc, (
        "The frontend initialises a Supabase JS client pointing at the APIM proxy URL. "
        "Used solely for session storage (IndexedDB) and realtime subscriptions. "
        "All mutating API calls go through FastAPI, not directly to PostgREST."
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. SIGNUP FLOW
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "4. Signup Flow — Sequence Diagram & Walkthrough", 1)

    heading(doc, "4.1 Signup Sequence Diagram", 2)
    seq_png = draw_signup_sequence()
    seq_buf = io.BytesIO(seq_png)
    doc.add_picture(seq_buf, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    heading(doc, "4.2 Step-by-Step Walkthrough", 2)
    steps_detail = [
        ("Step 1 — User fills the 3-step signup wizard (Frontend: Signup.tsx)",
         [
             "Step 0: Enter email + password. Frontend validates locally: RFC 5322 email, "
             "password >= 8 chars with 1 upper, 1 lower, 1 digit, 1 special char.",
             "Step 1: Choose subscription tier (Starter / Professional / Enterprise) + agent packs.",
             "Step 2: Optionally register a device (name, OS, app version).",
         ]),
        ("Step 2 — Frontend calls POST /api/v1/auth/signup (Backend: auth_proxy.py)",
         [
             "Request body: { email, password, metadata: { requested_plan, requested_model_id, "
             "device_name, device_profile } }",
             "Backend rate-limits to 5 signups/min per IP via slowapi.",
             "Backend validates email format and password strength.",
         ]),
        ("Step 3 — Backend proxies to Supabase GoTrue via APIM",
         [
             "POST {SUPABASE_URL}/auth/v1/signup with { email, password, data: metadata }",
             "Traffic: FastAPI → Azure APIM → Supabase GoTrue",
             "APIM adds subscription keys and enforces additional rate limits.",
         ]),
        ("Step 4 — Supabase creates auth.users row",
         [
             "GoTrue creates the user in auth.users with hashed password.",
             "Returns { access_token, refresh_token, expires_in, user: { id, email } }",
         ]),
        ("Step 5 — PostgreSQL trigger fires: on_auth_user_created()",
         [
             "Auto-creates app_users row",
             "Auto-creates tenants row if new organisation",
             "Creates user_roles row with default role = 'user'",
             "If requested_model_id in metadata → creates activated_models row",
             "If device_name in metadata → creates devices row with signed device JWT",
             "Seat limit enforced by trigger — returns error if tenant.seat_limit reached",
         ]),
        ("Step 6 — Backend writes audit record",
         [
             "Calls insert_auth_audit_row(user_id, 'signup', ...)",
             "Computes SHA-256 integrity hash over event fields",
             "Inserts into immutable auth_audit table (trigger blocks UPDATE/DELETE)",
         ]),
        ("Step 7 — Backend returns tokens to Frontend",
         [
             "Response: { access_token, refresh_token, expires_in, user, role }",
             "Frontend stores session via Supabase JS client (IndexedDB)",
         ]),
        ("Step 8 — Frontend updates display name",
         [
             "PATCH /api/v1/auth/profile/name with { full_name }",
             "Backend updates app_users.full_name via PostgREST PATCH",
         ]),
        ("Step 9 — Navigation to Dashboard",
         [
             "React Router navigates to / (Dashboard.tsx)",
             "All subsequent API calls include Authorization: Bearer <access_token>",
             "Token auto-refreshed via POST /api/v1/auth/refresh when near expiry",
         ]),
    ]
    for title, sub_steps in steps_detail:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        for s in sub_steps:
            bullet(doc, s, level=1)
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. INFRASTRUCTURE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "5. Infrastructure Overview Diagram", 1)
    infra_png = draw_infra_overview()
    infra_buf = io.BytesIO(infra_png)
    doc.add_picture(infra_buf, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    heading(doc, "5.1 Component Descriptions", 2)
    infra_rows = [
        ("User (Browser/Desktop)",  "Accesses the app via HTTPS. Connects to both Next.js frontend and FastAPI backend."),
        ("Next.js Frontend",        "Azure App Service (Linux Docker). Serves the React SPA. Port 443."),
        ("FastAPI Backend",         "Azure App Service (Linux Docker). REST API, auth proxy, extraction logic. Port 8080."),
        ("Azure APIM",              "API Management gateway. Proxies all Supabase traffic. Adds rate limiting and subscription auth."),
        ("Supabase GoTrue",         "Authentication service (signup, login, token refresh, session management)."),
        ("PostgreSQL (Supabase)",   "Primary database. Multi-tenant, Row-Level Security, pgvector for embeddings."),
        ("Azure ACR (neurapodacr)", "neurapodacr.azurecr.io — Docker image registry for frontend and backend containers."),
        ("GitHub Actions",          "CI/CD. Builds, lints, tests, SonarQube scan, pushes to ACR, deploys to App Service."),
        ("RunPod / GPU",            "On-demand GPU instances for LLM inference (llama, mistral) and fine-tuning jobs."),
        ("Supabase Storage",        "Object storage for uploaded PDFs, ACORD forms, POD templates."),
        ("SeaweedFS  [DEV only]",   "S3-compatible object store for GGUF model artifacts and fine-tuned LoRA adapters. "
                                    "Endpoint: http://20.40.61.106:8333. Implemented and tested on dev environment. "
                                    "Staging deployment pending — not yet promoted to staging/production."),
        ("DB Migrations",           "supabase/migrations/*.sql — applied via db-migrate.yml workflow."),
    ]
    add_table(doc, ["Component", "Description"], infra_rows, col_widths=[2.0, 4.5])
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Note — SeaweedFS storage:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    doc.add_paragraph(
        "SeaweedFS is currently implemented and running on the dev environment only "
        "(SEAWEEDFS_ENDPOINT configured in backend/.env). Fine-tuned LoRA adapters and GGUF "
        "model files produced by RunPod fine-tuning jobs are written here. "
        "Promotion to staging environment is planned for the next sprint."
    ).runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. BACKEND ROUTES
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "6. Backend Routes Reference", 1)
    route_groups = [
        ("Auth Proxy (auth_proxy.py)", [
            ("POST /api/v1/auth/signup",        "Create account (rate: 5/min). Proxies to GoTrue. Fires DB trigger."),
            ("POST /api/v1/auth/login",         "Email+password login (rate: 10/min). Returns JWT tokens."),
            ("POST /api/v1/auth/logout",        "Revoke session. Writes auth_audit row."),
            ("POST /api/v1/auth/refresh",       "Token refresh (rate: 30/min). Returns new access+refresh tokens."),
            ("PATCH /api/v1/auth/profile/name", "Update display name in app_users."),
        ]),
        ("Devices (device.py)", [
            ("POST /api/v1/devices/register",              "Register device, returns signed device JWT."),
            ("PUT  /api/v1/devices/heartbeat",             "Device keep-alive ping (PUT, not POST)."),
            ("GET  /api/v1/devices/models",                "List AI models available for this device."),
            ("POST /api/v1/devices/{device_id}/revoke",    "Revoke device token."),
            ("POST /api/v1/devices/link",                  "Link device to user account."),
            ("GET  /api/v1/admin/devices",                 "Admin: list all tenant devices."),
            ("GET  /api/v1/admin/devices/{device_id}",     "Admin: device detail."),
        ]),
        ("Device Pairing (device.py)", [
            ("POST /api/devices/pairing/start",                  "Initiate device pairing session."),
            ("GET  /api/devices/pairing/status/{pairing_id}",    "Poll pairing session status."),
            ("POST /api/devices/pairing/confirm",                "Confirm and complete pairing."),
        ]),
        ("ACORD Extraction (acord.py)", [
            ("POST /api/acord/parse",                        "Parse ACORD form structure from PDF."),
            ("POST /api/acord/extract",                      "Synchronous extract (returns immediately)."),
            ("POST /api/acord/extract/start",                "Async extract kick-off — returns job_id."),
            ("GET  /api/acord/extract/status/{job_id}",      "Poll async extraction job status."),
            ("GET  /api/acord/runs",                         "List extraction runs for tenant."),
            ("GET  /api/acord/runs/{run_id}",                "Single run detail."),
            ("POST /api/acord/runs/{run_id}/re-extract",     "Re-run extraction on existing run."),
            ("POST /api/acord/runs/{run_id}/submit",         "Submit run for human review."),
            ("POST /api/acord/runs/{run_id}/preview-training-jsonl", "Preview training data for run."),
            ("GET  /api/acord/admin/queue",                  "Admin: review queue listing."),
            ("GET  /api/acord/admin/queue/stats",            "Admin: queue statistics."),
            ("GET  /api/acord/admin/queue/{run_id}/detail",  "Admin: run detail with field breakdown."),
            ("PATCH /api/acord/admin/queue/{run_id}/detail", "Admin: update run detail fields."),
            ("POST /api/acord/admin/{run_id}/review",        "Admin: approve/reject single run."),
            ("POST /api/acord/admin/batch-review",           "Admin: bulk approve/reject runs."),
            ("GET  /api/acord/admin/jobs",                   "Admin: fine-tuning job list."),
            ("GET  /api/acord/admin/jobs/{job_id}",          "Admin: fine-tuning job detail."),
            ("GET  /api/acord/admin/runs/{run_id}/health-card", "Admin: extraction quality health card."),
        ]),
        ("POD Extraction (pods.py)", [
            ("POST /api/pods/extract",          "Extract custom form from PDF."),
            ("POST /api/pods/submit-review",    "Approve/reject POD extraction."),
        ]),
        ("Admin (admin.py)", [
            ("GET  /api/list-users",                                    "List all users (admin only)."),
            ("GET  /api/admin/dashboard-stats",                         "Aggregate dashboard statistics."),
            ("POST /api/admin-create-user",                             "Create user account directly."),
            ("POST /api/admin-set-user-role",                           "Set user role."),
            ("GET  /api/user-creation-requests",                        "List pending user creation requests."),
            ("POST /api/user-creation-requests/{request_id}/approve",   "Approve user creation request."),
            ("POST /api/user-creation-requests/{request_id}/reject",    "Reject user creation request."),
        ]),
        ("Webhooks (webhooks.py)", [
            ("GET   /api/v1/webhooks",                          "List webhooks."),
            ("POST  /api/v1/webhooks",                          "Register webhook (HMAC-signed delivery)."),
            ("POST  /webhooks",                                  "Legacy alias (no /api/v1 prefix)."),
            ("PATCH /api/v1/webhooks/{webhook_id}",             "Update webhook."),
            ("DELETE /api/v1/webhooks/{webhook_id}",            "Delete webhook."),
            ("POST  /api/v1/webhooks/{webhook_id}/rotate-secret","Rotate HMAC signing secret."),
            ("POST  /api/v1/webhooks/test-event",               "Fire test delivery to webhook URL."),
        ]),
        ("Model Registry (model_registry.py)", [
            ("GET /api/v1/model-registry",              "List available models."),
            ("POST /api/v1/model-registry/activate",    "Activate a model for the user."),
            ("POST /api/v1/model-registry/sync-mlflow", "Sync catalog from MLflow."),
        ]),
        ("Notifications (notifications.py)", [
            ("GET /api/v1/notifications",                "List notifications for current user."),
            ("PATCH /api/v1/notifications/{id}/read",    "Mark notification as read."),
        ]),
        ("User Data (user_data.py)", [
            ("GET /api/settings/profile",  "User profile + role."),
            ("PATCH /api/settings/profile","Update profile."),
        ]),
        ("Other", [
            ("GET /health",         "Health check / pod readiness probe."),
            ("GET /api/v1/activity","User activity log."),
        ]),
    ]
    for group_name, routes in route_groups:
        para(doc, group_name, bold=True)
        add_table(doc, ["Endpoint", "Description"], routes, col_widths=[2.8, 3.7])
        doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. FRONTEND PAGES
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "7. Frontend Pages Reference", 1)
    page_rows = [
        ("/auth",             "Auth.tsx",           "Login, password reset, OAuth SSO"),
        ("/signup",           "Signup.tsx",         "3-step signup wizard"),
        ("/",                 "Dashboard.tsx",      "Home overview — pods, activity, stats"),
        ("/playground",       "Playground.tsx",     "Real-time extraction sandbox"),
        ("/training",         "Training.tsx",       "Fine-tuning job monitor"),
        ("/devices",          "Devices.tsx",        "Registered device list (admin only)"),
        ("/devices/:id",      "DeviceDetails.tsx",  "Device info + heartbeat (admin only)"),
        ("/workflows",        "Workflows.tsx",      "Agent workflow definitions"),
        ("/schedules",        "AgentSchedules.tsx", "Scheduled workflow runs"),
        ("/agent-workflows",  "AgentWorkflows.tsx", "Advanced workflow editor"),
        ("/my-models",        "MyModels.tsx",       "User-activated AI models"),
        ("/admin",            "AdminDashboard.tsx", "Global admin panel"),
        ("/users",            "Users.tsx",          "User management (admin only)"),
        ("/pod/:podId",       "PodDashboard.tsx",   "RunPod instance dashboard"),
        ("/settings",         "Settings.tsx",       "User/tenant settings"),
        ("/admin/acord-queue","AdminAcordQueue.tsx","ACORD extraction review queue"),
        ("/admin/model-registry","ModelRegistry.tsx","Model catalog management"),
        ("/activity",         "Activity.tsx",       "User activity log"),
        ("/review-queue",     "ReviewQueue.tsx",    "General review/approval queue"),
        ("/documents",        "Documents.tsx",      "Uploaded document management"),
        ("/mailbox",          "Mailbox.tsx",        "In-app messages"),
    ]
    add_table(doc, ["Route", "Component", "Purpose"], page_rows, col_widths=[1.8, 1.9, 2.8])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. CI/CD
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "8. CI/CD Pipeline Summary", 1)
    para(doc, (
        "Pipelines are in .github/workflows/ and triggered on push to "
        "v*-dev, v*-staging, and v1 branches, or by manual workflow_dispatch."
    ))
    cicd_rows = [
        ("backend-ci.yml",        "Backend",  "Lint (flake8) → Test (pytest + coverage) → SonarQube → Docker build → Push neurapodacr.azurecr.io → Deploy App Service"),
        ("frontend-ci.yml",       "Frontend", "ESLint → Unit tests → SonarQube → Docker build → Push neurapodacr.azurecr.io → Deploy App Service"),
        ("deploy-app-service.yml","Both",     "Reusable deploy: syncs GitHub secrets → App Service env vars, pulls Docker image from ACR, restarts service"),
        ("deploy-env.yml",        "Both",     "Reusable Docker-on-VM deploy via SSH. Pulls image, stops old container, starts new, health checks, auto-rollback on failure."),
        ("db-migrate.yml",        "Database", "Applies supabase/migrations/*.sql via SSH tunnel. Triggered on migration path push or manual dispatch."),
        ("security-audit.yml",    "Both",     "Weekly cron + manual. Runs npm audit (fail on high/critical) and pip-audit on backend deps. Uploads reports as artifacts."),
        ("hotfix-pipeline.yml",   "Both",     "Expedited path targeting v1 (production). Builds and pushes images but requires manual workflow_dispatch to deploy. No staging gate."),
        ("release-train.yml",     "Both",     "Manually triggered. Cuts versioned release branch from v1-staging, tags it, and updates VERSION.md."),
    ]
    add_table(doc, ["Workflow", "Scope", "Steps"], cicd_rows, col_widths=[1.8, 0.8, 4.0])
    doc.add_paragraph()
    para(doc, "Deployment environments:", bold=True)
    bullet(doc, "dev     — auto-deploy on push to v*-dev branch")
    bullet(doc, "staging — auto-deploy on push to v*-staging branch")
    bullet(doc, "production — manual approval gate (GitHub Environments protection rule)")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. ENV VARS
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "9. Environment Variables Reference", 1)

    heading(doc, "9.1 Backend (.env)", 2)
    env_be = [
        ("SUPABASE_URL",                  "https://fideon-staging-apim.azure-api.net/supabase", "Supabase via APIM proxy"),
        ("SUPABASE_ANON_KEY",             "<anon_jwt>",                                 "Public JWT (frontend Supabase JS client)"),
        ("SUPABASE_SERVICE_ROLE_KEY",     "<service_jwt>",                              "Private JWT (bypasses RLS — backend only)"),
        ("PGVECTOR_DATABASE_URL",         "postgresql://...",                           "Direct PG connection for psycopg3 / pgvector"),
        ("DEVICE_JWT_SECRET",             "<high_entropy>",                             "HMAC key for device tokens"),
        ("RUNPOD_API_KEY",                "<key>",                                      "RunPod GPU inference"),
        ("RUNPOD_GENERATE_URL",           "https://<pod>-8000.proxy.runpod.net/generate","RunPod endpoint"),
        ("WEBHOOK_SECRET_ENCRYPTION_KEY", "<base64_32>",                               "Webhook HMAC signing key"),
    ]
    add_table(doc, ["Variable", "Example Value", "Purpose"], env_be, col_widths=[2.2, 2.0, 2.3])
    doc.add_paragraph()

    heading(doc, "9.2 Frontend (.env.local)", 2)
    env_fe = [
        ("NEXT_PUBLIC_SUPABASE_URL",             "https://fideon-staging-apim.azure-api.net/supabase", "Supabase via APIM"),
        ("NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY", "<anon_jwt>",                                  "Supabase JS client auth"),
        ("NEXT_PUBLIC_API_URL",                  "http://localhost:8000",                        "FastAPI backend base URL"),
        ("NEXT_PUBLIC_APP_BASE_URL",             "http://localhost:3000",                        "App base URL for OAuth redirects"),
    ]
    add_table(doc, ["Variable", "Example Value", "Purpose"], env_fe, col_widths=[2.8, 2.0, 1.7])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. SECURITY REVIEW
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "10. Security Review & Open Items", 1)
    para(doc, (
        "The following items were raised in a pre-production security and reliability review. "
        "Critical items must be addressed before the first production release. "
        "Warning items should be resolved before General Availability (GA)."
    ))

    heading(doc, "10.1 Critical — Must Fix Before Production", 2)

    # C1
    p = doc.add_paragraph()
    r = p.add_run("C1 · Service Role Key — No Scope Restriction  [Security · Auth]")
    r.bold = True; r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    para(doc, (
        "SUPABASE_SERVICE_ROLE_KEY bypasses all Row-Level Security policies. "
        "Any code path that obtains this key can read and write every tenant's data. "
        "Currently the key is server-side only and never sent to the frontend, but there is "
        "no restricted Postgres role limiting what tables or operations it can access."
    ))
    para(doc, "Interim mitigations in place:", bold=True)
    bullet(doc, "SUPABASE_SERVICE_ROLE_KEY stored as Azure App Service environment variable — never in source code.")
    bullet(doc, "Key is never returned in API responses, error messages, or client-facing logs.")
    bullet(doc, "All service-key queries on multi-tenant tables must include an explicit tenant_id filter — code-review rule.")
    para(doc, "Full fix required:", bold=True)
    bullet(doc, "Create a restricted Postgres role backend_service with column-level SELECT/INSERT/UPDATE grants.")
    bullet(doc, "Use backend_service credentials for routine reads; reserve SERVICE_ROLE_KEY for migrations only.")
    doc.add_paragraph()

    # C2
    p = doc.add_paragraph()
    r = p.add_run("C2 · PGVECTOR_DATABASE_URL — Bypasses APIM and RLS  [Security · Database]")
    r.bold = True; r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    para(doc, (
        "PGVECTOR_DATABASE_URL is a raw PostgreSQL connection string that bypasses Azure APIM, "
        "Supabase GoTrue, and all RLS policies. Vector queries can read any tenant's embeddings "
        "unless the application explicitly filters by tenant_id."
    ))
    para(doc, "Interim mitigations in place:", bold=True)
    bullet(doc, "Every pgvector query must include WHERE tenant_id = :tenant_id — code-review enforced rule.")
    bullet(doc, "PGVECTOR_DATABASE_URL uses sslmode=require — connection is encrypted in transit.")
    bullet(doc, "The connection string is server-side only and never returned to clients.")
    para(doc, "Full fix required:", bold=True)
    bullet(doc, "Create a vector_service Postgres role restricted to embeddings tables only.")
    bullet(doc, "Wrap pgvector calls in a FastAPI service layer so they flow through APIM + RLS.")
    doc.add_paragraph()

    # C3
    p = doc.add_paragraph()
    r = p.add_run("C3 · Frontend API URL HTTP Fallback  [Security · Frontend]")
    r.bold = True; r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    para(doc, (
        "src/lib/apiBaseUrl.ts falls back to http://127.0.0.1:8080 (plaintext HTTP) "
        "if NEXT_PUBLIC_API_URL is not set. In a misconfigured production deployment, "
        "JWT Bearer tokens would be sent unencrypted."
    ))
    para(doc, "Fix required:", bold=True)
    bullet(doc, "Throw an explicit error when NEXT_PUBLIC_API_URL is absent outside localhost.")
    bullet(doc, "Add NEXT_PUBLIC_API_URL to GitHub Actions secrets for staging and production environments.")
    doc.add_paragraph()

    # Warnings
    heading(doc, "10.2 Warnings — Should Fix Before GA", 2)

    warn_items = [
        (
            "W1 · No ORM — No Schema Migration Safety Net  [Reliability · Backend]",
            (
                "PostgREST HTTP calls are stringly-typed strings. A column rename or type change "
                "in a migration will silently break API calls at runtime with no compile-time error."
            ),
            [
                "Maintain a typed Python dataclass or TypedDict for each table used in "
                "postgrest_get/insert/patch calls. Update these types as part of every migration PR.",
                "Add integration tests against the real Supabase schema (not mocks) to catch "
                "schema/code drift before deployment.",
            ]
        ),
        (
            "W2 · No Retry / Circuit-Breaker for RunPod Cold Starts  [Reliability · LLM/GPU]",
            (
                "There is no documented retry policy, exponential backoff, or circuit-breaker "
                "when RunPod is unavailable. A sustained RunPod outage will surface as 504 "
                "Gateway Timeout to users with no graceful degradation."
            ),
            [
                "Implement exponential backoff with jitter in the LLM service layer (tenacity): "
                "3 retries, base 2s, max 30s.",
                "Add a circuit-breaker: after 5 consecutive RunPod failures, open the circuit "
                "and return a clear 503 with Retry-After header.",
                "Fall through to the litellm fallback chain (Groq → Gemini → OpenAI) when RunPod circuit is open.",
            ]
        ),
        (
            "W3 · Multi-Tenant Isolation Relies Solely on RLS  [Security · Multi-tenancy]",
            (
                "RLS is a good first layer but depends on the JWT being correctly scoped. "
                "If a service-key call omits a tenant_id filter by mistake, it will silently "
                "read all tenants' data. There is no application-layer tenant check."
            ),
            [
                "Add a FastAPI dependency (get_current_tenant) that extracts tenant_id from the "
                "verified JWT and injects it into every route handler.",
                "Routes that use service_headers() must explicitly receive tenant_id as a "
                "parameter — never derive it from user input.",
                "Write a test that signs a JWT for tenant A and asserts zero rows returned from tenant B.",
            ]
        ),
        (
            "W4 · No Documented Secret Rotation Strategy  [Security · Secrets Management]",
            (
                "Several long-lived secrets are referenced: DEVICE_JWT_SECRET, "
                "WEBHOOK_SECRET_ENCRYPTION_KEY, SUPABASE_SERVICE_ROLE_KEY, and multiple LLM API keys. "
                "No rotation cadence or breach-response plan is documented."
            ),
            [
                "Store all secrets in Azure Key Vault; reference them in App Service via Key Vault references.",
                "SUPABASE_SERVICE_ROLE_KEY and DEVICE_JWT_SECRET: rotate every 90 days.",
                "LLM API keys: rotate every 180 days or on provider recommendation.",
            ]
        ),
        (
            "W5 · No Observability / Distributed Tracing  [Observability]",
            (
                "The stack has multiple async hops: Browser → Next.js → FastAPI → APIM → Supabase/RunPod. "
                "There is no trace ID propagation, structured logging correlation, or APM tool."
            ),
            [
                "Add a request-id header (UUID) at the FastAPI entry point via middleware; "
                "propagate it in all outgoing httpx calls.",
                "Use structlog or python-json-logger for structured JSON logs — include "
                "trace_id, tenant_id, user_id, endpoint, duration_ms.",
                "Integrate OpenTelemetry and export traces to Azure Monitor / Application Insights.",
            ]
        ),
        (
            "W6 · CI/CD Branch Strategy — Production Gate  [DevOps · CI/CD]",
            (
                "The pipeline triggers on v*-dev, v*-staging, and v1. The v1 branch appears to be "
                "production but there is no documented branch protection rule or required reviewer "
                "count before merge to v1."
            ),
            [
                "Add branch protection on v1: require 2 approvals, require status checks "
                "(lint, test, SonarQube gate), no force-push.",
                "Add a GitHub Environment named 'production' with required reviewers.",
                "Document the release flow: feature → v*-dev → v*-staging (auto) → v1 (manual PR + approval).",
            ]
        ),
    ]

    for title_text, desc, fixes in warn_items:
        p = doc.add_paragraph()
        r = p.add_run(title_text)
        r.bold = True; r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        para(doc, desc)
        para(doc, "Recommended fixes:", bold=True)
        for fix in fixes:
            bullet(doc, fix, level=1)
        doc.add_paragraph()

    heading(doc, "10.3 Summary", 2)
    summary_rows = [
        ("C1", "Service role key — no scope restriction",   "CRITICAL", "Open", "Interim: server-side only + mandatory tenant_id filter rule"),
        ("C2", "pgvector bypasses APIM + RLS",              "CRITICAL", "Open", "Interim: mandatory tenant_id filter + sslmode=require"),
        ("C3", "Frontend HTTP fallback in production",      "CRITICAL", "Open", "5-line guard in apiBaseUrl.ts required before first staging deploy"),
        ("W1", "No ORM schema safety net",                  "WARNING",  "Open", "Typed Python dicts + integration tests against real schema"),
        ("W2", "No RunPod retry / circuit-breaker",         "WARNING",  "Open", "tenacity retry + circuit breaker, fall to litellm fallback chain"),
        ("W3", "App-layer tenant isolation missing",        "WARNING",  "Open", "get_current_tenant FastAPI dependency on all service-key routes"),
        ("W4", "No secret rotation strategy",               "WARNING",  "Open", "Azure Key Vault references + 90-day rotation cadence"),
        ("W5", "No distributed tracing / APM",              "WARNING",  "Open", "OpenTelemetry + Azure App Insights, request-id propagation"),
        ("W6", "Production branch gate unclear",            "WARNING",  "Open", "GitHub branch protection on v1: 2 approvals + required checks"),
    ]
    add_table(doc,
              ["ID", "Issue", "Severity", "Status", "Action"],
              summary_rows,
              col_widths=[0.35, 2.2, 0.85, 0.6, 2.55])

    return doc


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    out_path = "fideon_architecture_pipelineAndBackend.docx"
    print("Building document...")
    doc = build_doc()
    doc.save(out_path)
    print(f"Saved: {out_path}")
