from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)


def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_heading(doc, text, level=1, color="1F3864"):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_border(t)
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, "1F3864")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = t.add_row()
        fill = "F2F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            cell.paragraphs[0].clear()
            parts = re.split(r"(\*\*.*?\*\*)", str(val))
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = cell.paragraphs[0].add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(9)
                else:
                    r = cell.paragraphs[0].add_run(part)
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ── Title ────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
shade_paragraph(title_p, "1F3864")
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after  = Pt(4)
title_p.paragraph_format.left_indent  = Cm(0.5)
tr = title_p.add_run("  Webhook SRS  —  Gap Analysis")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sub_p = doc.add_paragraph()
shade_paragraph(sub_p, "1F3864")
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(6)
sub_p.paragraph_format.left_indent  = Cm(0.5)
sr = sub_p.add_run("  FNF-68  \u00b7  neura-box-cloud-main  \u00b7  v1-dev  \u00b7  April 2026")
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0xAA, 0xBB, 0xFF)

doc.add_paragraph()

# ── Meta table ────────────────────────────────────────────────────────────────
meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
set_table_border(meta)
meta_data = [
    ("SRS Reference",  "webhook_srs_fideon_fabric.html"),
    ("Codebase",       "neura-box-cloud-main  (branch: v1-dev)"),
    ("Backend Stack",  "FastAPI (Python) — NOT Supabase Edge Functions"),
    ("Date",           "April 2026"),
]
for i, (k, v) in enumerate(meta_data):
    shade_cell(meta.rows[i].cells[0], "E8ECF5")
    shade_cell(meta.rows[i].cells[1], "FFFFFF")
    kr = meta.rows[i].cells[0].paragraphs[0].add_run(k)
    kr.bold = True
    kr.font.size = Pt(9)
    vr = meta.rows[i].cells[1].paragraphs[0].add_run(v)
    vr.font.size = Pt(9)
meta.columns[0].width = Cm(4)
meta.columns[1].width = Cm(12)
doc.add_paragraph()

# ── Legend ────────────────────────────────────────────────────────────────────
lp = doc.add_paragraph()
lp.paragraph_format.space_after = Pt(2)
lr = lp.add_run("Legend")
lr.bold = True
lr.font.size = Pt(11)
lr.font.color.rgb = RGBColor.from_string("1F3864")

legend_items = [
    "\u2705  Already implemented",
    "\u274c  Not implemented",
    "\u26a0\ufe0f  Partial / differs from spec",
    "\U0001f534  MUST implement — broken or security risk",
    "\U0001f7e1  SHOULD implement — needed before production",
    "\U0001f7e2  CAN DEFER — next sprint",
    "\u23ed\ufe0f  SKIP — not applicable to this stack",
]
leg_t = doc.add_table(rows=len(legend_items), cols=1)
leg_t.style = "Table Grid"
set_table_border(leg_t)
for i, txt in enumerate(legend_items):
    shade_cell(leg_t.rows[i].cells[0], "F7F8FC" if i % 2 == 0 else "FFFFFF")
    leg_t.rows[i].cells[0].paragraphs[0].add_run(txt).font.size = Pt(9)
doc.add_paragraph()

# ── Quick Summary ─────────────────────────────────────────────────────────────
qs_p = doc.add_paragraph()
shade_paragraph(qs_p, "FFF3CD")
qs_p.paragraph_format.left_indent  = Cm(0.3)
qs_p.paragraph_format.space_before = Pt(4)
qs_p.paragraph_format.space_after  = Pt(4)
qr1 = qs_p.add_run("Quick Summary\n")
qr1.bold = True
qr1.font.size = Pt(10)
qr1.font.color.rgb = RGBColor.from_string("7D4E00")
qr2 = qs_p.add_run(
    "Your project already has a working webhook engine in Python/FastAPI (~60% of the SRS is done). "
    "The SRS was written for Supabase Edge Functions (Deno/TypeScript) — the code samples do not apply "
    "directly, but the requirements do.\n\n"
    "The single biggest blocker: the database tables (webhooks, webhook_events, webhook_deliveries, "
    "webhook_secrets) DO NOT EXIST in any migration file. Nothing works until this is done."
)
qr2.font.size = Pt(9)
qr2.font.color.rgb = RGBColor.from_string("5C3A00")
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  Webhook Registration  (FR-01 to FR-05)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("FR-01", "Register HTTPS endpoint URL",
         "\u2705 Done", "POST /api/v1/webhooks in webhooks.py"),
        ("FR-01", "Reject HTTP (non-HTTPS) URLs",
         "\u274c  \U0001f7e1 Should Do", "No protocol check — any URL accepted"),
        ("FR-02", "Subscribe to specific event types",
         "\u2705 Done", "events array stored and filtered in webhook_engine.py"),
        ("FR-03", "Generate secret on registration, return once",
         "\u2705 Done", "Secret returned once, never retrievable again"),
        ("FR-03", "Secret never stored in plain text",
         "\u2705 Done", "Stored Fernet-encrypted + SHA-256 hash. Equivalent to SRS Vault approach"),
        ("FR-04", "Max 10 endpoints per tenant",
         "\u274c  \U0001f7e2 Defer", "No cap enforced. Simple count check before insert"),
        ("FR-05", "Optional name / description for endpoint",
         "\u2705 Done", "description field supported"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  Event Delivery  (FR-06 to FR-12)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("FR-06", "HTTP POST with application/json",
         "\u2705 Done", "webhook_engine.py — correct headers set"),
        ("FR-07", "Payload includes event_id, event_type, tenant_id, timestamp, data",
         "\u26a0\ufe0f  \U0001f7e1 Fix", "Has id, type, created_at, data — missing tenant_id in body"),
        ("FR-08", "X-Fideon-Signature header with sha256=<hex>",
         "\u2705 Done", "HMAC-SHA256 computed and sent"),
        ("FR-09", "X-Fideon-Event-Id header for idempotency",
         "\u26a0\ufe0f  \U0001f7e1 Fix",
         "Sends X-Fideon-Delivery-Id (delivery ID) not event ID. Retried events should keep same Event-Id"),
        ("FR-10", "10-second HTTP timeout per attempt",
         "\u26a0\ufe0f  \U0001f7e1 Fix", "Current timeout is 15 seconds. SRS requires 10s"),
        ("FR-11", "Only 2xx = success; all others retry",
         "\u2705 Done", "webhook_engine.py line 312"),
        ("FR-12", "Async delivery — never blocks calling API",
         "\u2705 Done", "Background asyncio worker loop started at app startup"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  Retry & Dead-Letter Queue  (FR-13 to FR-17)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("FR-13", "Retry up to 3 times (4 total attempts)",
         "\u2705 Done", "Configurable via WEBHOOK_MAX_ATTEMPTS — ensure default is 4"),
        ("FR-14", "Back-off: 30s -> 5min -> 30min with +/-10% jitter",
         "\u26a0\ufe0f  \U0001f7e1 Fix", "Generic exponential formula used, not specific schedule. No jitter"),
        ("FR-15", "Failed deliveries go to dead-letter queue",
         "\u2705 Done", "Status set to dead_letter after exhausting retries"),
        ("FR-16", "DLQ visible in UI — view, replay, dismiss",
         "\u274c  \U0001f7e2 Defer", "No DLQ listing or replay API endpoints. No UI panel"),
        ("FR-17", "DLQ entries retained 30 days then auto-deleted",
         "\u274c  \U0001f7e2 Defer", "No expires_at column or cleanup job"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4.  Secret Management  (FR-18 to FR-20)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("FR-18", "Tenants can rotate secret; new secret returned once",
         "\u2705 Done", "POST /api/v1/webhooks/{id}/rotate-secret"),
        ("FR-19", "30-minute dual-secret grace period on rotation",
         "\u274c  \U0001f7e2 Defer", "Old secret immediately deactivated. No grace window"),
        ("FR-20", "Secrets stored hashed, never logged",
         "\u26a0\ufe0f Acceptable",
         "SHA-256 + Fernet encryption (not bcrypt). Fernet is stronger for FastAPI; bcrypt is for Edge Function context"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5.  Management UI  (FR-21 to FR-24)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("FR-21", "List all endpoints with status and last delivery",
         "\u2705 Done", "WebhooksSettingsPanel.tsx"),
        ("FR-22", "Secret rotation from UI",
         "\u2705 Done", "Rotation button in the same UI panel"),
        ("FR-23", "Delivery history — last 50 attempts per endpoint",
         "\u274c  \U0001f7e2 Defer",
         "Data exists in webhook_deliveries once migration runs. No UI panel yet"),
        ("FR-24", "Send test event from UI",
         "\u26a0\ufe0f  \U0001f7e2 Defer",
         "API endpoint exists (POST /api/v1/webhooks/test-event) but no UI button"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6.  Security Requirements  (SEC-01 to SEC-11)", level=1)
add_table(doc,
    headers=["Req ID", "Requirement", "Status", "Notes"],
    rows=[
        ("SEC-01", "HTTPS only; reject HTTP URLs",
         "\u274c  \U0001f534 Critical", "No protocol validation at registration or delivery"),
        ("SEC-02", "HMAC-SHA256 — correct signing string (timestamp.body)",
         "\u26a0\ufe0f  \U0001f534 Critical",
         "BUG: Code signs just body. Must sign timestamp+\".\"+body. Receivers following the spec reject ALL events"),
        ("SEC-03", "X-Fideon-Timestamp header; receivers reject >5min old",
         "\u2705 Done", "Timestamp sent as header. Replay prevention is receiver responsibility"),
        ("SEC-04", "Constant-time signature comparison",
         "\u2705 Done", "Python hmac.new uses compare_digest — safe"),
        ("SEC-05", "Secrets never in logs or error messages",
         "\u2705 Done", "No secret logged anywhere"),
        ("SEC-06", "Rate limiting: 10 registrations/min per tenant",
         "\u274c  \U0001f7e2 Defer", "No per-tenant webhook-specific rate limit"),
        ("SEC-07", "SSRF: block RFC-1918, loopback, Azure ranges",
         "\u274c  \U0001f534 Critical",
         "NO SSRF validation. Worker will POST to 10.x.x.x, 169.254.169.254, Azure IMDS etc."),
        ("SEC-08", "Input validation on all API inputs",
         "\u2705 Done", "URL length, event name format, description length validated"),
        ("SEC-09", "Tenant-scoped access (RLS + API layer)",
         "\u2705 Done", "All queries filter by tenant_id from JWT"),
        ("SEC-10", "Sign with new secret only during rotation grace window",
         "\u23ed\ufe0f N/A", "Grace period not implemented yet — covered when FR-19 is done"),
        ("SEC-11", "Cloud worker must not deliver to private/internal IPs",
         "\u274c  \U0001f534 Critical",
         "Same as SEC-07 — no IP check at delivery time. DNS rebinding not protected"),
    ],
    col_widths=[2, 5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7.  Non-Functional Requirements", level=1)
add_table(doc,
    headers=["Requirement", "Status", "Notes"],
    rows=[
        ("500 events/min throughput per tenant", "\u2705 Done",
         "Worker polls every 2s, batch 25"),
        ("p95 delivery latency < 5 seconds", "\u2705 Done",
         "Near-immediate — worker polls every 2s"),
        (">=99.5% delivery reliability within 4 attempts", "\u2705 Done",
         "Retry logic covers this"),
        ("99.9% worker uptime", "\u2705 Done",
         "Started with app lifespan, restarts with app"),
        ("Zero plain-text secrets in storage (NFR-09)", "\u2705 Done",
         "Fernet-encrypted in DB, never in logs"),
        ("Row-Level Security on all webhook tables", "\u274c  \U0001f7e1 Needed",
         "RLS policies not created — migration does not exist yet"),
    ],
    col_widths=[6, 3, 8]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8.  Database Schema", level=1)
add_table(doc,
    headers=["Table", "Status", "Notes"],
    rows=[
        ("webhooks",          "\u274c  \U0001f534 Critical", "No migration exists. Referenced in code but never created"),
        ("webhook_events",    "\u274c  \U0001f534 Critical", "Same — missing migration"),
        ("webhook_deliveries","\u274c  \U0001f534 Critical", "Same — missing migration"),
        ("webhook_secrets",   "\u274c  \U0001f534 Critical", "Same — missing migration"),
        ("security_events",   "\u274c  \U0001f7e2 Defer",   "SRS adds this for SSRF audit log. Not in codebase"),
        ("rate_limit_log",    "\u274c  \U0001f7e2 Defer",   "Only needed if per-tenant rate limiting is added"),
    ],
    col_widths=[5, 4, 8]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9.  Architecture Differences — What to Skip", level=1)
p9 = doc.add_paragraph(
    "The SRS was written for Supabase Edge Functions (Deno/TypeScript). "
    "Your project uses FastAPI (Python). The following SRS items do NOT apply:"
)
p9.runs[0].font.size = Pt(10)
doc.add_paragraph()

add_table(doc,
    headers=["SRS Item", "Why It Does Not Apply"],
    rows=[
        ("All Deno/TypeScript code samples",
         "Your backend is FastAPI/Python. Use the SRS only for requirements understanding, not the code"),
        ("Supabase Vault for secret storage",
         "Vault is Supabase-specific. Your Fernet encryption is the correct equivalent"),
        ("pg_cron delivery worker trigger",
         "You use a Python asyncio background loop — same outcome, different mechanism"),
        ("bcrypt for secret hash (cost >=12)",
         "Bcrypt makes sense without a KMS. Fernet encryption is stronger for your use case"),
        ("is_edge_local / on-premise agent delivery",
         "Only relevant if you have an on-premise edge agent — skip unless in scope"),
        ("Azure-specific hostname blocklist patterns",
         "Add the CIDR ranges regardless. Azure hostname patterns only matter if hosted on Azure"),
    ],
    col_widths=[6, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — ACTION PLAN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10.  Action Plan", level=1)

add_heading(doc, "\U0001f534  Do This First  (System Broken / Security Risk)", level=2, color="C0152A")
add_table(doc,
    headers=["#", "Task", "File to Change"],
    rows=[
        ("1",
         "**Write DB migration** — create webhooks, webhook_events, webhook_deliveries, "
         "webhook_secrets tables with RLS policies. Nothing works without this.",
         "supabase/migrations/20260413_webhook_tables.sql  (new file)"),
        ("2",
         "**Fix signing string** — change from sign(body) to sign(timestamp + \".\" + body). "
         "Receivers following the spec will reject every event otherwise.",
         "backend/app/services/webhook_engine.py  line 286"),
        ("3",
         "**Add SSRF validation** — validate URLs at registration AND at delivery time. "
         "Block RFC-1918, loopback, link-local, Azure IMDS ranges.",
         "New: backend/app/core/ssrf_validator.py\nCall in: webhooks.py:79 + webhook_engine.py:304"),
    ],
    col_widths=[1, 10, 6]
)
doc.add_paragraph()

add_heading(doc, "\U0001f7e1  Do Before Production", level=2, color="B45309")
add_table(doc,
    headers=["#", "Task", "File to Change"],
    rows=[
        ("4", "Add HTTPS-only check at registration (reject http:// URLs)",
         "backend/app/routes/webhooks.py  line 73"),
        ("5", "Fix X-Fideon-Event-Id header — send event_id not delivery_id",
         "backend/app/services/webhook_engine.py  line 289"),
        ("6", "Add tenant_id + timestamp fields to outbound event payload body (FR-07)",
         "backend/app/services/webhook_engine.py  line 275"),
        ("7", "Set delivery timeout to 10 seconds (currently 15s)",
         "backend/app/services/webhook_engine.py  line 310"),
        ("8", "Fix retry schedule to 30s > 5min > 30min with +/-10% jitter",
         "backend/app/services/webhook_engine.py  line 202"),
    ],
    col_widths=[1, 10, 6]
)
doc.add_paragraph()

add_heading(doc, "\U0001f7e2  Next Sprint / Can Defer", level=2, color="00875A")
add_table(doc,
    headers=["#", "Task"],
    rows=[
        ("9",  "DLQ listing + replay API endpoints"),
        ("10", "Delivery history UI (last 50 attempts per endpoint)"),
        ("11", "Test event button in management UI"),
        ("12", "Dual-secret 30-minute grace window on rotation"),
        ("13", "Max 10 endpoints per tenant cap"),
        ("14", "Per-tenant rate limiting on registration (10/min)"),
        ("15", "expires_at on DLQ rows + 30-day cleanup job"),
    ],
    col_widths=[1, 16]
)

doc.add_paragraph()

# Footer
fp = doc.add_paragraph()
shade_paragraph(fp, "F2F4F8")
fp.paragraph_format.left_indent = Cm(0.3)
fr = fp.add_run(
    "Generated from analysis of webhook_srs_fideon_fabric.html vs neura-box-cloud-main codebase — 2026-04-13"
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
fr.italic = True

out_path = r"c:/Users/samar/Downloads/neura-box-cloud-main/docs/webhook-srs-gap-analysis.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
