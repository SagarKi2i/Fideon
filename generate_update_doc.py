"""
generate_update_doc.py
Generates the update DOCX for sir — covers everything done in the Quantization folder.
"""

import io, os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

C = {
    "bg":      "#080C14",
    "purple":  "#AFA9EC", "purple_d": "#3C3489",
    "teal":    "#5DCAA5", "teal_d":   "#085041",
    "blue":    "#85B7EB", "blue_d":   "#0C447C",
    "amber":   "#EF9F27", "amber_d":  "#633806",
    "green":   "#4ADE80", "red":      "#EF4444",
    "ink":     "#EDF2FF", "ink2":     "#8DA4C4",
}

# ─────────────────────────────────────────────
# DIAGRAM — Storage architecture (sir's spec)
# ─────────────────────────────────────────────
def make_storage_arch_diagram():
    fig, ax = plt.subplots(figsize=(13, 5.5))
    fig.patch.set_facecolor(C["bg"])
    ax.set_facecolor(C["bg"])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")

    ax.text(0.5, 0.95, "Storage Architecture — Sir's Specification",
            ha="center", va="top", fontsize=11, fontweight="bold",
            color=C["ink"], fontfamily="DejaVu Sans")
    ax.text(0.5, 0.875, "Our code always targets SeaweedFS only · SeaweedFS handles cloud backend internally",
            ha="center", va="top", fontsize=8, color=C["ink2"], fontfamily="DejaVu Sans")

    def rbox(x, y, w, h, fc, ec, lines, fsizes=None):
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.012",
                             linewidth=1.3, edgecolor=ec, facecolor=fc, alpha=0.85, zorder=3)
        ax.add_patch(box)
        if isinstance(lines, str): lines = [lines]
        if fsizes is None: fsizes = [8] * len(lines)
        total = len(lines)
        for i, (line, fs) in enumerate(zip(lines, fsizes)):
            offset = (i - (total-1)/2) * 0.042
            ax.text(x+w/2, y+h/2 - offset, line, ha="center", va="center",
                    fontsize=fs, color=ec, fontfamily="DejaVu Sans",
                    fontweight="bold" if i == 0 else "normal", zorder=4)

    def arr(x1, y1, x2, y2, col="#475569", label=None):
        ax.annotate("", xy=(x2,y2), xytext=(x1,y1),
                    arrowprops=dict(arrowstyle="-|>", color=col, lw=1.5, mutation_scale=11), zorder=5)
        if label:
            ax.text((x1+x2)/2+0.01, (y1+y2)/2+0.025, label,
                    fontsize=6.5, color=C["ink2"], ha="center", fontfamily="DejaVu Sans", zorder=6)

    # Our code box
    rbox(0.02, 0.42, 0.22, 0.20, C["purple_d"], C["purple"],
         ["upload.py", "boto3 S3-compat", "single upload call"],
         [9, 7.5, 7])

    arr(0.24, 0.52, 0.34, 0.52, C["teal"], "S3-compat API\nboto3")

    # SeaweedFS box
    rbox(0.34, 0.35, 0.22, 0.35, C["teal_d"], C["teal"],
         ["SeaweedFS", "S3-compat layer", "presigned URLs", "SHA-256 verify"],
         [9.5, 7.5, 7, 7])

    # Three arrows down to cloud backends
    arr(0.40, 0.35, 0.18, 0.18, C["blue"],   "Azure")
    arr(0.45, 0.35, 0.45, 0.18, C["amber"],  "AWS")
    arr(0.50, 0.35, 0.72, 0.18, C["purple"], "GCP")

    # Cloud backend boxes
    rbox(0.04, 0.04, 0.22, 0.13, C["blue_d"], C["blue"],
         ["Azure Blob", "configured by infra"], [9, 7])
    rbox(0.34, 0.04, 0.22, 0.13, C["amber_d"], C["amber"],
         ["AWS S3", "configured by infra"], [9, 7])
    rbox(0.63, 0.04, 0.22, 0.13, C["purple_d"], C["purple"],
         ["GCP Workspace", "configured by infra"], [9, 7])

    # Supabase registry
    rbox(0.64, 0.42, 0.22, 0.20, C["amber_d"], C["amber"],
         ["Supabase Registry", "blob_url (SeaweedFS)", "sha256 · version",
          "docker_sha · rollback_ptr"],
         [9, 7, 7, 7])
    arr(0.56, 0.52, 0.64, 0.52, C["amber"], "registry insert\nafter upload")

    # Electron
    rbox(0.78, 0.62, 0.20, 0.14, C["teal_d"], C["teal"],
         ["Electron App", "downloads via presigned URL"], [9, 7])
    arr(0.75, 0.525, 0.88, 0.62, C["teal"], "presigned URL")

    # label: our code boundary
    bbox = FancyBboxPatch((0.01, 0.40), 0.24, 0.25,
                          boxstyle="round,pad=0.008", linewidth=1,
                          edgecolor=C["purple"], facecolor="none",
                          linestyle="--", alpha=0.5, zorder=2)
    ax.add_patch(bbox)
    ax.text(0.13, 0.665, "Our code boundary", ha="center", fontsize=6.5,
            color=C["purple"], fontfamily="DejaVu Sans")

    bbox2 = FancyBboxPatch((0.03, 0.01), 0.84, 0.32,
                           boxstyle="round,pad=0.008", linewidth=1,
                           edgecolor=C["ink2"], facecolor="none",
                           linestyle="--", alpha=0.3, zorder=2)
    ax.add_patch(bbox2)
    ax.text(0.45, 0.325, "Infra team's responsibility — SeaweedFS backend config",
            ha="center", fontsize=6.5, color=C["ink2"], fontfamily="DejaVu Sans")

    plt.tight_layout(pad=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C["bg"], edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# DIAGRAM — Quantization folder file map
# ─────────────────────────────────────────────
def make_file_map_diagram():
    fig, ax = plt.subplots(figsize=(13, 4.5))
    fig.patch.set_facecolor(C["bg"])
    ax.set_facecolor(C["bg"])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")

    ax.text(0.5, 0.96, "Quantization/ — File Map & Data Flow",
            ha="center", va="top", fontsize=11, fontweight="bold",
            color=C["ink"], fontfamily="DejaVu Sans")

    files = [
        (0.02, "config.env.example", C["blue"],   "Environment config template\nSeaweedFS · Supabase · HF_TOKEN"),
        (0.22, "setup.sh",           C["purple"],  "Run once on Vast.ai node\nBuilds llama.cpp with CUDA"),
        (0.42, "quantize.py",        C["teal"],    "merge_and_unload() → FP16 GGUF\nQ5_K_M · Q4_K_M · SHA-256 · manifest"),
        (0.62, "upload.py",          C["amber"],   "boto3 → SeaweedFS S3-compat\nSHA-256 verify · Supabase registry"),
        (0.82, "run_pipeline.sh",    C["green"],   "Master script\nquantize.py → upload.py"),
    ]

    for x, name, color, desc in files:
        box = FancyBboxPatch((x, 0.42), 0.18, 0.38,
                             boxstyle="round,pad=0.01", linewidth=1.3,
                             edgecolor=color, facecolor=color.replace("EC","20").replace("A5","20"),
                             alpha=0.25, zorder=2)
        ax.add_patch(box)
        # use a very dark semi-transparent fill
        box2 = FancyBboxPatch((x, 0.42), 0.18, 0.38,
                              boxstyle="round,pad=0.01", linewidth=1.3,
                              edgecolor=color,
                              facecolor="#0D1422", alpha=0.88, zorder=2)
        ax.add_patch(box2)
        ax.text(x+0.09, 0.73, name, ha="center", va="center",
                fontsize=7.5, color=color, fontweight="bold",
                fontfamily="DejaVu Sans", zorder=4)
        for i, line in enumerate(desc.split("\n")):
            ax.text(x+0.09, 0.60 - i*0.10, line, ha="center", va="center",
                    fontsize=6.3, color=C["ink2"], fontfamily="DejaVu Sans", zorder=4)

    # arrows between quantize and upload
    ax.annotate("", xy=(0.62, 0.615), xytext=(0.60, 0.615),
                arrowprops=dict(arrowstyle="-|>", color=C["teal"], lw=1.3, mutation_scale=9), zorder=5)
    ax.text(0.61, 0.655, "manifest.json\n+ *.gguf", ha="center", fontsize=6,
            color=C["teal"], fontfamily="DejaVu Sans", zorder=6)

    # flow label bottom
    ax.text(0.5, 0.30,
            "Vast.ai node: setup.sh  →  run_pipeline.sh  →  quantize.py  →  upload.py  →  node destroyed",
            ha="center", fontsize=8, color=C["ink2"], fontfamily="DejaVu Sans")
    ax.text(0.5, 0.18,
            "Output artifacts: model-q5km.gguf · model-q4km.gguf · model-fp16.gguf · manifest.json · *.sig",
            ha="center", fontsize=7.5, color=C["teal"], fontfamily="DejaVu Sans")
    ax.text(0.5, 0.07,
            "All artifacts stored in SeaweedFS → Azure Blob (internally) · registry row in Supabase",
            ha="center", fontsize=7.5, color=C["amber"], fontfamily="DejaVu Sans")

    plt.tight_layout(pad=0)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C["bg"], edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex="2B2B6B"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_para(doc, text, bold=False, italic=False, size=10, color_hex=None,
             space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_code(doc, text, size=8.5, color_hex="2DD4BF"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot); pPr.append(pBdr)

def make_table(doc, rows, header_color, col_widths=None, header_text_color="FFFFFF"):
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths is None:
        col_widths = [Inches(6.6 / ncols)] * ncols
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, (text, width) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[j]
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(header_text_color)
                set_cell_bg(cell, header_color)
            else:
                run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────
# BUILD DOCX
# ─────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.4)
        section.right_margin  = Cm(2.4)

    # ══════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Quantization Pipeline — Implementation Update")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x2B, 0x2B, 0x6B)

    add_para(doc, "What was built · Architecture decisions · Next steps",
             bold=True, size=12, color_hex="0C447C",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_para(doc, "Neura-Box Cloud  ·  April 2026  ·  Prepared for review",
             size=10, color_hex="8DA4C4",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_para(doc,
             "This document covers the Quantization/ folder created in the project — "
             "all files, what each does, key architecture decisions made (especially around "
             "SeaweedFS sitting on Azure Blob / AWS S3 / GCP), the corrected storage design, "
             "and the remaining steps to run the pipeline end-to-end.",
             size=10, color_hex="4C6280",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    add_divider(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1 — WHAT WAS BUILT
    # ══════════════════════════════════════════
    add_heading(doc, "1. What Was Built — Quantization/ Folder", level=1, color_hex="2B2B6B")
    add_para(doc,
             "A self-contained Quantization/ folder has been created in the project root. "
             "It is designed to be transferred to the Vast.ai node and run there. "
             "It covers the full pipeline from merging the LoRA adapter through to uploading "
             "signed GGUF files to SeaweedFS and updating the Supabase model registry.",
             size=10, space_after=10)

    # file map diagram
    fmap_buf = make_file_map_diagram()
    doc.add_picture(fmap_buf, width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 1 — Files in Quantization/ and their data flow",
             size=8, color_hex="8DA4C4", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=14)

    make_table(doc, [
        ("File",                    "Purpose",                                              "Run where"),
        ("config.env.example",      "Environment config template — copy to config.env",     "Local machine"),
        ("requirements.txt",        "Python dependencies — boto3, transformers, peft, etc.", "Vast.ai node"),
        ("setup.sh",                "One-time node setup: installs llama.cpp with CUDA",    "Vast.ai node"),
        ("quantize.py",             "Merge adapter → FP16 GGUF → Q5_K_M → Q4_K_M → sign", "Vast.ai node"),
        ("upload.py",               "Upload to SeaweedFS · verify SHA-256 · update Supabase","Vast.ai node"),
        ("run_pipeline.sh",         "Master script — runs quantize.py then upload.py",      "Vast.ai node"),
        ("STEPS.md",                "Step-by-step operator guide",                          "Reference"),
    ], header_color="2B2B6B",
       col_widths=[Inches(2.0), Inches(3.2), Inches(1.4)])

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 2 — QUANTIZE.PY
    # ══════════════════════════════════════════
    add_heading(doc, "2. quantize.py — What It Does", level=1, color_hex="085041")
    add_para(doc,
             "This is the core quantization script. It runs entirely on the Vast.ai GPU node "
             "and produces the GGUF artifacts that are later uploaded to SeaweedFS.",
             size=10, space_after=8)

    make_table(doc, [
        ("Step", "Action",                          "Detail"),
        ("1",    "Load base model in BF16",         "Qwen/Qwen2.5-14B-Instruct · ~28 GB into VRAM · device_map=auto"),
        ("2",    "Load LoRA adapter",               "PeftModel.from_pretrained(model, adapter_path)"),
        ("3",    "merge_and_unload()",              "W' = W + (α/r)·B·A · single merged model · no PEFT needed after this"),
        ("4",    "Save merged BF16 to disk",        "/workspace/merged_bf16/ · tokenizer saved alongside"),
        ("5",    "Free VRAM",                       "del model · gc.collect() · torch.cuda.empty_cache()"),
        ("6",    "convert_hf_to_gguf.py --outtype f16", "FP16 GGUF (~28 GB) · CI reference baseline"),
        ("7",    "llama-quantize Q5_K_M",           "~10.7 GB · production use · attn 6-bit · FFN 5-bit"),
        ("8",    "llama-quantize Q4_K_M",           "~8.0 GB · edge / air-gap use · attn 5-bit · FFN 4-bit"),
        ("9",    "SHA-256 per file",                "Computed locally after each quantize step"),
        ("10",   "GPG sign (optional)",             "gpg --detach-sign · *.sig file per GGUF"),
        ("11",   "Write manifest.json",             "schema_version · artifacts[] · sha256 · size_bytes · purpose"),
    ], header_color="085041",
       col_widths=[Inches(0.4), Inches(2.4), Inches(3.8)])

    add_para(doc, "GPU used: H200 / A100 SXM4. CUDA 12.1+. llama.cpp built with -DGGML_CUDA=ON.",
             size=9, color_hex="4C6280", space_before=4)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 3 — STORAGE ARCHITECTURE (KEY DECISION)
    # ══════════════════════════════════════════
    add_heading(doc, "3. Storage Architecture — Key Design Decision", level=1, color_hex="0C447C")
    add_para(doc,
             "This is the most important architectural decision confirmed during implementation. "
             "It directly follows sir's specification.",
             size=10, space_after=10)

    arch_buf = make_storage_arch_diagram()
    doc.add_picture(arch_buf, width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 2 — SeaweedFS sits on top of Azure Blob / AWS S3 / GCP Workspace",
             size=8, color_hex="8DA4C4", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=14)

    add_heading(doc, "3.1  The Rule", level=2, color_hex="0C447C")
    add_para(doc,
             "Our code (upload.py) uploads ONCE to SeaweedFS via its S3-compatible API using boto3. "
             "SeaweedFS internally persists data to the cloud backend — which could be Azure Blob, "
             "AWS S3, or GCP Workspace depending on how the infra team has configured SeaweedFS on "
             "the server. Our pipeline code never calls any cloud SDK directly.",
             size=10, space_after=8)

    make_table(doc, [
        ("What our code does",          "What SeaweedFS does internally (infra team)"),
        ("boto3 S3-compat PutObject",   "Persists to Azure Blob (if on Azure)"),
        ("boto3 generate_presigned_url","Persists to AWS S3 (if on AWS)"),
        ("Single endpoint in config.env","Persists to GCP Workspace (if on GCP)"),
        ("No Azure/AWS/GCP SDK needed", "SeaweedFS config file sets the backend"),
    ], header_color="0C447C",
       col_widths=[Inches(3.3), Inches(3.3)])

    add_heading(doc, "3.2  What Was Corrected", level=2, color_hex="0C447C")
    add_para(doc,
             "An earlier version of upload.py incorrectly called both Azure Blob SDK and SeaweedFS "
             "as separate upload targets. This was corrected after clarification:",
             size=10, space_after=6)

    make_table(doc, [
        ("",          "Before (wrong)",                           "After (correct)"),
        ("Upload 1",  "SeaweedFS — S3-compat boto3",             "SeaweedFS — S3-compat boto3"),
        ("Upload 2",  "Azure Blob — azure-storage-blob SDK",     "Nothing — removed entirely"),
        ("SDK deps",  "boto3 + azure-storage-blob",              "boto3 only"),
        ("config.env","AZURE_STORAGE_CONNECTION_STRING needed",  "Not needed — infra concern"),
        ("blob_url",  "Azure presigned URL",                     "SeaweedFS presigned URL"),
    ], header_color="0C447C",
       col_widths=[Inches(1.4), Inches(2.6), Inches(2.6)])

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 4 — UPLOAD.PY
    # ══════════════════════════════════════════
    add_heading(doc, "4. upload.py — What It Does", level=1, color_hex="633806")
    add_para(doc,
             "Runs after quantize.py. Reads manifest.json, verifies local SHA-256 for each file, "
             "uploads to SeaweedFS, verifies SHA-256 on SeaweedFS, then inserts a row into the "
             "Supabase model_registry table.",
             size=10, space_after=8)

    make_table(doc, [
        ("Step", "Action",                       "Detail"),
        ("1",    "Read manifest.json",           "Parses artifact list: filename · sha256 · size_bytes · purpose"),
        ("2",    "Local SHA-256 verify",         "Computed before upload — mismatch → exit(1), do not upload"),
        ("3",    "Upload to SeaweedFS",          "boto3 S3 PutObject · endpoint from SEAWEEDFS_ENDPOINT env var"),
        ("4",    "Remote SHA-256 verify",        "get_object → hash body → compare with manifest · mismatch → delete + error"),
        ("5",    "Generate presigned URL",       "s3.generate_presigned_url · 24hr expiry · this is blob_url in registry"),
        ("6",    "Upload *.sig files",           "GPG signatures uploaded alongside each GGUF"),
        ("7",    "Upload manifest.json",         "Uploaded to models/{version}/manifest/ path"),
        ("8",    "Supabase registry INSERT",     "model_version · sha256 · blob_url · docker_sha · rollback_ptr"),
    ], header_color="633806",
       col_widths=[Inches(0.4), Inches(2.0), Inches(4.2)])

    add_heading(doc, "4.1  Supabase model_registry Fields (Sir's Spec)", level=2, color_hex="633806")
    make_table(doc, [
        ("Field",         "Value",                                "Status"),
        ("model_version", "e.g. v1.0",                           "Set from MODEL_VERSION env var"),
        ("sha256",        "SHA-256 of the GGUF file",            "Computed and verified twice"),
        ("blob_url",      "SeaweedFS presigned URL (24hr)",      "Generated after upload"),
        ("docker_sha",    "null",                                 "Set later when Docker image is built"),
        ("rollback_ptr",  "null",                                 "Set by backend when promoting to active"),
        ("quant_level",   "q5_k_m · q4_k_m · fp16",             "From artifact in manifest.json"),
        ("is_available",  "true",                                 "Set on insert"),
    ], header_color="633806",
       col_widths=[Inches(1.6), Inches(3.0), Inches(2.0)])

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 5 — SUPABASE MIGRATION NEEDED
    # ══════════════════════════════════════════
    add_heading(doc, "5. Supabase Migration Required", level=1, color_hex="2B2B6B")
    add_para(doc,
             "Before the pipeline can run, the following columns must be added to the "
             "existing model_registry table in Supabase:",
             size=10, space_after=6)

    add_code(doc, "-- Run in Supabase SQL Editor")
    add_code(doc, "ALTER TABLE public.model_registry")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS blob_url      text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS sha256        text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS quant_level   text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS docker_sha    text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS rollback_ptr  text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS filename      text,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS size_bytes    bigint,")
    add_code(doc, "  ADD COLUMN IF NOT EXISTS is_available  boolean default true;")
    doc.add_paragraph()

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 6 — VAST.AI SETUP
    # ══════════════════════════════════════════
    add_heading(doc, "6. Vast.ai Node — Recommended Spec", level=1, color_hex="2B2B6B")
    make_table(doc, [
        ("Parameter",    "Value",                        "Reason"),
        ("GPU",          "A100 SXM4 40GB ($0.697/hr)",  "Covers training (16GB) + GPU quantization (32GB)"),
        ("Template",     "PyTorch (Vast)",               "PyTorch + CUDA pre-installed"),
        ("Container disk","200 GB",                      "Base model (28GB) + merged (28GB) + GGUFs + packages"),
        ("System RAM",   "≥ 64 GB",                      "CPU fallback for quantization if needed"),
        ("CUDA",         "12.1 or 12.6",                 "Required by Unsloth + llama.cpp GGML CUDA"),
        ("SSH",          "Enable",                       "Required to run setup.sh and transfer files"),
        ("Cost estimate","~$0.697/hr × ~0.5 hr = ~$0.35","Per quantization cycle after training"),
    ], header_color="2B2B6B",
       col_widths=[Inches(1.6), Inches(2.8), Inches(2.2)])

    add_para(doc,
             "Note: H200 (141GB VRAM, $2.37/hr) also available and confirmed compatible. "
             "A100 SXM4 40GB is recommended for cost efficiency — sufficient for the full pipeline.",
             size=9, color_hex="4C6280", space_before=4)

    add_divider(doc)

    # ══════════════════════════════════════════
    # SECTION 7 — NEXT STEPS
    # ══════════════════════════════════════════
    add_heading(doc, "7. Next Steps", level=1, color_hex="2B2B6B")

    steps = [
        ("SeaweedFS deployment",
         "Confirm SeaweedFS is deployed and accessible. "
         "Provide SEAWEEDFS_ENDPOINT, SEAWEEDFS_ACCESS_KEY, SEAWEEDFS_SECRET_KEY "
         "for the config.env file. This is the blocker before the pipeline can run."),
        ("Supabase migration",
         "Run the ALTER TABLE migration in Section 5 against the production Supabase project."),
        ("Fill config.env",
         "Copy config.env.example → config.env. Fill in SEAWEEDFS_*, SUPABASE_*, "
         "HF_TOKEN, MODEL_VERSION."),
        ("Confirm adapter availability",
         "The fine-tuned LoRA adapter (adapter_model.safetensors + adapter_config.json + tokenizer files) "
         "must be available to transfer to the Vast.ai node."),
        ("Rent Vast.ai node + run pipeline",
         "Rent A100 SXM4 40GB with 200GB disk and PyTorch (Vast) template. "
         "Transfer Quantization/ folder + adapter. Run setup.sh then run_pipeline.sh."),
        ("Verify & destroy node",
         "Confirm GGUF files appear in SeaweedFS and model_registry row is inserted in Supabase. "
         "Destroy the node immediately to stop billing."),
        ("Electron model-updater.ts",
         "Separate task: implement the download handler in Electron that reads blob_url from "
         "Supabase, streams the GGUF, verifies sha256, and calls ollama create."),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f"  {i}.  {title}")
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
        add_para(doc, f"     {desc}", size=9.5, color_hex="4C6280",
                 space_before=0, space_after=5)

    add_divider(doc)

    add_para(doc,
             "Key constraint: SeaweedFS must be running and accessible before any upload can happen. "
             "All other steps (quantize.py, setup.sh, Supabase migration) can be prepared in parallel. "
             "GPU billing on Vast.ai is approximately $0.35 per full quantization cycle on A100 SXM4.",
             size=9, color_hex="8DA4C4", space_before=8)

    out_path = os.path.join(OUT_DIR, "Quantization_Pipeline_Update.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_docx()
